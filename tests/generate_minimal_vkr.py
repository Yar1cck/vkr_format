"""Генератор минимального валидного .docx для загрузки в VKR.Format.

Документ содержит ровно то, что требует пайплайн, чтобы пройти без
критических структурных нарушений:

  • Титульный лист с маркером «ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА»
    (триггер detect_title_page_end → весь блок становится неприкосновенным).
  • СОДЕРЖАНИЕ, ВВЕДЕНИЕ, одна нумерованная глава, ЗАКЛЮЧЕНИЕ,
    СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ — обязательные разделы (structure.py
    REQUIRED) в правильном порядке, оформленные стилем Heading 1.
  • Одна ссылка [1] во введении + одна позиция в списке источников —
    citations/bibliography сходятся.
  • Базовое форматирование тела по rules_v1.yaml: Times New Roman 14,
    интервал 1.5, поля 30/15/20/20 мм, абзацный отступ 1.25 см.

Запуск:
    python -m tests.generate_minimal_vkr            # → minimal_vkr.docx в корне
    python -m tests.generate_minimal_vkr out.docx
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

FONT = "Times New Roman"


def _set_font(run, size: float = 14, bold: bool = False) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)


def _center(p) -> None:
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)


def _title_line(doc, text: str, size: float = 14, bold: bool = False) -> None:
    p = doc.add_paragraph()
    _center(p)
    _set_font(p.add_run(text), size=size, bold=bold)


def _heading(doc, text: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.page_break_before = True
    for r in list(p.runs):
        r.text = ""
    _set_font(p.add_run(text), size=16, bold=True)


def _body(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    _set_font(p.add_run(text), size=14)


def build(out_path: Path) -> None:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)

    # ── Титульный лист ────────────────────────────────────────────────
    _title_line(doc, "Министерство науки и высшего образования "
                      "Российской Федерации", size=12)
    _title_line(doc, "Федеральное государственное бюджетное образовательное "
                      "учреждение высшего образования", size=12)
    _title_line(doc, "«Московский государственный университет геодезии "
                      "и картографии» (МИИГАиК)", size=12, bold=True)
    doc.add_paragraph()
    _title_line(doc, "Факультет прикладной информатики", size=12)
    _title_line(doc, "Кафедра вычислительной техники и информационных "
                      "технологий", size=12)
    for _ in range(4):
        doc.add_paragraph()
    _title_line(doc, "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", size=16, bold=True)
    doc.add_paragraph()
    _title_line(doc, "на тему:", size=14)
    _title_line(doc, "«Минимальный шаблон выпускной квалификационной "
                      "работы»", size=14, bold=True)
    for _ in range(5):
        doc.add_paragraph()
    _title_line(doc, "Выполнил: студент группы ПИ-V-2  Иванов И. И.", size=12)
    _title_line(doc, "Руководитель: доц., к.т.н.  Петров П. П.", size=12)
    for _ in range(4):
        doc.add_paragraph()
    _title_line(doc, "Москва 2026", size=14)

    # ── СОДЕРЖАНИЕ ────────────────────────────────────────────────────
    _heading(doc, "СОДЕРЖАНИЕ")
    for line in (
        "ВВЕДЕНИЕ",
        "1 Основная часть",
        "ЗАКЛЮЧЕНИЕ",
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        _set_font(p.add_run(line), size=14)

    # ── ВВЕДЕНИЕ ──────────────────────────────────────────────────────
    _heading(doc, "ВВЕДЕНИЕ")
    _body(doc, "Настоящая работа представляет собой минимальный по объёму "
               "документ, структура которого в точности соответствует "
               "требованиям сервиса автоматического форматирования. "
               "Актуальность темы подтверждается источником [1].")
    _body(doc, "Целью работы является демонстрация корректной структуры "
               "выпускной квалификационной работы. Для достижения цели "
               "поставлены и решены соответствующие задачи.")

    # ── ОСНОВНАЯ ЧАСТЬ ────────────────────────────────────────────────
    _heading(doc, "1 Основная часть")
    _body(doc, "В основной части излагается содержание исследования. Раздел "
               "приведён в единственном экземпляре, поскольку документ "
               "является минимальным шаблоном и служит для проверки "
               "распознавания структуры, а не содержательной полноты.")
    _body(doc, "Текст оформлен шрифтом Times New Roman размером 14 пунктов "
               "с полуторным межстрочным интервалом и абзацным отступом "
               "1,25 см, что соответствует нормативным требованиям.")

    # ── ЗАКЛЮЧЕНИЕ ────────────────────────────────────────────────────
    _heading(doc, "ЗАКЛЮЧЕНИЕ")
    _body(doc, "В ходе работы продемонстрирована минимально достаточная "
               "структура выпускной квалификационной работы, пригодная для "
               "загрузки и обработки сервисом без критических структурных "
               "нарушений.")

    # ── СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ ──────────────────────────────
    _heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    _set_font(p.add_run(
        "1. Иванов И. И. Основы оформления выпускных квалификационных "
        "работ : учебное пособие / И. И. Иванов. — Москва : Издательство "
        "МИИГАиК, 2025. — 120 с."
    ), size=14)

    doc.save(str(out_path))


def main() -> int:
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("minimal_vkr.docx")
    build(out)
    print(f"Готово: {out.resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
