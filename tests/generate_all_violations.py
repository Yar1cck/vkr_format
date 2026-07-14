"""Генератор .docx, провоцирующего МАКСИМУМ типов нарушений пайплайна.

Цель — один документ, на котором можно прогнать сервис и увидеть/исправить
все группы замечаний: структура, заголовки, нумерация, ссылки, библиография,
формулы, подписи, таблицы, форматирование, объём.

Некоторые типы взаимоисключающи в пределах одного файла (например
bibliography_missing vs bibliography_numbering_gap, volume_below vs above) —
выбран вариант с большим покрытием.

Запуск:
    python -m tests.generate_all_violations            # → all_violations.docx
    python -m tests.generate_all_violations out.docx
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

FONT = "Times New Roman"


def _font(run, size=14, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._r.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), FONT)


def _center(doc, text, size=14, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    _font(p.add_run(text), size, bold)
    return p


def _h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.page_break_before = True
    for r in list(p.runs):
        r.text = ""
    _font(p.add_run(text), 16, True)
    return p


def _h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.first_line_indent = Cm(0)
    for r in list(p.runs):
        r.text = ""
    _font(p.add_run(text), 14, True)
    return p


def _body(doc, text, size=14, space_after=0.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    _font(p.add_run(text), size)
    return p


def _plain(doc, text, bold=False, center=False, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    _font(p.add_run(text), size, bold)
    return p


def build(out: Path) -> None:
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(3.0)
    s.right_margin = Cm(1.5)

    # ── Титульный лист (валиден, неприкосновенен) ─────────────────────
    _center(doc, "Министерство науки и высшего образования Российской "
                 "Федерации", 12)
    _center(doc, "ФГБОУ ВО «Московский государственный университет геодезии "
                 "и картографии» (МИИГАиК)", 12, True)
    doc.add_paragraph()
    _center(doc, "Факультет прикладной информатики", 12)
    _center(doc, "Кафедра вычислительной техники", 12)
    for _ in range(3):
        doc.add_paragraph()
    _center(doc, "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", 16, True)
    _center(doc, "на тему: «Документ с полным набором нарушений»", 14, True)
    for _ in range(3):
        doc.add_paragraph()
    _center(doc, "Выполнил: студент Иванов И. И.", 12)
    _center(doc, "Руководитель: Петров П. П.", 12)
    doc.add_paragraph()
    _center(doc, "Москва 2026", 14)

    # ── ВВЕДЕНИЕ (намеренно ДО оглавления → structural_order_violation) ─
    _h1(doc, "ВВЕДЕНИЕ")
    # citation_sequence_issue: [2] раньше [1];
    # citation_separator_missing: [3][4];
    # citation_range_suspicious: [1-30];
    # citation_without_source: [9] (нет записи 9);
    # figure_reference_missing: рис. 9; table_reference_missing: табл. 8
    _body(doc, "Актуальность подтверждается источником [2], затем [1]. "
               "Подряд идущие ссылки [3][4] оформлены без разделителя. "
               "Широкий диапазон [50-80] почти наверняка опечатка. "
               "Несуществующий источник [9] цитируется в тексте. "
               "См. рис. 9 и табл. 8, которых в документе нет.")
    _body(doc, "Текст введения набран увеличенным кеглем — провокация "
               "несогласованности размеров шрифта.", size=16)
    _body(doc, "Этот абзац имеет ненулевой интервал после — провокация "
               "нарушения межабзацного интервала.", space_after=10.0)

    # ── ОГЛАВЛЕНИЕ (forbidden → СОДЕРЖАНИЕ) ───────────────────────────
    _h1(doc, "ОГЛАВЛЕНИЕ")
    # Со страничными номерами, как реальное оглавление: иначе строка
    # «СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ» совпала бы с алиасом раздела
    # раньше настоящего заголовка и сломала границы библиографии.
    for line, pg in (("ВВЕДЕНИЕ", 3), ("1 Теоретическая часть", 5),
                     ("3 Практическая часть", 9),
                     ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 12)):
        _plain(doc, f"{line}\t{pg}")

    # ── Глава 1 (forbidden «теоретическая часть») ─────────────────────
    _h1(doc, "1 Теоретическая часть")
    _body(doc, "Глава с запрещённым названием. Ниже подзаголовки с разрывом "
               "нумерации и пропущенным звеном для восстановления иерархии.")
    _h2(doc, "1.1 Постановка задачи")
    _body(doc, "Содержимое подраздела 1.1. Здесь достаточно текста, чтобы "
               "соседний абзац считался телом, а не заголовком, и сработало "
               "восстановление пропущенного звена иерархии между 1.1 и 1.3.")
    # 1.2 НЕ оформлен стилем → кандидат на heading_recovered (gap-fill)
    _plain(doc, "1.2 Обзор существующих решений")
    _body(doc, "Текст подраздела 1.2, который формально не оформлен как "
               "заголовок, но заполняет дыру между 1.1 и 1.3 в нумерации, "
               "поэтому должен быть восстановлен алгоритмом автоматически.")
    _h2(doc, "1.3 Выводы по главе")
    _body(doc, "Завершающий текст первой главы с выводами по результатам "
               "теоретического анализа предметной области исследования.")

    # possible_missed_heading: короткая изолированная строка-кандидат
    # сразу после принятого заголовка, без оформления, без точки.
    _h2(doc, "1.4 Дополнительные сведения")
    _plain(doc, "Архитектура программного комплекса")
    _body(doc, "Длинный абзац тела, который идёт после неоформленной строки "
               "и заставляет алгоритм заподозрить пропущенный заголовок выше "
               "по тексту, так как строка короткая и без завершающей точки.")

    # heading_confirm: короткая изолированная строка без точки и без
    # сильного якоря (capped → LOW → семантический кандидат). Чтобы
    # сработал сигнал isolated, оба соседних абзаца должны быть длинными
    # (≥120 символов).
    _body(doc, "Предшествующий длинный абзац тела документа, который нужен "
               "для того, чтобы у следующей короткой строки сработал сигнал "
               "изолированности и она была предложена на подтверждение.")
    _plain(doc, "Особый случай применения метода", bold=True, center=True)
    _body(doc, "Последующий длинный абзац тела документа, также превышающий "
               "сто двадцать символов, чтобы строка между двумя длинными "
               "абзацами считалась изолированным кандидатом в заголовки.")

    # ── Формулы ───────────────────────────────────────────────────────
    # formula_without_explanation: центрированная формула без «где» после.
    _center(doc, "E = m·c² + ∑αᵢ", 14)
    _body(doc, "Сразу идёт обычный текст, а не пояснение со словом «где», "
               "поэтому формула считается без расшифровки символов.")
    # formula_numbering_gap: (1) затем (3), пропущен (2).
    _center(doc, "S = π·r² (1)", 14)
    _body(doc, "где S — площадь, r — радиус окружности.")
    _center(doc, "F = G·m₁·m₂ / r² (3)", 14)
    _body(doc, "где F — сила, G — гравитационная постоянная.")

    # ── Подписи рисунков/таблиц ───────────────────────────────────────
    # caption_renumbered: первый рисунок назван «Рисунок 4» → станет 1.
    # figure_not_referenced: на него нет ссылки в тексте.
    _plain(doc, "Рисунок 4 – Архитектура системы", center=True)
    _body(doc, "Описание к рисунку. Ссылки на этот рисунок в тексте намеренно "
               "отсутствуют, чтобы сработало замечание о неупомянутом рисунке.")

    # Таблица: подпись с неверным номером (caption_renumbered),
    # не упомянута (table_not_referenced), >15 строк (long_table_split),
    # ≥60% жирных ячеек (table_mass_bold_removed).
    _plain(doc, "Таблица 7 – Сравнительные характеристики решений", center=True)
    tbl = doc.add_table(rows=17, cols=3)
    tbl.style = "Table Grid"
    for ri, row in enumerate(tbl.rows):
        for ci, cell in enumerate(row.cells):
            cell.text = ""
            par = cell.paragraphs[0]
            run = par.add_run(f"Зн{ri}.{ci}" if ri else f"Колонка {ci+1}")
            _font(run, 12, bold=True)  # все ячейки жирные → mass bold
    _body(doc, "После таблицы идёт обычный текст. Ссылок вида «табл.» на эту "
               "таблицу в документе нет специально.")

    # ── Глава 3 (пропущена «2» → heading_number_conflict;
    #            «практическая часть» → forbidden) ─────────────────────
    _h1(doc, "3 Практическая часть")
    _body(doc, "Глава с пропущенным номером (ожидалась 2) и запрещённым "
               "названием. Здесь корректно цитируется источник [1] для "
               "проверки сходимости библиографии.")

    # ── ЗАКЛЮЧЕНИЕ намеренно ОТСУТСТВУЕТ → missing_structural_section ──

    # ── СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ ──────────────────────────────
    _h1(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    # 1 — полная (цитируется [1]); 2 — неполная, нет года, <20 симв.
    #   → bibliography_entry_incomplete (цитируется [2]);
    # 4 — пропущен 3 → bibliography_numbering_gap (цитируется [4]);
    # 6 — пропущен 5 и НИГДЕ не цитируется → source_not_cited.
    # Цитата [9] не имеет записи → citation_without_source.
    _plain(doc, "1. Иванов И. И. Теория систем : учебник / И. И. Иванов. — "
                "Москва : Наука, 2021. — 320 с.")
    _plain(doc, "2. Краткая запись")
    _plain(doc, "4. Петров П. П. Практикум по проектированию : учебное "
                "пособие / П. П. Петров. — Санкт-Петербург : Питер, 2022. — "
                "210 с.")
    _plain(doc, "6. Сидоров С. С. Методология исследований : монография / "
                "С. С. Сидоров. — Казань : КФУ, 2023. — 180 с.")

    doc.save(str(out))


def main() -> int:
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("all_violations.docx")
    build(out)
    print(f"Готово: {out.resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
