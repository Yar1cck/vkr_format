"""Замер качества детектора заголовков на корпусе ВКР (A-N1, ТЗ §7.9 / Р-04).

Два режима:

  manual      — эталон из JSON с ручной разметкой. Даёт настоящие F1/precision
                /recall на реальных документах, в т.ч. с потерянными стилями.
  word-style  — baseline по `paragraph.style.name == "Heading N"`. Работает на
                любом корпусе без подготовки, но осмыслен только на чистых
                документах со стилевой разметкой.

Запуск:
  python -m tests.quality.measure_heading_detection \\
      --corpus path/to/docx/ \\
      --annotations path/to/annotations.json \\
      --output reports/heading_quality.json

Подробности и формат разметки — в `tests/quality/README.md`.
"""

from __future__ import annotations

import argparse
import json
import sys
from collections import defaultdict
from collections.abc import Iterable
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document as WordDocument

from services.core.vkr_core.config.normative_loader import load_default_rules
from services.core.vkr_core.engine.detection import detect_headings
from services.core.vkr_core.engine.stats import collect_stats
from services.core.vkr_core.engine.title_page import detect_title_page_end
from services.core.vkr_core.engine.toc import detect_toc_section_indexes


@dataclass
class GoldLabel:
    """Эталонная разметка одного абзаца. `level` имеет смысл только если
    `is_heading=True`, для отрицательных примеров — None."""
    is_heading: bool
    level: int | None = None


@dataclass
class DocumentEvaluation:
    doc_id: str
    filename: str
    paragraphs_total: int
    tp: int = 0
    fp: int = 0
    fn: int = 0
    tn: int = 0
    level_correct: int = 0
    level_wrong: int = 0
    confusion_levels: dict[str, int] = field(default_factory=dict)
    fp_examples: list[str] = field(default_factory=list)
    fn_examples: list[str] = field(default_factory=list)


def _load_annotations(path: Path) -> dict[str, dict[int, GoldLabel]]:
    """Загружает разметку из JSON. Возвращает dict {filename: {paragraph_index: GoldLabel}}."""
    data = json.loads(path.read_text(encoding="utf-8"))
    result: dict[str, dict[int, GoldLabel]] = {}
    for doc in data.get("documents", []):
        filename = doc["filename"]
        labels: dict[int, GoldLabel] = {}
        for ann in doc.get("annotations", []):
            idx = int(ann["paragraph_index"])
            is_h = bool(ann["is_heading"])
            level = int(ann["level"]) if is_h and "level" in ann else None
            labels[idx] = GoldLabel(is_heading=is_h, level=level)
        result[filename] = labels
    return result


def _word_style_labels(doc: WordDocument) -> dict[int, GoldLabel]:
    """Эталон по native Word-стилям. Heading 1/2/3 (или Заголовок 1/2/3) →
    положительные примеры. Остальные абзацы НЕ помечаем (не считаем как
    отрицательные — в word-style mode мы измеряем согласие только на тех
    абзацах, где Word уверенно сказал «это heading»).
    """
    labels: dict[int, GoldLabel] = {}
    for i, p in enumerate(doc.paragraphs):
        name = (p.style.name or "").strip().lower() if p.style else ""
        for prefix in ("heading ", "заголовок "):
            if name.startswith(prefix):
                tail = name[len(prefix):].strip()
                if tail.isdigit():
                    level = max(1, min(int(tail), 3))
                    labels[i] = GoldLabel(is_heading=True, level=level)
                    break
    return labels


def _detect_headings_for_doc(docx_path: Path, rules: dict) -> dict[int, int]:
    """Возвращает {paragraph_index: level} того, что детектор посчитал заголовками."""
    doc = WordDocument(str(docx_path))
    stats = collect_stats(doc)
    title_end = detect_title_page_end(stats, rules)
    title_page_indexes = set(range(title_end))
    toc_indexes = detect_toc_section_indexes(doc, rules)
    skip = title_page_indexes | toc_indexes
    headings = detect_headings(stats, rules, skip)
    return {h.paragraph_index: h.level for h in headings}


def _get_skip_indexes(docx_path: Path, rules: dict) -> set[int]:
    """Возвращает set индексов, которые детектор исключает из анализа
    (титульный лист + секция TOC)."""
    doc = WordDocument(str(docx_path))
    stats = collect_stats(doc)
    title_end = detect_title_page_end(stats, rules)
    toc_indexes = detect_toc_section_indexes(doc, rules)
    return set(range(title_end)) | toc_indexes


def _evaluate_document(
    docx_path: Path,
    rules: dict,
    gold: dict[int, GoldLabel],
    mode: str,
) -> DocumentEvaluation:
    doc = WordDocument(str(docx_path))
    detected = _detect_headings_for_doc(docx_path, rules)
    paragraphs_total = len(doc.paragraphs)

    # В word-style mode исключаем из gold-набора абзацы, которые детектор
    # обоснованно не обрабатывает: титульный лист (неприкосновенен по ТЗ §6.4)
    # и секцию TOC (записи оглавления — не структурные заголовки контента).
    if mode == "word-style":
        skip = _get_skip_indexes(docx_path, rules)
        gold = {idx: lbl for idx, lbl in gold.items() if idx not in skip}

    ev = DocumentEvaluation(
        doc_id=docx_path.stem,
        filename=docx_path.name,
        paragraphs_total=paragraphs_total,
    )

    # В word-style mode «не размеченных» абзацев в gold-комплекте нет —
    # измеряем только в позициях, где gold = heading (recall) и в позициях,
    # где детектор сказал heading (precision). Все остальные абзацы
    # выпадают из выборки. В manual mode наоборот: каждый абзац имеет
    # эталон (явный или дефолтный «не heading»).
    if mode == "manual":
        indices = range(paragraphs_total)
    else:
        # word-style: union(gold ∪ detected)
        indices = sorted(set(gold.keys()) | set(detected.keys()))

    for idx in indices:
        gold_label = gold.get(idx, GoldLabel(is_heading=False))
        if mode == "word-style" and idx not in gold and idx not in detected:
            continue
        is_detected = idx in detected
        # В word-style mode для «не размеченных в gold, но обнаруженных
        # детектором» абзацев у нас нет ground truth — пропускаем.
        if mode == "word-style" and idx not in gold and is_detected:
            continue

        if gold_label.is_heading and is_detected:
            ev.tp += 1
            actual_level = detected[idx]
            if gold_label.level == actual_level:
                ev.level_correct += 1
            else:
                ev.level_wrong += 1
                key = f"{gold_label.level}->{actual_level}"
                ev.confusion_levels[key] = ev.confusion_levels.get(key, 0) + 1
        elif gold_label.is_heading and not is_detected:
            ev.fn += 1
            text = doc.paragraphs[idx].text.strip()[:80]
            ev.fn_examples.append(f"#{idx} L{gold_label.level}: {text}")
        elif not gold_label.is_heading and is_detected:
            ev.fp += 1
            text = doc.paragraphs[idx].text.strip()[:80]
            ev.fp_examples.append(f"#{idx} L{detected[idx]}: {text}")
        else:
            ev.tn += 1

    return ev


def _aggregate(evaluations: Iterable[DocumentEvaluation]) -> dict:
    tp = fp = fn = tn = 0
    lvl_correct = lvl_wrong = 0
    confusion: dict[str, int] = defaultdict(int)
    for ev in evaluations:
        tp += ev.tp; fp += ev.fp; fn += ev.fn; tn += ev.tn
        lvl_correct += ev.level_correct; lvl_wrong += ev.level_wrong
        for k, v in ev.confusion_levels.items():
            confusion[k] += v

    precision = tp / (tp + fp) if (tp + fp) else 0.0
    recall    = tp / (tp + fn) if (tp + fn) else 0.0
    f1 = 2 * precision * recall / (precision + recall) if (precision + recall) else 0.0
    level_total = lvl_correct + lvl_wrong
    level_accuracy = lvl_correct / level_total if level_total else 0.0

    return {
        "binary": {
            "tp": tp, "fp": fp, "fn": fn, "tn": tn,
            "precision": round(precision, 4),
            "recall":    round(recall, 4),
            "f1":        round(f1, 4),
        },
        "level": {
            "correct":  lvl_correct,
            "wrong":    lvl_wrong,
            "accuracy": round(level_accuracy, 4),
        },
        "confusion_levels": dict(confusion),
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Замер качества детектора заголовков")
    parser.add_argument("--corpus", required=True, type=Path,
                        help="Каталог с docx-файлами для оценки")
    parser.add_argument("--annotations", type=Path,
                        help="JSON с ручной разметкой (mode=manual)")
    parser.add_argument("--baseline", choices=["word-style"],
                        help="Использовать word-style baseline вместо ручной разметки")
    parser.add_argument("--output", type=Path, required=True,
                        help="Куда сохранить отчёт (JSON)")
    args = parser.parse_args()

    if bool(args.annotations) == bool(args.baseline):
        parser.error("Укажите ровно одно из: --annotations или --baseline")

    if not args.corpus.is_dir():
        parser.error(f"Корпус не найден: {args.corpus}")

    rules = load_default_rules()

    mode = "manual" if args.annotations else "word-style"
    annotations: dict[str, dict[int, GoldLabel]] = (
        _load_annotations(args.annotations) if args.annotations else {}
    )

    evaluations: list[DocumentEvaluation] = []
    for docx_path in sorted(args.corpus.glob("*.docx")):
        if mode == "manual":
            gold = annotations.get(docx_path.name)
            if gold is None:
                print(f"  [skip] {docx_path.name}: нет разметки в annotations.json")
                continue
        else:
            doc = WordDocument(str(docx_path))
            gold = _word_style_labels(doc)
            if not gold:
                print(f"  [skip] {docx_path.name}: нет heading-стилей (word-style mode бесполезен)")
                continue

        try:
            ev = _evaluate_document(docx_path, rules, gold, mode)
            evaluations.append(ev)
            print(f"  [eval] {docx_path.name}: TP={ev.tp} FP={ev.fp} FN={ev.fn} TN={ev.tn}")
        except Exception as exc:
            print(f"  [fail] {docx_path.name}: {exc}", file=sys.stderr)

    if not evaluations:
        print("Не удалось оценить ни одного документа", file=sys.stderr)
        return 2

    aggregated = _aggregate(evaluations)
    report = {
        "mode": mode,
        "documents_evaluated": len(evaluations),
        **aggregated,
        "per_document": [
            {
                "doc_id": ev.doc_id,
                "filename": ev.filename,
                "paragraphs_total": ev.paragraphs_total,
                "tp": ev.tp, "fp": ev.fp, "fn": ev.fn, "tn": ev.tn,
                "level_correct": ev.level_correct,
                "level_wrong": ev.level_wrong,
                "confusion_levels": ev.confusion_levels,
                "fp_examples": ev.fp_examples[:10],
                "fn_examples": ev.fn_examples[:10],
            }
            for ev in evaluations
        ],
    }

    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    print()
    print(f"Отчёт сохранён в {args.output}")
    print(f"  documents:  {len(evaluations)}")
    print(f"  precision:  {aggregated['binary']['precision']}")
    print(f"  recall:     {aggregated['binary']['recall']}")
    print(f"  F1:         {aggregated['binary']['f1']}")
    print(f"  level acc:  {aggregated['level']['accuracy']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
