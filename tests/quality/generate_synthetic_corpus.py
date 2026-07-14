"""Генератор разнообразного синтетического корпуса для quality-замеров.

Создаёт несколько ВКР разного объёма и качества:

    Чистые работы (для замера precision — ложных тревог быть не должно):
    clean_bachelor_short.docx   ~20 стр, бакалавр, без нарушений
    clean_bachelor_medium.docx  ~50 стр, бакалавр, без нарушений
    clean_bachelor_long.docx    ~90 стр, бакалавр, без нарушений

  «Грязные» работы (для замера recall — алгоритм должен найти):
    messy_citations.docx        Нарушения в цитировании
    messy_bibliography.docx     Битый список источников
    messy_headings.docx         Сломанные заголовки/нумерация
    messy_appendices.docx       Битые приложения
    messy_full.docx             Микс нарушений всех типов

  Edge cases:
    edge_long_body.docx         Длинная сплошная глава без подразделов (нагрузочный)

Эталон ожидаемых нарушений строится по тем же правилам, что и сами вставки,
и сохраняется в `_annotations.json` рядом с документами.

    python -m tests.quality.generate_synthetic_corpus \\
        --output tests/quality/corpus/synthetic/
"""

from __future__ import annotations

import argparse
import json
import sys
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

FONT = "Times New Roman"


# ─── Атомарные строители абзацев ─────────────────────────────────────────────

def _set_font(run, size: float = 14, bold: bool = False, italic: bool = False) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)


def _center(p) -> None:
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)


def _title_line(doc, text: str, size: float = 14, bold: bool = False) -> None:
    p = doc.add_paragraph()
    _center(p)
    _set_font(p.add_run(text), size=size, bold=bold)


def _heading(doc, text: str, size: float = 14, level: int = 1) -> int:
    """Добавляет заголовок и возвращает его paragraph_index."""
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    if level == 1:
        p.paragraph_format.page_break_before = True
    for r in list(p.runs):
        r.text = ""
    _set_font(p.add_run(text), size=size, bold=True)
    return len(doc.paragraphs) - 1


def _body(doc, text: str) -> int:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    _set_font(p.add_run(text), size=14)
    return len(doc.paragraphs) - 1


def _bib_entry(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    _set_font(p.add_run(text), size=14)


def _caption_figure(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    _set_font(p.add_run(text), size=14)


def _caption_table(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    _set_font(p.add_run(text), size=14)


def _setup_section(doc) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)


def _title_page(doc, *, work_title: str, title_label: str = "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА",
                degree: str = "бакалавра") -> None:
    # Используем 14pt везде на титуле, чтобы не триггерить font_size_inconsistency
    # (детектор не различает «титул» / «тело», если у пары размеров доля > 0).
    _title_line(doc, "Министерство науки и высшего образования Российской Федерации", size=14)
    _title_line(doc, "Федеральное государственное бюджетное образовательное учреждение "
                     "высшего образования", size=14)
    _title_line(doc, "«Московский государственный университет геодезии и картографии» (МИИГАиК)",
                size=14, bold=True)
    doc.add_paragraph()
    _title_line(doc, "Факультет прикладной информатики", size=14)
    _title_line(doc, "Кафедра вычислительной техники и информационных технологий", size=14)
    for _ in range(4):
        doc.add_paragraph()
    _title_line(doc, title_label, size=14, bold=True)
    doc.add_paragraph()
    _title_line(doc, f"({degree})", size=14)
    doc.add_paragraph()
    _title_line(doc, "на тему:", size=14)
    _title_line(doc, f"«{work_title}»", size=14, bold=True)
    for _ in range(5):
        doc.add_paragraph()
    _title_line(doc, "Выполнил: студент группы ПИ-V-2  Иванов И. И.", size=14)
    _title_line(doc, "Руководитель: доц., к.т.н.  Петров П. П.", size=14)
    for _ in range(4):
        doc.add_paragraph()
    _title_line(doc, "Москва 2026", size=14)


# ─── Шаблонные тексты для разной длины ───────────────────────────────────────

LOREM = (
    "В настоящем разделе подробно рассматривается актуальная задача, "
    "которая решается в рамках выпускной квалификационной работы. "
    "Автор приводит описание методологии, использованной для достижения "
    "поставленной цели, а также формулирует основные результаты "
    "проведённого анализа. Особое внимание уделяется обзору существующих "
    "подходов и их сравнительной характеристике. "
)


def _long_body(doc, paragraphs: int, with_refs: list[int] | None = None) -> None:
    """Добавляет N абзацев тела, опционально вставляя ссылки [N] в возрастающем порядке."""
    refs = sorted(with_refs) if with_refs else []
    ref_idx = 0
    for i in range(paragraphs):
        text = LOREM * 2
        # Ссылки идут строго по возрастанию — иначе детектор cuts citation_sequence_issue
        if refs and i % 3 == 0:
            ref = refs[ref_idx % len(refs)]
            text += f"Эта мысль подтверждается источником [{ref}]. "
            ref_idx += 1
        _body(doc, text)


# ─── Спецификации эталона ────────────────────────────────────────────────────

@dataclass
class CorpusDocument:
    filename: str
    expected_violations: list[dict] = field(default_factory=list)
    description: str = ""

    def add_expected(self, vtype: str, paragraph_index: int | None = None, note: str = "") -> None:
        rec = {"type": vtype, "expected": True, "_detected": True}
        if paragraph_index is not None:
            rec["paragraph_index"] = paragraph_index
        if note:
            rec["_note"] = note
        self.expected_violations.append(rec)


# ─── Билдеры конкретных документов ───────────────────────────────────────────

def build_clean_bachelor_short(out: Path) -> CorpusDocument:
    """Короткая чистая бакалаврская — ~20 страниц, без нарушений."""
    doc = Document()
    _setup_section(doc)
    _title_page(doc, work_title="Краткое исследование систем автоматизации")

    _heading(doc, "СОДЕРЖАНИЕ")
    for line in ("ВВЕДЕНИЕ", "1 Анализ предметной области", "2 Реализация",
                 "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        _set_font(p.add_run(line), size=14)

    _heading(doc, "ВВЕДЕНИЕ")
    _long_body(doc, paragraphs=15, with_refs=[1, 2])

    _heading(doc, "1 Анализ предметной области")
    _long_body(doc, paragraphs=40, with_refs=[1, 3])
    _heading(doc, "1.1 Существующие решения", level=2)
    _long_body(doc, paragraphs=25, with_refs=[2])
    _heading(doc, "1.2 Постановка задачи", level=2)
    _long_body(doc, paragraphs=20, with_refs=[3])

    _heading(doc, "2 Программная реализация")
    _long_body(doc, paragraphs=40, with_refs=[1, 2, 3])

    _heading(doc, "ЗАКЛЮЧЕНИЕ")
    _long_body(doc, paragraphs=5)

    _heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    _bib_entry(doc, "1. Иванов И. И. Основы автоматизации : учебник / И. И. Иванов. — Москва : МИИГАиК, 2025. — 200 с.")
    _bib_entry(doc, "2. Петров П. П. Системный анализ : учебное пособие / П. П. Петров. — Москва : Наука, 2024. — 150 с.")
    _bib_entry(doc, "3. Сидоров С. С. Информационные технологии : учебник / С. С. Сидоров. — Москва : Юрайт, 2023. — 300 с.")

    doc.save(str(out))
    return CorpusDocument(filename=out.name, description="Чистая бакалаврская, ~20 страниц")


def build_clean_bachelor_medium(out: Path) -> CorpusDocument:
    """Средняя чистая бакалаврская — ~50 страниц."""
    doc = Document()
    _setup_section(doc)
    _title_page(doc, work_title="Разработка информационной системы учёта документов")

    _heading(doc, "СОДЕРЖАНИЕ")
    for line in ("ВВЕДЕНИЕ", "1 Теоретическая часть", "2 Проектирование",
                 "3 Реализация", "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        _set_font(p.add_run(line), size=14)

    _heading(doc, "ВВЕДЕНИЕ")
    _long_body(doc, paragraphs=10, with_refs=[1, 2, 3])

    for ch_num, ch_title in [
        (1, "Анализ предметной области"),
        (2, "Проектирование информационной системы"),
        (3, "Программная реализация и тестирование"),
    ]:
        _heading(doc, f"{ch_num} {ch_title}")
        _long_body(doc, paragraphs=30, with_refs=[1, 2, 3, 4])
        _heading(doc, f"{ch_num}.1 Подраздел A", level=2)
        _long_body(doc, paragraphs=20, with_refs=[1, 3])
        _heading(doc, f"{ch_num}.2 Подраздел B", level=2)
        _long_body(doc, paragraphs=20, with_refs=[2, 4])
        _heading(doc, f"{ch_num}.3 Подраздел C", level=2)
        _long_body(doc, paragraphs=15, with_refs=[5])

    _heading(doc, "ЗАКЛЮЧЕНИЕ")
    _long_body(doc, paragraphs=8)

    _heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    for i, src in enumerate([
        "Основы программирования : учебник / А. А. Алексеев. — Москва : МИИГАиК, 2025. — 250 с.",
        "Базы данных : учебное пособие / Б. Б. Борисов. — Москва : Наука, 2024. — 180 с.",
        "Архитектура ПО : учебник / В. В. Васильев. — Москва : Юрайт, 2023. — 320 с.",
        "Тестирование систем : учебное пособие / Г. Г. Григорьев. — Москва : Лань, 2024. — 200 с.",
        "Безопасность данных : учебник / Д. Д. Дмитриев. — Москва : Питер, 2025. — 280 с.",
    ], start=1):
        _bib_entry(doc, f"{i}. {src}")

    doc.save(str(out))
    return CorpusDocument(filename=out.name, description="Чистая бакалаврская, ~50 страниц")


def build_clean_bachelor_long(out: Path) -> CorpusDocument:
    """Длинная чистая бакалаврская работа — ~90 страниц."""
    doc = Document()
    _setup_section(doc)
    _title_page(doc, work_title="Методы интеллектуального анализа геопространственных данных",
                degree="бакалавра")

    _heading(doc, "СОДЕРЖАНИЕ")
    for line in ("ВВЕДЕНИЕ", "1 Обзор литературы", "2 Методология",
                 "3 Эксперименты", "4 Анализ результатов",
                 "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        _set_font(p.add_run(line), size=14)

    _heading(doc, "ВВЕДЕНИЕ")
    _long_body(doc, paragraphs=15, with_refs=list(range(1, 8)))

    for ch_num, ch_title in [
        (1, "Обзор литературы"),
        (2, "Методология"),
        (3, "Эксперименты"),
        (4, "Анализ результатов"),
    ]:
        _heading(doc, f"{ch_num} {ch_title}")
        _long_body(doc, paragraphs=30, with_refs=list(range(1, 8)))
        _heading(doc, f"{ch_num}.1 Подраздел A", level=2)
        _long_body(doc, paragraphs=15, with_refs=[1, 3, 5])
        _heading(doc, f"{ch_num}.2 Подраздел B", level=2)
        _long_body(doc, paragraphs=15, with_refs=[2, 4, 6])
        _heading(doc, f"{ch_num}.3 Подраздел C", level=2)
        _long_body(doc, paragraphs=12, with_refs=[7, 8])

    _heading(doc, "ЗАКЛЮЧЕНИЕ")
    _long_body(doc, paragraphs=12)

    _heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    for i, src in enumerate([
        "Машинное обучение : учебник / А. А. Алексеев. — Москва : МИИГАиК, 2025. — 350 с.",
        "Геоинформационные системы : учебное пособие / Б. Б. Борисов. — Москва : Наука, 2024. — 280 с.",
        "Картография : учебник / В. В. Васильев. — Москва : Юрайт, 2023. — 420 с.",
        "Анализ пространственных данных : учебное пособие / Г. Г. Григорьев. — Москва : Лань, 2024. — 200 с.",
        "Базы знаний : учебник / Д. Д. Дмитриев. — Москва : Питер, 2025. — 380 с.",
        "Нейронные сети : учебник / Е. Е. Ершов. — Москва : МИИГАиК, 2024. — 300 с.",
        "Большие данные : учебное пособие / Ж. Ж. Жуков. — Москва : Наука, 2025. — 250 с.",
        "Spatial Statistics : textbook / I. Ivanov. — Moscow : MIIGAiK, 2024. — 400 p.",
    ], start=1):
        _bib_entry(doc, f"{i}. {src}")

    doc.save(str(out))
    return CorpusDocument(filename=out.name, description="Чистая бакалаврская, ~90 страниц")


def build_messy_citations(out: Path) -> CorpusDocument:
    """Цитирование сломано — ожидаются типы:
       citation_separator_missing, citation_range_suspicious,
       citation_sequence_issue, citation_without_source."""
    doc = Document()
    _setup_section(doc)
    _title_page(doc, work_title="Документ с битым цитированием")
    cd = CorpusDocument(filename=out.name, description="Сломанное цитирование")

    _heading(doc, "СОДЕРЖАНИЕ")
    for line in ("ВВЕДЕНИЕ", "1 Основная часть", "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        _set_font(p.add_run(line), size=14)

    _heading(doc, "ВВЕДЕНИЕ")
    _body(doc, "Тема актуальна, что подтверждают источники [1][2] — здесь забыта запятая.")
    cd.add_expected("citation_separator_missing", note="[1][2] без запятой")
    _body(doc, "В дальнейшем работа опирается на исследования [3].")

    _heading(doc, "1 Основная часть")
    _body(doc, "Существенный массив литературы [1-100] поддерживает данное утверждение.")
    cd.add_expected("citation_range_suspicious", note="[1-100] — слишком широкий диапазон")
    _body(doc, "Также важны работы [7] и [5] — порядок ссылок нарушен.")
    cd.add_expected("citation_sequence_issue", note="ссылки не по возрастанию")
    _body(doc, "Особое значение имеет источник [99], которого нет в библиографии.")
    cd.add_expected("citation_without_source", note="[99] нет в списке литературы")

    _heading(doc, "ЗАКЛЮЧЕНИЕ")
    _body(doc, "Заключение основано на изученных материалах.")

    _heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    for i in range(1, 8):
        _bib_entry(doc, f"{i}. Автор {i}. Работа № {i} : учебник / Автор {i}. — Москва : Изд-во, 2024. — 150 с.")

    doc.save(str(out))
    return cd


def build_messy_bibliography(out: Path) -> CorpusDocument:
    """Битая библиография — bibliography_numbering_gap, bibliography_entries_unrecognised,
       bibliography_entry_incomplete, source_not_cited."""
    doc = Document()
    _setup_section(doc)
    _title_page(doc, work_title="Документ с битой библиографией")
    cd = CorpusDocument(filename=out.name, description="Битый список источников")

    _heading(doc, "СОДЕРЖАНИЕ")
    for line in ("ВВЕДЕНИЕ", "1 Основная часть", "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        _set_font(p.add_run(line), size=14)

    _heading(doc, "ВВЕДЕНИЕ")
    _body(doc, "Опираясь на источники [1] и [3], можно утверждать следующее.")

    _heading(doc, "1 Основная часть")
    _long_body(doc, paragraphs=10, with_refs=[1, 3])

    _heading(doc, "ЗАКЛЮЧЕНИЕ")
    _body(doc, "Подводим итоги исследования.")

    _heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    _bib_entry(doc, "1. Иванов И. И. Полный источник : учебник / И. И. Иванов. — Москва : МИИГАиК, 2024. — 200 с.")
    _bib_entry(doc, "Просто строка без номера и формата")  # unrecognised
    cd.add_expected("bibliography_entries_unrecognised", note="строка без номера")
    _bib_entry(doc, "3. Сидоров С. С. (без года и издательства)")  # incomplete + создаст gap (2 пропущен)
    cd.add_expected("bibliography_entry_incomplete", note="нет года и места")
    cd.add_expected("bibliography_numbering_gap", note="пропущен номер 2")
    _bib_entry(doc, "4. Не упомянутый автор. Не упомянутая работа : учебник. — Москва, 2024. — 100 с.")
    cd.add_expected("source_not_cited", note="источник 4 не упомянут в тексте")

    doc.save(str(out))
    return cd


def build_messy_headings(out: Path) -> CorpusDocument:
    """Сломанные заголовки — heading_number_conflict, forbidden_heading,
       missing_structural_section."""
    doc = Document()
    _setup_section(doc)
    _title_page(doc, work_title="Документ со сломанными заголовками")
    cd = CorpusDocument(filename=out.name, description="Сломанные заголовки и нумерация")

    _heading(doc, "ОГЛАВЛЕНИЕ")  # forbidden — должно быть СОДЕРЖАНИЕ
    cd.add_expected("forbidden_heading", note="ОГЛАВЛЕНИЕ запрещено")
    for line in ("ВВЕДЕНИЕ", "1 Глава первая", "3 Глава третья",
                 "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        _set_font(p.add_run(line), size=14)

    _heading(doc, "ВВЕДЕНИЕ")
    _long_body(doc, paragraphs=5, with_refs=[1])

    _heading(doc, "1 Глава первая")
    _long_body(doc, paragraphs=10, with_refs=[1])

    _heading(doc, "3 Глава третья")  # пропущена 2
    cd.add_expected("heading_number_conflict", note="пропущена глава 2 (1 → 3)")
    _long_body(doc, paragraphs=10, with_refs=[1])

    # Намеренно НЕТ заключения → missing_structural_section
    cd.add_expected("missing_structural_section", note="нет ЗАКЛЮЧЕНИЯ")

    _heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    _bib_entry(doc, "1. Иванов И. И. Тест : учебник. — Москва, 2024. — 100 с.")

    doc.save(str(out))
    return cd


def build_messy_appendices(out: Path) -> CorpusDocument:
    """Битые приложения — appendix_letter_gap, appendix_not_referenced, appendix_letter_invalid."""
    doc = Document()
    _setup_section(doc)
    _title_page(doc, work_title="Документ со сломанными приложениями")
    cd = CorpusDocument(filename=out.name, description="Битые приложения")

    _heading(doc, "СОДЕРЖАНИЕ")
    for line in ("ВВЕДЕНИЕ", "1 Основная часть", "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        _set_font(p.add_run(line), size=14)

    _heading(doc, "ВВЕДЕНИЕ")
    _body(doc, "Во введении упомянуто приложение А. Приложение В содержит дополнительные данные.")
    # Приложение Б упомянем не будет → not_referenced

    _heading(doc, "1 Основная часть")
    _long_body(doc, paragraphs=10, with_refs=[1])

    _heading(doc, "ЗАКЛЮЧЕНИЕ")
    _body(doc, "Подводим итоги.")

    _heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    _bib_entry(doc, "1. Иванов И. И. Тест : учебник. — Москва, 2024. — 100 с.")

    _heading(doc, "Приложение А")
    _body(doc, "Содержимое приложения А.")
    _heading(doc, "Приложение В")  # пропущена Б
    cd.add_expected("appendix_letter_gap", note="пропущена буква Б (А → В)")
    _body(doc, "Содержимое приложения В.")
    _heading(doc, "Приложение З")  # З — запрещённая буква
    cd.add_expected("appendix_letter_invalid", note="З не входит в допустимый ряд")
    _body(doc, "Содержимое приложения З.")

    doc.save(str(out))
    return cd


def build_messy_full(out: Path) -> CorpusDocument:
    """Микс многих типов нарушений в одном документе — стресс-тест recall."""
    doc = Document()
    _setup_section(doc)
    _title_page(doc, work_title="Документ с разнообразными нарушениями")
    cd = CorpusDocument(filename=out.name, description="Микс нарушений всех модулей")

    _heading(doc, "ОГЛАВЛЕНИЕ")  # forbidden
    cd.add_expected("forbidden_heading")
    for line in ("ВВЕДЕНИЕ", "1 Глава первая", "2 Глава вторая",
                 "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        _set_font(p.add_run(line), size=14)

    _heading(doc, "ВВЕДЕНИЕ")
    _body(doc, "Опираемся на источники [1][2] без запятой.")
    cd.add_expected("citation_separator_missing")
    _body(doc, "Ссылка [50-90] на массив источников.")
    cd.add_expected("citation_range_suspicious")

    _heading(doc, "1 Глава первая")
    _long_body(doc, paragraphs=10, with_refs=[1])
    _caption_figure(doc, "Рисунок 1 — Первая диаграмма")
    cd.add_expected("figure_not_referenced", note="на Рисунок 1 нет ссылки в тексте")

    _heading(doc, "2 Глава вторая")
    _long_body(doc, paragraphs=10, with_refs=[3])
    _caption_table(doc, "Таблица 1 — Сравнение методов")
    cd.add_expected("table_not_referenced", note="на Таблицу 1 нет ссылки в тексте")

    _heading(doc, "ЗАКЛЮЧЕНИЕ")
    _body(doc, "Подведение итогов.")

    _heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    _bib_entry(doc, "1. Иванов И. И. Источник 1 : учебник. — Москва, 2024. — 100 с.")
    _bib_entry(doc, "Просто строка")  # unrecognised
    cd.add_expected("bibliography_entries_unrecognised")
    _bib_entry(doc, "3. Сидоров С. С. (неполная запись)")  # gap+incomplete
    cd.add_expected("bibliography_entry_incomplete")
    cd.add_expected("bibliography_numbering_gap")

    doc.save(str(out))
    return cd


def build_edge_long_body(out: Path) -> CorpusDocument:
    """Edge case: одна очень длинная глава без подразделов (нагрузочный, ~120 страниц)."""
    doc = Document()
    _setup_section(doc)
    _title_page(doc, work_title="Длинный документ без подразделов")

    _heading(doc, "СОДЕРЖАНИЕ")
    for line in ("ВВЕДЕНИЕ", "1 Основная часть", "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        _set_font(p.add_run(line), size=14)

    _heading(doc, "ВВЕДЕНИЕ")
    _long_body(doc, paragraphs=8, with_refs=[1, 2])

    _heading(doc, "1 Основная часть")
    _long_body(doc, paragraphs=120, with_refs=[1, 2, 3])  # очень длинная

    _heading(doc, "ЗАКЛЮЧЕНИЕ")
    _long_body(doc, paragraphs=8)

    _heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    for i in range(1, 4):
        _bib_entry(doc, f"{i}. Автор {i}. Работа № {i} : учебник. — Москва, 2024. — 200 с.")

    doc.save(str(out))
    return CorpusDocument(filename=out.name, description="Длинный документ (~120 страниц) без подразделов")


BUILDERS = [
    ("clean_bachelor_short.docx", build_clean_bachelor_short),
    ("clean_bachelor_medium.docx", build_clean_bachelor_medium),
    ("clean_bachelor_long.docx", build_clean_bachelor_long),
    ("messy_citations.docx", build_messy_citations),
    ("messy_bibliography.docx", build_messy_bibliography),
    ("messy_headings.docx", build_messy_headings),
    ("messy_appendices.docx", build_messy_appendices),
    ("messy_full.docx", build_messy_full),
    ("edge_long_body.docx", build_edge_long_body),
]


def main() -> int:
    parser = argparse.ArgumentParser(description="Генератор синтетического корпуса для quality-замеров")
    parser.add_argument("--output", required=True, type=Path,
                        help="Куда положить .docx и _annotations.json")
    args = parser.parse_args()

    args.output.mkdir(parents=True, exist_ok=True)
    docs: list[CorpusDocument] = []

    for filename, builder in BUILDERS:
        path = args.output / filename
        print(f"  [build] {filename}", file=sys.stderr)
        cd = builder(path)
        docs.append(cd)
        print(f"    {path.stat().st_size // 1024} KB, {len(cd.expected_violations)} ожидаемых нарушений", file=sys.stderr)

    annotations = {
        "_format_version": 1,
        "_documentation": [
            "Эталон ожидаемых нарушений, сгенерирован automatically",
            "вместе с .docx. Если изменили .docx — перегенерируйте.",
        ],
        "documents": [
            {
                "filename": doc.filename,
                "_description": doc.description,
                "expected_violations": doc.expected_violations,
            }
            for doc in docs
        ],
    }
    out_json = args.output / "_annotations.json"
    out_json.write_text(json.dumps(annotations, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\nКорпус: {args.output} ({len(docs)} файлов, {sum(len(d.expected_violations) for d in docs)} нарушений)")
    print(f"Эталон: {out_json}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
