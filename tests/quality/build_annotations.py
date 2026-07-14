"""Генератор эталонной разметки заголовков для всего корпуса (A-N1, ТЗ §7.9).

Эталон строится из АВТОРСКОЙ стилевой разметки Word (`Heading N`), а не из
вывода детектора — это независимый ground truth. Из набора вычитаются
skip-индексы (титульный лист + секция TOC), которые детектор обоснованно
не обрабатывает.

Особый случай — испорченные пайплайном файлы: если документ является
устаревшим артефактом старого (багованного) пайплайна, его стили доверия
не заслуживают. Для такого файла эталон берётся у его «чистого близнеца» —
того же исходного документа, обработанного исправленным пайплайном
(индексы абзацев совпадают, текст идентичен по позициям).

Запуск:
    python -m tests.quality.build_annotations \\
        --corpus debug/ --output tests/quality/corpus_annotations.json
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document as WordDocument

from services.core.vkr_core.config.normative_loader import load_default_rules
from services.core.vkr_core.engine.stats import collect_stats
from services.core.vkr_core.engine.title_page import detect_title_page_end
from services.core.vkr_core.engine.toc import detect_toc_section_indexes

# Файлы — устаревшие артефакты старого пайплайна; эталон берётся у близнеца
# (тот же исходный документ, корректно обработанный). Ключ — испорченный
# файл, значение — файл-донор корректной стилевой разметки.
_CORRUPTED_TWIN = {
    "processed.docx": "processed2.docx",
}


def _word_style_headings(doc: WordDocument) -> dict[int, int]:
    """{paragraph_index: level} по авторским стилям Heading N / Заголовок N."""
    result: dict[int, int] = {}
    for i, p in enumerate(doc.paragraphs):
        name = (p.style.name or "").strip().lower() if p.style else ""
        for prefix in ("heading ", "заголовок "):
            if name.startswith(prefix):
                tail = name[len(prefix):].strip()
                if tail.isdigit():
                    result[i] = max(1, min(int(tail), 3))
                break
    return result


def _skip_indexes(doc: WordDocument, rules: dict) -> set[int]:
    stats = collect_stats(doc)
    title_end = detect_title_page_end(stats, rules)
    toc_indexes = detect_toc_section_indexes(doc, rules)
    return set(range(title_end)) | toc_indexes


def _annotations_for(docx_path: Path, rules: dict, corpus: Path) -> list[dict]:
    """Истинные заголовки файла (после вычитания skip)."""
    twin = _CORRUPTED_TWIN.get(docx_path.name)
    source_path = corpus / twin if twin else docx_path

    src_doc = WordDocument(str(source_path))
    headings = _word_style_headings(src_doc)
    skip = _skip_indexes(src_doc, rules)

    anns: list[dict] = []
    for idx, level in sorted(headings.items()):
        if idx in skip:
            continue
        anns.append({"paragraph_index": idx, "is_heading": True, "level": level})
    return anns


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Построить эталонную разметку заголовков для корпуса"
    )
    parser.add_argument("--corpus", required=True, type=Path,
                        help="Каталог с docx-файлами")
    parser.add_argument("--output", required=True, type=Path,
                        help="Куда сохранить JSON с разметкой")
    args = parser.parse_args()

    if not args.corpus.is_dir():
        parser.error(f"Корпус не найден: {args.corpus}")

    rules = load_default_rules()
    documents = []
    for docx_path in sorted(args.corpus.glob("*.docx")):
        twin = _CORRUPTED_TWIN.get(docx_path.name)
        # Сырой вход без нативных Heading-стилей и без index-совпадающего
        # близнеца разметить нечем — ground truth физически отсутствует.
        # Пропускаем (как это делает и word-style режим).
        if twin is None and not _word_style_headings(WordDocument(str(docx_path))):
            print(f"  [skip] {docx_path.name}: нет нативных стилей и близнеца — нечем размечать")
            continue
        anns = _annotations_for(docx_path, rules, args.corpus)
        documents.append({
            "doc_id": docx_path.stem,
            "filename": docx_path.name,
            "_source": twin or docx_path.name,
            "annotations": anns,
        })
        note = f" (эталон от {twin})" if twin else ""
        print(f"  {docx_path.name}: {len(anns)} заголовков{note}")

    report = {
        "_format_version": 1,
        "_generated_by": "tests.quality.build_annotations",
        "documents": documents,
    }
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(
        json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(f"\nРазметка сохранена в {args.output} ({len(documents)} документов)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
