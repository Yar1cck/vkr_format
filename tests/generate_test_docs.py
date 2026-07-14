"""
Генерация тестовых .docx документов для проверки алгоритма определения заголовков.

Документ 1 (draft): черновик с нарушениями форматирования
Документ 2 (clean): эталонный чистовой вариант с правильным форматированием
"""

import copy
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def set_font(run, name="Times New Roman", size=14, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Кириллический шрифт (для Word)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def page_margins(doc, top=2.0, bottom=2.0, left=3.0, right=1.5):
    section = doc.sections[0]
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


# ──────────────────────────────────────────────
# ДОКУМЕНТ 1: ЧЕРНОВИК (с нарушениями)
# ──────────────────────────────────────────────

def make_draft():
    doc = Document()
    page_margins(doc)

    # --- Титул (грубый, без нормального форматирования) ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Министерство науки и высшего образования Российской Федерации")
    set_font(run, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ФГБОУ ВО «МИИГАиК»")
    set_font(run, size=12, bold=True)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Факультет информационных технологий")
    set_font(run, size=12)

    doc.add_paragraph()
    doc.add_paragraph()

    # Тема (маленький шрифт, не выровнен нормально)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Разработка системы автоматической проверки форматирования\n"
                    "выпускных квалификационных работ")
    set_font(run, size=13, bold=True)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Пояснительная записка")
    set_font(run, size=12)

    doc.add_paragraph()
    doc.add_paragraph()

    # Нарушение: автор выровнен по центру вместо правого края
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Студент: Иванов И.И.")
    set_font(run, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Руководитель: проф. Петров П.П.")
    set_font(run, size=12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Москва 2025")
    set_font(run, size=12)

    doc.add_page_break()

    # --- АННОТАЦИЯ ---
    # Нарушение: слово «АННОТАЦИЯ» строчными буквами, курсив
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Аннотация")
    set_font(run, size=14, bold=True, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        "В данной работе рассматривается разработка веб-сервиса для автоматической "
        "проверки форматирования выпускных квалификационных работ (ВКР) "
        "в соответствии с требованиями МИИГАиК. "
        "Система позволяет выявлять отклонения от норм оформления, сокращая "
        "время ручной проверки преподавателями."
    )
    set_font(run, size=14)
    p.paragraph_format.first_line_indent = Cm(1.25)

    doc.add_paragraph()

    # --- СОДЕРЖАНИЕ (без стиля «Heading») ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("СОДЕРЖАНИЕ")
    set_font(run, size=14, bold=True)

    # Нарушение: ручной список без табуляций (нет точек-лидеров)
    items = [
        "Введение ........................................................ 4",
        "1 Постановка задачи ........................................... 5",
        "1.1 Требования к системе ...................................... 6",
        "1.2 Анализ существующих решений ............................... 8",
        "2 Проектирование системы ...................................... 10",
        "2.1 Архитектура ................................................ 11",
        "2.2 Схема базы данных .......................................... 13",
        "3 Реализация ................................................... 15",
        "3.1 Бэкенд ..................................................... 15",
        "3.2 Фронтенд ................................................... 18",
        "Заключение ..................................................... 22",
        "Список использованных источников ............................... 23",
        "Приложение А ................................................... 25",
    ]
    for item in items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_font(run, size=14)

    doc.add_page_break()

    # --- ВВЕДЕНИЕ ---
    # Нарушение: «Введение» не по центру, не CAPS
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Введение")
    set_font(run, size=14, bold=True)

    body_text = (
        "Актуальность темы определяется необходимостью автоматизации рутинных "
        "операций проверки соответствия оформления ВКР нормативным требованиям. "
        "В настоящее время данная процедура выполняется вручную, что приводит к "
        "значительным временным затратам и субъективности оценки. "
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(body_text)
    set_font(run, size=14)
    p.paragraph_format.first_line_indent = Cm(1.25)
    # Нарушение: межстрочный интервал 1.0 вместо 1.5
    p.paragraph_format.line_spacing = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        "Цель работы — разработать веб-сервис, который принимает файл ВКР в "
        "формате .docx и возвращает структурированный отчёт о нарушениях форматирования."
    )
    set_font(run, size=14)
    p.paragraph_format.first_line_indent = Cm(1.25)

    doc.add_paragraph()

    # --- ГЛАВА 1 ---
    # Нарушение: глава не с новой страницы, заголовок с точкой в конце
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("1 ПОСТАНОВКА ЗАДАЧИ.")
    set_font(run, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        "Для достижения цели работы необходимо решить следующие задачи: "
        "разработать алгоритм анализа структуры документа, реализовать "
        "серверную часть с REST API, создать пользовательский интерфейс. "
        "Система должна обрабатывать документы объёмом до 200 страниц."
    )
    set_font(run, size=14)
    p.paragraph_format.first_line_indent = Cm(1.25)

    # --- ПОДРАЗДЕЛ 1.1 ---
    # Нарушение: подраздел с отступом слева, маленький шрифт
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run("1.1 Требования к системе")
    set_font(run, size=13, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        "Система должна обеспечивать загрузку файлов .docx, анализ следующих "
        "параметров форматирования: шрифт, кегль, межстрочный интервал, поля "
        "страницы, выравнивание текста, оформление заголовков всех уровней, "
        "нумерация страниц, оформление списков и таблиц."
    )
    set_font(run, size=14)
    p.paragraph_format.first_line_indent = Cm(1.25)

    # Нарушение: список оформлен через тире, а не нумерованный
    for item in [
        "— загрузка файла через веб-интерфейс или API;",
        "— обработка документа в фоновом режиме;",
        "— формирование отчёта с указанием номеров страниц нарушений;",
        "— экспорт отчёта в PDF.",
    ]:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_font(run, size=14)
        p.paragraph_format.left_indent = Cm(1.25)

    # --- ПОДРАЗДЕЛ 1.2 ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("1.2 Анализ существующих решений")
    set_font(run, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        "На рынке существуют частичные решения данной задачи. "
        "Сервис «Антиплагиат» проверяет уникальность текста, но не форматирование. "
        "Плагин для Microsoft Word позволяет задать стилевой шаблон, однако "
        "требует установки ПО на каждом рабочем месте и не формирует подробных отчётов."
    )
    set_font(run, size=14)
    p.paragraph_format.first_line_indent = Cm(1.25)

    doc.add_paragraph()

    # Таблица с нарушениями оформления (нет номера и названия)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["Решение", "Проверка форматирования", "Отчёт"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=12, bold=True)
    rows_data = [
        ("Антиплагиат", "Нет", "Только уникальность"),
        ("Word-шаблон", "Частично", "Нет"),
        ("Предлагаемая система", "Да", "Подробный PDF-отчёт"),
    ]
    for r_idx, row_data in enumerate(rows_data):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            run = row.cells[c_idx].paragraphs[0].add_run(val)
            set_font(run, size=12)

    doc.add_paragraph()

    # --- ГЛАВА 2 ---
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2 ПРОЕКТИРОВАНИЕ СИСТЕМЫ")
    set_font(run, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        "Система реализована в виде микросервисной архитектуры и включает "
        "следующие компоненты: API-сервис (FastAPI), воркер обработки (Celery), "
        "база данных (PostgreSQL), кэш (Redis), фронтенд (React)."
    )
    set_font(run, size=14)
    p.paragraph_format.first_line_indent = Cm(1.25)

    # --- 2.1 ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("2.1 Архитектура")
    set_font(run, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        "Взаимодействие компонентов осуществляется через брокер сообщений. "
        "Клиент загружает файл через фронтенд, API сохраняет его в объектное "
        "хранилище и ставит задачу в очередь. Воркер извлекает задачу, "
        "анализирует документ и записывает результат в БД. "
        "Клиент опрашивает API до получения готового отчёта."
    )
    set_font(run, size=14)
    p.paragraph_format.first_line_indent = Cm(1.25)

    # --- 2.2 ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("2.2 Схема базы данных")
    set_font(run, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        "База данных содержит таблицы: users, documents, violations, reports. "
        "Таблица violations связана с documents отношением «многие-к-одному». "
        "Для хранения JSON-структуры нарушений используется тип JSONB PostgreSQL."
    )
    set_font(run, size=14)
    p.paragraph_format.first_line_indent = Cm(1.25)

    # --- ГЛАВА 3 ---
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("3 РЕАЛИЗАЦИЯ")
    set_font(run, size=14, bold=True)

    # --- 3.1 ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("3.1 Бэкенд")
    set_font(run, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        "API реализован на FastAPI. Для разбора .docx используется библиотека "
        "python-docx. Каждый параграф документа анализируется на соответствие "
        "правилам: шрифт Times New Roman 14 пт, межстрочный интервал 1,5, "
        "отступы полей: левое 3 см, правое 1,5 см, верхнее и нижнее 2 см."
    )
    set_font(run, size=14)
    p.paragraph_format.first_line_indent = Cm(1.25)

    # Нарушение: блок кода внутри обычного параграфа (без форматирования кода)
    p = doc.add_paragraph()
    run = p.add_run(
        "Пример кода определения заголовка: if paragraph.style.name.startswith('Heading'): ..."
    )
    set_font(run, size=12, italic=True)
    p.paragraph_format.left_indent = Cm(1.25)

    # --- 3.2 ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("3.2 Фронтенд")
    set_font(run, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        "Пользовательский интерфейс реализован на React с использованием "
        "библиотеки компонентов. Для отображения PDF-предпросмотра применяется "
        "pdf.js. Состояние приложения управляется через Zustand."
    )
    set_font(run, size=14)
    p.paragraph_format.first_line_indent = Cm(1.25)

    doc.add_page_break()

    # --- ЗАКЛЮЧЕНИЕ ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ЗАКЛЮЧЕНИЕ")
    set_font(run, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        "В ходе выполнения работы была разработана система автоматической "
        "проверки форматирования ВКР. Реализован алгоритм анализа .docx-файлов "
        "с выявлением нарушений по 15 критериям. Система прошла тестирование "
        "на 30 реальных ВКР, точность определения нарушений составила 94%."
    )
    set_font(run, size=14)
    p.paragraph_format.first_line_indent = Cm(1.25)

    doc.add_paragraph()

    # --- СПИСОК ИСТОЧНИКОВ ---
    # Нарушение: заголовок не CAPS
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Список использованных источников")
    set_font(run, size=14, bold=True)

    sources = [
        "1. ГОСТ Р 7.0.11-2011. Диссертация и автореферат диссертации. — М.: Стандартинформ, 2012.",
        "2. МИИГАиК. Положение о выпускных квалификационных работах. — М.: МИИГАиК, 2023.",
        "3. Python Software Foundation. python-docx documentation [Электронный ресурс]. "
        "URL: https://python-docx.readthedocs.io (дата обращения: 01.03.2025).",
        "4. FastAPI. FastAPI documentation [Электронный ресурс]. "
        "URL: https://fastapi.tiangolo.com (дата обращения: 05.03.2025).",
        "5. Hochreiter S., Schmidhuber J. Long Short-Term Memory // Neural Computation. "
        "— 1997. — Vol. 9, № 8. — P. 1735–1780.",
    ]
    for src in sources:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(src)
        set_font(run, size=14)
        # Нарушение: отступ у источников как у обычного текста
        p.paragraph_format.first_line_indent = Cm(1.25)

    doc.add_page_break()

    # --- ПРИЛОЖЕНИЕ А ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ПРИЛОЖЕНИЕ А")
    set_font(run, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Листинг основного модуля анализа")
    set_font(run, size=14)

    p = doc.add_paragraph()
    run = p.add_run(
        "def analyze_paragraph(para):\n"
        "    violations = []\n"
        "    if para.style.name == 'Normal':\n"
        "        if para.runs and para.runs[0].font.size != Pt(14):\n"
        "            violations.append('Неверный кегль')\n"
        "    return violations"
    )
    set_font(run, "Courier New", size=11)

    path = os.path.join(OUTPUT_DIR, "test_draft.docx")
    doc.save(path)
    print(f"Черновик сохранён: {path}")
    return path


# ──────────────────────────────────────────────
# ДОКУМЕНТ 2: ЧИСТОВОЙ ВАРИАНТ (эталон)
# ──────────────────────────────────────────────

def make_clean():
    doc = Document()
    page_margins(doc)

    # Настройка стилей Normal
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    from docx.shared import Pt as Pt2
    style.paragraph_format.line_spacing = Pt2(21)  # 1.5 × 14

    # --- Титульный лист ---
    def center(text, size=12, bold=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, size=size, bold=bold)
        return p

    center("Министерство науки и высшего образования Российской Федерации", 12)
    center("Федеральное государственное бюджетное образовательное учреждение\n"
           "высшего образования", 12)
    center("«Московский государственный университет геодезии и картографии»\n"
           "(МИИГАиК)", 12, bold=True)
    doc.add_paragraph()
    center("Факультет информационных технологий", 12)
    center("Кафедра информационных технологий и систем управления", 12)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("РАЗРАБОТКА СИСТЕМЫ АВТОМАТИЧЕСКОЙ ПРОВЕРКИ\n"
                  "ФОРМАТИРОВАНИЯ ВЫПУСКНЫХ КВАЛИФИКАЦИОННЫХ РАБОТ")
    set_font(r, size=14, bold=True)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Пояснительная записка к выпускной квалификационной работе")
    set_font(r, size=14)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Студент и руководитель — по правому краю с пробелами (ГОСТ)
    for line in [
        "Студент группы ИТ4-31         Иванов И.И.  _________",
        "Научный руководитель,",
        "профессор, д.т.н.              Петров П.П.  _________",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(line)
        set_font(r, size=12)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    center("Москва 2025", 12)
    doc.add_page_break()

    # --- АННОТАЦИЯ ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("АННОТАЦИЯ")
    set_font(r, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(
        "В данной работе рассматривается разработка веб-сервиса для автоматической "
        "проверки форматирования выпускных квалификационных работ (ВКР) "
        "в соответствии с требованиями МИИГАиК. Система позволяет выявлять "
        "отклонения от норм оформления, сокращая время ручной проверки "
        "преподавателями. Ключевые слова: форматирование, ВКР, автоматизация, "
        "python-docx, FastAPI."
    )
    set_font(r, size=14)
    p.paragraph_format.first_line_indent = Cm(1.25)
    from docx.shared import Pt as Pt3
    p.paragraph_format.line_spacing = Pt3(21)

    doc.add_page_break()

    # --- СОДЕРЖАНИЕ ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("СОДЕРЖАНИЕ")
    set_font(r, size=14, bold=True)

    toc_items = [
        ("Введение", "4"),
        ("1 Постановка задачи", "5"),
        ("   1.1 Требования к системе", "6"),
        ("   1.2 Анализ существующих решений", "8"),
        ("2 Проектирование системы", "10"),
        ("   2.1 Архитектура", "11"),
        ("   2.2 Схема базы данных", "13"),
        ("3 Реализация", "15"),
        ("   3.1 Бэкенд", "15"),
        ("   3.2 Фронтенд", "18"),
        ("Заключение", "22"),
        ("Список использованных источников", "23"),
        ("Приложение А", "25"),
    ]
    for title, page in toc_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(title)
        set_font(r, size=14)
        # Табуляция с лидером
        tab_stop = OxmlElement('w:tab')
        tab_stop.set(qn('w:val'), 'right')
        tab_stop.set(qn('w:pos'), '8800')
        tab_stop.set(qn('w:leader'), 'dot')
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tabs.append(tab_stop)
        pPr.append(tabs)
        r2 = p.add_run("\t" + page)
        set_font(r2, size=14)

    doc.add_page_break()

    # --- ВВЕДЕНИЕ ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ВВЕДЕНИЕ")
    set_font(r, size=14, bold=True)

    def body(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(text)
        set_font(r, size=14)
        p.paragraph_format.first_line_indent = Cm(1.25)
        from docx.shared import Pt as Pt4
        p.paragraph_format.line_spacing = Pt4(21)
        return p

    body(
        "Актуальность темы определяется необходимостью автоматизации рутинных "
        "операций проверки соответствия оформления ВКР нормативным требованиям. "
        "В настоящее время данная процедура выполняется вручную, что приводит к "
        "значительным временным затратам и субъективности оценки."
    )

    body(
        "Цель работы — разработать веб-сервис, который принимает файл ВКР в "
        "формате .docx и возвращает структурированный отчёт о нарушениях форматирования."
    )

    body(
        "Для достижения указанной цели необходимо решить следующие задачи:\n"
        "1) разработать алгоритм анализа структуры документа;\n"
        "2) реализовать серверную часть с REST API;\n"
        "3) разработать пользовательский интерфейс;\n"
        "4) провести тестирование на реальных документах."
    )

    doc.add_page_break()

    # --- ГЛАВА 1 ---
    def heading1(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, size=14, bold=True)
        return p

    def heading2(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        set_font(r, size=14, bold=True)
        return p

    heading1("1 ПОСТАНОВКА ЗАДАЧИ")

    body(
        "Для достижения цели работы необходимо решить следующие задачи: "
        "разработать алгоритм анализа структуры документа, реализовать "
        "серверную часть с REST API, создать пользовательский интерфейс. "
        "Система должна обрабатывать документы объёмом до 200 страниц "
        "за время не более 60 секунд."
    )

    heading2("1.1 Требования к системе")

    body(
        "Система должна обеспечивать загрузку файлов .docx, анализ следующих "
        "параметров форматирования: шрифт, кегль, межстрочный интервал, поля "
        "страницы, выравнивание текста, оформление заголовков всех уровней, "
        "нумерация страниц, оформление списков и таблиц."
    )

    body("Функциональные требования:")

    # Правильный нумерованный список с отступом
    for i, item in enumerate([
        "загрузка файла через веб-интерфейс или API;",
        "обработка документа в фоновом режиме;",
        "формирование отчёта с указанием номеров страниц нарушений;",
        "экспорт отчёта в PDF.",
    ], 1):
        p = doc.add_paragraph()
        r = p.add_run(f"{i}) {item}")
        set_font(r, size=14)
        p.paragraph_format.left_indent = Cm(1.25)

    heading2("1.2 Анализ существующих решений")

    body(
        "На рынке существуют частичные решения данной задачи. "
        "Сервис «Антиплагиат» проверяет уникальность текста, но не форматирование. "
        "Плагин для Microsoft Word позволяет задать стилевой шаблон, однако "
        "требует установки ПО на каждом рабочем месте и не формирует подробных отчётов."
    )

    # Правильная таблица с подписью
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Таблица 1.1 — Сравнение существующих решений")
    set_font(r, size=14)

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["Решение", "Проверка форматирования", "Отчёт"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cell.paragraphs[0].add_run(h)
        set_font(r, size=12, bold=True)
    rows_data = [
        ("Антиплагиат", "Нет", "Только уникальность"),
        ("Word-шаблон", "Частично", "Нет"),
        ("Предлагаемая система", "Да", "Подробный PDF-отчёт"),
    ]
    for r_idx, row_data in enumerate(rows_data):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            row.cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = row.cells[c_idx].paragraphs[0].add_run(val)
            set_font(r, size=12)

    doc.add_paragraph()

    doc.add_page_break()

    # --- ГЛАВА 2 ---
    heading1("2 ПРОЕКТИРОВАНИЕ СИСТЕМЫ")

    body(
        "Система реализована в виде микросервисной архитектуры и включает "
        "следующие компоненты: API-сервис (FastAPI), воркер обработки (Celery), "
        "база данных (PostgreSQL), кэш (Redis), фронтенд (React). "
        "Выбор микросервисного подхода обусловлен требованиями к масштабируемости."
    )

    heading2("2.1 Архитектура")

    body(
        "Взаимодействие компонентов осуществляется через брокер сообщений. "
        "Клиент загружает файл через фронтенд, API сохраняет его в объектное "
        "хранилище и ставит задачу в очередь. Воркер извлекает задачу, "
        "анализирует документ и записывает результат в БД. "
        "Клиент опрашивает API до получения готового отчёта."
    )

    # Подпись к рисунку
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[Рисунок 2.1 — Схема архитектуры системы]")
    set_font(r, size=14, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Рисунок 2.1 — Схема архитектуры системы")
    set_font(r, size=14)

    heading2("2.2 Схема базы данных")

    body(
        "База данных содержит таблицы: users, documents, violations, reports. "
        "Таблица violations связана с documents отношением «многие-к-одному». "
        "Для хранения JSON-структуры нарушений используется тип JSONB PostgreSQL, "
        "что позволяет эффективно индексировать поля отчёта."
    )

    doc.add_page_break()

    # --- ГЛАВА 3 ---
    heading1("3 РЕАЛИЗАЦИЯ")

    heading2("3.1 Бэкенд")

    body(
        "API реализован на FastAPI версии 0.111. Для разбора .docx используется "
        "библиотека python-docx 1.1.2. Каждый параграф документа анализируется "
        "на соответствие правилам: шрифт Times New Roman 14 пт, "
        "межстрочный интервал 1,5, отступы полей: левое 3 см, правое 1,5 см, "
        "верхнее и нижнее 2 см."
    )

    heading2("3.2 Фронтенд")

    body(
        "Пользовательский интерфейс реализован на React 18 с использованием "
        "библиотеки компонентов Headless UI. Для отображения PDF-предпросмотра "
        "применяется pdf.js версии 4.2. Состояние приложения управляется "
        "через Zustand 4.5."
    )

    doc.add_page_break()

    # --- ЗАКЛЮЧЕНИЕ ---
    heading1("ЗАКЛЮЧЕНИЕ")

    body(
        "В ходе выполнения работы была разработана система автоматической "
        "проверки форматирования ВКР. Реализован алгоритм анализа .docx-файлов "
        "с выявлением нарушений по 15 критериям."
    )

    body(
        "Система прошла тестирование на 30 реальных ВКР, точность определения "
        "нарушений составила 94%. Время обработки документа объёмом 100 страниц "
        "не превышает 18 секунд, что соответствует требованиям."
    )

    doc.add_page_break()

    # --- СПИСОК ИСТОЧНИКОВ ---
    heading1("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")

    sources = [
        "1. ГОСТ Р 7.0.11-2011. Диссертация и автореферат диссертации. — М.: Стандартинформ, 2012. — 12 с.",
        "2. МИИГАиК. Положение о выпускных квалификационных работах. — М.: МИИГАиК, 2023. — 18 с.",
        "3. Python Software Foundation. python-docx documentation [Электронный ресурс]. "
        "URL: https://python-docx.readthedocs.io (дата обращения: 01.03.2025).",
        "4. FastAPI. FastAPI documentation [Электронный ресурс]. "
        "URL: https://fastapi.tiangolo.com (дата обращения: 05.03.2025).",
        "5. Hochreiter S., Schmidhuber J. Long Short-Term Memory // Neural Computation. "
        "— 1997. — Vol. 9, № 8. — P. 1735–1780.",
    ]
    for src in sources:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(src)
        set_font(r, size=14)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0)
        from docx.shared import Pt as Pt5
        p.paragraph_format.line_spacing = Pt5(21)

    doc.add_page_break()

    # --- ПРИЛОЖЕНИЕ А ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ПРИЛОЖЕНИЕ А")
    set_font(r, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("(обязательное)")
    set_font(r, size=14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Листинг основного модуля анализа")
    set_font(r, size=14, bold=True)

    code_lines = [
        "def analyze_paragraph(para: Paragraph) -> list[str]:",
        "    violations = []",
        "    if para.style.name == 'Normal':",
        "        for run in para.runs:",
        "            if run.font.size and run.font.size != Pt(14):",
        "                violations.append('Неверный кегль')",
        "    return violations",
    ]
    for line in code_lines:
        p = doc.add_paragraph()
        r = p.add_run(line)
        set_font(r, "Courier New", size=11)
        p.paragraph_format.left_indent = Cm(1.25)

    path = os.path.join(OUTPUT_DIR, "test_clean.docx")
    doc.save(path)
    print(f"Чистовой вариант сохранён: {path}")
    return path


if __name__ == "__main__":
    make_draft()
    make_clean()
    print("Готово.")
