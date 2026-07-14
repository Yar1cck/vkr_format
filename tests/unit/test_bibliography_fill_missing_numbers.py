"""Источникам без номера дописываются номера, продолжая существующую
последовательность; уже пронумерованные записи не трогаются.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document as WordDocument

from services.core.vkr_core.config.normative_loader import load_default_rules
from services.core.vkr_core.engine.bibliography import (
    apply_bibliography_numbering,
    collect_bibliography,
    validate_bibliography,
)
from services.core.vkr_core.engine.stats import collect_stats


def _build(tmp_path: Path) -> WordDocument:
    doc = WordDocument()
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("1. Иванов И. И. Анализ. – М.: Наука, 2020. – 200 с.")
    doc.add_paragraph("2. Петров П. П. Методы. – СПб.: Питер, 2021. – 150 с.")
    # Без номеров (как на «Фото 2»):
    doc.add_paragraph("Программа развития МИИГАиК на 2024-2027 гг. – М., 2024.")
    doc.add_paragraph("Информационно-аналитические материалы. – М., 2024.")
    doc.add_paragraph("Положение о ВКР МГТУ «СТАНКИН». – М., 2023.")
    path = tmp_path / "doc.docx"
    doc.save(path)
    return WordDocument(str(path))


def test_missing_numbers_are_filled_continuing_sequence(tmp_path: Path) -> None:
    doc = _build(tmp_path)
    rules = load_default_rules()
    stats = collect_stats(doc)

    fixed = apply_bibliography_numbering(doc, stats, rules, headings=[])

    assert fixed == [3, 4, 5]
    # Существующие номера не тронуты.
    assert doc.paragraphs[1].text.startswith("1. Иванов")
    assert doc.paragraphs[2].text.startswith("2. Петров")
    # Пропущенным дописаны 3, 4, 5.
    assert doc.paragraphs[3].text.startswith("3. Программа развития")
    assert doc.paragraphs[4].text.startswith("4. Информационно")
    assert doc.paragraphs[5].text.startswith("5. Положение о ВКР")

    # После правки список распознаётся сквозным 1..5.
    stats = collect_stats(doc)
    entries, bounds, malformed = collect_bibliography(stats, rules, headings=[])
    assert [e.number for e in entries] == [1, 2, 3, 4, 5]
    assert malformed == []


def test_fully_unnumbered_list_gets_numbered(tmp_path: Path) -> None:
    doc = WordDocument()
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("Иванов И. И. Анализ. – М.: Наука, 2020. – 200 с.")
    doc.add_paragraph("Петров П. П. Методы. – СПб.: Питер, 2021. – 150 с.")
    path = tmp_path / "doc.docx"
    doc.save(path)
    doc = WordDocument(str(path))

    rules = load_default_rules()
    stats = collect_stats(doc)
    fixed = apply_bibliography_numbering(doc, stats, rules, headings=[])

    assert fixed == [1, 2]
    assert doc.paragraphs[1].text.startswith("1. Иванов")
    assert doc.paragraphs[2].text.startswith("2. Петров")


def test_continuation_line_is_not_numbered(tmp_path: Path) -> None:
    doc = WordDocument()
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("Иванов И. И. Анализ данных и обработка. – М.: Наука,")
    # Строка-продолжение (строчная буква в начале) — не запись.
    doc.add_paragraph("2020. – 200 с.")
    path = tmp_path / "doc.docx"
    doc.save(path)
    doc = WordDocument(str(path))

    rules = load_default_rules()
    stats = collect_stats(doc)
    fixed = apply_bibliography_numbering(doc, stats, rules, headings=[])

    assert fixed == [1]
    assert doc.paragraphs[1].text.startswith("1. Иванов")
    assert doc.paragraphs[2].text.startswith("2020. – 200 с.")


def test_journal_issue_continuation_is_not_bibliography_number(tmp_path: Path) -> None:
    doc = WordDocument()
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("1. Иванов И. И. Анализ данных. – М.: Наука, 2020.")
    doc.add_paragraph("2. Петров П. П. Методы обработки. – СПб.: Питер, 2021.")
    # Продолжение журнальной записи после склейки/переноса. Это не источник №5.
    doc.add_paragraph("5. № 12-1 (81). С. 71-77. [Электронный ресурс]. URL: https://example.org")
    path = tmp_path / "doc.docx"
    doc.save(path)
    doc = WordDocument(str(path))

    rules = load_default_rules()
    stats = collect_stats(doc)
    entries, bounds, malformed = collect_bibliography(stats, rules, headings=[])
    violations = validate_bibliography(entries, citations=[], bounds=bounds, malformed=malformed)

    assert [e.number for e in entries] == [1, 2]
    assert not [v for v in violations if v.type == "bibliography_numbering_gap"]
