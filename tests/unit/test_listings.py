"""Тесты listings.py — обнаружение тела листингов кода (§7.3 ext)."""
from __future__ import annotations

from pathlib import Path

from docx import Document as WordDocument
from docx.shared import Pt

from services.core.vkr_core.engine.listings import (
    _NUMBERED_LINE_RE,
    detect_listing_code_indexes,
    renumber_listings,
)


def _build(tmp_path: Path, items) -> WordDocument:
    """items: list of (text, font_name | None)."""
    doc = WordDocument()
    for text, font in items:
        p = doc.add_paragraph()
        run = p.add_run(text)
        if font:
            run.font.name = font
            run.font.size = Pt(10)
    out = tmp_path / "doc.docx"
    doc.save(out)
    return WordDocument(str(out))


def test_numbered_line_regex() -> None:
    assert _NUMBERED_LINE_RE.match("1. print(x)")
    assert _NUMBERED_LINE_RE.match("12) return")
    assert _NUMBERED_LINE_RE.match("3 import os")
    assert not _NUMBERED_LINE_RE.match("обычный текст")


def test_detect_listing_code_via_monospace(tmp_path: Path) -> None:
    doc = _build(tmp_path, [
        ("Листинг 1 — пример кода", None),
        ("def foo():", "Courier New"),
        ("    return 42", "Courier New"),
        ("Обычный абзац после кода.", None),
    ])
    indexes = detect_listing_code_indexes(doc.paragraphs, skip_indexes=set())
    # 1, 2 — строки кода; 0 — подпись; 3 — обычный текст
    assert 1 in indexes and 2 in indexes
    assert 0 not in indexes and 3 not in indexes


def test_detect_listing_code_via_numbered_lines(tmp_path: Path) -> None:
    # Нумерованные строки без моношрифта определяются как код только если
    # содержимое после числа начинается с ключевого слова языка или {}.
    # «y = 5» без шрифта неотличимо от «где N = 5», поэтому не детектируется.
    doc = _build(tmp_path, [
        ("Листинг — без номера", None),
        ("1. def foo():", None),   # ключевое слово после числа
        ("2. return 42", None),    # ключевое слово после числа
        ("3. y = 5", None),        # присваивание без шрифта — не детектируется
        ("Обычный текст.", None),
    ])
    indexes = detect_listing_code_indexes(doc.paragraphs, skip_indexes=set())
    assert 1 in indexes and 2 in indexes
    assert 3 not in indexes  # y = 5 без моношрифта неотличимо от формулы в тексте
    assert 4 not in indexes


def test_detect_listing_stops_at_non_code(tmp_path: Path) -> None:
    doc = _build(tmp_path, [
        ("Листинг 1 — код", None),
        ("def foo():", "Courier New"),
        ("Дальше пошёл обычный текст.", None),
        ("def bar():", "Courier New"),  # отдельный блок (без подписи перед ним)
    ])
    indexes = detect_listing_code_indexes(doc.paragraphs, skip_indexes=set())
    # 1 — точно код (после подписи), 3 — НЕ должен попасть в индексы
    assert 1 in indexes
    assert 3 not in indexes


def test_skip_indexes_excluded(tmp_path: Path) -> None:
    doc = _build(tmp_path, [
        ("Листинг 1 — код", None),
        ("def foo():", "Courier New"),
    ])
    indexes = detect_listing_code_indexes(doc.paragraphs, skip_indexes={0, 1})
    assert indexes == set()


def test_renumber_listings_sequential(tmp_path: Path) -> None:
    doc = _build(tmp_path, [
        ("Листинг 5 — первый", None),
        ("def a(): pass", "Courier New"),
        ("Текст.", None),
        ("Листинг — второй", None),  # без номера
        ("def b(): pass", "Courier New"),
    ])
    listings, changes = renumber_listings(doc.paragraphs, skip_indexes=set(), apply=True)
    # Должна быть сквозная нумерация 1, 2 (а не 5)
    nums = sorted(int(str(c.new_number)) for c in listings if str(c.new_number).isdigit())
    assert nums == [1, 2]


def test_renumber_listings_by_chapter_and_appendix(tmp_path: Path) -> None:
    doc = _build(tmp_path, [
        ("1 Первая глава", None),
        ("Листинг 5 — первый", None),
        ("def a(): pass", "Courier New"),
        ("2 Вторая глава", None),
        ("Листинг — второй", None),
        ("def b(): pass", "Courier New"),
        ("Приложение А", None),
        ("Листинг 9 — приложение", None),
        ("def c(): pass", "Courier New"),
    ])
    listings, _changes = renumber_listings(
        doc.paragraphs,
        skip_indexes=set(),
        apply=True,
        chapter_breaks=[(0, 1), (3, 2)],
    )
    assert [str(item.new_number) for item in listings] == ["1.1", "2.1", "А.1"]


def test_prose_with_listing_word_is_not_treated_as_caption(tmp_path: Path) -> None:
    doc = _build(tmp_path, [
        ("Листинг 3.2 показывает структуру модуля.", None),
        ("Обычный абзац текста.", None),
        ("Листинг 8 — Реальный пример", None),
        ("def real():", "Courier New"),
    ])
    listings, _changes = renumber_listings(doc.paragraphs, skip_indexes=set(), apply=True)
    assert [item.paragraph_index for item in listings] == [2]
    assert doc.paragraphs[0].text == "Листинг 3.2 показывает структуру модуля."


def test_post_listing_paragraphs_not_captured(tmp_path: Path) -> None:
    """Абзацы сразу после блока кода не должны захватываться в листинг."""
    doc = _build(tmp_path, [
        ("Листинг 1 — сортировка", None),
        ("def sort(arr):", "Courier New"),
        ("    return sorted(arr)", "Courier New"),
        # Следующие абзацы — обычный текст с унаследованным или явным Courier New
        ("где N = количество элементов", "Courier New"),   # формула с «=»
        ("1. Иванов И.И. Алгоритмы. М.: 2020.", None),     # библиография
        ("2. Петров П.П. Структуры данных. М.: 2021.", None),
    ])
    indexes = detect_listing_code_indexes(doc.paragraphs, skip_indexes=set())
    assert 1 in indexes and 2 in indexes   # строки кода
    assert 3 not in indexes  # «где N = …» с Courier New — не код
    assert 4 not in indexes  # библиографическая ссылка — не код
    assert 5 not in indexes


def test_plain_code_after_listing_caption_is_detected(tmp_path: Path) -> None:
    doc = _build(tmp_path, [
        ("Листинг 3.1 — Обработчик запроса", None),
        ("def handle_request(data):", None),
        ("    return data['id']", None),
    ])
    listings, _changes = renumber_listings(
        doc.paragraphs,
        skip_indexes=set(),
        apply=True,
        chapter_breaks=[(0, 3)],
    )
    assert [str(item.new_number) for item in listings] == ["3.1"]
    assert doc.paragraphs[0].text == "Листинг 3.1 — Обработчик запроса"
