"""Автоперенос слов в ВКР недопустим: конвейер должен его отключать.

Раньше pipeline ВКЛЮЧАЛ автоперенос (под узкие ячейки таблиц), теперь явно
выставляет `w:autoHyphenation w:val="false"` и убирает consecutiveHyphenLimit.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document as WordDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from services.core.vkr_core.config.normative_loader import load_default_rules
from services.core.vkr_core.engine import process_document
from services.core.vkr_core.engine.formatter import disable_document_hyphenation


def test_disable_hyphenation_sets_false_and_drops_limit() -> None:
    doc = WordDocument()
    settings = doc.settings.element
    # Эмулируем шаблон с уже включённым автопереносом.
    settings.append(OxmlElement("w:autoHyphenation"))
    consec = OxmlElement("w:consecutiveHyphenLimit")
    consec.set(qn("w:val"), "3")
    settings.append(consec)

    disable_document_hyphenation(doc)

    auto = settings.find(qn("w:autoHyphenation"))
    assert auto is not None
    assert auto.get(qn("w:val")) == "false"
    assert settings.find(qn("w:consecutiveHyphenLimit")) is None


def test_pipeline_disables_hyphenation(tmp_path: Path) -> None:
    doc = WordDocument()
    doc.add_paragraph("Выпускная квалификационная работа")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Основной текст раздела с достаточно длинными словами.")
    source = tmp_path / "source.docx"
    doc.save(source)

    result = process_document(
        source_path=source,
        rules=load_default_rules(),
        check_only=False,
    )

    processed = WordDocument(str(result.processed_docx_path))
    settings = processed.settings.element
    auto = settings.find(qn("w:autoHyphenation"))
    # Либо явный false, либо элемента нет — оба варианта = перенос выключен.
    assert auto is None or auto.get(qn("w:val")) == "false"
    assert settings.find(qn("w:consecutiveHyphenLimit")) is None
