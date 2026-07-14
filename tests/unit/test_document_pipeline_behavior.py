from __future__ import annotations

from pathlib import Path

from docx import Document as WordDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

from services.core.vkr_core.config.normative_loader import load_default_rules
from services.core.vkr_core.engine import process_document
from services.core.vkr_core.engine.formatter import apply_body_style
from services.core.vkr_core.engine.stats import collect_stats


def _write_doc(path: Path, paragraphs: list[str]) -> Path:
    doc = WordDocument()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)
    return path


def _write_cover_like_doc(path: Path) -> Path:
    doc = WordDocument()
    for text in (
        "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ",
        "Московский государственный университет геодезии и картографии",
        "КУРСОВАЯ РАБОТА",
    ):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    intro = doc.add_paragraph("ВВЕДЕНИЕ")
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.runs[0].bold = True
    doc.add_paragraph("Основной текст раздела.")
    doc.save(path)
    return path


def test_title_page_is_preserved(tmp_path: Path) -> None:
    source = _write_doc(
        tmp_path / "source.docx",
        [
            "Выпускная квалификационная работа",
            "На правах рукописи",
            "",
            "ВВЕДЕНИЕ",
            "Текст раздела.",
        ],
    )
    result = process_document(
        source_path=source,
        rules=load_default_rules(),
        check_only=False,
    )
    processed = WordDocument(str(result.processed_docx_path))
    texts = [p.text for p in processed.paragraphs]
    assert "На правах рукописи" in texts


def test_check_only_does_not_apply_renumbering(tmp_path: Path) -> None:
    source = _write_doc(
        tmp_path / "source_check_only.docx",
        [
            "Выпускная квалификационная работа",
            "",
            "ВВЕДЕНИЕ",
            "Текст [5] и [9]",
            "Рисунок 7 – Схема",
            "Таблица 5 – Данные",
        ],
    )
    result = process_document(
        source_path=source,
        rules=load_default_rules(),
        check_only=True,
    )
    types = {v.type for v in result.violations}
    assert "figures_renumbered" not in types
    assert "tables_renumbered" not in types

    processed = WordDocument(str(result.processed_docx_path))
    texts = [p.text for p in processed.paragraphs]
    # check_only не мутирует документ — подписи остаются как во вводе («–»).
    assert "Рисунок 7 – Схема" in texts
    assert "Таблица 5 – Данные" in texts


def test_captions_are_renumbered_in_full_mode(tmp_path: Path) -> None:
    source = _write_doc(
        tmp_path / "source_renumber.docx",
        [
            "Выпускная квалификационная работа",
            "",
            "ВВЕДЕНИЕ",
            "Текст до рисунков. См. рис. 1 и рис. 2.",
            "Рисунок 7 – Схема",
            "Между рисунками.",
            "Рисунок 9 – Вторая схема",
        ],
    )
    result = process_document(
        source_path=source,
        rules=load_default_rules(),
        check_only=False,
    )
    processed = WordDocument(str(result.processed_docx_path))
    texts = [p.text for p in processed.paragraphs]
    assert any(t.startswith("Рисунок 1") for t in texts)
    assert any(t.startswith("Рисунок 2") for t in texts)


def test_cover_paragraphs_are_not_reclassified_as_headings(tmp_path: Path) -> None:
    source = _write_cover_like_doc(tmp_path / "source_cover.docx")
    result = process_document(
        source_path=source,
        rules=load_default_rules(),
        check_only=False,
    )
    bad = [
        v
        for v in result.violations
        if v.section_title
        in {
            "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ",
            "Московский государственный университет геодезии и картографии",
            "КУРСОВАЯ РАБОТА",
        }
    ]
    assert not bad


def test_missing_references_section_raises_violation(tmp_path: Path) -> None:
    source = _write_doc(
        tmp_path / "source_no_refs.docx",
        [
            "Выпускная квалификационная работа",
            "",
            "ВВЕДЕНИЕ",
            "Текст работы без раздела источников.",
            "ЗАКЛЮЧЕНИЕ",
            "Итоги.",
        ],
    )
    result = process_document(
        source_path=source,
        rules=load_default_rules(),
        check_only=True,
    )
    types = {v.type for v in result.violations}
    assert "bibliography_missing" in types or "missing_structural_section" in types


def test_bibliography_malformed_number_is_flagged(tmp_path: Path) -> None:
    """Запись с битым номером-диапазоном «18-32» больше не теряется молча —
    выдаётся нарушение bibliography_entry_malformed (п.6.16 МИИГАиК)."""
    source = _write_doc(
        tmp_path / "bib_malformed.docx",
        [
            "Выпускная квалификационная работа",
            "",
            "ВВЕДЕНИЕ",
            "В работе [1] и [2] показано решение.",
            "ЗАКЛЮЧЕНИЕ",
            "Итоги.",
            "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
            "1. Иванов И. И. Первый источник. — Москва : Изд-во, 2024. — 100 с.",
            "18-32 Петров П. П. Битый диапазон. — Москва : Изд-во, 2023. — 50 с.",
            "2. Сидоров С. С. Второй источник. — Москва : Изд-во, 2025. — 80 с.",
        ],
    )
    result = process_document(
        source_path=source,
        rules=load_default_rules(),
        check_only=True,
    )
    types = {v.type for v in result.violations}
    assert "bibliography_entry_malformed" in types


def test_table_reference_sentence_gets_body_style() -> None:
    doc = WordDocument()
    paragraph = doc.add_paragraph(
        "Таблица 1.1 показывает, что часть требований имеет измеримый характер."
    )
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rules = load_default_rules()
    stats = collect_stats(doc)

    apply_body_style(doc, rules, stats, skip_indexes=set(), heading_indexes=set())

    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert abs(
        paragraph.paragraph_format.first_line_indent.cm
        - rules["body_text_style"]["first_line_indent_cm"]
    ) < 0.01


def test_check_only_flags_centered_body_table_reference(tmp_path: Path) -> None:
    doc = WordDocument()
    doc.add_paragraph("ВВЕДЕНИЕ")
    paragraph = doc.add_paragraph(
        "Таблица 1.1 показывает, что часть требований имеет измеримый характер."
    )
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source = tmp_path / "centered_body.docx"
    doc.save(source)

    result = process_document(
        source_path=source,
        rules=load_default_rules(),
        check_only=True,
    )

    assert any(v.type == "paragraph_spacing" for v in result.violations)
