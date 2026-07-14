"""Регресс: ручное оглавление голыми заголовками не должно рушить структуру.

minimal_vkr.docx: после «СОДЕРЖАНИЕ» идут голые строки ВВЕДЕНИЕ /
1 Основная часть / ЗАКЛЮЧЕНИЕ / СПИСОК… без точек и номеров страниц.
Раньше detect_toc_section_indexes возвращал пусто (первая строка —
структурный алиас → граница сразу), строки оглавления уходили в
заголовки, реальные разделы дублировались/демотировались.

Признак ручного оглавления — строки дословно повторяются ниже.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document as WordDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from services.core.vkr_core.config.normative_loader import load_default_rules
from services.core.vkr_core.engine.detection import detect_headings_full
from services.core.vkr_core.engine.stats import collect_stats
from services.core.vkr_core.engine.title_page import detect_title_page_end
from services.core.vkr_core.engine.toc import (
    _is_toc_line,
    _norm,
    detect_toc_section_indexes,
    insert_toc,
)


def _set_outline(paragraph, level: int = 0) -> None:
    """Проставляет w:outlineLvl — как делает formatter перед вставкой TOC."""
    p = paragraph._p
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    for old in list(pPr.findall(qn("w:outlineLvl"))):
        pPr.remove(old)
    olvl = OxmlElement("w:outlineLvl")
    olvl.set(qn("w:val"), str(level))
    pPr.append(olvl)


def _heading(doc, text: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    for r in list(p.runs):
        r.text = ""
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    _set_outline(p, 0)


def _bare(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(text).font.size = Pt(14)


def _body(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text).font.size = Pt(14)


def _build(tmp_path: Path) -> WordDocument:
    doc = WordDocument()
    doc.add_paragraph().add_run("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА")
    doc.add_paragraph().add_run("Москва 2026")
    _heading(doc, "СОДЕРЖАНИЕ")
    # Ручное оглавление — голые названия разделов (повторяются ниже).
    for t in ("ВВЕДЕНИЕ", "1 Основная часть", "ЗАКЛЮЧЕНИЕ",
              "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"):
        _bare(doc, t)
    _heading(doc, "ВВЕДЕНИЕ")
    _body(doc, "Вводный текст работы достаточной длины для тела документа. " * 3)
    _heading(doc, "1 Основная часть")
    _body(doc, "Основной текст исследования достаточной длины для тела. " * 3)
    _heading(doc, "ЗАКЛЮЧЕНИЕ")
    _body(doc, "Заключительный текст работы достаточной длины для тела. " * 3)
    _heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    _body(doc, "1. Иванов И. И. Пособие. — М.: МИИГАиК, 2025. — 120 с.")
    path = tmp_path / "doc.docx"
    doc.save(path)
    return WordDocument(str(path))


def test_manual_toc_region_detected_and_headings_intact(tmp_path: Path) -> None:
    doc = _build(tmp_path)
    rules = load_default_rules()
    stats = collect_stats(doc)
    title_end = detect_title_page_end(stats, rules)

    toc = detect_toc_section_indexes(doc, rules)
    # «СОДЕРЖАНИЕ» = idx 2; ручное оглавление = idx 3..6.
    assert toc == {3, 4, 5, 6}, toc

    heads, soft = detect_headings_full(stats, rules, set(range(title_end)) | toc)
    htexts = [h.text.strip() for h in heads]
    # Реальные разделы приняты как заголовки (каждый РОВНО один раз).
    for name in ("ВВЕДЕНИЕ", "1 Основная часть", "ЗАКЛЮЧЕНИЕ",
                 "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"):
        assert htexts.count(name) == 1, (name, htexts)
    # Никаких possible_missed_heading / heading_confirm на реальные разделы.
    assert not [
        v for v in soft
        if v.type in ("possible_missed_heading", "heading_confirm")
    ], [(v.type, v.section_title) for v in soft]


def test_insert_toc_removes_bare_manual_toc(tmp_path: Path) -> None:
    doc = _build(tmp_path)
    rules = load_default_rules()
    assert insert_toc(doc, rules, title_end_index=0) is True

    names = {"введение", "1 основная часть", "заключение",
             "список использованных источников"}
    for p in doc.paragraphs:
        if _norm(p.text) not in names:
            continue
        is_heading = (p.style.name or "").startswith("Heading")
        is_toc_entry = "PAGEREF" in p._p.xml or "_VkrToc_" in p._p.xml
        # Каждое вхождение названия раздела — либо реальный заголовок, либо
        # сгенерированная строка оглавления. Голой строки ручного
        # оглавления (Normal без PAGEREF) остаться не должно.
        assert is_heading or is_toc_entry, (
            f"остался висячий дубль ручного оглавления: {p.text!r} "
            f"(style={p.style.name})"
        )


def _build_toc_with_dotted_page_numbers(tmp_path: Path) -> WordDocument:
    doc = WordDocument()
    _heading(doc, "СОДЕРЖАНИЕ")
    # Ручное оглавление из реального файла: часть номеров страниц записана
    # как "5." / "32 .", а последняя строка совпадает со структурным разделом.
    for text in (
        "ВВЕДЕНИЕ\t4",
        "1 Анализ предметной области\t5.",
        "3 Разработка и тестирование\t32 .",
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ\t44.",
    ):
        _bare(doc, text)
    _heading(doc, "ВВЕДЕНИЕ")
    _body(doc, "Вводный текст работы достаточной длины для тела документа. " * 3)
    _heading(doc, "1 Анализ предметной области")
    _body(doc, "Основной текст исследования достаточной длины для тела. " * 3)
    _heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    _body(doc, "1. Иванов И. И. Пособие. — М.: МИИГАиК, 2025. — 120 с.")
    path = tmp_path / "toc_dotted_pages.docx"
    doc.save(path)
    return WordDocument(str(path))


def test_toc_lines_accept_page_number_with_trailing_dot() -> None:
    assert _is_toc_line("1 Анализ предметной области\t5.")
    assert _is_toc_line("3 Разработка и тестирование\t32 .")
    assert _is_toc_line("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ\t44.")


def test_dotted_page_toc_region_keeps_references_out_of_headings(tmp_path: Path) -> None:
    doc = _build_toc_with_dotted_page_numbers(tmp_path)
    rules = load_default_rules()
    stats = collect_stats(doc)
    title_end = detect_title_page_end(stats, rules)

    toc = detect_toc_section_indexes(doc, rules)
    assert toc == {1, 2, 3, 4}, toc

    heads, _ = detect_headings_full(stats, rules, set(range(title_end)) | toc)
    htexts = [h.text.strip() for h in heads]
    assert htexts.count("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ") == 1
    assert htexts.index("ВВЕДЕНИЕ") < htexts.index("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")


def test_insert_toc_removes_dotted_page_manual_toc(tmp_path: Path) -> None:
    doc = _build_toc_with_dotted_page_numbers(tmp_path)
    rules = load_default_rules()

    assert insert_toc(doc, rules, title_end_index=0) is True

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text or "\t" not in text:
            continue
        is_generated_toc_entry = "PAGEREF" in p._p.xml or "_VkrToc_" in p._p.xml
        assert is_generated_toc_entry, f"осталась старая ручная строка TOC: {text!r}"
