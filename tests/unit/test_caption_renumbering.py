from __future__ import annotations

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

from services.core.vkr_core.engine.captions import renumber_figures, renumber_tables


def test_renumber_figure_preserves_drawing_nodes() -> None:
    doc = Document()
    p = doc.add_paragraph("Рисунок 5 — Схема")
    run_with_drawing = p.add_run()
    run_with_drawing._r.append(OxmlElement("w:drawing"))

    figures, _changes = renumber_figures(doc.paragraphs, skip_indexes=set(), apply=True)

    assert figures
    assert doc.paragraphs[0].text.startswith("Рисунок 1 —")
    assert doc.paragraphs[0]._p.findall(".//" + qn("w:drawing"))


def test_table_caption_indent_is_forced_to_zero() -> None:
    doc = Document()
    caption = doc.add_paragraph("Таблица 7 — Пример данных")
    caption.paragraph_format.first_line_indent = Cm(1.25)
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "x"

    tables, _changes = renumber_tables(
        doc.paragraphs,
        skip_indexes=set(),
        apply=True,
        indent_cm=1.25,
    )

    assert tables
    assert caption.paragraph_format.first_line_indent == Cm(0)


def test_table_caption_without_dash_is_normalized_before_table() -> None:
    doc = Document()
    caption = doc.add_paragraph("Таблица 3.2 Показатели качества")
    caption.paragraph_format.first_line_indent = Cm(1.25)
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "x"

    tables, changes = renumber_tables(
        doc.paragraphs,
        skip_indexes=set(),
        apply=True,
        chapter_breaks=[(0, 3)],
    )

    assert [str(table_caption.new_number) for table_caption in tables] == ["3.1"]
    assert caption.text == "Таблица 3.1 — Показатели качества"
    assert caption.paragraph_format.first_line_indent == Cm(0)
    assert changes
