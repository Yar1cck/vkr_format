"""Регресс: призрачный дубль заголовка «СОДЕРЖАНИЕ» в оглавлении.

Воспроизводит механизм бага: после цикла «вставка TOC → LibreOffice
запекает поля и стирает _VkrToc_ → rebuild уходит в полную переврезку»
в документе остаётся ВТОРОЙ абзац «СОДЕРЖАНИЕ» с outline-уровнем
(внизу, после списка источников). Раньше он рисовался призрачной
строкой в конце оглавления и на следующем перестроении брался как
обычная строка TOC (дефект самонакапливался).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document as WordDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from services.core.vkr_core.config.normative_loader import load_default_rules
from services.core.vkr_core.engine.toc import (
    _find_headings_for_toc,
    _make_toc_entry_p,
    _norm,
    insert_toc,
    rebuild_toc_full,
)


def _set_outline(paragraph, level: int) -> None:
    """Проставляет w:outlineLvl (0-based) — имитирует заголовок."""
    p = paragraph._p
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    for old in list(pPr.findall(qn("w:outlineLvl"))):
        pPr.remove(old)
    olvl = OxmlElement("w:outlineLvl")
    olvl.set(qn("w:val"), str(level))
    pPr.append(olvl)


def _build_doc_with_ghost(tmp_path: Path) -> WordDocument:
    doc = WordDocument()
    h_contents = doc.add_paragraph("СОДЕРЖАНИЕ")
    _set_outline(h_contents, 0)
    h_intro = doc.add_paragraph("Введение")
    _set_outline(h_intro, 0)
    doc.add_paragraph("Тело введения. " * 20)
    h_biblio = doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    _set_outline(h_biblio, 0)
    doc.add_paragraph("1. Источник один. 2. Источник два.")
    # Осиротевший дубль заголовка оглавления — то, чего «нигде нет».
    ghost = doc.add_paragraph("СОДЕРЖАНИЕ")
    _set_outline(ghost, 0)
    path = tmp_path / "doc.docx"
    doc.save(path)
    return WordDocument(str(path))


def _count_contents_paragraphs(doc: WordDocument) -> int:
    return sum(1 for p in doc.paragraphs if _norm(p.text) == "содержание")


def test_insert_toc_drops_ghost_contents(tmp_path: Path) -> None:
    rules = load_default_rules()
    doc = _build_doc_with_ghost(tmp_path)
    assert _count_contents_paragraphs(doc) == 2  # якорь + призрак

    assert insert_toc(doc, rules, title_end_index=0) is True

    # Призрак удалён — остался единственный заголовок оглавления.
    assert _count_contents_paragraphs(doc) == 1
    # «СОДЕРЖАНИЕ» не попало в строки оглавления как заголовок.
    entries = _find_headings_for_toc(doc, 0, 0)
    assert all(_norm(t) != "содержание" for _, _, t in entries)
    titles = {_norm(t) for _, _, t in entries}
    assert "введение" in titles


def test_insert_toc_keeps_heading_glued_to_contents(tmp_path: Path) -> None:
    """Регресс: заголовок сразу после «СОДЕРЖАНИЕ» без разделителя (как после
    unwrap нативного Word TOC) не должен приниматься за фиктивную строку
    оглавления и удаляться — даже если после нормализации регистра он
    перестал быть капсом («ТЕРМИНЫ, ОПРЕДЕЛЕНИЯ И СОКРАЩЕНИЯ» →
    «Термины, определения и сокращения») и не входит в hardcoded список
    структурных алиасов."""
    rules = load_default_rules()
    doc = WordDocument()
    h_contents = doc.add_paragraph("СОДЕРЖАНИЕ")
    _set_outline(h_contents, 0)
    h_terms = doc.add_paragraph("Термины, определения и сокращения")
    h_terms.style = doc.styles["Heading 1"]
    doc.add_paragraph("В настоящей работе применяются следующие термины и определения.")
    h_intro = doc.add_paragraph("Введение")
    _set_outline(h_intro, 0)
    doc.add_paragraph("Тело введения. " * 20)
    path = tmp_path / "doc.docx"
    doc.save(path)
    doc = WordDocument(str(path))

    assert insert_toc(doc, rules, title_end_index=0) is True

    texts = {_norm(p.text) for p in doc.paragraphs}
    assert "термины, определения и сокращения" in texts


def test_rebuild_is_idempotent_no_ghost_growth(tmp_path: Path) -> None:
    """Повторные rebuild'ы не плодят дубли заголовка оглавления."""
    rules = load_default_rules()
    doc = _build_doc_with_ghost(tmp_path)
    insert_toc(doc, rules, title_end_index=0)
    for _ in range(3):
        rebuild_toc_full(doc, rules)
    assert _count_contents_paragraphs(doc) == 1


def test_generated_toc_entries_use_one_and_half_line_spacing() -> None:
    entry = _make_toc_entry_p("Введение", "_VkrToc_1", 9000)
    ppr = entry.find(qn("w:pPr"))
    spacing = ppr.find(qn("w:spacing"))

    assert spacing.get(qn("w:line")) == "360"
    assert spacing.get(qn("w:lineRule")) == "auto"
    assert spacing.get(qn("w:before")) == "0"
    assert spacing.get(qn("w:after")) == "0"


def test_generated_toc_entries_have_no_paragraph_indent() -> None:
    """Абзацный отступ в оглавлении всегда 0 — не по уровню/рангу заголовка."""
    ind = _make_toc_entry_p("Введение", "_VkrToc_1", 9000).find(qn("w:pPr")).find(qn("w:ind"))
    assert ind.get(qn("w:left")) == "0"
    assert ind.get(qn("w:firstLine")) == "0"
