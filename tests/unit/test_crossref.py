"""Регрессия: перенумерация и якоря ссылок на рисунки/таблицы.

Защищает от повторного всплытия багов: ссылки в тексте оставались со
старыми номерами; перечисление «рисунки N и M» ловилось наполовину;
проза «(Рисунок 2.1) поясняет …» считалась подписью; якорей не было.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from services.core.vkr_core.engine.captions import Caption, renumber_figures, validate_references
from services.core.vkr_core.engine.crossref import apply_crossrefs, iter_reference_spans
from services.core.vkr_core.engine.stats import DocStats
from services.core.vkr_core.engine.violations import SEVERITY_INFO


def test_iter_references_expands_enumerations():
    refs = [(kind, num) for kind, num, _ in
            iter_reference_spans("см. рисунки 2.3 и 2.4, а также табл. 5, 6")]
    assert ("figure", "2.3") in refs
    assert ("figure", "2.4") in refs  # раньше терялся
    assert ("table", "5") in refs
    assert ("table", "6") in refs


def test_prose_with_figure_word_is_not_a_caption():
    doc = Document()
    p = doc.add_paragraph("Рисунок 2 показывает общую архитектуру системы.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    figures, _ = renumber_figures(doc.paragraphs, set(), apply=True)
    assert figures == []  # проза не подпись — фантомного рисунка нет


def test_intext_references_renumbered_and_anchored():
    doc = Document()
    img = doc.add_paragraph()
    img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img.add_run()
    run._r.append(doc.element.makeelement(qn("w:drawing"), {}))
    cap = doc.add_paragraph("Рисунок 2.5 – Архитектура системы")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Подробнее это видно на рисунке 2.5 далее.")

    figures, _ = renumber_figures(doc.paragraphs, set(), apply=True)
    assert figures and figures[0].new_number == 1

    changed = apply_crossrefs(doc, figures, [], set())
    assert changed == 1

    body_after = doc.paragraphs[2]
    assert "рисунке 1" in body_after.text  # перенумеровано 2.5 -> 1
    assert "2.5" not in body_after.text
    hyperlinks = body_after._p.findall(qn("w:hyperlink"))
    assert len(hyperlinks) == 1
    assert hyperlinks[0].get(qn("w:anchor")) == "_Ref_fig1"

    bm = [
        b.get(qn("w:name"))
        for b in cap._p.findall(qn("w:bookmarkStart"))
    ]
    assert "_Ref_fig1" in bm


def test_correctly_numbered_caption_reports_no_renumber_violation():
    """«Таблица 1» с верным номером не должна давать карточку «1 → 1»."""
    from services.core.vkr_core.engine.captions import renumber_tables

    doc = Document()
    p = doc.add_paragraph("Таблица 1  —  Классификация процессов")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_table(1, 1)  # подпись должна стоять перед реальной таблицей
    tables, changes = renumber_tables(doc.paragraphs, set(), apply=True)
    assert tables and tables[0].new_number == 1
    assert changes == []  # номер верный — нарушения нет
    # Но оформление всё равно нормализовано (двойные пробелы убраны).
    assert doc.paragraphs[0].text == "Таблица 1 — Классификация процессов"


def test_table_reference_sentence_is_not_caption_even_when_centered_before_table():
    from services.core.vkr_core.engine.captions import renumber_tables

    doc = Document()
    p = doc.add_paragraph("Таблица 1.1 показывает состав требований.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_table(1, 1)

    tables, changes = renumber_tables(doc.paragraphs, set(), apply=True)

    assert tables == []
    assert changes == []
    assert doc.paragraphs[0].text == "Таблица 1.1 показывает состав требований."


def test_wrongly_numbered_caption_still_reported():
    """Иерархическая «Таблица 2.1» → «Таблица 1» — это реальная
    перенумерация, её сообщать надо."""
    from services.core.vkr_core.engine.captions import renumber_tables

    doc = Document()
    p = doc.add_paragraph("Таблица 2.1 – Параметры")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_table(1, 1)  # подпись должна стоять перед реальной таблицей
    _tables, changes = renumber_tables(doc.paragraphs, set(), apply=True)
    assert len(changes) == 1


def test_three_level_caption_number_parsed_whole_not_leaked_into_title():
    """«Таблица 2.3.1» — трёхуровневый номер (глава.раздел.номер) должен
    разбираться целиком: раньше числовая группа поддерживала только один
    уровень дроби, и «.1» утекал в начало нового заголовка подписи
    («Таблица 1 — 1 – Заголовок» вместо «Таблица 1 — Заголовок»)."""
    from services.core.vkr_core.engine.captions import renumber_figures, renumber_tables

    doc = Document()
    doc.add_paragraph("Таблица 2.3.1 – Заголовок таблицы")
    doc.add_table(1, 1)
    tables, changes = renumber_tables(doc.paragraphs, set(), apply=True)
    assert tables[0].original_number == "2.3.1"
    assert tables[0].title == "Заголовок таблицы"
    assert changes == [(0, "Таблица 2.3.1 – Заголовок таблицы", "Таблица 1 — Заголовок таблицы")]

    doc2 = Document()
    doc2.add_paragraph("Рисунок 1.2.3 – Схема архитектуры")
    figures, fchanges = renumber_figures(doc2.paragraphs, set(), apply=True)
    assert figures[0].original_number == "1.2.3"
    assert figures[0].title == "Схема архитектуры"
    assert fchanges == [(0, "Рисунок 1.2.3 – Схема архитектуры", "Рисунок 1 — Схема архитектуры")]


def test_intext_reference_to_three_level_number_resolves_no_false_violation():
    """Ссылка «таблице 2.3.1» на подпись с тем же трёхуровневым номером
    должна засчитываться — раньше ссылка усекалась до «2.3» и подпись
    ложно считалась неупомянутой."""
    from services.core.vkr_core.engine.captions import renumber_tables
    from services.core.vkr_core.engine.stats import collect_stats

    doc = Document()
    doc.add_paragraph("Как показано в таблице 2.3.1, результаты подтверждают гипотезу.")
    doc.add_paragraph("Таблица 2.3.1 – Заголовок таблицы")
    doc.add_table(1, 1)
    tables, _ = renumber_tables(doc.paragraphs, set(), apply=True)
    stats = collect_stats(doc)
    violations = validate_references(stats, [], tables, set())
    assert violations == []


def test_unreferenced_figure_is_info_with_advice():
    from types import SimpleNamespace

    stats = DocStats()
    stats.paragraphs = [
        SimpleNamespace(index=0, is_empty=False, stripped="Текст без ссылок на рисунки."),
        SimpleNamespace(index=1, is_empty=False, stripped="Рисунок 1 — Схема"),
    ]
    figures = [Caption(1, "1", 1, "Схема", "figure")]

    violations = validate_references(stats, figures=figures, tables=[], skip_indexes=set())

    fr = next(v for v in violations if v.type == "figure_not_referenced")
    assert fr.severity == SEVERITY_INFO
    assert "Совет:" in fr.description
    assert "см. рисунок 1" in fr.description


def test_heading_line_spacing_is_always_one_and_half():
    """Межстрочный интервал заголовков всегда 1.5 — даже у раздела 1-го уровня
    и у длинных, переносящихся на вторую строку заголовков. Зазор до/после
    заголовка — 6pt (120 twips) с обеих сторон, одинаковый на всех уровнях."""
    from docx.oxml.ns import qn as _qn

    from services.core.vkr_core.engine.formatter import _force_heading_direct_format

    def _spacing(text, level, structural=False):
        d = Document()
        p = d.add_paragraph()
        p.add_run(text)
        _force_heading_direct_format(p, structural=structural, level=level)
        sp = p._p.find(_qn("w:pPr")).find(_qn("w:spacing"))
        return sp.get(_qn("w:line")), sp.get(_qn("w:before")), sp.get(_qn("w:after"))

    line_h1, before_h1, after_h1 = _spacing("1 Короткий раздел", 1)
    assert line_h1 == "360"  # 1.5
    assert before_h1 == "120"  # 6pt
    assert after_h1 == "120"

    long = (
        "1 Анализ предметной области, архитектурных требований и средств "
        "автоматизации развертывания информационной системы LibSpace"
    )
    line_long, before_long, after_long = _spacing(long, 1)
    assert line_long == "360"
    assert before_long == "120"
    assert after_long == "120"

    line_sub, before_sub, after_sub = _spacing("1.1 Короткий подраздел", 2)
    assert line_sub == "360"
    assert before_sub == "120"
    assert after_sub == "120"

    line_struct, before_struct, after_struct = _spacing("ВВЕДЕНИЕ", 1, structural=True)
    assert line_struct == "360"
    assert before_struct == "120"
    assert after_struct == "120"


def test_heading_bold_depends_on_level():
    """Раздел (1) и подраздел (1.1) — полужирные; пункт (1.1.1) — обычным начертанием."""
    from services.core.vkr_core.engine.formatter import _force_heading_direct_format

    def _is_bold(text, level, structural=False):
        d = Document()
        p = d.add_paragraph()
        p.add_run(text)
        _force_heading_direct_format(p, structural=structural, level=level)
        return p.runs[0].font.bold

    assert _is_bold("1 Раздел", 1) is True
    assert _is_bold("1.1 Подраздел", 2) is True
    assert _is_bold("1.1.1 Пункт", 3) is False
    assert _is_bold("ВВЕДЕНИЕ", 1, structural=True) is True


def test_appendix_label_not_all_caps():
    """«ПРИЛОЖЕНИЕ А» → «Приложение А»: капс не по норме, буква-идентификатор
    остаётся заглавной. Остальные структурные заголовки капс сохраняют."""
    from services.core.vkr_core.engine.formatter import _normalise_heading_text

    assert _normalise_heading_text("ПРИЛОЖЕНИЕ А") == "Приложение А"
    assert _normalise_heading_text("приложение б") == "Приложение Б"
    assert _normalise_heading_text("Приложение В (обязательное)") == "Приложение В (обязательное)"
    # Другие структурные заголовки по-прежнему приводятся к капсу.
    assert _normalise_heading_text("введение") == "ВВЕДЕНИЕ"
    assert _normalise_heading_text("заключение") == "ЗАКЛЮЧЕНИЕ"


def test_appendix_title_has_no_own_page_break():
    """Название приложения, случайно распознанное как отдельный заголовок
    (например, студент применил стиль «Заголовок 1» и к лейблу, и к названию),
    не должно получать свой pageBreakBefore — иначе оно отрывается от лейбла
    «Приложение X» и уезжает на следующую страницу."""
    from services.core.vkr_core.config.normative_loader import load_default_rules
    from services.core.vkr_core.engine.detection import detect_headings
    from services.core.vkr_core.engine.formatter import apply_headings, style_appendix_titles
    from services.core.vkr_core.engine.stats import collect_stats

    rules = load_default_rules()
    doc = Document()
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("1. Источник один.")
    label = doc.add_paragraph("Приложение А")
    label.style = doc.styles["Heading 1"]
    title = doc.add_paragraph("Исходные данные и результаты пилотного запуска")
    title.style = doc.styles["Heading 1"]
    doc.add_paragraph("Текст приложения.")

    stats = collect_stats(doc)
    headings = detect_headings(stats, rules, set())
    apply_headings(doc, headings, stats)
    style_appendix_titles(doc)

    def _has_page_break(p):
        pPr = p._p.find(qn("w:pPr"))
        return pPr is not None and pPr.find(qn("w:pageBreakBefore")) is not None

    assert _has_page_break(doc.paragraphs[2]) is True  # лейбл — с новой страницы
    assert _has_page_break(doc.paragraphs[3]) is False  # название — та же страница


def test_appendix_label_gets_no_extra_inline_page_break():
    """Регресс: insert_l1_page_breaks добавлял <w:br type="page"/> в конец
    лейбла «Приложение X», потому что название приложения (в этот момент
    пайплайна ещё не разжаловано style_appendix_titles) считалось отдельным
    L1-заголовком, а лейбл — его «предыдущим абзацем». В итоге лейбл
    оставался один на странице, а название уезжало на следующую."""
    from services.core.vkr_core.config.normative_loader import load_default_rules
    from services.core.vkr_core.engine.detection import detect_headings
    from services.core.vkr_core.engine.formatter import apply_headings, insert_l1_page_breaks
    from services.core.vkr_core.engine.stats import collect_stats

    rules = load_default_rules()
    doc = Document()
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("1. Источник один.")
    label = doc.add_paragraph("Приложение А")
    label.style = doc.styles["Heading 1"]
    title = doc.add_paragraph("Исходные данные и результаты пилотного запуска")
    title.style = doc.styles["Heading 1"]
    doc.add_paragraph("Текст приложения.")

    stats = collect_stats(doc)
    headings = detect_headings(stats, rules, set())
    applied = apply_headings(doc, headings, stats)
    insert_l1_page_breaks(doc, headings, applied)

    label_p = doc.paragraphs[2]._p
    assert not any(br.get(qn("w:type")) == "page" for br in label_p.iter(qn("w:br")))


def test_landscape_section_stays_landscape():
    """Секция, сделанную студентом альбомной (широкая таблица/приложение),
    apply_page_settings не должен переворачивать обратно в книжную — только
    книжные секции приводятся к нормативным размерам А4."""
    from docx.enum.section import WD_ORIENT, WD_SECTION
    from docx.shared import Mm

    from services.core.vkr_core.config.normative_loader import load_default_rules
    from services.core.vkr_core.engine.formatter import apply_page_settings

    rules = load_default_rules()
    doc = Document()
    doc.add_paragraph("Портретная страница")

    landscape_section = doc.add_section(WD_SECTION.NEW_PAGE)
    landscape_section.orientation = WD_ORIENT.LANDSCAPE
    landscape_section.page_width = Mm(297)
    landscape_section.page_height = Mm(210)
    doc.add_paragraph("Альбомная страница с широкой таблицей")

    apply_page_settings(doc, rules, skip_first_section=False)

    portrait_section, landscape_section = doc.sections
    assert portrait_section.orientation == WD_ORIENT.PORTRAIT
    assert portrait_section.page_width < portrait_section.page_height

    assert landscape_section.orientation == WD_ORIENT.LANDSCAPE
    assert landscape_section.page_width > landscape_section.page_height


def test_crossref_is_idempotent():
    doc = Document()
    cap = doc.add_paragraph("Рисунок 1 – Схема")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Как показано на рисунке 1, всё работает.")
    figures, _ = renumber_figures(doc.paragraphs, set(), apply=True)
    apply_crossrefs(doc, figures, [], set())
    apply_crossrefs(doc, figures, [], set())  # второй прогон
    body = doc.paragraphs[1]
    assert body.text.count("рисунке 1") == 1
    assert len(body._p.findall(qn("w:hyperlink"))) == 1
    assert len(cap._p.findall(qn("w:bookmarkStart"))) == 1


def test_set_keep_flags_writes_keepnext_keeplines():
    """Хелпер пишет <w:keepNext/> и <w:keepLines/> в правильном порядке."""
    from services.core.vkr_core.engine.captions import _set_keep_flags

    doc = Document()
    p = doc.add_paragraph("Рисунок 1 — Схема")
    _set_keep_flags(p._p, keep_next=True, keep_lines=True)
    pPr = p._p.find(qn("w:pPr"))
    kids = [c.tag for c in pPr]
    assert qn("w:keepNext") in kids
    assert qn("w:keepLines") in kids
    # keepNext раньше keepLines (порядок дочерних элементов w:pPr по схеме).
    assert kids.index(qn("w:keepNext")) < kids.index(qn("w:keepLines"))
    # Идемпотентность: повторный вызов не плодит дубли.
    _set_keep_flags(p._p, keep_next=True, keep_lines=False)
    pPr = p._p.find(qn("w:pPr"))
    assert len(pPr.findall(qn("w:keepNext"))) == 1
    assert len(pPr.findall(qn("w:keepLines"))) == 0


def test_consecutive_headings_all_use_same_line_spacing():
    """«1» → «1.1» → «1.1.1»: у всех уровней одинаковый межстрочный интервал 1.5
    (двойной интервал у раздела 1-го уровня не применяется) и одинаковый
    зазор до/после — 6pt (120 twips)."""
    from services.core.vkr_core.engine.formatter import _force_heading_direct_format

    doc = Document()
    h1 = doc.add_paragraph("1 Основная часть")
    h2 = doc.add_paragraph("1.1 Постановка задачи")
    h3 = doc.add_paragraph("1.1.1 Анализ требований")
    doc.add_paragraph("Обычный текст раздела.")
    _force_heading_direct_format(h1, structural=False, level=1)
    _force_heading_direct_format(h2, structural=False, level=2)
    _force_heading_direct_format(h3, structural=False, level=3)

    for h in (h1, h2, h3):
        sp = h._p.find(qn("w:pPr")).find(qn("w:spacing"))
        assert sp.get(qn("w:line")) == "360"
        assert sp.get(qn("w:before")) == "120"
        assert sp.get(qn("w:after")) == "120"
