"""Регресс: «Продолжение таблицы N» вставляется и не плодит пустых строк.

Точку переноса теперь даёт UNO-измерение движка LibreOffice
(engine.uno_layout.measure_table_row_pages → {table_idx: [page_per_row]}),
а split_long_tables_to_pages принимает этот словарь. Здесь мы передаём
измерение синтетически (тесты не зовут LibreOffice).

Если UNO-измерения нет (пустой dict) — модуль не режет таблицу: надпись
«Продолжение таблицы N» создаётся только по реальному переносу. Надпись
несёт pageBreakBefore (отдельного пустого абзаца-разрыва нет).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document as WordDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from services.core.vkr_core.engine.captions import renumber_tables
from services.core.vkr_core.engine.table_continuation import (
    _LABEL_PREFIX,
    split_long_tables_to_pages,
    split_overflowing_fragments,
)
from services.core.vkr_core.engine.tables import apply_continuation_table_styles
from services.core.vkr_core.engine.toc import _squash


def test_long_table_gets_continuation_label(tmp_path: Path) -> None:
    doc = WordDocument()
    doc.add_paragraph("Текст перед таблицей достаточной длины. " * 2)
    doc.add_paragraph("Таблица 1 — Тестовая длинная таблица")
    t = doc.add_table(rows=6, cols=1)
    t.cell(0, 0).text = "Заголовок столбца"
    rows = [
        "Альфа строка один",
        "Бета строка два",
        "Гамма строка три",
        "Дельта строка четыре",
        "Эпсилон строка пять",
    ]
    for i, txt in enumerate(rows, start=1):
        t.cell(i, 0).text = txt
    doc.add_paragraph("Текст после таблицы.")
    p = tmp_path / "doc.docx"
    doc.save(p)
    doc = WordDocument(str(p))

    caps, _ = renumber_tables(doc.paragraphs, set(), apply=False)
    assert caps, "подпись таблицы должна распознаться"

    # UNO-измерение: таблица 0, 6 строк (шапка+5 данных). Шапка+строки 1-2
    # на стр 1, строки 3-5 на стр 2 → переносится после 2-й строки данных.
    row_pages = {0: [1, 1, 1, 2, 2, 2]}
    assert split_long_tables_to_pages(doc, row_pages, caps) is True

    labels = [p for p in doc.paragraphs if p.text.strip().startswith(_LABEL_PREFIX)]
    assert labels, "надпись «Продолжение таблицы N» должна быть вставлена"
    assert all("1" in p.text for p in labels)
    assert len(doc.tables) >= 2
    # Разрыв страницы — НА надписи (pageBreakBefore), без отдельного
    # пустого абзаца-разрыва (он давал лишнюю пустую страницу).
    for lp in labels:
        pPr = lp._p.find(qn("w:pPr"))
        assert pPr is not None and pPr.find(qn("w:pageBreakBefore")) is not None
        ind = pPr.find(qn("w:ind"))
        assert ind is not None
        assert ind.get(qn("w:firstLine")) == "0"
        assert ind.get(qn("w:left")) == "0"
        assert lp.text.strip() == "Продолжение таблицы 1"
    body = doc.element.body
    for el in body.iter(qn("w:p")):
        if _squash("".join(el.itertext())).startswith(_squash(_LABEL_PREFIX)):
            prev = el.getprevious()
            assert prev is not None and prev.tag == qn("w:tbl"), (
                "перед надписью должна идти таблица, а не пустой разрыв"
            )
            # Регресс: между фрагментом и надписью не должен зависать
            # пустой <w:p> (его раньше оставлял ensure_separator_after_-
            # tables_and_figures, что съезжало надпись вниз на странице).
            assert "".join(prev.itertext()).strip(), (
                "перед надписью не должно быть пустой строки"
            )


def test_uno_split_three_fragments(tmp_path: Path) -> None:
    """UNO-измерение [1,1,1,2,2,3]: данные распределены 2/2/1 по трём
    страницам → ровно три фрагмента (без рваных «по 1 строке»)."""
    doc = WordDocument()
    doc.add_paragraph("Таблица 1 — Плотная на три страницы")
    t = doc.add_table(rows=6, cols=1)
    t.cell(0, 0).text = "Шапка"
    for i, txt in enumerate(
        ["один", "два", "три", "четыре", "пять"], start=1
    ):
        t.cell(i, 0).text = txt
    p = tmp_path / "three.docx"
    doc.save(p)
    doc = WordDocument(str(p))

    caps, _ = renumber_tables(doc.paragraphs, set(), apply=False)
    # шапка стр1; данные: стр 1,1,2,2,3 → группы [2,2,1]
    row_pages = {0: [1, 1, 1, 2, 2, 3]}
    assert split_long_tables_to_pages(doc, row_pages, caps) is True
    labels = [p for p in doc.paragraphs if p.text.strip().startswith(_LABEL_PREFIX)]
    # два переноса → две надписи, три фрагмента таблицы
    assert len(labels) == 2
    assert len(doc.tables) == 3


def test_short_table_not_split(tmp_path: Path) -> None:
    doc = WordDocument()
    doc.add_paragraph("Таблица 1 — Короткая")
    t = doc.add_table(rows=3, cols=1)
    t.cell(0, 0).text = "Шапка"
    t.cell(1, 0).text = "Первая уникальная"
    t.cell(2, 0).text = "Вторая уникальная"
    p = tmp_path / "s.docx"
    doc.save(p)
    doc = WordDocument(str(p))
    caps, _ = renumber_tables(doc.paragraphs, set(), apply=False)
    # вся таблица на одной странице → UNO даёт один уровень → не режем
    row_pages = {0: [1, 1, 1]}
    assert split_long_tables_to_pages(doc, row_pages, caps) is False
    assert not [p for p in doc.paragraphs if p.text.strip().startswith(_LABEL_PREFIX)]


def test_no_split_when_uno_empty_even_for_long_table(tmp_path: Path) -> None:
    """Без точного UNO-измерения не вставляем «Продолжение таблицы»:
    эвристика по объёму может создать ложный перенос на той же странице."""
    doc = WordDocument()
    doc.add_paragraph("Таблица 1 — Длинная без UNO-измерения")
    n_data = 35
    t = doc.add_table(rows=n_data + 1, cols=2)
    t.cell(0, 0).text = "№"
    t.cell(0, 1).text = "Значение"
    for i in range(1, n_data + 1):
        t.cell(i, 0).text = str(i)
        t.cell(i, 1).text = "x"
    doc.add_paragraph("Текст после.")
    p = tmp_path / "fallback.docx"
    doc.save(p)
    doc = WordDocument(str(p))

    caps, _ = renumber_tables(doc.paragraphs, set(), apply=False)
    assert split_long_tables_to_pages(doc, {}, caps) is False

    labels = [p for p in doc.paragraphs if p.text.strip().startswith(_LABEL_PREFIX)]
    assert labels == []
    assert len(doc.tables) == 1


def test_fallback_skips_short_table_without_measurement(tmp_path: Path) -> None:
    """Без UNO-измерения короткая таблица (< порога) не режется."""
    doc = WordDocument()
    doc.add_paragraph("Таблица 1 — Короткая без UNO")
    n_data = 20
    t = doc.add_table(rows=n_data + 1, cols=2)
    t.cell(0, 0).text = "№"
    t.cell(0, 1).text = "V"
    for i in range(1, n_data + 1):
        t.cell(i, 0).text = str(i)
        t.cell(i, 1).text = "x"
    p = tmp_path / "short_no_uno.docx"
    doc.save(p)
    doc = WordDocument(str(p))

    caps, _ = renumber_tables(doc.paragraphs, set(), apply=False)
    assert split_long_tables_to_pages(doc, {}, caps) is False
    assert not [
        p for p in doc.paragraphs if p.text.strip().startswith(_LABEL_PREFIX)
    ]


def test_no_text_volume_split_without_uno(tmp_path: Path) -> None:
    """Даже плотную таблицу не режем без UNO: продолжение должно отражать
    реальную страницу переноса, а не оценку по символам."""
    doc = WordDocument()
    doc.add_paragraph("Таблица 1 — Сравнение фреймворков")
    body = "Длинное описание характеристики с пояснением деталей применения. " * 4
    t = doc.add_table(rows=9, cols=4)
    t.cell(0, 0).text = "Критерий"
    for j, h in enumerate(["A", "B", "C"], start=1):
        t.cell(0, j).text = "Фреймворк " + h
    criteria = ["Производительность", "Масштабируемость", "Документация",
                "Сообщество", "Лицензия", "Поддержка", "Зрелость", "Цена"]
    for i, c in enumerate(criteria, start=1):
        t.cell(i, 0).text = c
        for j in range(1, 4):
            t.cell(i, j).text = body
    p = tmp_path / "dense.docx"
    doc.save(p)
    doc = WordDocument(str(p))

    caps, _ = renumber_tables(doc.paragraphs, set(), apply=False)
    assert split_long_tables_to_pages(doc, {}, caps) is False
    labels = [p for p in doc.paragraphs if p.text.strip().startswith(_LABEL_PREFIX)]
    assert labels == []


def test_overflowing_fragment_split_by_break(tmp_path: Path) -> None:
    """split_overflowing_fragments режет цельную таблицу по первой точке
    переноса; номер берётся из подписи «Таблица N»."""
    doc = WordDocument()
    doc.add_paragraph("Таблица 1 — Длинная для дорезки")
    t = doc.add_table(rows=8, cols=1)
    t.cell(0, 0).text = "Шапка"
    for i, txt in enumerate(
        ["a", "b", "c", "d", "e", "f", "g"], start=1
    ):
        t.cell(i, 0).text = txt
    p = tmp_path / "overflow.docx"
    doc.save(p)
    doc = WordDocument(str(p))
    # шапка стр1; данные 1,2,3 на стр1; 4,5,6,7 на стр2 → перенос после 3-й
    row_pages = {0: [1, 1, 1, 1, 2, 2, 2, 2]}
    assert split_overflowing_fragments(doc, row_pages) is True
    labels = [p for p in doc.paragraphs if p.text.strip().startswith(_LABEL_PREFIX)]
    assert len(labels) == 1
    assert "1" in labels[0].text
    assert len(doc.tables) == 2


def test_overflow_rolled_into_next_fragment(tmp_path: Path) -> None:
    """Балансировка: если следующий фрагмент той же таблицы существует,
    лишние строки прокатываются в него, а не создают одинокий хвост.

    Готовим документ из двух фрагментов: tbl1 (шапка+3 данных) +
    «Продолжение таблицы 1» + tbl2 (шапка+2 данных). UNO «говорит», что
    у tbl1 переносится 3-я строка → она должна переехать в начало tbl2,
    а НЕ создать третий фрагмент. Итог: 2 фрагмента, tbl1 ужался.
    """
    from services.core.vkr_core.engine.table_continuation import _make_label_para

    doc = WordDocument()
    doc.add_paragraph("Таблица 1 — С двумя фрагментами")
    t1 = doc.add_table(rows=4, cols=1)  # шапка + 3 данных
    t1.cell(0, 0).text = "Шапка"
    for i, v in enumerate(["a", "b", "c"], start=1):
        t1.cell(i, 0).text = v
    # надпись + второй фрагмент
    label = _make_label_para("1")
    t1._tbl.addnext(label)
    t2 = doc.add_table(rows=3, cols=1)  # шапка + 2 данных
    t2.cell(0, 0).text = "Шапка"
    t2.cell(1, 0).text = "d"
    t2.cell(2, 0).text = "e"
    # переносим t2 после надписи
    label.addnext(t2._tbl)
    p = tmp_path / "two_frag.docx"
    doc.save(p)
    doc = WordDocument(str(p))

    n_tables_before = len(doc.tables)
    # tbl1 (idx 0): данные 1,2 на стр1, 3-я на стр2 → перенос после 2-й.
    # tbl2 (idx 1): помещается.
    row_pages = {0: [1, 1, 1, 2], 1: [1, 1, 1]}
    assert split_overflowing_fragments(doc, row_pages) is True
    # Число таблиц НЕ выросло — строка прокатилась, новый фрагмент не создан.
    assert len(doc.tables) == n_tables_before
    # tbl1 теперь 1+2 строки (шапка + a,b), tbl2 — 1+3 (шапка + c,d,e).
    assert len(doc.tables[0].rows) == 3
    assert len(doc.tables[1].rows) == 4


def test_existing_continuation_label_is_normalized(tmp_path: Path) -> None:
    from services.core.vkr_core.engine.table_continuation import (
        _normalize_continuation_labels,
    )

    doc = WordDocument()
    doc.add_paragraph("Продолжение таблицы2.1")
    p = tmp_path / "label.docx"
    doc.save(p)
    doc = WordDocument(str(p))

    _normalize_continuation_labels(doc)

    label = doc.paragraphs[0]
    assert label.text == "Продолжение таблицы 2.1"
    pPr = label._p.find(qn("w:pPr"))
    assert pPr is not None and pPr.find(qn("w:pageBreakBefore")) is not None
    ind = pPr.find(qn("w:ind"))
    assert ind is not None and ind.get(qn("w:firstLine")) == "0"


def test_continuation_table_gets_same_gost_style(tmp_path: Path) -> None:
    doc = WordDocument()
    doc.add_paragraph("Таблица 2.1 — Исходная таблица")
    base = doc.add_table(rows=2, cols=2)
    base.cell(0, 0).text = "Код"
    base.cell(0, 1).text = "Описание"
    base.cell(1, 0).text = "F-01"
    base.cell(1, 1).text = "Первая строка"
    doc.add_paragraph("Продолжение таблицы2.1")
    cont = doc.add_table(rows=2, cols=2)
    cont.cell(0, 0).text = "Код"
    cont.cell(0, 1).text = "Описание"
    cont.cell(1, 0).text = "F-02"
    cont.cell(1, 1).text = "Вторая строка"

    for row in cont.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "BDD7EE")
            tcPr.append(shd)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(18)
                    run.bold = False

    p = tmp_path / "continuation_style.docx"
    doc.save(p)
    doc = WordDocument(str(p))

    assert apply_continuation_table_styles(doc, 0) == 1

    styled = doc.tables[1]
    header_p = styled.cell(0, 0).paragraphs[0]
    data_p = styled.cell(1, 1).paragraphs[0]
    assert header_p.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert data_p.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert header_p.runs[0].bold is True
    assert header_p.runs[0].font.size.pt == 12
    assert data_p.runs[0].font.size.pt == 12
    assert styled.cell(0, 0)._tc.get_or_add_tcPr().find(qn("w:shd")) is None
    borders = styled._tbl.find(qn("w:tblPr")).find(qn("w:tblBorders"))
    assert borders is not None and borders.find(qn("w:insideH")) is not None


def test_no_empty_paragraph_before_label_after_split(tmp_path: Path) -> None:
    """Регресс: между фрагментом-таблицей и надписью не должно оставаться
    пустого <w:p> — даже если он был ДО резки (ensure_separator)."""
    doc = WordDocument()
    doc.add_paragraph("Таблица 1 — Тестовая для регресса")
    t = doc.add_table(rows=7, cols=1)
    t.cell(0, 0).text = "Заголовок"
    cells = ["Альфа", "Бета", "Гамма", "Дельта", "Эпсилон", "Зета"]
    for i, txt in enumerate(cells, start=1):
        t.cell(i, 0).text = txt
    # Пустой абзац СРАЗУ ПОСЛЕ таблицы — то, что вставляет
    # ensure_separator_after_tables_and_figures в реальном пайплайне.
    doc.add_paragraph("")
    doc.add_paragraph("Текст после.")
    p = tmp_path / "with_separator.docx"
    doc.save(p)
    doc = WordDocument(str(p))

    caps, _ = renumber_tables(doc.paragraphs, set(), apply=False)
    # шапка+строки 1-3 на стр1, строки 4-6 на стр2
    row_pages = {0: [1, 1, 1, 1, 2, 2, 2]}
    assert split_long_tables_to_pages(doc, row_pages, caps) is True

    body = doc.element.body
    for el in body.iter(qn("w:p")):
        if _squash("".join(el.itertext())).startswith(_squash(_LABEL_PREFIX)):
            prev = el.getprevious()
            assert prev is not None, "перед надписью что-то должно быть"
            assert prev.tag == qn("w:tbl"), (
                f"перед надписью должна идти таблица, нашли <{prev.tag}>"
            )
