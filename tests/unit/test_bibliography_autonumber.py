"""Регресс: список источников с автонумерацией Word (w:numPr).

Видимое «1.» рисует поле нумерации, в тексте абзаца цифры нет, поэтому
буквальный BIB_ENTRY_RE такие записи не ловил и выдавалось ложное
предупреждение 7.6.B2 «записи не распознаны». collect_bibliography
теперь распознаёт нумерованные абзацы внутри раздела источников и
синтезирует сквозные номера (1, 2, 3, …).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document as WordDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from services.core.vkr_core.config.normative_loader import load_default_rules
from services.core.vkr_core.engine.bibliography import (
    collect_bibliography,
    validate_bibliography,
)
from services.core.vkr_core.engine.stats import collect_stats


def _make_autonumbered(paragraph, num_id: int = 7) -> None:
    """Проставляет w:numPr (автонумерованный список Word)."""
    p = paragraph._p
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(nid)
    pPr.append(numPr)


def _build(tmp_path: Path) -> WordDocument:
    doc = WordDocument()
    doc.add_paragraph("Текст работы со ссылкой [1] и [2].")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    # Записи — пункты автосписка Word: в тексте номера НЕТ.
    for src in (
        "Иванов И. И. Анализ данных. – М.: Наука, 2020. – 200 с.",
        "Петров П. П. Методы обработки. – СПб.: Питер, 2021. – 150 с.",
        "Сидоров С. С. Алгоритмы. – М.: ДМК, 2019. – 320 с.",
    ):
        p = doc.add_paragraph(src)
        _make_autonumbered(p)
    path = tmp_path / "doc.docx"
    doc.save(path)
    return WordDocument(str(path))


def test_autonumbered_bibliography_is_recognised(tmp_path: Path) -> None:
    rules = load_default_rules()
    doc = _build(tmp_path)
    stats = collect_stats(doc)

    entries, bounds, malformed = collect_bibliography(stats, rules, headings=[])

    assert bounds is not None
    assert [e.number for e in entries] == [1, 2, 3]
    assert "Иванов" in entries[0].text
    assert malformed == []

    violations = validate_bibliography(entries, citations=[], bounds=bounds, malformed=malformed)
    types = {v.type for v in violations}
    # Ложное «записи не распознаны» больше не возникает.
    assert "bibliography_entries_unrecognised" not in types
