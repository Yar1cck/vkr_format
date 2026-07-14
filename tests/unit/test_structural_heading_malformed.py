"""Структурный заголовок, оформленный как строка оглавления (отточие + номер
страницы), детектируется и в полном режиме приводится к чистому виду.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document as WordDocument

from services.core.vkr_core.config.normative_loader import load_default_rules
from services.core.vkr_core.engine import process_document
from services.core.vkr_core.engine.stats import collect_stats
from services.core.vkr_core.engine.structural_format import (
    apply_structural_heading_cleanup,
    build_violations,
    find_malformed_structural_headings,
)

MALFORMED = "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ" + "\t" + "." * 30 + "43"


def test_find_and_cleanup_malformed_structural_heading(tmp_path: Path) -> None:
    doc = WordDocument()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    doc.add_paragraph(MALFORMED)
    doc.add_paragraph("Иванов И. И. Анализ. – М.: Наука, 2020. – 200 с.")
    path = tmp_path / "doc.docx"
    doc.save(path)
    doc = WordDocument(str(path))

    rules = load_default_rules()
    stats = collect_stats(doc)
    matches = find_malformed_structural_headings(stats, rules, skip_indexes=set())

    assert len(matches) == 1
    assert matches[0].key == "references"
    assert matches[0].cleaned_text == "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"

    apply_structural_heading_cleanup(doc, matches)
    assert doc.paragraphs[2].text == "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"


def test_clean_heading_is_not_flagged(tmp_path: Path) -> None:
    doc = WordDocument()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    path = tmp_path / "clean.docx"
    doc.save(path)
    doc = WordDocument(str(path))

    rules = load_default_rules()
    stats = collect_stats(doc)
    assert find_malformed_structural_headings(stats, rules, skip_indexes=set()) == []


def test_build_violations_status_depends_on_fix() -> None:
    from services.core.vkr_core.engine.structural_format import MalformedHeading
    from services.core.vkr_core.models.enums import ViolationStatus

    m = [MalformedHeading(2, "references", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", MALFORMED)]
    fixed = build_violations(m, fixed=True)
    not_fixed = build_violations(m, fixed=False)
    assert fixed[0].type == "structural_heading_malformed"
    assert fixed[0].status == ViolationStatus.auto_fixed
    assert not_fixed[0].status == ViolationStatus.manual_required


def test_pipeline_cleans_malformed_references_heading(tmp_path: Path) -> None:
    doc = WordDocument()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения со ссылкой [1].")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    doc.add_paragraph("Текст заключения.")
    doc.add_paragraph(MALFORMED)
    doc.add_paragraph("Иванов И. И. Анализ. – М.: Наука, 2020. – 200 с.")
    source = tmp_path / "source.docx"
    doc.save(source)

    result = process_document(
        source_path=source,
        rules=load_default_rules(),
        check_only=False,
    )
    types = {v.type for v in result.violations}
    assert "structural_heading_malformed" in types
    # Раздел источников больше не считается отсутствующим.
    assert "bibliography_missing" not in types

    processed = WordDocument(str(result.processed_docx_path))
    texts = [p.text for p in processed.paragraphs]
    assert "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ" in texts
    # Хвоста-отточия с номером страницы в теле быть не должно.
    assert not any("....." in t and t.rstrip().endswith("43") for t in texts)
