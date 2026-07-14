from __future__ import annotations

from dataclasses import replace
from pathlib import Path

from docx import Document as WordDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from services.core.vkr_core.config.normative_loader import load_default_rules
from services.core.vkr_core.engine.detection import (
    HEADING_SOURCE_HIGH,
    HEADING_SOURCE_MEDIUM,
    HEADING_SOURCE_RECOVERED,
    detect_headings,
    detect_headings_full,
)
from services.core.vkr_core.engine.formatter import strip_chapter_word_everywhere
from services.core.vkr_core.engine.heading_scoring import (
    TIER_HIGH,
    TIER_LOW,
    TIER_MEDIUM,
    TIER_NONE,
    score_paragraph,
)
from services.core.vkr_core.engine.stats import DocStats, ParagraphStats, collect_stats
from services.core.vkr_core.engine.title_page import detect_title_page_end

# ──────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────

def _build(tmp_path: Path, paragraphs) -> WordDocument:
    """Build a minimal DOCX from a list of specs.

    Each spec is either a plain text string OR a tuple
    (text, bold, centered, size).
    """
    doc = WordDocument()
    for spec in paragraphs:
        if isinstance(spec, tuple):
            text, bold, centered, size = spec
        else:
            text, bold, centered, size = spec, False, False, None
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    path = tmp_path / "doc.docx"
    doc.save(path)
    return WordDocument(str(path))


def _detect(doc: WordDocument):
    rules = load_default_rules()
    stats = collect_stats(doc)
    title_end = detect_title_page_end(stats, rules)
    headings = detect_headings(stats, rules, set(range(title_end)))
    return stats, headings


def _detect_full(doc: WordDocument):
    rules = load_default_rules()
    stats = collect_stats(doc)
    title_end = detect_title_page_end(stats, rules)
    return detect_headings_full(stats, rules, set(range(title_end)))


def _make_ps(
    index: int,
    text: str,
    *,
    bold_ratio: float = 0.0,
    centered: bool = False,
    font_size: float | None = None,
    first_line_indent_cm: float | None = 0.0,
    style_name: str = "Normal",
    numpr_ilvl: int | None = None,
    outline_level: int | None = None,
    has_page_break_before: bool = False,
    has_tab: bool = False,
    in_table: bool = False,
) -> ParagraphStats:
    from docx.enum.text import WD_ALIGN_PARAGRAPH as A
    stripped = text.strip()
    import re as _re
    tokens_match = _re.match(r"^(\d+(?:\.\d+){0,3})[.)]?\s+", stripped)
    tokens = (
        tuple(int(x) for x in tokens_match.group(1).split("."))
        if tokens_match else tuple()
    )
    derived = tokens_match.group(1) if tokens_match else None
    return ParagraphStats(
        index=index,
        text=text,
        stripped=stripped,
        length=len(stripped),
        word_count=len(stripped.split()),
        alignment=A.CENTER if centered else A.LEFT,
        first_line_indent_cm=first_line_indent_cm,
        left_indent_cm=0.0,
        max_font_size_pt=font_size,
        modal_font_size_pt=font_size,
        font_family="Times New Roman",
        bold_ratio=bold_ratio,
        italic_ratio=0.0,
        ends_with_period=stripped.endswith((".", "!", "?")),
        is_upper=stripped.isupper() and bool(stripped),
        is_empty=not stripped,
        in_table=in_table,
        style_name=style_name,
        numbered=bool(tokens),
        numbering_tokens=tokens,
        has_tab=has_tab,
        numpr_ilvl=numpr_ilvl,
        numpr_num_id=1 if numpr_ilvl is not None else None,
        outline_level=outline_level,
        has_page_break_before=has_page_break_before,
        derived_number=derived,
    )


def _make_stats(paragraphs: list[ParagraphStats], body_font: float = 14.0) -> DocStats:
    stats = DocStats()
    stats.paragraphs = paragraphs
    stats.modal_font_size_pt = body_font
    stats.modal_font_family = "Times New Roman"
    return stats


# ──────────────────────────────────────────────────────────────────────
# Integration tests (python-docx based)
# ──────────────────────────────────────────────────────────────────────

def test_numbered_pattern_is_unconditional_heading(tmp_path: Path) -> None:
    doc = _build(
        tmp_path,
        [
            ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", True, True, 16),
            ("", False, False, None),
            ("ВВЕДЕНИЕ", True, True, 14),
            "Тело " * 40,
            ("1 Анализ методов", True, True, 14),
            "Тело " * 40,
        ],
    )
    _, headings = _detect(doc)
    titles = [h.text for h in headings]
    assert any("1 Анализ методов" in t for t in titles)


def test_body_paragraph_is_not_classified_as_heading(tmp_path: Path) -> None:
    doc = _build(
        tmp_path,
        [
            ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", True, True, 16),
            ("", False, False, None),
            ("ВВЕДЕНИЕ", True, True, 14),
            "В тексте приведены результаты [1] и формула a=b. " * 5,
        ],
    )
    _, headings = _detect(doc)
    for h in headings:
        assert "результаты" not in h.text.lower()


def test_forbidden_heading_is_detected(tmp_path: Path) -> None:
    """Genuinely forbidden headings ("Теоретическая часть") are hard-rejected.

    Note: «ОГЛАВЛЕНИЕ» is NOT forbidden — it's a structural-keyword alias that
    will be renamed to «СОДЕРЖАНИЕ» by apply_forbidden_replacements. Here we
    verify the hard-reject path for truly banned phrases.
    """
    doc = _build(
        tmp_path,
        [
            ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", True, True, 16),
            ("", False, False, None),
            ("ВВЕДЕНИЕ", True, True, 14),
            "Тело " * 30,
            ("Теоретическая часть", True, True, 14),
            "Тело " * 30,
        ],
    )
    _, headings = _detect(doc)
    assert all("теоретическая часть" not in h.title_only.lower() for h in headings)


def test_three_consecutive_subsections_all_detected(tmp_path: Path) -> None:
    # The original bug: 1.1.1, 1.1.2, 1.1.3 in sequence — middle one used to be
    # dropped as a "list item". Score-based detection must not do that.
    doc = _build(
        tmp_path,
        [
            ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", True, True, 16),
            ("", False, False, None),
            ("ВВЕДЕНИЕ", True, True, 14),
            "Тело " * 40,
            ("1.1.1 Первый подраздел", True, False, 14),
            "Тело первого подраздела. " * 15,
            ("1.1.2 Проблемы организации работы репетитора", True, False, 14),
            "Тело второго подраздела. " * 15,
            ("1.1.3 Третий подраздел", True, False, 14),
            "Тело третьего подраздела. " * 15,
        ],
    )
    _, headings = _detect(doc)
    titles = [h.text for h in headings]
    assert any("1.1.1 Первый" in t for t in titles)
    assert any("1.1.2 Проблемы организации" in t for t in titles)
    assert any("1.1.3 Третий" in t for t in titles)


def test_chapter_word_is_heading(tmp_path: Path) -> None:
    doc = _build(
        tmp_path,
        [
            ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", True, True, 16),
            ("", False, False, None),
            ("ВВЕДЕНИЕ", True, True, 14),
            "Тело " * 30,
            ("Глава 1 Обзор литературы", True, False, 14),
            "Тело " * 30,
        ],
    )
    _, headings = _detect(doc)
    assert any("Глава 1" in h.text for h in headings)


def test_roman_chapter_word_is_heading(tmp_path: Path) -> None:
    doc = _build(
        tmp_path,
        [
            ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", True, True, 16),
            ("", False, False, None),
            ("ВВЕДЕНИЕ", True, True, 14),
            "Тело " * 30,
            ("Глава I. Обзор литературы", True, False, 14),
            "Тело " * 30,
        ],
    )
    _, headings = _detect(doc)
    hit = [h for h in headings if "Глава I" in h.text]
    assert hit
    assert hit[0].level == 1
    assert hit[0].title_only == "Обзор литературы"


def test_strip_chapter_word_everywhere_normalizes_chapter_heading_text() -> None:
    doc = WordDocument()
    p1 = doc.add_paragraph("ГЛАВА 3 Архитектура сервиса")
    p2 = doc.add_paragraph("Глава IV. Экспериментальная часть")

    changed = strip_chapter_word_everywhere(doc, skip_indexes=set())

    assert changed == 2
    assert p1.text == "3 Архитектура сервиса"
    assert p2.text == "4 Экспериментальная часть"


def test_compact_numbered_heading_is_detected(tmp_path: Path) -> None:
    doc = _build(
        tmp_path,
        [
            ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", True, True, 16),
            ("", False, False, None),
            ("ВВЕДЕНИЕ", True, True, 14),
            "Тело " * 30,
            ("1.1.Анализ предметной области", True, False, 14),
            "Тело " * 30,
        ],
    )
    _, headings = _detect(doc)
    hit = [h for h in headings if "1.1.Анализ" in h.text]
    assert hit
    assert hit[0].level == 2
    assert hit[0].derived_number == "1.1"
    assert hit[0].title_only == "Анализ предметной области"


def test_appendix_is_heading(tmp_path: Path) -> None:
    doc = _build(
        tmp_path,
        [
            ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", True, True, 16),
            ("", False, False, None),
            ("ВВЕДЕНИЕ", True, True, 14),
            "Тело " * 30,
            ("Приложение А Исходные данные", True, False, 14),
            "Таблица " * 20,
        ],
    )
    _, headings = _detect(doc)
    assert any("Приложение А" in h.text for h in headings)


def test_structural_keywords_detected(tmp_path: Path) -> None:
    doc = _build(
        tmp_path,
        [
            ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", True, True, 16),
            ("", False, False, None),
            ("ВВЕДЕНИЕ", True, True, 14),
            "Тело " * 30,
            ("ЗАКЛЮЧЕНИЕ", True, True, 14),
            "Тело " * 30,
            ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", True, True, 14),
            "1. Иванов И. И. Название. — М., 2020." * 3,
        ],
    )
    _, headings = _detect(doc)
    texts = {h.text for h in headings}
    assert "ВВЕДЕНИЕ" in texts
    assert "ЗАКЛЮЧЕНИЕ" in texts


def test_caption_is_hard_rejected(tmp_path: Path) -> None:
    doc = _build(
        tmp_path,
        [
            ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", True, True, 16),
            ("", False, False, None),
            ("ВВЕДЕНИЕ", True, True, 14),
            "Тело " * 40,
            ("Рисунок 3.1 — Диаграмма потоков данных", True, True, 16),
            "Тело " * 40,
        ],
    )
    _, headings = _detect(doc)
    assert all("рисунок" not in h.text.lower() for h in headings)


def test_too_long_paragraph_rejected(tmp_path: Path) -> None:
    long_text = "Очень длинный абзац текста "[:30] * 20  # ≥ 200 chars
    doc = _build(
        tmp_path,
        [
            ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", True, True, 16),
            ("", False, False, None),
            ("ВВЕДЕНИЕ", True, True, 14),
            "Тело " * 30,
            (long_text, True, True, 16),
            "Тело " * 30,
        ],
    )
    _, headings = _detect(doc)
    # The long-text paragraph is not accepted as a heading.
    for h in headings:
        assert len(h.text) < 200


# ──────────────────────────────────────────────────────────────────────
# Unit tests (pure scoring)
# ──────────────────────────────────────────────────────────────────────

def test_score_numpr_heading_without_text_number() -> None:
    """Word auto-numbered heading: paragraph.text lacks "1.1.2" because numPr
    generates it. Score must still classify as heading via numpr_ilvl +
    formatting signals."""
    target = _make_ps(
        2,
        "Проблемы организации работы репетитора",
        bold_ratio=1.0,
        font_size=14.0,
        numpr_ilvl=2,
    )
    long_body = "Тело " * 30
    paragraphs = [
        _make_ps(0, "Предыдущий длинный абзац. " + long_body, font_size=14.0),
        _make_ps(1, "", font_size=None),
        target,
        _make_ps(3, "", font_size=None),
        _make_ps(4, "Следующий длинный абзац. " + long_body, font_size=14.0),
    ]
    stats = _make_stats(paragraphs)
    score = score_paragraph(target, stats, {}, set())
    assert score.tier in (TIER_HIGH, TIER_MEDIUM), (
        f"expected HIGH/MEDIUM, got {score.tier} (score={score.score}, signals={score.signals})"
    )
    assert score.level == 3


def test_score_outline_level_alone_is_heading() -> None:
    """Paragraph with only outline_level=1 and no other signals must be HIGH."""
    target = _make_ps(1, "Раздел без форматирования", outline_level=1)
    paragraphs = [
        _make_ps(0, "Тело " * 30),
        target,
        _make_ps(2, "Тело " * 30),
    ]
    stats = _make_stats(paragraphs)
    score = score_paragraph(target, stats, {}, set())
    assert score.tier == TIER_HIGH
    assert score.level == 2  # outline_level=1 → level 2


def test_score_numbered_list_with_hanging_indent_rejected() -> None:
    """Numbered list item ("1. Item") with hanging indent and body formatting
    must be hard-rejected, regardless of literal number."""
    target = _make_ps(
        2,
        "1. Первый пункт списка",
        bold_ratio=0.0,
        font_size=14.0,
        first_line_indent_cm=-0.5,
    )
    stats = _make_stats([
        _make_ps(0, "Тело абзаца перед списком. " * 10),
        _make_ps(1, "", font_size=None),
        target,
        _make_ps(3, "2. Второй пункт списка", first_line_indent_cm=-0.5),
        _make_ps(4, "3. Третий пункт списка", first_line_indent_cm=-0.5),
    ])
    score = score_paragraph(target, stats, {}, set())
    assert score.hard_reject_reason == "hanging_indent"


def test_score_compact_numbered_list_sequence_rejected() -> None:
    target = _make_ps(1, "1.Первый пункт списка", font_size=14.0)
    stats = _make_stats([
        _make_ps(0, "Тело абзаца перед списком. " * 10),
        target,
        _make_ps(2, "2.Второй пункт списка", font_size=14.0),
        _make_ps(3, "3.Третий пункт списка", font_size=14.0),
    ])
    score = score_paragraph(target, stats, {}, set())
    assert score.hard_reject_reason == "list_sequence"


def test_score_dense_text_rejected() -> None:
    """A long paragraph with many commas must not reach heading tier even if
    it starts with a number."""
    text = (
        "1. В ходе исследования были рассмотрены различные аспекты, "
        "включая методы сбора данных, их анализ, обобщение результатов"
    )
    target = _make_ps(1, text, font_size=14.0)
    stats = _make_stats([
        _make_ps(0, "Тело. " * 30),
        target,
        _make_ps(2, "Тело. " * 30),
    ])
    score = score_paragraph(target, stats, {}, set())
    assert score.tier in (TIER_NONE, TIER_LOW), (
        f"dense text should not be accepted, got {score.tier}"
    )


def test_score_semantic_heading_is_low() -> None:
    """Short isolated line with capital start and no period — semantic candidate."""
    target = _make_ps(
        1,
        "Методология исследования",
        bold_ratio=0.0,
        font_size=14.0,
    )
    stats = _make_stats([
        _make_ps(0, "Длинный абзац перед. " * 10),
        target,
        _make_ps(2, "Длинный абзац после. " * 10),
    ])
    score = score_paragraph(target, stats, {}, set())
    # isolated (3) + short (1) + no_period (1) = 5 → LOW
    assert score.tier == TIER_LOW


def test_score_in_table_is_rejected() -> None:
    target = _make_ps(1, "1.1 Раздел", in_table=True, bold_ratio=1.0, font_size=14.0)
    stats = _make_stats([target])
    score = score_paragraph(target, stats, {}, set())
    assert score.hard_reject_reason == "in_table"


def test_score_title_page_paragraph_rejected() -> None:
    target = _make_ps(3, "1 Раздел", bold_ratio=1.0, font_size=14.0)
    stats = _make_stats([target])
    score = score_paragraph(target, stats, {}, {3})
    assert score.hard_reject_reason == "title_page"


def test_score_forbidden_heading_rejected() -> None:
    target = _make_ps(
        1, "Теоретическая часть", bold_ratio=1.0, font_size=14.0, centered=True,
    )
    stats = _make_stats([target])
    rules = {"forbidden_heading_map": {"теоретическая часть": None}}
    score = score_paragraph(target, stats, rules, set())
    assert score.hard_reject_reason == "forbidden"


def test_native_heading_with_trailing_colon_is_detected() -> None:
    target = _make_ps(
        1,
        "1.1 Нормативная основа и процесс подготовки ВКР:",
        style_name="Heading 2",
    )
    stats = _make_stats([
        _make_ps(0, "Тело " * 30),
        target,
        _make_ps(2, "Тело " * 30),
    ])
    score = score_paragraph(target, stats, {}, set())
    assert score.tier == TIER_HIGH
    assert score.level == 2


def test_native_heading_with_long_period_title_is_detected() -> None:
    target = _make_ps(
        1,
        "3 Разработка и тестирование веб-сервиса автоматизированной проверки и оформления ВКР.",
        style_name="Heading 1",
    )
    stats = _make_stats([
        _make_ps(0, "Тело " * 30),
        target,
        _make_ps(2, "Тело " * 30),
    ])
    score = score_paragraph(target, stats, {}, set())
    assert score.tier == TIER_HIGH
    assert score.level == 1


def test_toc_like_heading_style_with_tab_page_number_is_rejected() -> None:
    target = _make_ps(
        1,
        "3 Разработка и тестирование веб-сервиса\t32 .",
        style_name="Heading 1",
        has_tab=True,
    )
    stats = _make_stats([target])
    score = score_paragraph(target, stats, {}, set())
    assert score.hard_reject_reason == "toc_entry"


def test_visible_multilevel_number_overrides_wrong_heading_one_level() -> None:
    target = _make_ps(
        0,
        "1.3 Подраздел с точкой.",
        style_name="Heading 1",
        outline_level=0,
    )
    stats = _make_stats([target])

    score = score_paragraph(target, stats, load_default_rules(), set())

    assert score.tier == TIER_HIGH
    assert score.level == 2
    assert "visible_number_level:1.3" in score.signals


def test_numbered_period_text_without_context_is_not_auto_heading() -> None:
    target = _make_ps(0, "1.1 Обычный нумерованный абзац.")
    stats = _make_stats([target])

    score = score_paragraph(target, stats, load_default_rules(), set())

    assert score.tier == TIER_NONE


def test_chapter_word_compact_line_not_rejected_as_list_item() -> None:
    target = _make_ps(
        1,
        "Глава 3 Архитектура сервиса",
        numpr_ilvl=0,
        style_name="List Paragraph",
    )
    stats = _make_stats([
        _make_ps(0, "Тело " * 30),
        target,
        _make_ps(2, "Тело " * 30),
    ])

    score = score_paragraph(target, stats, load_default_rules(), set())

    assert score.hard_reject_reason is None
    assert score.tier == TIER_HIGH
    assert "chapter_word" in score.signals


def test_sentence_with_word_glava_is_not_heading() -> None:
    target = _make_ps(
        1,
        "В этой главе 3 рассмотрены особенности реализации подсистемы.",
        numpr_ilvl=0,
        style_name="List Paragraph",
    )
    stats = _make_stats([
        _make_ps(0, "Тело " * 30),
        target,
        _make_ps(2, "Тело " * 30),
    ])

    score = score_paragraph(target, stats, load_default_rules(), set())

    assert score.tier in (TIER_NONE, TIER_LOW)


# ──────────────────────────────────────────────────────────────────────
# Hierarchy tests
# ──────────────────────────────────────────────────────────────────────

def test_gap_recovery_promotes_low_candidate(tmp_path: Path) -> None:
    """HIGH 1.1.1, LOW 1.1.2 (unformatted), HIGH 1.1.3 → 1.1.2 is recovered."""
    doc = _build(
        tmp_path,
        [
            ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", True, True, 16),
            ("", False, False, None),
            ("ВВЕДЕНИЕ", True, True, 14),
            "Тело " * 40,
            ("1.1.1 Сформированный подраздел", True, False, 14),
            "Тело " * 25,
            # 1.1.2 without formatting (no bold, no centering) — LOW tier
            ("1.1.2 Слабо оформленный подраздел", False, False, 14),
            "Тело " * 25,
            ("1.1.3 Третий подраздел", True, False, 14),
            "Тело " * 25,
        ],
    )
    headings, soft = _detect_full(doc)
    titles = [h.text for h in headings]
    assert any("1.1.2 Слабо" in t for t in titles), (
        f"1.1.2 must be recovered; got {titles}"
    )


def test_first_child_heading_recovered_only_with_context(tmp_path: Path) -> None:
    doc = _build(
        tmp_path,
        [
            ("1 Анализ предметной области", True, True, 14),
            ("1.1 Нормативная основа.", False, False, 16),
            "Основной текст подраздела. " * 20,
            ("1.2 Следующий подраздел", True, False, 14),
            "Основной текст следующего подраздела. " * 20,
        ],
    )
    rules = load_default_rules()
    stats = collect_stats(doc)

    headings, soft = detect_headings_full(stats, rules, set())

    hit = [h for h in headings if h.text.startswith("1.1 Нормативная")]
    assert hit, f"1.1 should be contextually recovered; got {[h.text for h in headings]}"
    assert hit[0].level == 2
    assert any(
        v.type == "heading_recovered" and v.paragraph_index == hit[0].paragraph_index
        for v in soft
    )


def test_medium_tier_emits_confirm_violation(tmp_path: Path) -> None:
    """A short bold-only paragraph without number should land in MEDIUM and
    produce a manual_required 'confirm heading' violation."""
    doc = _build(
        tmp_path,
        [
            ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", True, True, 16),
            ("", False, False, None),
            ("ВВЕДЕНИЕ", True, True, 14),
            "Тело перед. " * 25,
            ("Методология исследования", True, False, 14),
            "Тело после. " * 25,
        ],
    )
    headings, soft = _detect_full(doc)
    confirm = [v for v in soft if v.type == "heading_confirm"]
    # Either the paragraph is accepted as heading (then MEDIUM with confirm
    # violation) or it's LOW-semantic (possible_missed_heading). Both surface
    # the paragraph to the user.
    assert confirm or any(v.type == "possible_missed_heading" for v in soft)


def test_soft_violations_use_manual_required_status(tmp_path: Path) -> None:
    from services.core.vkr_core.models.enums import ViolationStatus

    doc = _build(
        tmp_path,
        [
            ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", True, True, 16),
            ("", False, False, None),
            ("ВВЕДЕНИЕ", True, True, 14),
            "Тело перед. " * 25,
            ("Короткая строка без форматирования", False, False, 14),
            "Тело после. " * 25,
        ],
    )
    _, soft = _detect_full(doc)
    for v in soft:
        if v.type in ("heading_confirm", "possible_missed_heading"):
            assert v.status == ViolationStatus.manual_required


def test_toc_title_never_flagged_as_possible_heading(tmp_path: Path) -> None:
    """«СОДЕРЖАНИЕ»/«ОГЛАВЛЕНИЕ» — раздел, которым владеет конвейер. Он не
    должен попадать в heading_confirm / possible_missed_heading: иначе в
    превью ложно подчёркивается слово «содержание» в обычном тексте."""
    for toc_word in ("СОДЕРЖАНИЕ", "ОГЛАВЛЕНИЕ"):
        doc = _build(
            tmp_path,
            [
                ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", True, True, 16),
                ("", False, False, None),
                ("ВВЕДЕНИЕ", True, True, 14),
                (toc_word, True, True, 14),
                "Научный руководитель отвечает за содержание ВКР. " * 8,
            ],
        )
        _, soft = _detect_full(doc)
        offenders = [
            v for v in soft
            if v.type in ("heading_confirm", "possible_missed_heading")
            and (v.section_title or "").strip().lower() == toc_word.lower()
        ]
        assert not offenders, f"{toc_word} ошибочно помечен: {offenders}"
