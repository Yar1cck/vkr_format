"""Тесты formulas.py — детекция формул и валидация (§7.2)."""
from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH

from services.core.vkr_core.engine.formulas import (
    detect_formulas,
    validate_formulas,
)
from services.core.vkr_core.engine.stats import DocStats, ParagraphStats


def _ps(idx: int, text: str, *, centred: bool = False, in_table: bool = False) -> ParagraphStats:
    return ParagraphStats(
        index=idx, text=text, stripped=text.strip(), length=len(text),
        word_count=len(text.split()),
        alignment=WD_ALIGN_PARAGRAPH.CENTER if centred else WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_cm=None, left_indent_cm=None,
        max_font_size_pt=14.0, modal_font_size_pt=14.0, font_family="Times New Roman",
        bold_ratio=0.0, italic_ratio=0.0,
        ends_with_period=text.endswith("."), is_upper=text.isupper(),
        is_empty=not text.strip(), in_table=in_table, style_name="Normal",
        numbered=False, numbering_tokens=(), has_tab=False,
    )


def _stats(*items) -> DocStats:
    paragraphs = []
    for i, item in enumerate(items):
        if isinstance(item, tuple):
            text, *flags = item
            centred = "centred" in flags
            in_table = "in_table" in flags
            paragraphs.append(_ps(i, text, centred=centred, in_table=in_table))
        else:
            paragraphs.append(_ps(i, item))
    return DocStats(paragraphs=paragraphs)


def test_detect_centred_formula() -> None:
    stats = _stats(("y = α · x + β (1)", "centred"))
    formulas = detect_formulas(stats, set())
    assert len(formulas) == 1
    assert formulas[0].number == "1"


def test_detect_left_aligned_with_number() -> None:
    # Не по центру, но с номером — считается формулой
    stats = _stats("E = m × c² (2)")
    formulas = detect_formulas(stats, set())
    assert len(formulas) == 1
    assert formulas[0].number == "2"


def test_no_formula_without_equals() -> None:
    stats = _stats(("α + β + γ", "centred"))
    assert detect_formulas(stats, set()) == []


def test_no_formula_when_in_table() -> None:
    stats = _stats(("y = α · x (1)", "centred", "in_table"))
    assert detect_formulas(stats, set()) == []


def test_no_formula_when_too_long() -> None:
    long_text = "y = " + "α " * 70 + "(1)"
    stats = _stats((long_text, "centred"))
    assert detect_formulas(stats, set()) == []


def test_skip_indexes_excluded() -> None:
    stats = _stats(("y = α · x (1)", "centred"))
    assert detect_formulas(stats, {0}) == []


def test_validate_formula_without_explanation() -> None:
    stats = _stats(
        ("y = α · x (1)", "centred"),
        "Обычный текст без где.",
    )
    formulas = detect_formulas(stats, set())
    violations = validate_formulas(stats, formulas)
    assert any(v.type == "formula_without_explanation" for v in violations)


def test_validate_formula_with_explanation_no_violation() -> None:
    stats = _stats(
        ("y = α · x + β (1)", "centred"),
        "где α — коэффициент, β — смещение.",
    )
    formulas = detect_formulas(stats, set())
    violations = validate_formulas(stats, formulas)
    assert not any(v.type == "formula_without_explanation" for v in violations)


def test_validate_numbering_gap() -> None:
    stats = _stats(
        ("y = α (1)", "centred"),
        "где α — параметр.",
        ("z = β (3)", "centred"),  # пропущен (2)
        "где β — другой параметр.",
    )
    formulas = detect_formulas(stats, set())
    violations = validate_formulas(stats, formulas)
    gaps = [v for v in violations if v.type == "formula_numbering_gap"]
    assert len(gaps) == 1


def test_validate_numbering_continuous_no_gap() -> None:
    stats = _stats(
        ("y = α (1)", "centred"),
        "где α — параметр.",
        ("z = β (2)", "centred"),
        "где β — другой параметр.",
    )
    formulas = detect_formulas(stats, set())
    violations = validate_formulas(stats, formulas)
    assert not any(v.type == "formula_numbering_gap" for v in violations)


# ── F4: номер формулы должен быть справа (правый таб) ──────────────────────

def _ps_with_tab(idx: int, text: str, *, has_tab: bool) -> ParagraphStats:
    return ParagraphStats(
        index=idx, text=text, stripped=text.strip(), length=len(text),
        word_count=len(text.split()),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_cm=None, left_indent_cm=None,
        max_font_size_pt=14.0, modal_font_size_pt=14.0, font_family="Times New Roman",
        bold_ratio=0.0, italic_ratio=0.0,
        ends_with_period=False, is_upper=False,
        is_empty=False, in_table=False, style_name="Normal",
        numbered=False, numbering_tokens=(), has_tab=has_tab,
    )


def test_formula_number_not_right_violation() -> None:
    # Номер есть, таб отсутствует → нарушение
    ps = _ps_with_tab(0, "y = α · x (1)", has_tab=False)
    stats = DocStats(paragraphs=[ps, _ps(1, "где α — параметр.")])
    formulas = detect_formulas(stats, set())
    assert formulas[0].has_tab is False
    violations = validate_formulas(stats, formulas)
    assert any(v.type == "formula_number_not_right" for v in violations)


def test_formula_number_right_no_violation() -> None:
    # Номер есть и таб есть → нарушения нет
    ps = _ps_with_tab(0, "y = α · x\t(1)", has_tab=True)
    stats = DocStats(paragraphs=[ps, _ps(1, "где α — параметр.")])
    formulas = detect_formulas(stats, set())
    assert formulas[0].has_tab is True
    violations = validate_formulas(stats, formulas)
    assert not any(v.type == "formula_number_not_right" for v in violations)


def test_formula_without_number_no_right_violation() -> None:
    # Центрированная формула без номера → F4 не применяется
    ps = _ps_with_tab(0, "y = α · x + β", has_tab=False)
    stats = DocStats(paragraphs=[ps, _ps(1, "где α — параметр.")])
    formulas = detect_formulas(stats, set())
    violations = validate_formulas(stats, formulas)
    assert not any(v.type == "formula_number_not_right" for v in violations)


def test_fix_formula_number_alignment() -> None:
    from docx import Document as WordDocument

    from services.core.vkr_core.config.normative_loader import load_default_rules
    from services.core.vkr_core.engine.formatter import fix_formula_number_alignment

    doc = WordDocument()
    para = doc.add_paragraph("y = α · x (1)")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rules = load_default_rules()
    formula = type("F", (), {"paragraph_index": 0, "number": "1", "has_tab": False})()
    n = fix_formula_number_alignment(doc, [formula], rules)
    assert n == 1

    # После фикса в абзаце должен быть элемент <w:tab>
    from docx.oxml.ns import qn
    p_elem = doc.paragraphs[0]._p
    tab_elements = p_elem.findall(f".//{qn('w:tab')}")
    assert tab_elements, "ожидался <w:tab/> в XML абзаца"

    # Текст после фикса содержит номер
    assert "(1)" in doc.paragraphs[0].text
