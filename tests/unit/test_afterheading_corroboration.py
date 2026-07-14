"""Регресс: after_heading сам по себе — недостаточное основание.

Абзац, стоящий первым после заголовка, флагается как
possible_missed_heading только если у него есть ≥1 настоящего признака
заголовка (жирный / по центру / ВСЕ ЗАГЛАВНЫЕ / структурное слово /
номер-префикс / native-стиль). Голый фрагмент прозы — не «возможный
заголовок».
"""
from __future__ import annotations

from pathlib import Path

from docx import Document as WordDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from services.core.vkr_core.config.normative_loader import load_default_rules
from services.core.vkr_core.engine.detection import detect_headings_full
from services.core.vkr_core.engine.stats import collect_stats
from services.core.vkr_core.engine.title_page import detect_title_page_end


def _build(tmp_path: Path, paragraphs) -> WordDocument:
    doc = WordDocument()
    for spec in paragraphs:
        if isinstance(spec, tuple):
            text, bold, centered, size = spec
        else:
            text, bold, centered, size = spec, False, False, None
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        p.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
        )
    path = tmp_path / "doc.docx"
    doc.save(path)
    return WordDocument(str(path))


def _soft(doc: WordDocument):
    rules = load_default_rules()
    stats = collect_stats(doc)
    title_end = detect_title_page_end(stats, rules)
    _, soft = detect_headings_full(stats, rules, set(range(title_end)))
    return soft


def _pmh(soft):
    return [v for v in soft if v.type == "possible_missed_heading"]


def test_plain_prose_after_heading_is_not_flagged(tmp_path: Path) -> None:
    doc = _build(
        tmp_path,
        [
            ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", True, True, 16),
            ("", False, False, None),
            ("ВВЕДЕНИЕ", True, True, 14),
            # Кандидат: короткий, с заглавной, без точки, но БЕЗ форматирования.
            ("Научный руководитель курирует подготовку работы", False, False, 14),
            "Тело после введения с достаточной длиной абзаца. " * 6,
        ],
    )
    pmh = _pmh(_soft(doc))
    assert all(
        "научный руководитель" not in (v.section_title or "").lower() for v in pmh
    ), f"голый фрагмент прозы не должен флагаться: {[v.section_title for v in pmh]}"


def test_formatted_line_after_heading_is_flagged_with_signal(tmp_path: Path) -> None:
    doc = _build(
        tmp_path,
        [
            ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", True, True, 16),
            ("", False, False, None),
            ("ВВЕДЕНИЕ", True, True, 14),
            # Тот же по позиции кандидат, но ЖИРНЫЙ — есть признак заголовка.
            ("Методология исследования", True, False, 14),
            "Тело после с достаточной длиной обычного абзаца тела. " * 6,
        ],
    )
    pmh = _pmh(_soft(doc))
    hit = [v for v in pmh if "методология" in (v.section_title or "").lower()]
    assert hit, "жирная строка после заголовка должна оставаться кандидатом"
    sigs = hit[0].detector_signals or []
    assert "after_heading" in sigs
    assert any(s in sigs for s in ("bold", "bold_partial", "all_caps", "centered")), (
        f"должен быть подтверждающий сигнал, а есть: {sigs}"
    )


def test_plain_title_after_bare_chapter_is_flagged(tmp_path: Path) -> None:
    doc = _build(
        tmp_path,
        [
            ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", True, True, 16),
            ("", False, False, None),
            ("ВВЕДЕНИЕ", True, True, 14),
            "Тело перед главой с достаточной длиной обычного абзаца. " * 6,
            ("ГЛАВА 1", True, True, 14),
            ("Анализ предметной области", False, False, 14),
            "Тело после с достаточной длиной обычного абзаца тела. " * 6,
        ],
    )
    pmh = _pmh(_soft(doc))
    hit = [v for v in pmh if "анализ предметной области" in (v.section_title or "").lower()]
    assert hit, "название после отдельной строки ГЛАВА N должно быть подсказкой"
    assert "after_bare_chapter" in (hit[0].detector_signals or [])


def test_contents_title_after_heading_still_suppressed(tmp_path: Path) -> None:
    doc = _build(
        tmp_path,
        [
            ("ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", True, True, 16),
            ("", False, False, None),
            ("ВВЕДЕНИЕ", True, True, 14),
            ("СОДЕРЖАНИЕ", True, True, 14),
            "Тело после с достаточной длиной обычного абзаца тела. " * 6,
        ],
    )
    pmh = _pmh(_soft(doc))
    assert all(
        (v.section_title or "").strip().lower() != "содержание" for v in pmh
    )
