from __future__ import annotations

from docx import Document

from services.core.vkr_core.config.normative_loader import load_default_rules
from services.core.vkr_core.engine.captions import renumber_figures
from services.core.vkr_core.engine.pipeline import _caption_chapter_breaks_for_mode


def test_default_caption_numbering_uses_chapter_breaks() -> None:
    rules = load_default_rules()
    breaks = [(0, 1), (2, 2)]

    assert _caption_chapter_breaks_for_mode(rules, "figure", breaks) == breaks
    assert _caption_chapter_breaks_for_mode(rules, "table", breaks) == breaks
    assert _caption_chapter_breaks_for_mode(rules, "listing", breaks) == breaks


def test_continuous_caption_numbering_disables_chapter_breaks() -> None:
    rules = load_default_rules()
    rules["numbering_rules"]["figure_mode"] = "continuous"

    assert _caption_chapter_breaks_for_mode(rules, "figure", [(0, 1)]) is None


def test_renumber_figures_by_default_config_per_chapter() -> None:
    doc = Document()
    doc.add_paragraph("1 Первая глава")
    doc.add_paragraph("Рисунок 1 — Первый")
    doc.add_paragraph("2 Вторая глава")
    doc.add_paragraph("Рисунок 2 — Второй")

    breaks = _caption_chapter_breaks_for_mode(
        load_default_rules(),
        "figure",
        [(0, 1), (2, 2)],
    )
    figures, _changes = renumber_figures(doc.paragraphs, set(), apply=False, chapter_breaks=breaks)

    assert [figure.new_number for figure in figures] == ["1.1", "2.1"]
