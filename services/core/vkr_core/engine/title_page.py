"""Обнаружение титульного листа / задания.

Титульные листы переформатировать нельзя: ТЗ §6.1 выводит первую секцию
из-под действия правил.
"""

from __future__ import annotations

import re

from docx.enum.text import WD_ALIGN_PARAGRAPH

from services.core.vkr_core.engine.stats import DocStats

DEFAULT_TITLE_MARKERS = (
    "выпускная квалификационная работа",
    "курсовая работа",
    "дипломная работа",
    "бакалаврская работа",
    "вкр",
)
CONTEXT_MARKERS = ("университет", "факультет", "кафедра", "по дисциплине", "тема")
H1_RE = re.compile(r"^\d+[.)]?\s+\S")
# «Глава N», «Раздел N», «Chapter N» — тоже начало тела документа.
_CHAPTER_WORD_BOUNDARY_RE = re.compile(
    r"^(?:глава|раздел|часть|chapter|section|part)\s+\d+", re.IGNORECASE
)


def _norm(text: str) -> str:
    return " ".join(text.strip().lower().split())


def detect_title_page_end(stats: DocStats, rules: dict) -> int:
    """Возвращает индекс абзаца, с которого начинается тело документа.

    Всё с индексом < возвращаемого значения относится к титульному листу
    или странице задания и должно сохраняться в неизменном виде.
    """
    structural = rules.get("structural_elements", {})
    markers = {
        _norm(item)
        for item in structural.get("title_page", []) + structural.get("task", [])
    }
    markers.update(_norm(item) for item in DEFAULT_TITLE_MARKERS)

    scan_limit = min(len(stats.paragraphs), 60)
    has_title_marker = False
    centered_short = 0
    context_hits = 0

    for i in range(scan_limit):
        ps = stats.paragraphs[i]
        if ps.is_empty:
            continue
        normalized = _norm(ps.stripped)
        if any(marker in normalized for marker in markers if marker):
            has_title_marker = True
            break
        if ps.alignment == WD_ALIGN_PARAGRAPH.CENTER and ps.length <= 120:
            centered_short += 1
        if any(marker in normalized for marker in CONTEXT_MARKERS):
            context_hits += 1

    if not has_title_marker:
        if centered_short >= 4 and context_hits >= 2:
            has_title_marker = True
        else:
            return 0

    boundary = {
        _norm(item) for item in structural.get("contents", []) + structural.get("introduction", [])
    }
    boundary.discard("")

    for i, ps in enumerate(stats.paragraphs):
        if ps.is_empty:
            continue
        if _norm(ps.stripped) in boundary:
            return i
        if H1_RE.match(ps.stripped):
            return i
        if _CHAPTER_WORD_BOUNDARY_RE.match(ps.stripped):
            return i
        # Заголовок уровня 1, обнаруженный по метаданным Word (numPr или
        # outline level) — у текста может не быть текстового номера, если
        # включена автонумерация.
        if ps.outline_level == 0 or ps.numpr_ilvl == 0:
            return i

    return min(len(stats.paragraphs), 25)
