"""DocStats — проход сбора статистики (ТЗ §6.5.1).

Собирает признаки абзацев (размер шрифта, доля жирного, выравнивание,
длина) и распределения по документу (модальный шрифт тела, кластеры
шрифтов), которые использует детектор заголовков и другие детерминированные
проверки.
"""

from __future__ import annotations

import statistics
from collections import Counter
from dataclasses import dataclass, field

from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


@dataclass
class ParagraphStats:
    index: int
    text: str
    stripped: str
    length: int
    word_count: int
    alignment: WD_ALIGN_PARAGRAPH | None
    first_line_indent_cm: float | None
    left_indent_cm: float | None
    max_font_size_pt: float | None
    modal_font_size_pt: float | None
    font_family: str | None
    bold_ratio: float
    italic_ratio: float
    ends_with_period: bool
    is_upper: bool
    is_empty: bool
    in_table: bool
    style_name: str
    numbered: bool
    numbering_tokens: tuple[int, ...]
    has_tab: bool
    # Вспомогательные сигналы из Word-XML, используемые детектором заголовков
    numpr_ilvl: int | None = None
    numpr_num_id: int | None = None
    outline_level: int | None = None
    has_page_break_before: bool = False
    derived_number: str | None = None


@dataclass
class DocStats:
    paragraphs: list[ParagraphStats] = field(default_factory=list)
    modal_font_size_pt: float = 14.0
    modal_font_family: str = "Times New Roman"
    font_size_distribution: Counter = field(default_factory=Counter)
    total_paragraphs_non_empty: int = 0
    modal_line_spacing: float = 1.5
    # Нормализованные тексты абзацев, лежащих *внутри* блоков w:sdt и
    # поэтому невидимых для doc.paragraphs. Использует validate_structure,
    # чтобы автогенерированное Word'ом оглавление со словом "СОДЕРЖАНИЕ"
    # внутри SDT не считалось отсутствующим.
    sdt_paragraph_texts: frozenset[str] = field(default_factory=frozenset)
    # Размеры групп numId: {num_id: количество абзацев с этим num_id}.
    # Большая группа (≥4 элементов) — почти всегда обычный список, а не
    # последовательность заголовков. heading_scoring использует это, чтобы
    # отбрасывать пункты списка с outline_level от шаблона Word.
    numpr_group_sizes: dict[int, int] = field(default_factory=dict)


def _run_font_size_pt(run) -> float | None:
    try:
        if run.font.size is not None:
            return float(run.font.size.pt)
    except Exception:
        return None
    return None


def _style_font_size(paragraph: Paragraph) -> float | None:
    try:
        if paragraph.style and paragraph.style.font and paragraph.style.font.size:
            return float(paragraph.style.font.size.pt)
    except Exception:
        return None
    return None


def _style_font_name(paragraph: Paragraph) -> str | None:
    try:
        if paragraph.style and paragraph.style.font and paragraph.style.font.name:
            return str(paragraph.style.font.name)
    except Exception:
        return None
    return None


def _collect_font_sizes(paragraph: Paragraph) -> list[float]:
    sizes: list[float] = []
    for run in paragraph.runs:
        size = _run_font_size_pt(run)
        if size is not None:
            sizes.append(size)
    style_size = _style_font_size(paragraph)
    if style_size is not None and not sizes:
        sizes.append(style_size)
    return sizes


def _font_family(paragraph: Paragraph) -> str | None:
    for run in paragraph.runs:
        try:
            if run.font.name:
                return str(run.font.name)
        except Exception:
            continue
    return _style_font_name(paragraph)


def _bold_ratio(paragraph: Paragraph) -> float:
    printable = [run for run in paragraph.runs if run.text.strip()]
    if not printable:
        return 0.0
    bold_chars = sum(len(run.text) for run in printable if run.bold)
    total_chars = sum(len(run.text) for run in printable)
    return bold_chars / total_chars if total_chars else 0.0


def _italic_ratio(paragraph: Paragraph) -> float:
    printable = [run for run in paragraph.runs if run.text.strip()]
    if not printable:
        return 0.0
    italic_chars = sum(len(run.text) for run in printable if run.italic)
    total_chars = sum(len(run.text) for run in printable)
    return italic_chars / total_chars if total_chars else 0.0


def _first_line_indent_cm(paragraph: Paragraph) -> float | None:
    try:
        fli = paragraph.paragraph_format.first_line_indent
        if fli is None:
            return None
        return float(fli.cm)
    except Exception:
        return None


def _left_indent_cm(paragraph: Paragraph) -> float | None:
    try:
        li = paragraph.paragraph_format.left_indent
        if li is None:
            return None
        return float(li.cm)
    except Exception:
        return None


def _in_table(paragraph: Paragraph) -> bool:
    try:
        parent = paragraph._element.getparent()
        while parent is not None:
            if parent.tag.endswith("}tc"):
                return True
            parent = parent.getparent()
    except Exception:
        pass
    return False


# Агрессивная Unicode-зачистка пробелов: Python str.strip() оставляет
# NBSP, en/em/thin-пробелы и нулевой ширины. Заголовки часто подцепляют
# эти символы из "Вставка символа" Word или вставленного текста, и без
# их удаления regex ниже не сматчит числовой префикс.
_WS_STRIP_RE = __import__("re").compile(
    r"^[\s\xa0 -​  　]+|[\s\xa0 -​  　]+$"
)


def _unicode_strip(text: str) -> str:
    return _WS_STRIP_RE.sub("", text)


_NUMBERING_RE = __import__("re").compile(r"^(\d+(?:\.\d+){0,3})[.)]?\s+")


def _numbering_tokens(text: str) -> tuple[int, ...]:
    match = _NUMBERING_RE.match(text)
    if not match:
        return tuple()
    try:
        return tuple(int(p) for p in match.group(1).split("."))
    except ValueError:
        return tuple()


def _extract_numpr(paragraph: Paragraph) -> tuple[int | None, int | None]:
    """Читает (ilvl, num_id) из w:numPr. Возвращает (None, None), если нет."""
    try:
        pPr = paragraph._p.find(qn("w:pPr"))
        if pPr is None:
            return None, None
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            return None, None
        ilvl_el = numPr.find(qn("w:ilvl"))
        num_id_el = numPr.find(qn("w:numId"))
        ilvl = int(ilvl_el.get(qn("w:val"))) if ilvl_el is not None else None
        num_id = int(num_id_el.get(qn("w:val"))) if num_id_el is not None else None
        return ilvl, num_id
    except Exception:
        return None, None


def _extract_outline_level(paragraph: Paragraph) -> int | None:
    """Читает w:outlineLvl из pPr абзаца или наследует из стиля.

    Word отражает иерархию заголовков через outlineLvl: 0 → H1, 1 → H2,
    2 → H3. Значения больше 2 — это тело текста, считаем "не заголовок".
    """
    try:
        pPr = paragraph._p.find(qn("w:pPr"))
        if pPr is not None:
            olvl = pPr.find(qn("w:outlineLvl"))
            if olvl is not None:
                val = olvl.get(qn("w:val"))
                if val is not None:
                    return int(val)
        # Наследуется из стиля
        style = paragraph.style
        if style is not None:
            style_el = getattr(style, "element", None)
            if style_el is not None:
                style_pPr = style_el.find(qn("w:pPr"))
                if style_pPr is not None:
                    olvl = style_pPr.find(qn("w:outlineLvl"))
                    if olvl is not None:
                        val = olvl.get(qn("w:val"))
                        if val is not None:
                            return int(val)
    except Exception:
        pass
    return None


def _has_page_break_before(paragraph: Paragraph) -> bool:
    """Проверяет наличие w:pageBreakBefore в pPr (присутствует и непустое значение)."""
    try:
        pPr = paragraph._p.find(qn("w:pPr"))
        if pPr is None:
            return False
        pb = pPr.find(qn("w:pageBreakBefore"))
        if pb is None:
            return False
        val = pb.get(qn("w:val"))
        # Отсутствующий val = ON. Явные "0"/"false"/"off" = OFF.
        if val is None or val.lower() not in ("0", "false", "off"):
            return True
    except Exception:
        pass
    return False


def _resolve_derived_numbers(paragraphs: list[ParagraphStats]) -> None:
    """Заполняет `derived_number` для каждого абзаца.

    Приоритет — текстовый числовой префикс. Иначе, если есть numPr с
    ilvl ≤ 2, имитируем счётчики по numId и выдаём номер вроде "1.1.2".
    Получаемая строка — только подсказка для последующего скоринга; точный
    разбор формата numbering.xml НЕ выполняется.

    Если между двумя абзацами с одинаковым num_id есть содержательный
    body-абзац (не numPr и не пустой), счётчик сбрасывается. Без этого
    разные списки в исходнике, использующие один num_id, получают сквозную
    нумерацию через весь документ — и в обработанном файле "5", "6", "7",
    ..., "29" вместо "1", "2", "3" в каждом разделе.
    """
    numpr_counters: dict[int, list[int]] = {}
    last_seen_index: dict[int, int] = {}
    for ps in paragraphs:
        if ps.numbering_tokens:
            ps.derived_number = ".".join(str(t) for t in ps.numbering_tokens)
            continue
        if ps.is_empty:
            continue
        if ps.numpr_ilvl is None or ps.numpr_num_id is None:
            continue
        ilvl = ps.numpr_ilvl
        if ilvl < 0 or ilvl > 8:
            continue

        num_id = ps.numpr_num_id
        last_idx = last_seen_index.get(num_id)
        if last_idx is not None:
            broken = any(
                (not paragraphs[j].is_empty)
                and (paragraphs[j].numpr_num_id != num_id)
                for j in range(last_idx + 1, ps.index)
            )
            if broken:
                # Между двумя пунктами одного num_id попался обычный абзац —
                # это разные списки, склеенные одним num_id. Сбрасываем счётчик.
                numpr_counters[num_id] = [0] * 9

        counters = numpr_counters.setdefault(num_id, [0] * 9)
        counters[ilvl] += 1
        # Если родительский счётчик ещё 0 — поднимаем его до 1: когда
        # документ перепрыгивает сразу на подраздел "1.1" без явной
        # родительской "1 Глава" с тем же num_id, мы хотим, чтобы
        # derived_number было "1.1", а не "0.1".
        for i in range(ilvl):
            if counters[i] == 0:
                counters[i] = 1
        for i in range(ilvl + 1, 9):
            counters[i] = 0
        ps.derived_number = ".".join(str(counters[i]) for i in range(ilvl + 1))
        last_seen_index[num_id] = ps.index





def _paragraph_stats(paragraph: Paragraph, index: int) -> ParagraphStats:
    text = paragraph.text
    stripped = _unicode_strip(text)
    sizes = _collect_font_sizes(paragraph)
    max_size = max(sizes) if sizes else None
    modal_size = None
    if sizes:
        counter = Counter(round(s, 1) for s in sizes)
        modal_size = counter.most_common(1)[0][0]

    style_name = paragraph.style.name if paragraph.style else ""
    tokens = _numbering_tokens(stripped)
    ilvl, num_id = _extract_numpr(paragraph)

    return ParagraphStats(
        index=index,
        text=text,
        stripped=stripped,
        length=len(stripped),
        word_count=len(stripped.split()),
        alignment=paragraph.alignment,
        first_line_indent_cm=_first_line_indent_cm(paragraph),
        left_indent_cm=_left_indent_cm(paragraph),
        max_font_size_pt=max_size,
        modal_font_size_pt=modal_size,
        font_family=_font_family(paragraph),
        bold_ratio=_bold_ratio(paragraph),
        italic_ratio=_italic_ratio(paragraph),
        ends_with_period=stripped.endswith((".", "!", "?")),
        is_upper=stripped.isupper() and bool(stripped),
        is_empty=not stripped,
        in_table=_in_table(paragraph),
        style_name=style_name or "",
        numbered=bool(tokens),
        numbering_tokens=tokens,
        has_tab="\t" in text,
        numpr_ilvl=ilvl,
        numpr_num_id=num_id,
        outline_level=_extract_outline_level(paragraph),
        has_page_break_before=_has_page_break_before(paragraph),
        derived_number=None,  # заполняется в _resolve_derived_numbers
    )


def _collect_sdt_texts(doc: DocxDocument) -> frozenset[str]:
    """Возвращает нормализованные тексты абзацев внутри контейнеров w:sdt.

    Word оборачивает автогенерируемые оглавления в SDT-блоки, поэтому такие
    абзацы (включая заголовок "СОДЕРЖАНИЕ") невидимы в doc.paragraphs.
    Этот хелпер обходит XML тела документа напрямую.
    """
    texts: set[str] = set()
    try:
        body = doc.element.body
        for sdt in body.iter(qn("w:sdt")):
            sdt_content = sdt.find(qn("w:sdtContent"))
            if sdt_content is None:
                continue
            for p in sdt_content.iter(qn("w:p")):
                # Собираем текст из всех w:t-элементов внутри абзаца.
                raw = "".join(
                    el.text or ""
                    for el in p.iter(qn("w:t"))
                )
                stripped = raw.strip()
                if stripped:
                    texts.add(" ".join(stripped.lower().split()))
    except Exception:
        pass
    return frozenset(texts)


def collect_stats(doc: DocxDocument) -> DocStats:
    stats = DocStats()

    body_sizes: list[float] = []
    body_families: list[str] = []
    body_spacings: list[float] = []

    paragraphs = list(doc.paragraphs)
    stats.paragraphs = [_paragraph_stats(p, i) for i, p in enumerate(paragraphs)]
    _resolve_derived_numbers(stats.paragraphs)
    stats.sdt_paragraph_texts = _collect_sdt_texts(doc)

    # Размеры групп по num_id — нужны heading_scoring'у, чтобы отличать
    # пункты длинного списка от настоящих заголовков, унаследовавших
    # outline_level от стиля.
    group_counter: Counter = Counter()
    for ps in stats.paragraphs:
        if ps.numpr_num_id is not None and not ps.is_empty:
            group_counter[ps.numpr_num_id] += 1
    stats.numpr_group_sizes = dict(group_counter)

    # stats.paragraphs и doc.paragraphs формируются в одном проходе ниже —
    # длины равны по построению. strict=False оставляем, чтобы lint не
    # ругался и тест не падал, если в будущем добавится side-loop.
    for ps, paragraph in zip(stats.paragraphs, paragraphs, strict=False):
        if ps.is_empty:
            continue
        stats.total_paragraphs_non_empty += 1

        if ps.modal_font_size_pt is not None:
            stats.font_size_distribution[round(ps.modal_font_size_pt, 1)] += 1

        # Эвристика для абзаца тела: длинный, без нумерации, не жирный
        # (первая буква может быть и строчной, и заглавной)
        if ps.length >= 120 and ps.bold_ratio < 0.2 and not ps.numbered:
            if ps.modal_font_size_pt:
                body_sizes.append(ps.modal_font_size_pt)
            if ps.font_family:
                body_families.append(ps.font_family)
            try:
                ls = paragraph.paragraph_format.line_spacing
                if isinstance(ls, (int, float)):
                    body_spacings.append(float(ls))
            except Exception:
                pass

    if body_sizes:
        stats.modal_font_size_pt = statistics.mode(body_sizes) if len(body_sizes) > 1 else body_sizes[0]
    elif stats.font_size_distribution:
        stats.modal_font_size_pt = stats.font_size_distribution.most_common(1)[0][0]

    if body_families:
        stats.modal_font_family = Counter(body_families).most_common(1)[0][0]

    if body_spacings:
        try:
            stats.modal_line_spacing = statistics.mode(body_spacings)
        except statistics.StatisticsError:
            stats.modal_line_spacing = body_spacings[0]

    return stats
