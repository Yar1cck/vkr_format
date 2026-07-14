"""Детекция структурного заголовка, оформленного как строка оглавления.

Типовой дефект: заголовок обязательного раздела (например, «СПИСОК
ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ») в теле документа набран с отточием и номером
страницы — `СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ.................43` — то есть как
строка оглавления, скопированная в тело. Детектор заголовков
(`heading_scoring.py`) такой абзац hard-reject'ит как `toc_entry` (таб/отточие +
число в конце), поэтому раздел молча «исчезает», а пользователь видит лишь
«раздел не найден».

Этот модуль ловит такие абзацы (вне титула и вне ручного оглавления), выдаёт
понятное нарушение и в режиме полной обработки срезает хвост-отточие, оставляя
чистый текст заголовка — дальше его штатно подхватывает детектор и оформляет
`apply_headings`.
"""

from __future__ import annotations

import re
from dataclasses import dataclass

from docx.document import Document as DocxDocument

from services.core.vkr_core.engine.stats import DocStats
from services.core.vkr_core.engine.violations import (
    SEVERITY_WARNING,
    PipelineViolation,
)
from services.core.vkr_core.models.enums import ViolationStatus

# Хвост «строки оглавления»: ведущий таб / 2+ пробелов / отточие, затем номер
# страницы (1–4 цифры) в самом конце строки. Достаточно отделяющего разделителя
# (таб, длинный пробел или серия точек), чтобы не срезать осмысленное число у
# обычного заголовка без отточия.
_TOC_TAIL_RE = re.compile(r"(?:\t+|\s{2,}|\.{2,}[.\s]*)\s*\d{1,4}\s*$")


def _norm(text: str) -> str:
    return " ".join(text.strip().lower().split())


@dataclass
class MalformedHeading:
    paragraph_index: int
    key: str
    cleaned_text: str
    raw_text: str


def find_malformed_structural_headings(
    stats: DocStats,
    rules: dict,
    skip_indexes: set[int],
) -> list[MalformedHeading]:
    """Находит структурные заголовки, оформленные как строка оглавления."""
    structural = rules.get("structural_elements", {})
    alias_to_key: dict[str, str] = {}
    for key, aliases in structural.items():
        for alias in aliases:
            alias_to_key[_norm(alias)] = key
    if not alias_to_key:
        return []

    found: list[MalformedHeading] = []
    for ps in stats.paragraphs:
        if ps.index in skip_indexes or ps.is_empty:
            continue
        tail = _TOC_TAIL_RE.search(ps.stripped)
        if not tail:
            continue
        cleaned = ps.stripped[: tail.start()].strip()
        key = alias_to_key.get(_norm(cleaned))
        if key is None:
            continue
        found.append(
            MalformedHeading(
                paragraph_index=ps.index,
                key=key,
                cleaned_text=cleaned,
                raw_text=ps.stripped,
            )
        )
    return found


def build_violations(
    matches: list[MalformedHeading],
    fixed: bool,
) -> list[PipelineViolation]:
    status = ViolationStatus.auto_fixed if fixed else ViolationStatus.manual_required
    tail = (
        "оформлен как обычный заголовок."
        if fixed
        else "оформите как обычный заголовок (по центру, без отточия и номера страницы)."
    )
    violations: list[PipelineViolation] = []
    for m in matches:
        violations.append(
            PipelineViolation(
                type="structural_heading_malformed",
                rule_reference="п. 6.11",
                description=(
                    f"Структурный заголовок «{m.cleaned_text}» оформлен как строка "
                    f"оглавления (отточие и номер страницы) — {tail}"
                ),
                status=status,
                severity=SEVERITY_WARNING,
                paragraph_index=m.paragraph_index,
                original_text=m.raw_text,
                fixed_text=m.cleaned_text if fixed else None,
            )
        )
    return violations


def apply_structural_heading_cleanup(
    doc: DocxDocument,
    matches: list[MalformedHeading],
) -> None:
    """Срезает хвост-отточие у найденных абзацев, оставляя чистый заголовок.

    Меняет только содержимое runs (число абзацев не трогает), поэтому индексы
    абзацев остаются валидными. Кастомные таб-стопы удаляются, чтобы при
    последующем оформлении не осталось точечного лидера.
    """
    paragraphs = doc.paragraphs
    for m in matches:
        if m.paragraph_index >= len(paragraphs):
            continue
        paragraph = paragraphs[m.paragraph_index]
        # Очищаем все runs, кладём один с чистым текстом.
        for run in list(paragraph.runs):
            run._element.getparent().remove(run._element)
        paragraph.add_run(m.cleaned_text)
        try:
            paragraph.paragraph_format.tab_stops.clear_all()
        except Exception:
            pass
