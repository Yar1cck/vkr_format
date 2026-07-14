"""Детерминированный оркестратор обработки (ТЗ §6 + §7).

Вся тяжёлая работа выполняется в чистых python-модулях; этот файл их
связывает. Конвейер строго последовательный:

    1.  Сбор DocStats.
    2.  Определение границы титульного листа — всё выше неприкосновенно.
    3.  Запуск детектора заголовков (5 проходов).
    4.  Запуск алгоритмов §7 на чётко определённом наборе абзацев:
         числительные, формулы, подписи, запрещённые заголовки, цитаты,
         библиография, структура, объём.
    5.  В режиме "full" — применение детерминированного форматирования
         (страница, тело, заголовки, оглавление). В check-only документ
         остаётся нетронутым.
    6.  Возврат обработанного файла + агрегированный список нарушений.
"""

from __future__ import annotations

import re as _re
import shutil
import tempfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document as WordDocument
from docx.oxml.ns import qn

from services.core.vkr_core.engine.appendix import validate_appendices
from services.core.vkr_core.engine.autofix_guard import may_apply_heading
from services.core.vkr_core.engine.bibliography import (
    apply_bibliography_numbering,
    collect_bibliography,
    find_bibliography_start,
    validate_bibliography,
)
from services.core.vkr_core.engine.captions import (
    keep_figures_with_captions,
    move_table_captions_before_tables,
    renumber_figures,
    renumber_tables,
    validate_references,
)
from services.core.vkr_core.engine.citations import (
    check_citation_separators,
    check_combined_citations,
    collect_citations,
    detect_suspicious_citation_ranges,
    fix_citation_separators,
    fix_combined_citations,
    validate_citation_order,
)
from services.core.vkr_core.engine.crossref import apply_crossrefs
from services.core.vkr_core.engine.detection import detect_headings, detect_headings_full
from services.core.vkr_core.engine.forbidden import (
    apply_forbidden_replacements,
    apply_section_title_canonicalization,
    check_forbidden,
    check_section_titles,
)
from services.core.vkr_core.engine.formatter import (
    _strip_paragraph_outline_level,
    apply_body_style,
    apply_footer_page_number,
    apply_headings,
    apply_page_settings,
    check_paragraph_spacing,
    disable_document_hyphenation,
    ensure_separator_after_tables_and_figures,
    fix_formula_number_alignment,
    insert_l1_page_breaks,
    insert_pretoc_section_break,
    normalize_spaces,
    remove_empty_paragraphs,
    strip_chapter_word_everywhere,
    style_appendix_titles,
)
from services.core.vkr_core.engine.formulas import detect_formulas, validate_formulas
from services.core.vkr_core.engine.heading_numbers import validate_heading_numbers
from services.core.vkr_core.engine.instrumentation import timed
from services.core.vkr_core.engine.listings import (
    detect_listing_code_indexes,
    renumber_listings,
)
from services.core.vkr_core.engine.stats import collect_stats
from services.core.vkr_core.engine.structural_format import (
    apply_structural_heading_cleanup,
    find_malformed_structural_headings,
)
from services.core.vkr_core.engine.structural_format import (
    build_violations as build_structural_format_violations,
)
from services.core.vkr_core.engine.structure import validate_structure
from services.core.vkr_core.engine.table_continuation import _normalize_continuation_labels
from services.core.vkr_core.engine.tables import (
    apply_continuation_table_styles,
    apply_table_styles,
    ensure_table_pagebreak_safety,
)
from services.core.vkr_core.engine.title_page import detect_title_page_end
from services.core.vkr_core.engine.toc import (
    _unwrap_sdt_toc,
    apply_table_continuation_only,
    detect_toc_section_indexes,
    finalize_toc,
    insert_toc,
)
from services.core.vkr_core.engine.violations import (
    SEVERITY_INFO,
    PipelineViolation,
)
from services.core.vkr_core.engine.volume import estimate_pages, validate_volume
from services.core.vkr_core.models.enums import ViolationStatus

_H1_NUM_PREFIX = _re.compile(r"^(\d+)\b")


def _chapter_breaks_from_headings(headings, paragraphs) -> list[tuple[int, int]]:
    """Нумерованные H1-заголовки → [(paragraph_index, chapter_num), ...]. Для check_only."""
    result = []
    for h in sorted(headings, key=lambda x: x.paragraph_index):
        if h.level != 1:
            continue
        idx = h.paragraph_index
        if idx >= len(paragraphs):
            continue
        # Сначала пробуем derived_number (уже разобранный номер — работает для «ГЛАВА N»,
        # у которых сырой текст начинается не с цифры).
        if h.derived_number:
            m = _H1_NUM_PREFIX.match(h.derived_number)
            if m:
                result.append((idx, int(m.group(1))))
                continue
        m = _H1_NUM_PREFIX.match(paragraphs[idx].text.strip())
        if m:
            result.append((idx, int(m.group(1))))
    return result


def _chapter_breaks_from_xml(paragraphs) -> list[tuple[int, int]]:
    """Нумерованные H1-заголовки → [(paragraph_index, chapter_num), ...]. Читает XML после apply_headings."""
    result = []
    for idx, p in enumerate(paragraphs):
        pPr = p._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        lvl = pPr.find(qn("w:outlineLvl"))
        if lvl is None or lvl.get(qn("w:val")) != "0":
            continue
        m = _H1_NUM_PREFIX.match(p.text.strip())
        if m:
            result.append((idx, int(m.group(1))))
    return result


def _caption_chapter_breaks_for_mode(
    rules: dict,
    kind: str,
    chapter_breaks: list[tuple[int, int]],
) -> list[tuple[int, int]] | None:
    mode = (rules.get("numbering_rules") or {}).get(f"{kind}_mode", "by_chapter")
    return chapter_breaks if mode == "by_chapter" else None


@dataclass
class PipelineResult:
    processed_docx_path: Path
    violations: list[PipelineViolation]
    volume_pages: int
    toc_warning: str | None = None


def _append_caption_violations(
    violations: list[PipelineViolation],
    changes: list[tuple[int, str, str]],
) -> None:
    """Список (paragraph_index, было, стало) → PipelineViolation. Рисунки, таблицы, листинги."""
    for para_idx, orig_text, new_text in changes:
        violations.append(
            PipelineViolation(
                type="caption_renumbered",
                rule_reference="п. 6.12",
                description="Сквозная нумерация по порядку появления в тексте (см. Было/Стало).",
                status=ViolationStatus.manual_required,
                severity=SEVERITY_INFO,
                paragraph_index=para_idx,
                original_text=orig_text,
                fixed_text=new_text,
            )
        )


def process_document(
    source_path: Path,
    rules: dict,
    check_only: bool,
) -> PipelineResult:
    working_dir = Path(tempfile.mkdtemp(prefix="vkr-pipeline-"))
    src_copy = working_dir / "input.docx"
    shutil.copy2(source_path, src_copy)

    doc = WordDocument(str(src_copy))

    # SDT до статистики: detect_title_page_end иначе принимает эти абзацы за титул и сдвигает индексы.
    _unwrap_sdt_toc(doc, rules)

    violations: list[PipelineViolation] = []

    with timed("collect_stats"):
        stats = collect_stats(doc)

    with timed("detect_title_page"):
        title_end_index = detect_title_page_end(stats, rules)
        title_page_indexes = set(range(title_end_index))

    with timed("detect_toc_section"):
        toc_section_indexes = detect_toc_section_indexes(doc, rules)
        skip_indexes = title_page_indexes | toc_section_indexes

    with timed("structural_format_cleanup"):
        struct_malformed = find_malformed_structural_headings(stats, rules, skip_indexes)
        if struct_malformed and not check_only:
            apply_structural_heading_cleanup(doc, struct_malformed)
            stats = collect_stats(doc)
        violations.extend(build_structural_format_violations(struct_malformed, fixed=not check_only))

    with timed("detect_headings"):
        headings, soft_violations = detect_headings_full(stats, rules, skip_indexes)
        violations.extend(soft_violations)
        heading_indexes = {h.paragraph_index for h in headings}

    with timed("forbidden_and_section_titles"):
        violations.extend(check_forbidden(headings, rules))
        violations.extend(check_section_titles(headings, rules))

    # §6.5 нумерация заголовков — в конце: insert_toc и remove_empty сдвигают индексы.

    with timed("citations"):
        # Останавливаем сбор цитирований на начале раздела источников:
        # перекрёстные ссылки внутри библиографических записей не учитываются
        # при проверке порядка [N] в тексте.
        _bib_start = find_bibliography_start(stats, rules, headings)
        citations = collect_citations(stats, title_page_indexes, max_index=_bib_start)
        numbering_rules = rules.get("numbering_rules", {})
        if numbering_rules.get("validate_citation_sequence", True):
            violations.extend(validate_citation_order(citations))
        violations.extend(check_citation_separators(stats, skip_indexes))
        violations.extend(check_combined_citations(stats, skip_indexes))
        violations.extend(detect_suspicious_citation_ranges(stats, skip_indexes))

    with timed("bibliography"):
        bib_entries, bib_bounds, bib_malformed = collect_bibliography(stats, rules, headings)
        if not check_only and bib_bounds is not None:
            fixed_idx = apply_bibliography_numbering(doc, stats, rules, headings)
            if fixed_idx:
                stats = collect_stats(doc)
                bib_entries, bib_bounds, bib_malformed = collect_bibliography(stats, rules, headings)
                violations.append(
                    PipelineViolation(
                        type="bibliography_entries_autonumbered",
                        rule_reference="п. 6.16",
                        description=(
                            f"Автоматически пронумерованы источники без номера: {len(fixed_idx)}."
                        ),
                        status=ViolationStatus.auto_fixed,
                        severity=SEVERITY_INFO,
                    )
                )
        violations.extend(
            validate_bibliography(bib_entries, citations, bib_bounds, bib_malformed)
        )

    with timed("structure_and_appendices"):
        violations.extend(validate_structure(stats, rules, headings, title_page_indexes))
        violations.extend(validate_appendices(headings, stats, title_page_indexes))

    with timed("formulas"):
        formulas = detect_formulas(stats, title_page_indexes | heading_indexes)
        violations.extend(validate_formulas(stats, formulas))

    with timed("volume"):
        # Ячейки таблиц не входят в doc.paragraphs — добавляем их отдельно,
        # иначе документы с большими таблицами занижают оценку объёма.
        table_chars = sum(
            len(cell.text)
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        )
        pages = estimate_pages(stats, title_page_indexes, extra_chars=table_chars)
        violations.extend(validate_volume(pages, rules))

    # Смешанные размеры шрифта: титул исключаем (16pt/11pt там допустимы).
    from collections import Counter as _Counter
    body_size = rules.get("body_text_style", {}).get("font_size_pt", 14)
    size_dist: _Counter[float] = _Counter()
    for ps in stats.paragraphs:
        if ps.index not in title_page_indexes and not ps.is_empty and ps.modal_font_size_pt is not None:
            size_dist[round(ps.modal_font_size_pt, 1)] += 1
    if size_dist:
        distinct = sorted(size_dist.items(), key=lambda kv: -kv[1])
        non_standard = [(sz, cnt) for sz, cnt in distinct if abs(sz - body_size) > 0.1]
        if non_standard:
            total = sum(size_dist.values()) or 1
            parts = [f"{sz}pt — {cnt} абзац. ({cnt * 100 // total}%)" for sz, cnt in distinct]
            violations.append(
                PipelineViolation(
                    type="font_size_inconsistency",
                    rule_reference="Общие требования",
                    description=(
                        "В документе встречаются разные размеры шрифта. "
                        "Распределение по абзацам: " + "; ".join(parts) + "."
                    ),
                    status=ViolationStatus.auto_fixed,
                    severity=SEVERITY_INFO,
                )
            )

    with timed("paragraph_spacing"):
        violations.extend(
            check_paragraph_spacing(doc, stats, skip_indexes, heading_indexes, rules)
        )

    toc_warning: str | None = None
    processed_path = working_dir / "processed.docx"

    if check_only:
        with timed("check_only_finalize"):
            _co_breaks = _chapter_breaks_from_headings(headings, doc.paragraphs)
            figure_breaks = _caption_chapter_breaks_for_mode(rules, "figure", _co_breaks)
            table_breaks = _caption_chapter_breaks_for_mode(rules, "table", _co_breaks)
            figures, _fig_changes = renumber_figures(
                doc.paragraphs, title_page_indexes, apply=False, chapter_breaks=figure_breaks
            )
            tables, _tbl_changes = renumber_tables(
                doc.paragraphs, title_page_indexes, apply=False, chapter_breaks=table_breaks
            )
            violations.extend(validate_references(stats, figures, tables, title_page_indexes))
            violations.extend(validate_heading_numbers(headings, stats))
            doc.save(str(processed_path))

        return PipelineResult(
            processed_docx_path=processed_path,
            violations=violations,
            volume_pages=pages,
            toc_warning=None,
        )

    # --- Дальше — путь полной обработки ---

    with timed("apply_headings"):
        safe_headings = [h for h in headings if may_apply_heading(stats.paragraphs[h.paragraph_index])]

        apply_forbidden_replacements(doc, headings, rules)
        apply_section_title_canonicalization(doc, headings, rules)

        applied_headings = apply_headings(doc, safe_headings, stats)
        heading_indexes = applied_headings | heading_indexes

    # Body-абзацы на стиле "Заголовок N" (hard-reject): переключаем на Normal до сбора heading_indexes —
    # иначе попадут в TOC через outline_level и пропустят apply_body_style.
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx in title_page_indexes or idx in heading_indexes:
            continue
        style_name = (paragraph.style.name or "").strip().lower() if paragraph.style else ""
        if style_name.startswith(("heading", "заголовок", "подзаголовок")):
            try:
                paragraph.style = "Normal"
            except KeyError:
                pass

    for idx, paragraph in enumerate(doc.paragraphs):
        if idx in title_page_indexes:
            continue
        if paragraph.style and paragraph.style.name.lower().startswith(("heading", "заголовок")):
            heading_indexes.add(idx)

    with timed("apply_page_and_body"):
        apply_page_settings(doc, rules, skip_first_section=bool(title_page_indexes))
        disable_document_hyphenation(doc)

        # Строки кода исключаем: justify растянет их, Times New Roman перебьёт моноширину.
        listing_code_indexes = detect_listing_code_indexes(doc.paragraphs, title_page_indexes)
        apply_body_style(
            doc, rules, stats,
            title_page_indexes | listing_code_indexes,
            heading_indexes,
        )

        # Разрывы страниц перед L1 — после apply_body_style (вставка абзацев сдвигает индексы stats).
        insert_l1_page_breaks(doc, safe_headings, heading_indexes)

    # Снимаем outlineLvl у body-абзацев — иначе шаблонный outlineLvl=0 попадает в Word TOC.
    # heading_indexes устарел после insert_l1_page_breaks → проверяем по XML (≤2 → заголовок).
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx in title_page_indexes:
            continue
        pPr = paragraph._p.find(qn("w:pPr"))
        if pPr is not None:
            olvl = pPr.find(qn("w:outlineLvl"))
            if olvl is not None:
                try:
                    if int(olvl.get(qn("w:val"), "8")) <= 2:
                        continue  # настоящий заголовок — outlineLvl оставляем
                except (ValueError, TypeError):
                    pass
        _strip_paragraph_outline_level(paragraph)

    # Нужно до _chapter_breaks_from_xml: иначе «ГЛАВА 3 ...» ещё не начинается
    # с числа, chapter_breaks не строятся, а подписи/листинги съезжают к 1.
    strip_chapter_word_everywhere(doc, title_page_indexes)

    with timed("captions_and_crossrefs"):
        move_table_captions_before_tables(doc)
        chapter_breaks = _chapter_breaks_from_xml(doc.paragraphs)
        figure_breaks = _caption_chapter_breaks_for_mode(rules, "figure", chapter_breaks)
        table_breaks = _caption_chapter_breaks_for_mode(rules, "table", chapter_breaks)
        listing_breaks = _caption_chapter_breaks_for_mode(rules, "listing", chapter_breaks)
        figures, figure_changes = renumber_figures(
            doc.paragraphs, title_page_indexes, apply=True, chapter_breaks=figure_breaks
        )
        tables, table_changes = renumber_tables(
            doc.paragraphs,
            title_page_indexes,
            apply=True,
            indent_cm=None,
            chapter_breaks=table_breaks,
        )
        keep_figures_with_captions(doc, figures, title_page_indexes)
        _append_caption_violations(violations, figure_changes)
        _append_caption_violations(violations, table_changes)
        _listings, listing_changes = renumber_listings(
            doc.paragraphs,
            title_page_indexes,
            apply=True,
            chapter_breaks=listing_breaks,
        )
        _append_caption_violations(violations, listing_changes)

        style_appendix_titles(doc)
        apply_crossrefs(doc, figures, tables, title_page_indexes)

        stats_after = collect_stats(doc)
        violations.extend(validate_references(stats_after, figures, tables, title_page_indexes))

    with timed("table_styles"):
        violations.extend(apply_table_styles(doc, tables))
        apply_continuation_table_styles(doc, title_end_index)
        ensure_table_pagebreak_safety(doc, title_end_index)

    # §7.5: нарушения зафиксированы выше, здесь только мутация.
    fix_combined_citations(doc, skip_indexes)
    fix_citation_separators(doc, skip_indexes)

    # §6.14: номера формул — правый край страницы.
    fix_formula_number_alignment(doc, formulas, rules)

    # Индексы кода пересчитываем заново: insert_l1_page_breaks сдвинул нумерацию.
    listing_code_indexes_late = detect_listing_code_indexes(
        doc.paragraphs, title_page_indexes
    )
    normalize_spaces(
        doc, title_page_indexes | listing_code_indexes_late, heading_indexes
    )
    _normalize_continuation_labels(doc)
    apply_continuation_table_styles(doc, title_end_index)
    remove_empty_paragraphs(doc, title_page_indexes)
    ensure_separator_after_tables_and_figures(doc, title_page_indexes)

    with timed("footer_and_toc_insert"):
        _has_pretoc_section = insert_pretoc_section_break(doc, rules, title_end_index)
        apply_footer_page_number(doc, with_blank_page_section=_has_pretoc_section)

        inserted = insert_toc(doc, rules, title_end_index)
        if inserted:
            toc_warning = rules.get("report_templates", {}).get(
                "warning_toc",
                "Содержание обновлено автоматически. Проверьте номера страниц в Word/LibreOffice (F9).",
            )

        doc.save(str(processed_path))

    with timed("finalize_toc"):
        # PAGEREF через LO-макрос + нормализация шрифта TOC. Нет LO → заглушки и предупреждение F9.
        if inserted:
            if not finalize_toc(processed_path, tables):
                toc_warning = rules.get("report_templates", {}).get(
                    "warning_toc_lo_failed",
                    "Не удалось обновить поля оглавления через LibreOffice "
                    "(возможно, soffice недоступен или превышен таймаут). "
                    "Откройте файл в Word/LibreOffice и нажмите F9 для "
                    "пересчёта номеров страниц.",
                )
        else:
            # Документ без TOC: finalize_toc не вызывается, но резку таблиц
            # всё равно делаем (она требует LO-рендер, поэтому живёт здесь).
            apply_table_continuation_only(processed_path, tables)

    with timed("final_validate_heading_numbers"):
        # §6.5 нумерация заголовков: перечитываем после finalize_toc — LO вставил строки TOC и сдвинул индексы.
        final_doc = WordDocument(str(processed_path))

        final_stats = collect_stats(final_doc)
        final_title_end = detect_title_page_end(final_stats, rules)
        final_title_page_indexes = set(range(final_title_end))
        final_toc_indexes = detect_toc_section_indexes(final_doc, rules)
        final_skip_indexes = final_title_page_indexes | final_toc_indexes
        final_headings = detect_headings(final_stats, rules, final_skip_indexes)
        violations.extend(validate_heading_numbers(final_headings, final_stats))

    return PipelineResult(
        processed_docx_path=processed_path,
        violations=violations,
        volume_pages=pages,
        toc_warning=toc_warning,
    )
