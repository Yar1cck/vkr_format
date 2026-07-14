"""Скоринговое обнаружение заголовков (ТЗ §6.5).

Каждый абзац получает heading-скор, вычисленный по трём классам сигналов:
  - сильные: родной стиль, numPr, outline_level, числовой префикс,
    структурные ключевые слова. Один сильный сигнал уводит абзац в HIGH.
  - слабые: шрифт, жирный, выравнивание, изолированность. Складываются
    в MEDIUM/LOW.
  - отрицательные: оформление списка, подписи, конечная пунктуация.
    Либо штрафуют скор, либо hard-reject'ят абзац.

Уровни:
  HIGH   (≥ 10) — уверенный заголовок; подтверждения пользователя не нужно.
  MEDIUM (6-9)  — вероятный заголовок; выдаём manual_required для подтверждения.
  LOW    (3-5)  — семантический кандидат; принимается только после
    восстановления иерархии, иначе превращается в нарушение
    "возможно пропущенный заголовок".
  NONE   (< 3)  — не заголовок.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field

from docx.enum.text import WD_ALIGN_PARAGRAPH

from services.core.vkr_core.engine.heading_prefix import parse_heading_prefix
from services.core.vkr_core.engine.stats import DocStats, ParagraphStats

TIER_NONE = "none"
TIER_LOW = "low"
TIER_MEDIUM = "medium"
TIER_HIGH = "high"

TIER_MEDIUM_THRESHOLD = 6
TIER_LOW_THRESHOLD = 3

_TIER_THRESHOLDS = [
    (TIER_HIGH, 10),
    (TIER_MEDIUM, TIER_MEDIUM_THRESHOLD),
    (TIER_LOW, TIER_LOW_THRESHOLD),
]


@dataclass
class HeadingScore:
    index: int
    score: int
    tier: str
    level: int
    derived_number: str | None
    signals: list[str] = field(default_factory=list)
    hard_reject_reason: str | None = None


_APPENDIX_RE = re.compile(r"^приложение\s+[а-яё]", re.IGNORECASE)
_CAPTION_RE = re.compile(
    r"^(?:рисунок|рис\.|таблица|табл\.|листинг)\s+\d+(?:\.\d+)*"
    r"(?:\s*[:–—-]\s*|\.\s+)\S"
    r"|^(?:продолжение|окончание)\s+таблицы\s+\d+(?:\.\d+)*\b"
    r"|^формула\s+\(?\d+(?:\.\d+)*\)?\b",
    re.IGNORECASE,
)
_INTERNAL_SENTENCE_RE = re.compile(r"[а-яa-z][.!?]\s+[А-ЯA-Z]")

# Сигналы, перебивающие hard-reject по hanging-indent'у.
# numPr и литеральный номер намеренно исключены — в отдельности недостаточны.
_STRONG_ANCHOR_SIGNALS = (
    "native_style", "outline_level",
    "chapter_word", "appendix", "structural_keyword",
)


def _norm(text: str) -> str:
    return " ".join(text.strip().lower().split())


def _native_heading_level(style_name: str) -> int | None:
    lowered = style_name.strip().lower()
    for prefix in ("heading ", "заголовок "):
        if lowered.startswith(prefix):
            tail = lowered[len(prefix):].strip()
            if tail.isdigit():
                return max(1, min(int(tail), 3))
    match = re.match(r"(heading|заголовок)\s*(\d+)$", lowered)
    if match:
        return max(1, min(int(match.group(2)), 3))
    return None


def _structural_keyword_level(normalized: str, rules: dict) -> int | None:
    structural = rules.get("structural_elements", {})
    for key in ("contents", "introduction", "conclusion", "references", "terms"):
        for alias in structural.get(key, []):
            if _norm(alias) == normalized:
                return 1
    return None


def _is_forbidden(normalized: str, rules: dict) -> bool:
    """True, если заголовок запрещён нормативной базой БЕЗ замены.

    Запрещённые с заменой (например, "оглавление" → "СОДЕРЖАНИЕ") сюда НЕ
    попадают: их обрабатывает apply_forbidden_replacements уже после
    обнаружения. Источник — rules.forbidden_heading_map (БД), никаких
    хардкод-baseline'ов в коде нет.
    """
    forbidden_map = rules.get("forbidden_heading_map") or {}
    for raw_key, replacement in forbidden_map.items():
        if replacement is None and _norm(raw_key) == normalized:
            return True
    return False


def _prev_body_length(stats: DocStats, index: int) -> int:
    for i in range(index - 1, -1, -1):
        ps = stats.paragraphs[i]
        if not ps.is_empty:
            return ps.length
    return 0


def _next_body_length(stats: DocStats, index: int) -> int:
    for i in range(index + 1, len(stats.paragraphs)):
        ps = stats.paragraphs[i]
        if not ps.is_empty:
            return ps.length
    return 0


def _visible_numbering_tokens(ps: ParagraphStats) -> tuple[int, ...]:
    if ps.numbering_tokens:
        return ps.numbering_tokens
    prefix = parse_heading_prefix(ps.stripped)
    if prefix is not None and prefix.kind == "number":
        return prefix.tokens
    return ()


def _prev_numbering_tokens(stats: DocStats, index: int) -> tuple[int, ...]:
    for i in range(index - 1, -1, -1):
        ps = stats.paragraphs[i]
        if not ps.is_empty:
            return _visible_numbering_tokens(ps)
    return ()


def _next_numbering_tokens(stats: DocStats, index: int) -> tuple[int, ...]:
    for i in range(index + 1, len(stats.paragraphs)):
        ps = stats.paragraphs[i]
        if not ps.is_empty:
            return _visible_numbering_tokens(ps)
    return ()


def _dense_text(stripped: str) -> bool:
    return len(stripped) >= 100 and stripped.count(",") >= 3


def _classify_tier(score: int) -> str:
    for tier, threshold in _TIER_THRESHOLDS:
        if score >= threshold:
            return tier
    return TIER_NONE


def _heading_detection_config(rules: dict) -> tuple[int, int, float, int]:
    cfg = rules.get("heading_detection", {}) or {}
    max_length = int(cfg.get("max_length", 200))
    max_words = int(cfg.get("max_words", 14))
    font_jump = float(cfg.get("font_jump_pt", 1.0))
    neighbour_long_threshold = int(cfg.get("neighbour_long_threshold", 120))
    return max_length, max_words, max(font_jump, 0.1), neighbour_long_threshold


def _reject(index: int, reason: str) -> HeadingScore:
    return HeadingScore(
        index=index,
        score=0,
        tier=TIER_NONE,
        level=2,
        derived_number=None,
        signals=[],
        hard_reject_reason=reason,
    )


def score_paragraph(
    ps: ParagraphStats,
    stats: DocStats,
    rules: dict,
    title_page_indexes: set[int],
) -> HeadingScore:
    """Вычисляет HeadingScore для одного абзаца."""
    max_length, max_words, font_jump, neighbour_long_threshold = _heading_detection_config(rules)
    if ps.is_empty:
        # Пустой Heading-стиль — placeholder из шаблона. Word явно пометил, доверяем.
        if ps.index not in title_page_indexes and not ps.in_table:
            native_level = _native_heading_level(ps.style_name)
            if native_level is not None:
                return HeadingScore(
                    index=ps.index,
                    score=10,
                    tier=TIER_HIGH,
                    level=native_level,
                    derived_number=None,
                    signals=["native_style_empty"],
                )
        return _reject(ps.index, "empty")
    if ps.in_table:
        return _reject(ps.index, "in_table")
    if ps.index in title_page_indexes:
        return _reject(ps.index, "title_page")
    if _CAPTION_RE.match(ps.stripped):
        return _reject(ps.index, "caption")
    if ps.length > max_length:
        return _reject(ps.index, "too_long")

    native_level = _native_heading_level(ps.style_name)

    # Строки TOC по стилю или таб-стопу. Ручные строки отсекаются через toc_section_indexes в pipeline.py до детектора.
    style_lower = ps.style_name.strip().lower()
    if style_lower.startswith(("toc", "оглавление", "table of contents")):
        return _reject(ps.index, "toc_entry")
    # LO-сгенерированная строка: таб-стоп + цифра страницы в конце.
    if (ps.has_tab or "\t" in ps.stripped) and re.search(r"\d+\s*\.?\s*$", ps.stripped):
        return _reject(ps.index, "toc_entry")

    normalized = _norm(ps.stripped)
    visible_prefix = parse_heading_prefix(ps.stripped)
    visible_number_tokens = (
        visible_prefix.tokens
        if visible_prefix is not None and visible_prefix.kind == "number"
        else ps.numbering_tokens
    )
    visible_derived_number = (
        visible_prefix.number if visible_prefix is not None else ps.derived_number
    )
    strict_chapter_candidate = (
        visible_prefix is not None
        and visible_prefix.kind == "chapter_word"
        and ps.word_count <= max_words
        and not ps.ends_with_period
        and not ps.stripped.endswith(":")
        and not _dense_text(ps.stripped)
    )
    if _is_forbidden(normalized, rules):
        return _reject(ps.index, "forbidden")

    # Числовой префикс + плотная проза — пункт списка, не заголовок.
    if visible_number_tokens and _dense_text(ps.stripped):
        return _reject(ps.index, "dense_numbered_body")

    # Пункты w:numPr, унаследовавшие outline_level от шаблонного стиля.
    # Признаки списка: ≥4 элементов с одним numId, или точка/точка-с-запятой в конце.
    # Без этого правила outline_level=0 из шаблона превращает каждый пункт в Heading 1.
    if ps.numpr_num_id is not None and ps.numpr_ilvl is not None:
        group_size = stats.numpr_group_sizes.get(ps.numpr_num_id, 0)
        # List Paragraph / Абзац списка — однозначный пункт перечисления.
        not_native = _native_heading_level(ps.style_name) is None
        style_is_list = ps.style_name.strip().lower() in (
            "list paragraph", "listparagraph", "абзац списка",
        )
        if (
            group_size >= 4
            or ps.ends_with_period
            or (not_native and ps.stripped.endswith(";"))
            or (not_native and style_is_list)
        ) and not strict_chapter_candidate:
            return _reject(ps.index, "list_item")

    # Body-абзац с outline_level=0 из шаблонного стиля: длинный + точка в конце + не жирный.
    # Без этого правила outline_level=0 даёт +10 и каждый body-абзац становится Heading 1.
    if (
        ps.ends_with_period
        and ps.length > 60
        and ps.bold_ratio < 0.3
        and native_level is None
        and not (visible_prefix is not None and visible_prefix.kind == "chapter_word")
        and not strict_chapter_candidate
    ):
        return _reject(ps.index, "body_text")

    # Двоеточие в конце — lead-in перед списком, не заголовок.
    if (
        ps.stripped.endswith(":")
        and native_level is None
        and not (visible_prefix is not None and visible_prefix.kind in {"chapter_word", "number"})
        and not (ps.outline_level is not None and 0 <= ps.outline_level <= 2)
    ):
        return _reject(ps.index, "lead_in_colon")

    signals: list[str] = []
    score = 0
    level: int | None = None

    # --- СИЛЬНЫЕ сигналы ---
    if native_level is not None:
        score += 10
        level = native_level
        signals.append(f"native_style:{native_level}")

    if ps.outline_level is not None and 0 <= ps.outline_level <= 2:
        score += 10
        if level is None:
            level = ps.outline_level + 1
        signals.append(f"outline_level:{ps.outline_level}")

    if ps.numpr_ilvl is not None and 0 <= ps.numpr_ilvl <= 2:
        # numPr слабее outline_level — встречается и в списках, HIGH только в комбинации.
        score += 4
        if level is None:
            level = ps.numpr_ilvl + 1
        signals.append(f"numpr_ilvl:{ps.numpr_ilvl}")

    if visible_prefix is not None and visible_prefix.kind == "chapter_word":
        score += 10
        # «Глава N.» — L1. «Раздел 2.1» получает уровень из номера.
        word = (visible_prefix.word or "").lower()
        if word in ("раздел", "section") and len(visible_prefix.tokens) > 1:
            level = visible_prefix.level
        else:
            level = 1
        signals.append("chapter_word")

    if _APPENDIX_RE.match(normalized):
        score += 10
        level = 1  # «Приложение X» — всегда L1 независимо от стиля
        signals.append("appendix")

    kw_level = _structural_keyword_level(normalized, rules)
    if kw_level is not None:
        score += 10
        level = kw_level  # Введение/Заключение/… — всегда L1 независимо от стиля
        signals.append("structural_keyword")

    if (
        visible_number_tokens
        and ps.length <= max_length
        and ps.word_count <= max_words
        and not ps.ends_with_period
    ):
        score += 6
        # Текстовый номер "2.1.1" надёжнее outline_level/numPr — перекрываем уровень.
        # Иначе шаблонный outline_level=0 сводит все вложенные заголовки к L1.
        level = max(1, min(len(visible_number_tokens), 3))
        signals.append(f"literal_number:{visible_derived_number}")

    # Если Word уже явно пометил абзац заголовком, видимый многоуровневый
    # номер точнее шаблонного Heading 1 / outlineLvl=0. Это НЕ добавляет
    # баллов: обычный "1.1 Текст." без заголовочного якоря не становится
    # заголовком только из-за номера.
    if (
        visible_prefix is not None
        and visible_prefix.kind == "number"
        and len(visible_number_tokens) >= 2
        and any(
            s.split(":", 1)[0] in {"native_style", "outline_level", "numpr_ilvl"}
            for s in signals
        )
    ):
        level = visible_prefix.level
        signals.append(f"visible_number_level:{visible_derived_number}")

    # --- СЛАБЫЕ сигналы ---
    body_size = stats.modal_font_size_pt or 14.0
    font_size = ps.max_font_size_pt or ps.modal_font_size_pt
    if font_size is not None:
        if font_size >= body_size + (font_jump * 2):
            score += 4
            signals.append(f"font_large:{font_size:.1f}")
        elif font_size >= body_size + font_jump:
            score += 2
            signals.append(f"font_small:{font_size:.1f}")

    if ps.bold_ratio >= 0.8:
        score += 3
        signals.append(f"bold:{ps.bold_ratio:.2f}")
    elif ps.bold_ratio >= 0.5:
        score += 2
        signals.append(f"bold_partial:{ps.bold_ratio:.2f}")

    if ps.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        score += 3
        signals.append("centered")

    if ps.is_upper and ps.length <= 80:
        score += 3
        signals.append("all_caps")

    if ps.length <= min(100, max_length) and ps.word_count <= max_words:
        score += 1
        signals.append("short")

    prev_len = _prev_body_length(stats, ps.index)
    next_len = _next_body_length(stats, ps.index)
    if prev_len >= neighbour_long_threshold and next_len >= neighbour_long_threshold:
        score += 3
        signals.append("isolated")

    if not ps.ends_with_period:
        score += 1
        signals.append("no_period")

    if ps.has_page_break_before:
        score += 3
        signals.append("page_break_before")

    # --- ОТРИЦАТЕЛЬНЫЕ сигналы ---
    has_strong_anchor = any(
        s.split(":", 1)[0] in _STRONG_ANCHOR_SIGNALS for s in signals
    )

    # Hanging indent — признак авто-списка Word. Hard-reject, если нет сильного якоря.
    if ps.first_line_indent_cm is not None and ps.first_line_indent_cm < -0.1:
        if not has_strong_anchor:
            return _reject(ps.index, "hanging_indent")
        score -= 5
        signals.append("hanging_indent_penalty")

    if ps.stripped and ps.stripped[0].islower() and not has_strong_anchor:
        score -= 5
        signals.append("lowercase_start")

    if _INTERNAL_SENTENCE_RE.search(ps.stripped):
        score -= 4
        signals.append("internal_sentence")

    if ps.stripped.endswith(":"):
        score -= 2
        signals.append("ends_with_colon")

    if _dense_text(ps.stripped):
        score -= 3
        signals.append("dense_text")

    if level is None:
        level = 2

    # Заголовок с точкой в конце — только при семантическом якоре или многоуровневом номере.
    # outline_level/numpr_ilvl из шаблона точку не перебивают.
    if ps.ends_with_period:
        _PERIOD_SAFE = {"native_style", "chapter_word",
                        "structural_keyword", "appendix"}
        period_safe = any(
            s.split(":", 1)[0] in _PERIOD_SAFE for s in signals
        )
        multi_level_number = len(visible_number_tokens) >= 2
        if not period_safe and not multi_level_number:
            return _reject(ps.index, "ends_with_period")

    # Без сильного сигнала скор только от слабых (bold+short+isolated) — ограничиваем LOW.
    # Жирный лид-ин без номера и без стиля — не раздел.
    _STRONG = ("native_style", "outline_level", "chapter_word", "appendix",
               "structural_keyword", "literal_number", "numpr_ilvl")
    has_strong = any(s.split(":", 1)[0] in _STRONG for s in signals)
    if not has_strong and score >= TIER_MEDIUM_THRESHOLD:
        score = TIER_LOW_THRESHOLD  # cap → LOW
        signals.append("no_strong_signal_capped")

    # Несколько коротких жирных абзацев подряд без якоря — inline-перечисление, не заголовки.
    # Требуем оба соседа с теми же признаками, чтобы не задеть одиночные лид-ины.
    if not has_strong and ps.bold_ratio >= 0.8 and ps.length <= min(80, max_length):
        enum_neighbors = 0
        for direction in (-1, 1):
            i = ps.index + direction
            while 0 <= i < len(stats.paragraphs):
                np = stats.paragraphs[i]
                if not np.is_empty:
                    if (np.bold_ratio >= 0.8 and np.length <= 80
                            and not np.ends_with_period):
                        enum_neighbors += 1
                    break
                i += direction
        if enum_neighbors >= 2:
            return _reject(ps.index, "inline_enum")

    # Нумерованный список с одиночным префиксом (1, 2, 3…).
    # Кейс A: единственный сильный — literal_number, сосед N±1 → список.
    # Кейс B: есть native_style/outline_level/numpr_ilvl без семантического якоря,
    # сосед тоже короткий и без точки → список, стиль мог назначить форматтер.

    _STRUCT_ANCHORS_SET = {"native_style", "outline_level", "chapter_word",
                           "appendix", "structural_keyword", "numpr_ilvl"}
    _SEMANTIC_STRUCTURAL = {"chapter_word", "appendix", "structural_keyword"}
    _NON_SEMANTIC_STRUCTURAL = {"native_style", "outline_level", "numpr_ilvl"}

    has_semantic_structural = any(
        s.split(":", 1)[0] in _SEMANTIC_STRUCTURAL for s in signals
    )

    if visible_number_tokens and len(visible_number_tokens) == 1 and not has_semantic_structural:
        n = visible_number_tokens[0]
        prev_tok = _prev_numbering_tokens(stats, ps.index)
        next_tok = _next_numbering_tokens(stats, ps.index)

        only_literal_strong = (
            has_strong
            and not any(s.split(":", 1)[0] in _STRUCT_ANCHORS_SET for s in signals)
        )

        if only_literal_strong:
            # Кейс A: одного соседа N±1 достаточно.
            if prev_tok == (n - 1,) or next_tok == (n + 1,):
                return _reject(ps.index, "list_sequence")
        elif any(s.split(":", 1)[0] in _NON_SEMANTIC_STRUCTURAL for s in signals):
            # Кейс B: сосед должен сам быть коротким и без точки (список, не тело).
            def _neighbor_is_listlike(tok_matches: bool, neighbor_idx_range) -> bool:
                if not tok_matches:
                    return False
                for i in neighbor_idx_range:
                    np = stats.paragraphs[i]
                    if not np.is_empty:
                        return np.length <= 120 and not np.ends_with_period
                return False

            prev_match = prev_tok == (n - 1,)
            next_match = next_tok == (n + 1,)
            if (_neighbor_is_listlike(prev_match, range(ps.index - 1, -1, -1))
                    or _neighbor_is_listlike(next_match, range(ps.index + 1, len(stats.paragraphs)))):
                return _reject(ps.index, "list_sequence")

    return HeadingScore(
        index=ps.index,
        score=score,
        tier=_classify_tier(score),
        level=max(1, min(level, 3)),
        derived_number=visible_derived_number,
        signals=signals,
        hard_reject_reason=None,
    )


def score_all(
    stats: DocStats,
    rules: dict,
    title_page_indexes: set[int],
) -> list[HeadingScore]:
    """Возвращает скоры всех абзацев, которые не были hard-reject'нуты И
    набрали хотя бы уровень LOW. Hard-rejected абзацы в список не попадают."""
    results: list[HeadingScore] = []
    for ps in stats.paragraphs:
        score = score_paragraph(ps, stats, rules, title_page_indexes)
        if score.hard_reject_reason is not None:
            continue
        if score.tier == TIER_NONE:
            continue
        results.append(score)
    return results
