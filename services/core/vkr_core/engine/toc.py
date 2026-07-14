r"""Вставка и пост-обработка оглавления (ТЗ §6.6).

Генерируем оглавление вручную: для каждого заголовка (outlineLvl 0–2)
создаём абзац «текст + TAB + PAGEREF-поле». LibreOffice headless при
конвертации обновляет PAGEREF, проставляя реальные номера страниц.

Почему не Word TOC-поле (`TOC \o ...`):
  При нажатии «Обновить поле» Word перезаписывает прямое форматирование
  TOC-абзацев иерархическими w:ind (TOC2 — отступ, TOC3 — больший отступ)
  независимо от определений стилей. Статические PAGEREF-абзацы обновляют
  ТОЛЬКО номера страниц, не трогая структуру и отступы.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

_logger = logging.getLogger(__name__)

# Лимит итераций резки таблиц. Сходится за ≈ макс. глубину переноса; страховка от зависания.
_MAX_TABLE_SPLIT_ITERS = 6


def _norm(text: str) -> str:
    return " ".join(text.strip().lower().split())


def _has_explicit_page_break(paragraph: Paragraph) -> bool:
    for br in paragraph._element.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


_NATIVE_HEADING_STYLE_RE = re.compile(r"^(heading|заголовок)\s*[1-3]$")


def _looks_like_real_heading(paragraph: Paragraph) -> bool:
    """True, если абзац точно настоящий заголовок — по родному стилю Word
    («Heading N»/«Заголовок N») или по прямому outlineLvl 0–2 (уже
    проставленному apply_headings).

    Нужен как страховка от эвристики «текстовое оглавление без лидеров»:
    после нормализации регистра apply_headings настоящий заголовок («ТЕРМИНЫ
    И ОПРЕДЕЛЕНИЯ» → «Термины и определения») перестаёт быть капсом и больше
    не ловится текстовым признаком «капс = структурный заголовок» — тогда его
    ошибочно принимают за фиктивную строку оглавления и удаляют.
    """
    style_name = (paragraph.style.name or "").strip().lower() if paragraph.style else ""
    if _NATIVE_HEADING_STYLE_RE.match(style_name):
        return True
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return False
    olvl = pPr.find(qn("w:outlineLvl"))
    if olvl is None:
        return False
    try:
        return 0 <= int(olvl.get(qn("w:val"), "8")) <= 2
    except (ValueError, TypeError):
        return False


_DEFAULT_CONTENTS_NORMS = {"содержание", "оглавление"}


def _contents_norms(rules: dict | None) -> set[str]:
    """Нормализованные алиасы заголовка оглавления (СОДЕРЖАНИЕ/ОГЛАВЛЕНИЕ
    + синонимы из rules). rules может быть None — тогда только дефолтные."""
    norms = set(_DEFAULT_CONTENTS_NORMS)
    if rules:
        for item in rules.get("structural_elements", {}).get("contents", []):
            n = _norm(item)
            if n:
                norms.add(n)
    return norms


def _find_contents_heading(doc: DocxDocument, rules: dict) -> Paragraph | None:
    aliases = _contents_norms(rules)
    for paragraph in doc.paragraphs:
        if _norm(paragraph.text) in aliases:
            return paragraph
    return None


def _canonicalize_contents_headings(
    doc: DocxDocument, rules: dict | None, anchor: Paragraph
) -> None:
    """Удаляет дубли абзаца «СОДЕРЖАНИЕ», оставляя только anchor.

    После finalize→LO→rebuild в документе может остаться второй «СОДЕРЖАНИЕ»,
    который попадает в TOC призрачной строкой и накапливается при каждом цикле.
    """
    norms = _contents_norms(rules)
    for para in list(doc.paragraphs):
        if para._p is anchor._p:
            continue
        if _norm(para.text) not in norms:
            continue
        p = para._p
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)


def _paragraph_index(doc: DocxDocument, target: Paragraph) -> int | None:
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph._p is target._p:
            return idx
    return None


def _manual_toc_by_repetition(paragraphs, heading_idx: int) -> set[int]:
    """Ручное оглавление голыми заголовками (без лидеров и номеров страниц).

    Признак: каждая строка блока дословно повторяется ниже реальным заголовком.
    Возвращает пусто, если блок короче 2 строк — тогда работают остальные фазы.
    """
    norm_indexes: dict[str, list[int]] = {}
    for idx, paragraph in enumerate(paragraphs):
        key = _norm(paragraph.text)
        if key:
            norm_indexes.setdefault(key, []).append(idx)

    manual: set[int] = set()
    for idx in range(heading_idx + 1, len(paragraphs)):
        text = paragraphs[idx].text.strip()
        if not text:
            break  # пустая строка — оглавление кончилось
        later = norm_indexes.get(_norm(text), ())
        if any(j > idx for j in later):
            manual.add(idx)
        else:
            break  # ниже не повторяется — это уже реальный раздел
    return manual if len(manual) >= 2 else set()


def _is_toc_field_paragraph(paragraph: Paragraph) -> bool:
    xml = paragraph._p.xml
    return " TOC " in xml or "TOC \\o" in xml


def _sdt_is_toc(sdt_elem) -> bool:
    """True если <w:sdt> является оглавлением Word (Quick TOC gallery).

    Не использует .xml (который есть только у python-docx-зарегистрированных
    классов), а ищет маркеры через lxml iter().
    """
    # docPartGallery с "Table of Contents" или "Оглавление"
    sdt_pr = sdt_elem.find(qn("w:sdtPr"))
    if sdt_pr is not None:
        for elem in sdt_pr.iter(qn("w:docPartGallery")):
            val = (elem.get(qn("w:val")) or "").lower()
            if "contents" in val or "оглавлени" in val:
                return True
    # instrText с полем TOC внутри sdtContent
    for instr in sdt_elem.iter(qn("w:instrText")):
        text = instr.text or ""
        if " TOC " in text or "TOC \\o" in text:
            return True
    return False


# Паттерны конца строки ручного оглавления:
#   "Введение ......... 3"   — точечные лидеры
#   "Введение\t3" / "\t3."  — TAB + номер
#   "Введение   3" / " 3."  — 3+ пробела + номер
#   "Введение — 3" / "— 3." — тире + номер
_TOC_LINE_RE = re.compile(
    r"(?:"
    r"\.{3,}\s*\d+\s*\.?\s*$"
    r"|\t\s*\d{1,3}\s*\.?\s*$"
    r"|\s{3,}\d{1,3}\s*\.?\s*$"
    r"|[—–-]\s*\d{1,3}\s*\.?\s*$"
    r")"
)


def _is_toc_line(text: str) -> bool:
    return bool(_TOC_LINE_RE.search(text))


def _is_toc_entry(paragraph: Paragraph) -> bool:
    """True, если абзац является записью оглавления (любого вида).

    Ловит три случая Word-оглавления:
      1. Абзац с instrText «TOC \\o ...» — первый абзац поля оглавления.
      2. Абзац с TOC-стилем (TOC 1/2/3) — результирующие строки Word-TOC.
      3. Короткий абзац с PAGEREF-полем — Word-TOC без явных стилей.
    И один случай ручного оглавления:
      4. Строка с точечными лидерами / TAB+страница / тире+страница.
    """
    if _is_toc_field_paragraph(paragraph):
        return True
    if paragraph.style:
        name = (paragraph.style.name or "").strip().lower()
        if re.match(r"^toc\s*\d*$", name) or re.match(r"^(оглавление|содержание)\s*\d*$", name):
            return True
    text = paragraph.text.strip()
    if text and _is_toc_line(text):
        return True
    # Word-TOC result без видимого таба: text = "Заголовок3" (tab — w:tab, не w:t).
    # Распознаём по наличию PAGEREF-поля при коротком тексте.
    if text and len(text) < 200 and "PAGEREF" in paragraph._p.xml:
        return True
    return False


def _remove_existing_toc(doc: DocxDocument, rules: dict, title_end_index: int) -> None:
    heading = _find_contents_heading(doc, rules)
    if heading is None:
        return
    heading_idx = _paragraph_index(doc, heading)
    if heading_idx is None or heading_idx < title_end_index:
        return

    # Удаляем <w:sdt> целиком — doc.paragraphs не видит параграфы внутри него.
    body = doc.element.body
    heading_elem = heading._p
    after_heading = False
    for child in list(body):
        if child is heading_elem:
            after_heading = True
            continue
        if not after_heading:
            continue
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag != "sdt":
            continue
        if _sdt_is_toc(child):
            body.remove(child)

    paragraphs = doc.paragraphs
    # Пересчитываем heading_idx — sdt мог стоять перед СОДЕРЖАНИЕ.
    heading_idx = _paragraph_index(doc, heading)
    if heading_idx is None:
        return

    to_delete: list[Paragraph] = []
    found_toc_content = False

    # Фаза 1: TOC-поля, TOC-стили, PAGEREF в коротком абзаце, ручные строки с лидерами.
    for idx, paragraph in enumerate(paragraphs):
        if idx <= heading_idx:
            continue
        if _is_toc_entry(paragraph):
            to_delete.append(paragraph)
            found_toc_content = True
            continue
        if not paragraph.text.strip():
            if _has_explicit_page_break(paragraph):
                break  # разрыв страницы — конец TOC
            to_delete.append(paragraph)  # пустой абзац — удаляем условно
            continue
        # Непустая строка, не похожая на запись TOC — конец секции.
        if found_toc_content:
            break
        # Ещё не нашли ни одной TOC-строки — прерываем; может быть фаза 2.
        break

    if not found_toc_content:
        to_delete.clear()

        # Фаза 1.5: ручное оглавление голыми заголовками (без лидеров).
        _manual = _manual_toc_by_repetition(paragraphs, heading_idx)
        if _manual:
            for paragraph in (paragraphs[i] for i in sorted(_manual)):
                parent = paragraph._element.getparent()
                if parent is not None:
                    parent.remove(paragraph._element)
            return

        # alias_set из structural_elements — чтобы «Термины и определения» останавливало удаление.
        structural = rules.get("structural_elements", {})
        _p2_alias_set: set[str] = set()
        for _key, _val in structural.items():
            if _key == "appendix_regex":
                continue
            for _alias in (_val or []):
                _n = _norm(_alias)
                if _n:
                    _p2_alias_set.add(_n)
        _p2_appendix_re = None
        _appendix_pat = structural.get("appendix_regex")
        if _appendix_pat:
            try:
                _p2_appendix_re = re.compile(_appendix_pat, re.IGNORECASE)
            except re.error:
                pass

        def _p2_is_structural(paragraph: Paragraph) -> bool:
            s = paragraph.text.strip()
            if not s:
                return False
            if _looks_like_real_heading(paragraph):
                return True
            if _norm(s) in _p2_alias_set:
                return True
            if _CHAPTER_START_RE.match(s):
                return True
            if _p2_appendix_re is not None and _p2_appendix_re.match(s):
                return True
            letters = [c for c in s if c.isalpha()]
            if letters and len(s) <= 80 and all(c.isupper() for c in letters):
                return True
            return False

        # Граница: первый структурный заголовок, тело или таблица — берём min.
        _table_boundary: int | None = None
        _p_idx = -1
        _seen_heading = False
        for _child in doc.element.body.iterchildren():
            _tag = _child.tag.split("}")[-1] if "}" in _child.tag else _child.tag
            if _tag == "p":
                _p_idx += 1
                if _p_idx == heading_idx:
                    _seen_heading = True
            elif _tag == "tbl" and _seen_heading:
                _table_boundary = _p_idx
                break

        boundary: int | None = None
        for idx in range(heading_idx + 1, len(paragraphs)):
            if _p2_is_structural(paragraphs[idx]) or _is_body_paragraph(paragraphs[idx].text.strip()):
                boundary = idx
                break

        if _table_boundary is not None:
            boundary = _table_boundary if boundary is None else min(boundary, _table_boundary)

        if boundary is not None:
            for idx in range(heading_idx + 1, boundary):
                para = paragraphs[idx]
                if _p2_is_structural(para):
                    break
                if not _has_explicit_page_break(para):
                    to_delete.append(para)

    # Удаляем собранные абзацы (getparent может быть None для уже удалённых).
    for paragraph in to_delete:
        parent = paragraph._element.getparent()
        if parent is not None:
            parent.remove(paragraph._element)


_CHAPTER_START_RE = re.compile(
    r"^(?:глава|раздел|часть|chapter|section|part)\s+\d+", re.IGNORECASE
)
_NUM_PREFIX_TOC_RE = re.compile(r"^\d+[\.\d]*[\.\)]\s")


def _is_body_paragraph(text: str) -> bool:
    """True, если абзац похож на тело документа, а не на запись оглавления.

    Тело: длинный (>= 80 символов), не все заглавные, не начинается
    с «Глава N» или числового префикса, содержит строчные буквы.
    """
    if len(text) < 80:
        return False
    stripped = text.strip()
    letters = [c for c in stripped if c.isalpha()]
    if not letters:
        return False
    if all(c.isupper() for c in letters):
        return False  # всё заглавными — скорее заголовок
    if _CHAPTER_START_RE.match(stripped):
        return False
    if _NUM_PREFIX_TOC_RE.match(stripped):
        return False
    return sum(1 for c in stripped if c.islower()) >= 10


def detect_toc_section_indexes(doc: DocxDocument, rules: dict) -> set[int]:
    """Возвращает множество индексов параграфов ручного оглавления.

    Фаза 1 — классическое оглавление: строки с точечными лидерами
    (``\\.{3,}\\s*\\d+\\s*$``) или TOC-поля Word.

    Фаза 2 — текстовое оглавление без точек (студент просто перечислил
    названия глав): ищем первый абзац тела документа (_is_body_paragraph),
    затем смотрим назад до первого непустого абзаца — это первый настоящий
    заголовок. Всё между СОДЕРЖАНИЕ и этим заголовком добавляется в
    toc_indexes, сам заголовок — нет.
    """
    heading = _find_contents_heading(doc, rules)
    if heading is None:
        return set()
    heading_idx = _paragraph_index(doc, heading)
    if heading_idx is None:
        return set()

    paragraphs = doc.paragraphs
    toc_indexes: set[int] = set()

    # Фаза 1: точечные лидеры, TAB/пробелы + страница, тире + страница,
    # TOC-поле Word, TOC-стили, PAGEREF в коротком абзаце.
    for idx, paragraph in enumerate(paragraphs):
        if idx <= heading_idx:
            continue
        if _is_toc_entry(paragraph):
            toc_indexes.add(idx)
            continue
        if paragraph.text.strip() and toc_indexes:
            break

    if toc_indexes:
        return toc_indexes

    # Фаза 1.5: ручное оглавление голыми заголовками (повторяются ниже).
    manual_toc = _manual_toc_by_repetition(paragraphs, heading_idx)
    if manual_toc:
        return manual_toc

    # Фаза 2: текстовое оглавление без точечных лидеров (в т.ч. ручной TOC,
    # сделанный табом с w:leader="dot" — в paragraph.text ни точек, ни таба,
    # поэтому Фаза 1 его не видит).
    #
    # Граница TOC-региона — ПЕРВЫЙ структурный заголовок (раздел из
    # rules.structural_elements / «Глава N» / приложение / короткий заголовок
    # капсом), ПЕРВЫЙ абзац тела ИЛИ первая таблица после «СОДЕРЖАНИЕ» — что
    # встретится раньше. Раньше границей считался «ближайший непустой перед
    # первым длинным абзацем тела»; это вырезало раздел «ТЕРМИНЫ И
    # ОПРЕДЕЛЕНИЯ», стоящий до «Введения» и состоящий из таблицы (его текст
    # не body, а первый body-абзац оказывался уже во «Введении»).
    structural = rules.get("structural_elements", {})
    alias_set: set[str] = set()
    for key, val in structural.items():
        if key == "appendix_regex":
            continue
        for alias in (val or []):
            normalized = _norm(alias)
            if normalized:
                alias_set.add(normalized)
    appendix_re = None
    appendix_pat = structural.get("appendix_regex")
    if appendix_pat:
        try:
            appendix_re = re.compile(appendix_pat, re.IGNORECASE)
        except re.error:
            appendix_re = None

    def _is_structural_boundary(paragraph: Paragraph) -> bool:
        s = paragraph.text.strip()
        if not s:
            return False
        if _looks_like_real_heading(paragraph):
            return True
        if _norm(s) in alias_set:
            return True
        if _CHAPTER_START_RE.match(s):
            return True
        if appendix_re is not None and appendix_re.match(s):
            return True
        # Короткий заголовок, набранный капсом ("ТЕРМИНЫ И ОПРЕДЕЛЕНИЯ").
        letters = [c for c in s if c.isalpha()]
        if letters and len(s) <= 80 and all(c.isupper() for c in letters):
            return True
        return False

    # Индекс абзаца, идущего непосредственно перед первой таблицей после
    # «СОДЕРЖАНИЕ»: таблица в ручном оглавлении не встречается, значит TOC
    # уже закончился и предшествующий абзац — настоящий заголовок раздела.
    table_boundary: int | None = None
    p_idx = -1
    seen_heading = False
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            p_idx += 1
            if p_idx == heading_idx:
                seen_heading = True
        elif tag == "tbl" and seen_heading:
            table_boundary = p_idx  # последний абзац перед таблицей
            break

    boundary: int | None = None
    for idx in range(heading_idx + 1, len(paragraphs)):
        if _is_structural_boundary(paragraphs[idx]) or _is_body_paragraph(paragraphs[idx].text.strip()):
            boundary = idx
            break

    if table_boundary is not None:
        boundary = table_boundary if boundary is None else min(boundary, table_boundary)

    if boundary is None:
        return toc_indexes  # граница не найдена — ничего не трогаем

    for idx in range(heading_idx + 1, boundary):
        # Жёсткая страховка: структурный заголовок никогда не считается TOC.
        if _is_structural_boundary(paragraphs[idx]):
            continue
        toc_indexes.add(idx)

    return toc_indexes


_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


def _find_headings_for_toc(
    doc: DocxDocument,
    title_end_index: int,
    contents_idx: int,
    contents_norms: set[str] | None = None,
) -> list[tuple[int, int, str]]:
    """Возвращает [(para_idx, level, text)] для заголовков с outlineLvl 0–2.

    Все абзацы «СОДЕРЖАНИЕ»/«ОГЛАВЛЕНИЕ» исключаем — не только якорный,
    чтобы призрачный дубль не попал в TOC отдельной строкой.
    """
    cnorms = contents_norms if contents_norms is not None else _DEFAULT_CONTENTS_NORMS
    result = []
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx < title_end_index or idx == contents_idx:
            continue
        if _norm(paragraph.text) in cnorms:
            continue
        pPr = paragraph._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        olvl = pPr.find(qn("w:outlineLvl"))
        if olvl is None:
            continue
        try:
            level = int(olvl.get(qn("w:val"), "8")) + 1  # 0-based → 1-based
        except (ValueError, TypeError):
            continue
        if 1 <= level <= 3:
            text = paragraph.text.strip()
            if text:
                result.append((idx, level, text))
    return result


def _ensure_bookmark(paragraph: Paragraph, name: str, bm_id: int) -> None:
    """Вставляет bookmarkStart/End сразу после pPr (или в начало абзаца)."""
    p = paragraph._p
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(bm_id))
    bm_start.set(qn("w:name"), name)
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bm_id))
    pPr = p.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(bm_end)
        pPr.addnext(bm_start)
    else:
        p.insert(0, bm_end)
        p.insert(0, bm_start)


_TOC_LINE_SPACING_TWIPS = "360"  # 1.5 интервала при 14pt


def _make_toc_entry_p(
    text: str,
    bookmark_name: str,
    tab_twips: int,
) -> object:
    """Строит <w:p> для строки оглавления: текст + TAB + PAGEREF.

    Форматирование напрямую в pPr/rPr — не через стиль, чтобы не портилось при Update Field.
    """
    p = OxmlElement("w:p")

    # --- pPr ---
    pPr = OxmlElement("w:pPr")
    p.append(pPr)

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), _TOC_LINE_SPACING_TWIPS)
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)

    # Без иерархического отступа: все уровни оглавления от левого края,
    # различаются только текстом, а не абзацным отступом «по ранжиру».
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:firstLine"), "0")
    pPr.append(ind)

    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), str(tab_twips))
    tabs.append(tab)
    pPr.append(tabs)

    # Выравнивание по левому краю (jc = left или просто отсутствует)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "left")
    pPr.append(jc)

    def _rPr() -> object:
        rp = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rFonts.set(qn("w:cs"), "Times New Roman")
        rp.append(rFonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "28")       # 14pt = 28 half-points
        rp.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), "28")
        rp.append(szCs)
        return rp

    # --- текстовый run ---
    r_text = OxmlElement("w:r")
    r_text.append(_rPr())
    t = OxmlElement("w:t")
    t.set(_XML_SPACE, "preserve")
    t.text = text
    r_text.append(t)
    p.append(r_text)

    # --- TAB ---
    r_tab = OxmlElement("w:r")
    r_tab.append(_rPr())
    r_tab.append(OxmlElement("w:tab"))
    p.append(r_tab)

    # --- PAGEREF field: begin → instrText → separate → placeholder → end ---
    r_begin = OxmlElement("w:r")
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    r_begin.append(fc_begin)
    p.append(r_begin)

    r_instr = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(_XML_SPACE, "preserve")
    instr.text = f" PAGEREF {bookmark_name} \\h "
    r_instr.append(instr)
    p.append(r_instr)

    r_sep = OxmlElement("w:r")
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    r_sep.append(fc_sep)
    p.append(r_sep)

    r_ph = OxmlElement("w:r")
    r_ph.append(_rPr())
    t_ph = OxmlElement("w:t")
    t_ph.text = "1"        # placeholder, будет заменён LibreOffice/Word
    r_ph.append(t_ph)
    p.append(r_ph)

    r_end = OxmlElement("w:r")
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    r_end.append(fc_end)
    p.append(r_end)

    return p


def _set_jc_xml(paragraph, w_val: str) -> None:
    """Принудительно задаёт w:jc в pPr абзаца, обходя наследование стиля."""
    pPr = paragraph._p.get_or_add_pPr()
    for existing in list(pPr.findall(qn("w:jc"))):
        pPr.remove(existing)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), w_val)
    pPr.append(jc)


def _unwrap_sdt_toc(doc: DocxDocument, rules: dict) -> None:
    """Удаляет <w:sdt> Quick TOC и заменяет его абзацем «СОДЕРЖАНИЕ».

    Word Quick TOC прячет весь TOC в <w:sdt> — doc.paragraphs его не видит,
    поэтому _find_contents_heading не находит заголовок.
    """
    body = doc.element.body
    aliases = {
        _norm(item)
        for item in rules.get("structural_elements", {}).get("contents", [])
    }
    aliases.discard("")

    for sdt in list(body.findall(qn("w:sdt"))):
        if not _sdt_is_toc(sdt):
            continue
        # Создаём новый абзац «СОДЕРЖАНИЕ» на месте SDT.
        # insert_toc полностью перепишет его (bold, TNR 14pt, по центру).
        new_p = OxmlElement("w:p")
        r_elem = OxmlElement("w:r")
        t_elem = OxmlElement("w:t")
        t_elem.text = "СОДЕРЖАНИЕ"
        r_elem.append(t_elem)
        new_p.append(r_elem)
        sdt.addprevious(new_p)
        body.remove(sdt)
        break


def insert_toc(doc: DocxDocument, rules: dict, title_end_index: int) -> bool:
    heading = _find_contents_heading(doc, rules)
    if heading is None:
        return False
    heading_idx = _paragraph_index(doc, heading)
    if heading_idx is None or heading_idx < title_end_index:
        return False

    _remove_existing_toc(doc, rules, title_end_index)

    # heading.text = "..." пересоздаёт run без форматирования — работаем через XML.
    p = heading._p
    for r in list(p.findall(qn("w:r"))):
        p.remove(r)
    for h in list(p.findall(qn("w:hyperlink"))):
        p.remove(h)
    run = heading.add_run("СОДЕРЖАНИЕ")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.all_caps = False

    # Выравнивание через XML — надёжнее питоновского сеттера.
    _set_jc_xml(heading, "center")
    heading.paragraph_format.line_spacing = 1.5

    # Ширина текстового блока для tab-stop (правый край, лидер из точек).
    try:
        section = doc.sections[0]
        text_width_twips = int(
            (section.page_width - section.left_margin - section.right_margin)
            * 1440 / 914400
        )
    except Exception:
        text_width_twips = 9354  # 165 мм — A4 с полями ГОСТ

    # Схлопываем дубли «СОДЕРЖАНИЕ» до сбора заголовков.
    _canonicalize_contents_headings(doc, rules, heading)
    heading_idx = _paragraph_index(doc, heading)
    if heading_idx is None:
        return False

    # Находим все заголовки документа (по outlineLvl, уровни 1–3).
    cnorms = _contents_norms(rules)
    toc_entries = _find_headings_for_toc(doc, title_end_index, heading_idx, cnorms)
    if not toc_entries:
        return False

    # Закладки на заголовки и строки TOC: текст + TAB + PAGEREF.
    last_p = heading._p
    for para_idx, _level, text in toc_entries:
        heading_para = doc.paragraphs[para_idx]
        bm_name = f"_VkrToc_{para_idx}"
        _ensure_bookmark(heading_para, bm_name, bm_id=para_idx + 10000)

        entry_p = _make_toc_entry_p(text, bm_name, text_width_twips)
        last_p.addnext(entry_p)
        last_p = entry_p

    return True


def normalize_toc_indents(doc: DocxDocument) -> None:
    """Удаляет прямые w:ind у абзацев TOC 1/2/3.

    LO при пересчёте TOC может дописать иерархические w:ind, перекрывая стиль.
    """
    toc_style_names = {"toc 1", "toc 2", "toc 3"}
    for paragraph in doc.paragraphs:
        style_name = (paragraph.style.name or "").strip().lower() if paragraph.style else ""
        if style_name not in toc_style_names:
            continue
        pPr = paragraph._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        for ind in list(pPr.findall(qn("w:ind"))):
            pPr.remove(ind)


def rebuild_toc_full(doc: DocxDocument, rules: dict | None = None) -> None:
    """Полностью перестраивает строки оглавления по текущим заголовкам.

    Вызывается после любой ручной правки заголовка (fix-heading-number,
    promote-to-heading, demote-heading-to-text/list, revert heading_recovered),
    чтобы TOC отражал актуальную структуру документа — добавлял новые
    заголовки, удалял разжалованные и обновлял изменённые номера/тексты.
    """
    # 1. Найти абзац-заголовок оглавления: последний непустой абзац перед
    #    первой строкой с _VkrToc_ instrText.
    toc_heading: Paragraph | None = None
    last_nonempty: Paragraph | None = None
    for para in doc.paragraphs:
        has_vkr = any(
            "_VkrToc_" in (e.text or "")
            for e in para._p.iter(qn("w:instrText"))
        )
        if has_vkr:
            toc_heading = last_nonempty
            break
        if para.text.strip():
            last_nonempty = para

    if toc_heading is None:
        # _VkrToc_ маркеры отсутствуют — LO мог стереть instrText при convert-to docx.
        if rules is not None:
            # Чистим старые _VkrToc_* закладки перед insert_toc — иначе конфликт ID.
            for _para in doc.paragraphs:
                _p = _para._p
                for _bm in list(_p.findall(qn("w:bookmarkStart"))):
                    if not _bm.get(qn("w:name"), "").startswith("_VkrToc_"):
                        continue
                    _bid = _bm.get(qn("w:id"), "")
                    _p.remove(_bm)
                    for _be in list(_p.findall(qn("w:bookmarkEnd"))):
                        if _be.get(qn("w:id"), "") == _bid:
                            _p.remove(_be)
                            break
            insert_toc(doc, rules, title_end_index=0)
        return  # нет TOC — нечего перестраивать

    # 2. Удаляем _VkrToc_ записи и пустые абзацы между ними.
    heading_elem = toc_heading._p
    found_heading = False
    found_toc = False
    to_delete: list[Paragraph] = []
    for para in doc.paragraphs:
        if para._p is heading_elem:
            found_heading = True
            continue
        if not found_heading:
            continue
        has_vkr = any(
            "_VkrToc_" in (e.text or "")
            for e in para._p.iter(qn("w:instrText"))
        )
        if has_vkr:
            to_delete.append(para)
            found_toc = True
        elif not para.text.strip() and found_toc:
            to_delete.append(para)
        elif para.text.strip():
            break

    for para in to_delete:
        parent = para._element.getparent()
        if parent is not None:
            parent.remove(para._element)

    # 3. Чистим _VkrToc_* закладки с заголовков — чтобы не дублировались ID.
    for para in doc.paragraphs:
        p = para._p
        for bm_start in list(p.findall(qn("w:bookmarkStart"))):
            if not bm_start.get(qn("w:name"), "").startswith("_VkrToc_"):
                continue
            bm_id = bm_start.get(qn("w:id"), "")
            p.remove(bm_start)
            for bm_end in list(p.findall(qn("w:bookmarkEnd"))):
                if bm_end.get(qn("w:id"), "") == bm_id:
                    p.remove(bm_end)
                    break

    # 3b. Схлопываем дубли «СОДЕРЖАНИЕ» — иначе дефект накапливается при каждом rebuild.
    _canonicalize_contents_headings(doc, rules, toc_heading)

    # 4. Пересчитываем индекс TOC-heading после удаления строк.
    contents_idx = _paragraph_index(doc, toc_heading)
    if contents_idx is None:
        return

    # 5. Ширина text-блока для tab-стопа с лидером.
    try:
        section = doc.sections[0]
        text_width_twips = int(
            (section.page_width - section.left_margin - section.right_margin)
            * 1440 / 914400
        )
    except Exception:
        text_width_twips = 9354  # 165 мм — A4 с полями ГОСТ

    # 6. Найти все текущие заголовки (outlineLvl 0–2) после TOC.
    toc_entries = _find_headings_for_toc(
        doc, contents_idx + 1, contents_idx, _contents_norms(rules)
    )

    # 7. Вставить новые TOC-записи с обновлёнными закладками.
    last_p = toc_heading._p
    for para_idx, _level, text in toc_entries:
        heading_para = doc.paragraphs[para_idx]
        bm_name = f"_VkrToc_{para_idx}"
        _ensure_bookmark(heading_para, bm_name, bm_id=para_idx + 10000)
        entry_p = _make_toc_entry_p(text, bm_name, text_width_twips)
        last_p.addnext(entry_p)
        last_p = entry_p


_TNR = "Times New Roman"


def _force_tnr_rpr(rPr) -> None:
    """Выставляет в готовом w:rPr шрифт TNR 14pt (перетирая существующее)."""
    for tag in ("w:rFonts", "w:sz", "w:szCs"):
        for el in list(rPr.findall(qn(tag))):
            rPr.remove(el)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), _TNR)
    rFonts.set(qn("w:hAnsi"), _TNR)
    rFonts.set(qn("w:cs"), _TNR)
    rPr.insert(0, rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "28")
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "28")
    rPr.append(szCs)


def normalize_toc_fonts(doc: DocxDocument) -> None:
    """Выставляет TNR 14pt на все runs строк оглавления и на w:pPr/w:rPr.

    LO генерирует результат PAGEREF своим шрифтом, игнорируя rPr result-run.
    """
    for para in doc.paragraphs:
        p = para._p
        is_entry = any(
            "_VkrToc_" in (e.text or "")
            for e in p.iter(qn("w:instrText"))
        )
        if not is_entry:
            continue
        # w:pPr/w:rPr — базовый шрифт абзаца, наследуется перегенерированными LO runs.
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p.insert(0, pPr)
        para_rPr = pPr.find(qn("w:rPr"))
        if para_rPr is None:
            para_rPr = OxmlElement("w:rPr")
            pPr.append(para_rPr)
        _force_tnr_rpr(para_rPr)
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), _TOC_LINE_SPACING_TWIPS)
        spacing.set(qn("w:lineRule"), "auto")
        # Каждый run абзаца.
        for r in p.findall(qn("w:r")):
            rPr = r.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                r.insert(0, rPr)
            _force_tnr_rpr(rPr)


def _squash(text: str) -> str:
    """Нормализует строку для сопоставления: только буквы/цифры и пробелы."""
    s = re.sub(r"[^0-9a-zа-яё]+", " ", (text or "").lower())
    return " ".join(s.split())


def _toc_entry_result_runs(doc: DocxDocument) -> list[tuple[str, object]]:
    """[(текст_заголовка, w:t результата PAGEREF)] по порядку TOC.

    Текст — из первого w:t абзаца (до TAB). Result-run — между fldChar separate и end.
    """
    out: list[tuple[str, object]] = []
    for para in doc.paragraphs:
        p = para._p
        if not any(
            "_VkrToc_" in (e.text or "") for e in p.iter(qn("w:instrText"))
        ):
            continue
        first_t = None
        result_t = None
        state = "before"  # before → separate → done
        for r in p.findall(qn("w:r")):
            fld = r.find(qn("w:fldChar"))
            if fld is not None:
                ft = fld.get(qn("w:fldCharType"))
                if ft == "separate":
                    state = "separate"
                elif ft == "end":
                    state = "done"
                continue
            t = r.find(qn("w:t"))
            if t is None:
                continue
            if first_t is None and state == "before":
                first_t = t
            if state == "separate" and result_t is None:
                result_t = t
        if result_t is None:
            continue
        heading_text = (first_t.text or "") if first_t is not None else ""
        out.append((heading_text, result_t))
    return out


def _staticize_toc_fields(doc: DocxDocument, nums: list[int | None]) -> int:
    """Заменяет PAGEREF-поля TOC статическим текстом с верным номером.

    LO при --convert-to pdf пересчитывает PAGEREF своей ранней раскладкой.
    Убираем плумбинг поля (begin/instrText/separate/end), оставляем один run с нашим числом.
    Возвращает число оформленных записей.
    """
    count = 0
    i = 0
    for para in doc.paragraphs:
        p = para._p
        if not any(
            "_VkrToc_" in (e.text or "") for e in p.iter(qn("w:instrText"))
        ):
            continue
        if i >= len(nums):
            break
        num = nums[i]
        i += 1
        if num is None:
            continue  # не нашли страницу — оставляем поле как есть

        runs = p.findall(qn("w:r"))
        sep_i = end_i = None
        result_run = None
        state = "before"
        for ri, r in enumerate(runs):
            fld = r.find(qn("w:fldChar"))
            if fld is not None:
                ft = fld.get(qn("w:fldCharType"))
                if ft == "separate":
                    sep_i = ri
                    state = "separate"
                elif ft == "end":
                    end_i = ri
                continue
            if r.find(qn("w:instrText")) is not None:
                continue
            if state == "separate" and result_run is None and r.find(qn("w:t")) is not None:
                result_run = r
        if sep_i is None or end_i is None:
            continue  # это не наша PAGEREF-запись

        if result_run is None:
            result_run = OxmlElement("w:r")
            runs[end_i].addprevious(result_run)
        rPr = result_run.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            result_run.insert(0, rPr)
        _force_tnr_rpr(rPr)
        t = result_run.find(qn("w:t"))
        if t is None:
            t = OxmlElement("w:t")
            result_run.append(t)
        t.text = str(num)

        # vkrPageNum — атрибут переживает сохранение docx; _read_baked_toc_nums его читает.
        p.set("vkrPageNum", str(num))

        # Удаляем fldChar, instrText и все run'ы между separate и end кроме result_run.
        for ri, r in enumerate(runs):
            if r is result_run:
                continue
            fld = r.find(qn("w:fldChar"))
            is_instr = r.find(qn("w:instrText")) is not None
            between = sep_i < ri < end_i
            if fld is not None or is_instr or between:
                if r.getparent() is not None:
                    p.remove(r)
        count += 1
    return count


def _read_baked_toc_nums(doc: DocxDocument) -> list[int | None]:
    """Читает запечённые номера страниц из атрибута vkrPageNum или result-run.

    Вызывать ДО rebuild_toc_full — он удаляет старые записи.
    """
    # Попытка 1: instrText ещё не удалён → берём из result-run.
    entries = _toc_entry_result_runs(doc)
    if entries:
        nums: list[int | None] = []
        for _, result_t in entries:
            try:
                nums.append(int((result_t.text or "").strip()))
            except (ValueError, TypeError):
                nums.append(None)
        return nums

    # Попытка 2: TOC уже статицизирован → читаем из атрибута vkrPageNum.
    nums = []
    for para in doc.paragraphs:
        val = para._p.get("vkrPageNum")
        if val is not None:
            try:
                nums.append(int(val))
            except (ValueError, TypeError):
                nums.append(None)
    return nums


def _pdftotext(pdf_path: Path) -> str | None:
    """Извлекает текст PDF через poppler `pdftotext -layout`. None при сбое."""
    import subprocess

    try:
        proc = subprocess.run(
            ["pdftotext", "-layout", str(pdf_path), "-"],
            capture_output=True,
            timeout=120,
        )
        if proc.returncode != 0:
            return None
        return proc.stdout.decode("utf-8", errors="replace")
    except Exception:
        _logger.warning("pdftotext недоступен/упал", exc_info=True)
        return None


_FF = "\x0c"  # form feed — разделитель страниц в выводе pdftotext
_FOOTER_NUM_RE = re.compile(r"^\s*(\d{1,4})\s*$")
_TOC_LEADER_RE = re.compile(r"[.…]{3,}\s*\d{1,4}\s*$")


def _page_is_toc(page_text: str) -> bool:
    """Страница относится к региону оглавления (её пропускаем при поиске
    реального места заголовка в теле)."""
    sq = _squash(page_text)
    if sq.startswith("содержание") or sq.startswith("оглавление"):
        return True
    return len(_TOC_LEADER_RE.findall(page_text)) >= 4


def _footer_page_number(page_text: str) -> int | None:
    """Номер страницы, НАПЕЧАТАННЫЙ в колонтитуле (последняя строка-одиночное
    число на странице — колонтитул идёт ниже всего контента)."""
    found: int | None = None
    for ln in page_text.splitlines():
        m = _FOOTER_NUM_RE.match(ln)
        if m:
            found = int(m.group(1))
    return found


def _line_matches_heading(line_sq: str, esq: str) -> bool:
    if not line_sq or not esq:
        return False
    if line_sq == esq or line_sq.startswith(esq):
        return True
    # Перенос длинного заголовка на 2 строки PDF: первая строка — префикс.
    if len(line_sq) >= 12 and esq.startswith(line_sq):
        return True
    # Короткая строка-заголовок, целиком содержащая ожидаемый текст.
    if esq in line_sq and len(line_sq) <= len(esq) + 12:
        return True
    return False


def _heading_page_numbers_from_pdf(
    pdf_text: str, expected: list[str]
) -> list[int | None]:
    """Реальные номера страниц заголовков из отрендеренного PDF.

    LO резолвит PAGEREF до финальной раскладки — числа занижены.
    Находим физическую страницу каждого заголовка и берём номер из колонтитула.
    Возвращает список длиной len(expected).
    """
    result: list[int | None] = [None] * len(expected)
    pages = pdf_text.split(_FF)
    if len(pages) < 2:
        return result

    toc_flags = [_page_is_toc(p) for p in pages]
    # Предрасчёт нормализованных строк каждой страницы.
    page_lines_sq = [[_squash(ln) for ln in p.splitlines()] for p in pages]

    phys: list[int | None] = [None] * len(expected)
    last_page = 0
    for i, etext in enumerate(expected):
        esq = _squash(etext)
        if not esq:
            continue
        for pgi in range(last_page, len(pages)):
            if toc_flags[pgi]:
                continue
            hit = any(_line_matches_heading(ls, esq) for ls in page_lines_sq[pgi])
            if not hit and len(esq) > 20:
                # Префиксный фолбэк (как в PdfPane.findMatches).
                for cut in (40, 30, 20):
                    if len(esq) <= cut:
                        continue
                    pref = esq[:cut]
                    if any(ls.startswith(pref) for ls in page_lines_sq[pgi]):
                        hit = True
                        break
            if hit:
                phys[i] = pgi + 1  # 1-based физическая страница
                last_page = pgi
                break

    # Смещение «колонтитул − физическая страница» — самое частое значение (мода).
    offsets: dict[int, int] = {}
    for i, pg in enumerate(phys):
        if pg is None:
            continue
        fn = _footer_page_number(pages[pg - 1])
        if fn is not None:
            off = fn - pg
            offsets[off] = offsets.get(off, 0) + 1
    offset = max(offsets, key=offsets.get) if offsets else 0

    prev = 0
    for i, pg in enumerate(phys):
        if pg is None:
            continue
        val = pg + offset
        if val < prev:        # номера TOC не убывают
            val = prev
        if val < 1:
            val = 1
        result[i] = val
        prev = val
    return result


def _finalize_and_render(
    docx_path: Path,
    table_captions=None,
    skip_table_split: bool = False,
    known_page_nums: list[int | None] | None = None,
) -> tuple[bool, Path | None]:
    """Финализирует TOC: нормализует шрифты/отступы, определяет номера страниц, запекает.

    skip_table_split=True — пропустить резку таблиц (при ручных правках уже нарезаны).
    known_page_nums — если передан, R1 (измерительный рендер) пропускается.
    Возвращает (ok, pdf_path|None). pdf_path НЕ удаляется — очистка на вызывающей стороне.
    """
    from docx import Document as WordDocument

    from services.core.vkr_core.engine.preview import convert_docx_to_pdf

    doc = WordDocument(str(docx_path))
    entries_initial = _toc_entry_result_runs(doc)
    if not entries_initial:
        return (False, None)

    # 1. Нормализация ДО рендера — чтобы единственный PDF был с верным TNR.
    normalize_toc_indents(doc)
    normalize_toc_fonts(doc)
    doc.save(str(docx_path))

    import shutil as _shutil

    # 2. Номера страниц: если known_page_nums совпадает — R1 пропускаем.
    #    R1 нужен только чтобы найти реальные страницы заголовков (раскладка тела у LO корректна).
    pdf_text: str | None = None
    nums: list[int | None] | None = None

    if (
        known_page_nums is not None
        and skip_table_split  # резка таблиц не изменит число записей
        and len(known_page_nums) == len(entries_initial)
    ):
        nums = known_page_nums
    else:
        try:
            r1 = convert_docx_to_pdf(docx_path, fast=True)  # измерительный
            pdf_text = _pdftotext(r1)
        except Exception:
            _logger.warning("Не удалось отрендерить PDF для номеров TOC", exc_info=True)
            return (False, None)
        finally:
            try:
                _shutil.rmtree(r1.parent, ignore_errors=True)
            except Exception:
                pass

    # 2b. Резка таблиц по UNO-измерению. Меняет пагинацию → нужен свежий рендер для TOC.
    #     skip_table_split=True — пропускаем: таблицы уже нарезаны, UNO занимает десятки секунд.
    if not skip_table_split:
        try:
            from services.core.vkr_core.engine.captions import renumber_tables
            from services.core.vkr_core.engine.table_continuation import (
                _remerge_continuations,
                split_long_tables_to_pages,
                split_overflowing_fragments,
            )
            from services.core.vkr_core.engine.tables import apply_continuation_table_styles
            from services.core.vkr_core.engine.uno_layout import (
                measure_table_row_pages,
            )

            any_split = False
            # UNO-измерение + итеративная резка. Сходится, когда ни одна таблица больше не переносится.
            probe = measure_table_row_pages(docx_path)
            if probe:
                # Склеиваем фрагменты (повторный finalize), чтобы резать от цельных таблиц.
                mdoc = WordDocument(str(docx_path))
                _remerge_continuations(mdoc)
                mdoc.save(str(docx_path))
                for _ in range(_MAX_TABLE_SPLIT_ITERS):
                    rp = measure_table_row_pages(docx_path)
                    fdoc = WordDocument(str(docx_path))
                    if not split_overflowing_fragments(fdoc, rp):
                        break
                    apply_continuation_table_styles(fdoc, 0)
                    fdoc.save(str(docx_path))
                    any_split = True
            else:
                sdoc = WordDocument(str(docx_path))
                caps, _ = renumber_tables(sdoc.paragraphs, set(), apply=False)
                if caps and split_long_tables_to_pages(sdoc, {}, caps):
                    apply_continuation_table_styles(sdoc, 0)
                    sdoc.save(str(docx_path))
                    any_split = True
                else:
                    _logger.info(
                        "table_continuation: UNO-измерение недоступно, "
                        "fallback не нашёл таблиц для резки"
                    )

            if any_split:
                try:
                    r1b = convert_docx_to_pdf(docx_path, fast=True)  # измерит.
                    pdf_text = _pdftotext(r1b)
                except Exception:
                    _logger.warning(
                        "Не удалось перерендерить PDF после резки таблиц",
                        exc_info=True,
                    )
                finally:
                    try:
                        _shutil.rmtree(r1b.parent, ignore_errors=True)
                    except Exception:
                        pass
        except Exception:
            _logger.warning("Резка таблиц не удалась", exc_info=True)

    # 3. Реальные номера страниц + статицизация полей TOC (иначе LO снова подставит заниженные).
    doc2 = WordDocument(str(docx_path))
    entries = _toc_entry_result_runs(doc2)
    if nums is None:
        # pdf_text получен из R1 (не было known_page_nums или count не совпал)
        nums = _heading_page_numbers_from_pdf(pdf_text or "", [t for t, _ in entries])
    elif len(nums) != len(entries):
        # Количество записей изменилось после резки таблиц — пересчитать нельзя.
        nums = _heading_page_numbers_from_pdf(pdf_text or "", [t for t, _ in entries])
    baked = _staticize_toc_fields(doc2, nums)
    if not baked:
        return (False, None)
    doc2.save(str(docx_path))

    # 4. R2 — финальный рендер со статическим оглавлением. Это и есть превью.
    try:
        r2 = convert_docx_to_pdf(docx_path)
    except Exception:
        _logger.warning("Не удалось отрендерить финальный PDF", exc_info=True)
        return (True, None)
    return (True, r2)


def finalize_toc(docx_path: Path, table_captions=None) -> bool:
    """Финализирует оглавление (для pipeline): один рендер, PDF удаляется."""
    ok, pdf_path = _finalize_and_render(docx_path, table_captions)
    if pdf_path is not None:
        import shutil as _shutil

        _shutil.rmtree(pdf_path.parent, ignore_errors=True)
    return ok


def finalize_toc_with_preview(
    docx_path: Path,
    table_captions=None,
    skip_table_split: bool = False,
    known_page_nums: list[int | None] | None = None,
) -> tuple[bool, Path | None]:
    """Как finalize_toc, но ВОЗВРАЩАЕТ отрендеренный PDF для переиспользования
    как превью (без второго рендера). Очистку pdf_path.parent делает вызвавший.
    Если TOC нет — (False, None), вызывающий рендерит превью сам.

    known_page_nums — если передан, R1 (измерительный рендер) пропускается.
    """
    return _finalize_and_render(
        docx_path, table_captions,
        skip_table_split=skip_table_split,
        known_page_nums=known_page_nums,
    )


def apply_table_continuation_only(docx_path: Path, table_captions=None) -> bool:
    """Только резка таблиц по страницам — без TOC.

    Вызывается для документов без оглавления, где finalize_toc не запускается.
    Возвращает True, если хотя бы одна таблица была разрезана.
    """
    from docx import Document as WordDocument

    _logger.info("table_continuation: standalone run for non-TOC document")
    any_split = False
    try:
        from services.core.vkr_core.engine.captions import renumber_tables
        from services.core.vkr_core.engine.table_continuation import (
            _remerge_continuations,
            split_long_tables_to_pages,
            split_overflowing_fragments,
        )
        from services.core.vkr_core.engine.tables import apply_continuation_table_styles
        from services.core.vkr_core.engine.uno_layout import measure_table_row_pages

        probe = measure_table_row_pages(docx_path)
        if probe:
            mdoc = WordDocument(str(docx_path))
            _remerge_continuations(mdoc)
            mdoc.save(str(docx_path))
            for _ in range(_MAX_TABLE_SPLIT_ITERS):
                rp = measure_table_row_pages(docx_path)
                fdoc = WordDocument(str(docx_path))
                if not split_overflowing_fragments(fdoc, rp):
                    break
                apply_continuation_table_styles(fdoc, 0)
                fdoc.save(str(docx_path))
                any_split = True
        else:
            sdoc = WordDocument(str(docx_path))
            caps = table_captions
            if caps is None:
                caps, _ = renumber_tables(sdoc.paragraphs, set(), apply=False)
            if caps and split_long_tables_to_pages(sdoc, {}, caps):
                apply_continuation_table_styles(sdoc, 0)
                sdoc.save(str(docx_path))
                any_split = True
            else:
                _logger.info(
                    "table_continuation: UNO недоступно, heuristic не нашёл таблиц для резки"
                )
    except Exception:
        _logger.warning("table_continuation standalone: ошибка при резке таблиц", exc_info=True)
    return any_split
