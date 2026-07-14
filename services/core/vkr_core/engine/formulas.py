"""Обнаружение формул (ТЗ §7.2).

Детерминированные правила:
  F1. Абзац считается формулой, если содержит знак "=" и хотя бы один
      символ из подобранного набора (греческие буквы, операторы), а сам
      абзац короткий (< 80 символов) и либо выровнен по центру, либо
      заканчивается номером формулы в скобках.
  F2. За каждой формулой (в следующем непустом абзаце) должно идти
      пояснение, начинающееся со слова "где".
  F3. Формулы должны быть пронумерованы сквозной нумерацией по документу
      и упоминаться в тексте как (n) до их появления.
"""

from __future__ import annotations

import re
from dataclasses import dataclass

from docx.enum.text import WD_ALIGN_PARAGRAPH

from services.core.vkr_core.engine.stats import DocStats, ParagraphStats
from services.core.vkr_core.engine.violations import (
    SEVERITY_WARNING,
    PipelineViolation,
)
from services.core.vkr_core.models.enums import ViolationStatus

FORMULA_SYMBOLS = re.compile(r"[=±×÷∑∫√αβγδεζηθικλμνξοπρστυφχψωΑΒΓΔΕΗΘΛΠΣΦΨΩ]")
FORMULA_NUMBER_RE = re.compile(r"\((\d+(?:\.\d+)?)\)\s*$")
WHERE_RE = re.compile(r"^\s*где\b", re.IGNORECASE)


@dataclass
class FormulaOccurrence:
    paragraph_index: int
    number: str | None
    text: str
    has_tab: bool = False


def detect_formulas(stats: DocStats, skip_indexes: set[int]) -> list[FormulaOccurrence]:
    found: list[FormulaOccurrence] = []
    for ps in stats.paragraphs:
        if ps.index in skip_indexes or ps.is_empty or ps.in_table:
            continue
        text = ps.stripped
        if "=" not in text:
            continue
        if not FORMULA_SYMBOLS.search(text):
            continue
        if ps.length > 120:
            continue
        # Должна быть либо по центру, либо явно помечена номером формулы.
        number_match = FORMULA_NUMBER_RE.search(text)
        is_centred = ps.alignment == WD_ALIGN_PARAGRAPH.CENTER
        if not is_centred and not number_match:
            continue
        number = number_match.group(1) if number_match else None
        found.append(FormulaOccurrence(
            paragraph_index=ps.index,
            number=number,
            text=text,
            has_tab=ps.has_tab,
        ))
    return found


def _next_non_empty(stats: DocStats, from_index: int) -> ParagraphStats | None:
    for i in range(from_index + 1, len(stats.paragraphs)):
        ps = stats.paragraphs[i]
        if not ps.is_empty:
            return ps
    return None


def validate_formulas(
    stats: DocStats,
    formulas: list[FormulaOccurrence],
) -> list[PipelineViolation]:
    violations: list[PipelineViolation] = []

    # F4 — номер формулы должен быть выровнен по правому краю страницы
    # (отделён от формулы правым табом). Если has_tab=False, номер «прилип» к тексту.
    for formula in formulas:
        if formula.number is not None and not formula.has_tab:
            violations.append(
                PipelineViolation(
                    type="formula_number_not_right",
                    rule_reference="п. 6.14",
                    description=(
                        f"Номер формулы ({formula.number}) должен быть выровнен "
                        "по правому краю страницы через табуляцию."
                    ),
                    status=ViolationStatus.auto_fixed,
                    severity=SEVERITY_WARNING,
                    paragraph_index=formula.paragraph_index,
                    original_text=formula.text,
                )
            )

    # F2 — проверка наличия пояснения.
    for formula in formulas:
        successor = _next_non_empty(stats, formula.paragraph_index)
        if successor is None or not WHERE_RE.match(successor.stripped):
            violations.append(
                PipelineViolation(
                    type="formula_without_explanation",
                    rule_reference="п. 6.14",
                    description=(
                        "Формула не сопровождается пояснением, начинающимся со слова «где»."
                    ),
                    status=ViolationStatus.manual_required,
                    severity=SEVERITY_WARNING,
                    paragraph_index=formula.paragraph_index,
                    original_text=formula.text,
                )
            )

    # F3 — непрерывность нумерации.
    numbered = [f for f in formulas if f.number and f.number.isdigit()]
    numbered.sort(key=lambda f: f.paragraph_index)
    expected = 1
    for f in numbered:
        actual = int(f.number)
        if actual != expected:
            violations.append(
                PipelineViolation(
                    type="formula_numbering_gap",
                    rule_reference="п. 6.14",
                    description=(
                        f"Нарушена сквозная нумерация формул: ожидалась ({expected}), "
                        f"найдена ({actual})."
                    ),
                    status=ViolationStatus.manual_required,
                    severity=SEVERITY_WARNING,
                    paragraph_index=f.paragraph_index,
                    original_text=f.text,
                )
            )
            expected = actual + 1
        else:
            expected += 1

    return violations
