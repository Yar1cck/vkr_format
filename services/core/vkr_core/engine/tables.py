"""Стили таблиц основной части документа (ГОСТ 7.32-2017).

Применяется к "обычным" таблицам — тем, перед которыми идёт абзац-подпись
вида "Таблица N – ...", а также к их фрагментам после "Продолжение таблицы N".
Декоративные таблицы (table-layout титульного листа, бланки задания и т.п.) не
трогаем — у них нет подписи.

Правила:
  T1. Текст ячеек — Times New Roman 12 pt (тело документа — 14 pt).
  T2. Выравнивание (ГОСТ 7.32-2017, п. 6.6):
        шапка (первая строка) — по центру горизонтально и вертикально;
        данные, числовые      — по правому краю горизонтально, по центру по вертикали;
        данные, текстовые     — по левому краю горизонтально, по центру по вертикали.
  T3. Первая строка помечается `<w:tblHeader/>` — Word и LibreOffice сами
      повторят шапку на новой странице при переносе таблицы.
  T4. На все строки ставится `<w:cantSplit/>` — отдельная строка не будет
      разрываться между страницами (по ГОСТ строка таблицы целая).
  T5. Если в таблице больше TABLE_LONG_THRESHOLD строк, добавляется
      info-нарушение: формат flow не позволяет точно предсказать точку
      разрыва (см. CLAUDE.md), поэтому честно предупреждаем пользователя,
      что Word повторит шапку — без обманной надписи "Продолжение таблицы N",
      которая в реальном Word'е может уехать на 1–2 строки.
  T6. Если ≥ MASS_BOLD_THRESHOLD доли непустых runs в таблице жирные —
      считаем, что студент выделил всю таблицу жирным. Убираем жирность
      из строк данных (шапка не трогается — там bold уместен).
"""

from __future__ import annotations

import re

from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table

from services.core.vkr_core.engine.captions import Caption, _set_keep_flags
from services.core.vkr_core.engine.violations import (
    SEVERITY_INFO,
    SEVERITY_WARNING,
    PipelineViolation,
)
from services.core.vkr_core.models.enums import ViolationStatus

TABLE_FONT_NAME = "Times New Roman"
TABLE_FONT_SIZE_PT = 12
TABLE_LONG_THRESHOLD = 15
MASS_BOLD_THRESHOLD = 0.6
_NUMERIC_THRESHOLD = 0.7

# Значение числовой ячейки: опциональный знак, цифры, разделители, единицы.
_NUMERIC_CELL_RE = re.compile(r"^[+\-±]?\s*\d[\d\s.,]*[%°]?\s*$")
_CONTINUATION_LABEL_RE = re.compile(
    r"^\s*(?:продолжение|окончание)\s+таблицы\s*\d+(?:\.\d+)?\b",
    re.IGNORECASE,
)


def _is_numeric_cell(text: str) -> bool:
    return bool(_NUMERIC_CELL_RE.match(text.strip()))


def _column_is_numeric(table, col_idx: int) -> bool:
    """True если ≥70 % непустых ячеек данных (не шапки) содержат числа."""
    total = numeric = 0
    for row in table.rows[1:]:
        cells = row.cells
        if col_idx >= len(cells):
            continue
        t = cells[col_idx].text.strip()
        if not t:
            continue
        total += 1
        if _is_numeric_cell(t):
            numeric += 1
    return total > 0 and (numeric / total) >= _NUMERIC_THRESHOLD


def _ensure_trPr(tr_element):
    trPr = tr_element.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr_element.insert(0, trPr)
    return trPr


def _set_table_autofit(table) -> None:
    """Переключает таблицу в режим autofit to window.

    С layout=fixed узкие ячейки (tcW в dxa) переносят контент побуквенно:
    слово не помещается в 0.8 cm — каждая буква на отдельной строке.
    В режиме autofit Word растягивает столбцы под контент не выходя за
    ширину текстового блока, что исключает однобуквенные переносы.

    tblW=pct:5000 = 100 % текстового блока; существующие tcW (dxa/pct) в
    режиме autofit трактуются как относительные веса — пропорции сохраняются.
    """
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # tblW → 100 % ширины текстового блока
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), "5000")

    # tblLayout → autofit (убирает жёсткий fixed-режим)
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "autofit")

    # tblJc=center: корректное значение по ГОСТ для таблиц. В сочетании с
    # pct:5000 Word отображает её корректно. LibreOffice игнорирует pct при
    # jc=center и сжимает таблицу до контентной ширины — это дефект LO-рендера,
    # не влияет на открытие в Word. Убираем только явный jc=left/right,
    # который мог прийти из шаблона и сдвигает таблицу от поля.
    tblJc = tblPr.find(qn("w:jc"))
    if tblJc is not None:
        val = tblJc.get(qn("w:val"), "")
        if val != "center":
            tblPr.remove(tblJc)

    # tblInd: отступ таблицы от левого поля. Обнуляем, чтобы таблица
    # прилегала к левому краю текстового блока (как и требует ГОСТ).
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is not None:
        tblPr.remove(tblInd)


def _set_tr_flag(tr_element, flag_name: str) -> None:
    trPr = _ensure_trPr(tr_element)
    if trPr.find(qn(f"w:{flag_name}")) is None:
        trPr.append(OxmlElement(f"w:{flag_name}"))


def _set_table_borders(table) -> None:
    """Единая сетка 0.5 pt для всех фрагментов одной таблицы."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "auto")


def _clear_cell_shading(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for shd in list(tcPr.findall(qn("w:shd"))):
        tcPr.remove(shd)


def _set_cell_borders(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "auto")


def _set_run_table_font(run) -> None:
    run.font.name = TABLE_FONT_NAME
    run.font.size = Pt(TABLE_FONT_SIZE_PT)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), TABLE_FONT_NAME)


def _count_bold_runs(table) -> tuple[int, int]:
    """Возвращает (всего непустых runs, из них явно жирных)."""
    total = bold = 0
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text.strip():
                        total += 1
                        if run.bold:
                            bold += 1
    return total, bold


def _normalize_cell_spacing(paragraph) -> None:
    """Снимает space_before/after и autospacing-флаги через XML.

    python-docx сеттеры иногда не перебивают w:beforeAutospacing /
    w:afterAutospacing из стиля Normal — поэтому патчим напрямую.
    """
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.attrib.pop(qn("w:beforeAutospacing"), None)
    spacing.attrib.pop(qn("w:afterAutospacing"), None)


def _apply_cell_text_style(
    cell,
    alignment: WD_ALIGN_PARAGRAPH,
    strip_bold: bool = False,
    force_bold: bool | None = None,
) -> None:
    _clear_cell_shading(cell)
    _set_cell_borders(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        # Явный 0 (не None) — иначе наследуется отступ абзаца из стиля "Normal".
        pf.first_line_indent = 0
        pf.left_indent = 0
        # Одинарный интервал: ячейки наследуют 1.5/2.0 из Normal — текст
        # расплывается и строки ячейки выглядят как отдельные абзацы.
        pf.line_spacing = 1.0
        _normalize_cell_spacing(paragraph)
        # Убираем принудительный разрыв страницы внутри ячейки: студент мог
        # поставить его через Word UI (Вставка → Разрыв страницы), что
        # переносит строку таблицы на новую страницу, оставляя предыдущую
        # почти пустой.
        p_xml = paragraph._p
        pPr_xml = p_xml.get_or_add_pPr()
        # Запрещаем автоперенос слов в ячейках — иначе LO может оставить
        # одну букву слова в конце строки («б-» на одной строке, «ank» на следующей).
        if pPr_xml.find(qn("w:suppressAutoHyphens")) is None:
            sah = OxmlElement("w:suppressAutoHyphens")
            sah.set(qn("w:val"), "true")
            pPr_xml.append(sah)
        pbb = pPr_xml.find(qn("w:pageBreakBefore"))
        if pbb is not None:
            pPr_xml.remove(pbb)
        # Ручные разрывы страниц могут также лежать как w:br[@w:type="page"]
        # внутри w:r — удаляем и их.
        for r_elem in list(p_xml.findall(qn("w:r"))):
            for br in r_elem.findall(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    r_elem.remove(br)
            if not list(r_elem) or (
                len(list(r_elem)) == 1
                and r_elem.find(qn("w:rPr")) is not None
            ):
                p_xml.remove(r_elem)
        for run in paragraph.runs:
            _set_run_table_font(run)
            if force_bold is not None:
                run.bold = force_bold
            if strip_bold and run.bold:
                run.bold = False


def _apply_gost_table_style(table) -> bool:
    """Применяет полный стиль проекта к одной таблице или её фрагменту.

    Возвращает True, если в таблице была массовая жирность и она снята со
    строк данных. Нарезанные фрагменты используют этот же helper, чтобы шапка,
    сетка, шрифты и отступы не расходились между частями одной таблицы.
    """
    rows = list(table.rows)
    if not rows:
        return False

    _set_table_autofit(table)
    _set_table_borders(table)
    total_runs, bold_runs = _count_bold_runs(table)
    mass_bold = total_runs > 0 and (bold_runs / total_runs) >= MASS_BOLD_THRESHOLD

    for cell in rows[0].cells:
        _apply_cell_text_style(
            cell,
            WD_ALIGN_PARAGRAPH.CENTER,
            force_bold=True,
        )
    _set_tr_flag(rows[0]._tr, "tblHeader")

    col_count = max((len(row.cells) for row in rows[1:]), default=0)
    col_alignments = [
        WD_ALIGN_PARAGRAPH.RIGHT if _column_is_numeric(table, ci) else WD_ALIGN_PARAGRAPH.LEFT
        for ci in range(col_count)
    ]
    for row in rows[1:]:
        for ci, cell in enumerate(row.cells):
            align = col_alignments[ci] if ci < len(col_alignments) else WD_ALIGN_PARAGRAPH.LEFT
            _apply_cell_text_style(cell, align, strip_bold=mass_bold)

    for row in rows:
        _set_tr_flag(row._tr, "cantSplit")

    return mass_bold


def _find_captioned_tables(doc, table_captions: list[Caption]) -> list[tuple[Table, Caption]]:
    """Сопоставляет объекты Table с подписями. Подпись засчитывается, если
    она — последний непустой абзац перед таблицей (пустые абзацы между
    подписью и таблицей допускаются)."""
    caption_by_index = {cap.paragraph_index: cap for cap in table_captions}
    paragraphs = doc.paragraphs
    body = doc.element.body

    para_idx = -1
    table_idx = -1
    pending: Caption | None = None
    result: list[tuple[Table, Caption]] = []

    for elem in body.iterchildren():
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag == "p":
            para_idx += 1
            if para_idx in caption_by_index:
                pending = caption_by_index[para_idx]
            elif para_idx < len(paragraphs) and paragraphs[para_idx].text.strip():
                # любой непустой не-caption абзац разрывает связь подпись → таблица
                pending = None
        elif tag == "tbl":
            table_idx += 1
            if pending is not None and table_idx < len(doc.tables):
                result.append((doc.tables[table_idx], pending))
            pending = None

    return result


def _previous_nonempty_paragraph_text(tbl_elem) -> str | None:
    el = tbl_elem.getprevious()
    while el is not None:
        tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if tag == "p":
            text = "".join(el.itertext()).strip()
            if text:
                return text
        elif tag == "tbl":
            return None
        el = el.getprevious()
    return None


def _looks_decorative(table) -> bool:
    """Бланк-форма (титул/задание): ячейки пустые или только подчёркивания/
    знаки пунктуации. У таких меньше 2 «содержательных» ячеек."""
    meaningful = 0
    for row in table.rows:
        for cell in row.cells:
            t = cell.text.strip()
            if t and re.sub(r"[\s_().:;,–—\-]", "", t):
                meaningful += 1
                if meaningful >= 2:
                    return False
    return True


def ensure_table_pagebreak_safety(doc, title_end_index: int) -> None:
    """Гарантирует корректный перенос ЛЮБОЙ содержательной таблицы тела.

    apply_table_styles ставит tblHeader/cantSplit только подписанным
    «Таблица N» таблицам. Но в теле бывают таблицы без подписи (например
    глоссарий «ТЕРМИНЫ И ОПРЕДЕЛЕНИЯ» на 25 строк) — без повтора шапки и
    без cantSplit они переносятся через страницу неверно (шапка не
    дублируется, строка рвётся посередине). Здесь добавляем эти флаги всем
    таблицам тела (после титула/задания, не бланкам). Меняем ТОЛЬКО
    поведение разрыва — шрифты/выравнивание не трогаем (идемпотентно к
    apply_table_styles: _set_tr_flag не дублирует уже стоящий флаг)."""
    body = doc.element.body
    pidx = -1
    tidx = -1
    for elem in body.iterchildren():
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag == "p":
            pidx += 1
        elif tag == "tbl":
            tidx += 1
            if pidx < title_end_index:
                continue  # титул/задание — не трогаем
            if tidx >= len(doc.tables):
                continue
            table = doc.tables[tidx]
            rows = list(table.rows)
            if len(rows) < 2 or _looks_decorative(table):
                continue
            _set_table_autofit(table)
            _set_tr_flag(rows[0]._tr, "tblHeader")
            for row in rows:
                _set_tr_flag(row._tr, "cantSplit")


def apply_continuation_table_styles(doc, title_end_index: int = 0) -> int:
    """Применяет тот же табличный стиль к фрагментам после продолжений.

    `apply_table_styles()` находит таблицы по исходной подписи "Таблица N".
    Если документ уже был когда-то нарезан или дорезается после UNO-рендера,
    вторые и последующие фрагменты стоят после "Продолжение таблицы N" и не
    попадают в список подписанных таблиц. Их нужно стилизовать отдельным
    проходом, иначе верхняя и нижняя части одной таблицы выглядят по-разному.
    """
    body = doc.element.body
    pidx = -1
    tidx = -1
    styled = 0
    for elem in body.iterchildren():
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag == "p":
            pidx += 1
        elif tag == "tbl":
            tidx += 1
            if pidx < title_end_index or tidx >= len(doc.tables):
                continue
            text = _previous_nonempty_paragraph_text(elem)
            if not text or not _CONTINUATION_LABEL_RE.match(text):
                continue
            table = doc.tables[tidx]
            if _looks_decorative(table):
                continue
            _apply_gost_table_style(table)
            styled += 1
    return styled


def apply_table_styles(doc, table_captions: list[Caption]) -> list[PipelineViolation]:
    """Применяет ГОСТ-стили ко всем подписанным таблицам. Возвращает список
    info-нарушений о таблицах, которые могут переноситься через границу
    страницы (Word повторит шапку автоматически)."""
    violations: list[PipelineViolation] = []
    captioned = _find_captioned_tables(doc, table_captions)

    for table, caption in captioned:
        rows = list(table.rows)
        if not rows:
            continue

        mass_bold = _apply_gost_table_style(table)

        # Подпись таблицы не должна отрываться от самой таблицы: keepNext на
        # абзац подписи и на пустые абзацы-разделители между ней и таблицей.
        paragraphs = doc.paragraphs
        if caption.paragraph_index < len(paragraphs):
            cap_p = paragraphs[caption.paragraph_index]
            _set_keep_flags(cap_p._p, keep_next=True, keep_lines=True)
            sib = cap_p._p.getnext()
            while sib is not None and sib.tag == qn("w:p"):
                if "".join(sib.itertext()).strip():
                    break  # непустой абзац — связь подпись→таблица разорвана
                _set_keep_flags(sib, keep_next=True, keep_lines=False)
                sib = sib.getnext()

        if mass_bold:
            violations.append(
                PipelineViolation(
                    type="table_mass_bold_removed",
                    rule_reference="п. 6.13",
                    description=(
                        f"Таблица {caption.new_number}: большинство ячеек содержали жирный текст. "
                        f"По ГОСТ 7.32-2017 жирное начертание в строках данных не предусмотрено — "
                        f"жирность убрана. Шапка таблицы оставлена без изменений."
                    ),
                    status=ViolationStatus.auto_fixed,
                    severity=SEVERITY_WARNING,
                    paragraph_index=caption.paragraph_index,
                    section_title=f"Таблица {caption.new_number} – {caption.title}",
                )
            )

        if len(rows) > TABLE_LONG_THRESHOLD:
            violations.append(
                PipelineViolation(
                    type="long_table_split_warning",
                    rule_reference="п. 6.13",
                    description=(
                        f"Таблица {caption.new_number} содержит {len(rows)} строк "
                        f"и может занимать несколько страниц. Шапка таблицы помечена "
                        f"как повторяющийся заголовок: Word и LibreOffice автоматически "
                        f"продублируют её на каждой следующей странице. Перед сдачей "
                        f"откройте файл и визуально проверьте, что разрыв таблицы "
                        f"оформлен корректно."
                    ),
                    status=ViolationStatus.auto_fixed,
                    severity=SEVERITY_INFO,
                    paragraph_index=caption.paragraph_index,
                    section_title=f"Таблица {caption.new_number} – {caption.title}",
                )
            )

    return violations
