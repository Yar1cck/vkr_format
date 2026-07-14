"""ГОСТ «Продолжение таблицы N» на страницах-продолжениях.

Точку разрыва НЕ угадываем: документ уже рендерится в PDF, поэтому по
реальному рендеру известно, какая строка таблицы на какой физической
странице. Зная это, физически разрезаем `<w:tbl>` на фрагменты и сами
ставим жёсткий разрыв страницы + надпись «Продолжение таблицы N» + клон
строки-шапки. Место разрыва не предсказывается, а диктуется нами.

Детерминированно, без итерационного цикла рендеров: для каждой
страницы-продолжения резервируем одну строку данных (надпись по высоте
≤ строки данных ⇒ переполнения не будет).

Применяется ТОЛЬКО к нумерованным подписанным таблицам («Таблица N …»).
Бесподписные (глоссарий ТиО) не трогаем — номер выдумывать нельзя.
"""

from __future__ import annotations

import copy
import logging
import re

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from services.core.vkr_core.engine.captions import _set_keep_flags
from services.core.vkr_core.engine.tables import (
    _find_captioned_tables,
    _looks_decorative,
    apply_continuation_table_styles,
)
from services.core.vkr_core.engine.toc import _squash

_logger = logging.getLogger(__name__)

_LABEL_PREFIX = "Продолжение таблицы"
# squash("Продолжение таблицы") — для идемпотентного детекта уже вставленных
# надписей при повторном finalize.
_LABEL_SQUASH = _squash(_LABEL_PREFIX)

# ── идемпотентность: склейка ранее созданных фрагментов ──────────────────────

def _is_label_para(el) -> bool:
    return (
        el is not None
        and el.tag == qn("w:p")
        and _squash("".join(el.itertext())).startswith(_LABEL_SQUASH)
    )


def _is_empty_paragraph(el) -> bool:
    """True, если элемент — <w:p> без видимого текста (только пробелы/пусто)."""
    return (
        el is not None
        and el.tag == qn("w:p")
        and not "".join(el.itertext()).strip()
    )


def _drop_empty_paragraphs_between(left, right) -> None:
    """Удаляет все пустые <w:p> между left и right (не включая их самих).

    Используется в двух местах:
      * в _split_table — между фрагментом-предыдущей-таблицей (anchor) и
        новой таблицей-фрагментом могла остаться пустая строка от
        ensure_separator_after_tables_and_figures (она вставляется после
        каждой таблицы ещё в pipeline). Без чистки эта пустая строка
        попадает на страницу-продолжение ПЕРЕД «Продолжение таблицы N»
        и сдвигает её на одну строку вниз.
      * в _remerge_continuations — между двумя склеиваемыми таблицами,
        чтобы при повторных правках мусор не накапливался.
    """
    sibling = left.getnext()
    while sibling is not None and sibling is not right:
        next_sibling = sibling.getnext()
        if _is_empty_paragraph(sibling):
            parent = sibling.getparent()
            if parent is not None:
                parent.remove(sibling)
        sibling = next_sibling


def _remerge_continuations(doc) -> None:
    """Склеивает ранее разрезанные фрагменты обратно в одну таблицу.

    Нужно при повторном finalize (ручная правка перечитывает уже
    разрезанный processed.docx) — иначе фрагменты бы множились. Цепочка
    после резки: tbl1 → p_label (несёт pageBreakBefore) → tbl2 (клон шапки
    + строки). Узнаём блок по абзацу-надписи между двумя <w:tbl>.
    """
    body = doc.element.body
    changed = True
    while changed:
        changed = False
        for label in list(body):
            if not _is_label_para(label):
                continue
            tbl1 = label.getprevious()
            tbl2 = label.getnext()
            if (
                tbl1 is None
                or tbl1.tag != qn("w:tbl")
                or tbl2 is None
                or tbl2.tag != qn("w:tbl")
            ):
                continue
            trs2 = tbl2.findall(qn("w:tr"))
            # первый <w:tr> tbl2 — клон шапки, его отбрасываем
            for tr in trs2[1:]:
                tbl2.remove(tr)
                tbl1.append(tr)
            body.remove(tbl2)
            body.remove(label)
            # Между tbl1 и его следующим (теперь — содержимое после старого
            # tbl2) могла остаться пустая <w:p> ИЛИ её там не было: на всякий
            # случай чистим хвост, чтобы повторные finalize не накапливали
            # мусор.
            next_sibling = tbl1.getnext()
            if next_sibling is not None:
                _drop_empty_paragraphs_between(tbl1, next_sibling)
            changed = True
            break


# ── измерение пагинации: строки таблицы → страницы (через UNO) ───────────────


def _page_groups(row_pages: list[int | None]) -> list[int] | None:
    """Список количеств строк данных по физическим страницам, по порядку.

    None у строки → берём страницу предыдущей известной строки. Если
    подряд None в начале или совсем ничего не нашли — возвращаем None
    (ненадёжно, таблицу не режем)."""
    filled: list[int] = []
    last: int | None = None
    for p in row_pages:
        if p is None:
            if last is None:
                return None
            filled.append(last)
        else:
            if last is not None and p < last:
                p = last  # курсор не назад
            filled.append(p)
            last = p
    if last is None:
        return None
    groups: list[int] = []
    cur = filled[0]
    cnt = 0
    for p in filled:
        if p == cur:
            cnt += 1
        else:
            groups.append(cnt)
            cur = p
            cnt = 1
    groups.append(cnt)
    return groups


# ── OXML: клон шапки, разрез таблицы, вставка надписи ────────────────────────

def _clone_header(header_tr):
    tr = copy.deepcopy(header_tr)
    trPr = tr.find(qn("w:trPr"))
    if trPr is not None:
        for h in list(trPr.findall(qn("w:tblHeader"))):
            trPr.remove(h)  # повтор шапки делаем вручную, авто-повтор не нужен
    return tr


def _make_label_para(number: str):
    p = OxmlElement("w:p")
    _set_label_ppr(p)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), "Times New Roman")
    rf.set(qn("w:hAnsi"), "Times New Roman")
    rf.set(qn("w:cs"), "Times New Roman")
    rf.set(qn("w:eastAsia"), "Times New Roman")
    rPr.append(rf)
    for tag in ("w:sz", "w:szCs"):
        s = OxmlElement(tag)
        s.set(qn("w:val"), "28")  # 14pt
        rPr.append(s)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = f"{_LABEL_PREFIX} {number}"
    r.append(t)
    p.append(r)
    return p


def _set_label_ppr(p_elem) -> None:
    pPr = p_elem.get_or_add_pPr()
    for tag in ("w:pStyle", "w:pageBreakBefore", "w:spacing", "w:ind", "w:jc"):
        for existing in list(pPr.findall(qn(tag))):
            pPr.remove(existing)

    # Надпись держится со следующим фрагментом и сама начинает новую страницу.
    _set_keep_flags(p_elem, keep_next=True, keep_lines=True)

    pbb = OxmlElement("w:pageBreakBefore")
    pPr.append(pbb)

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "360")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)

    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:right"), "0")
    ind.set(qn("w:firstLine"), "0")
    pPr.append(ind)

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "left")
    pPr.append(jc)


def _normalize_continuation_labels(doc) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        match = _CONT_NUM_RE.match(text)
        if not match:
            continue
        normalized = f"{_LABEL_PREFIX} {match.group(1)}"
        if text != normalized:
            paragraph.text = normalized
        _set_label_ppr(paragraph._p)
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)


def _split_table(table, sizes: list[int], number: str) -> None:
    """Разрезает <w:tbl> на len(sizes) фрагментов по размерам sizes (строки
    данных). Первый — без надписи; остальные = разрыв + надпись + клон шапки
    + свои строки."""
    tbl = table._tbl
    trs = tbl.findall(qn("w:tr"))
    if len(trs) < 2:
        return
    header = trs[0]
    data = trs[1:]
    tblPr = tbl.find(qn("w:tblPr"))
    tblGrid = tbl.find(qn("w:tblGrid"))

    # offsets границ
    bounds = []
    acc = 0
    for s in sizes:
        bounds.append((acc, acc + s))
        acc += s
    bounds[-1] = (bounds[-1][0], len(data))  # хвост всё забирает

    # фрагмент 1 — оставляем в исходной таблице data[0:b1]; лишние убираем
    first_end = bounds[0][1]
    for tr in data[first_end:]:
        tbl.remove(tr)

    anchor = tbl  # после кого вставлять следующий блок
    for fi in range(1, len(bounds)):
        s, e = bounds[fi]
        new_tbl = OxmlElement("w:tbl")
        if tblPr is not None:
            new_tbl.append(copy.deepcopy(tblPr))
        if tblGrid is not None:
            new_tbl.append(copy.deepcopy(tblGrid))
        new_tbl.append(_clone_header(header))
        for tr in data[s:e]:
            new_tbl.append(tr)  # перенос элемента (move)
        anchor.addnext(new_tbl)
        # Между anchor и new_tbl мог остаться пустой <w:p> от
        # ensure_separator_after_tables_and_figures — иначе он попадёт на
        # страницу-продолжение ПЕРЕД «Продолжение таблицы N» как лишняя
        # пустая строка. Удаляем перед вставкой надписи, чтобы p_label
        # встал ровно после anchor.
        _drop_empty_paragraphs_between(anchor, new_tbl)
        p_label = _make_label_para(number)  # несёт pageBreakBefore
        new_tbl.addprevious(p_label)
        anchor = new_tbl


def _has_straddling_vmerge(table) -> bool:
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            vm = tc.find(qn("w:tcPr"))
            if vm is None:
                continue
            m = vm.find(qn("w:vMerge"))
            if m is not None:
                return True
    return False


def _sizes_from_uno_pages(pages: list[int] | None) -> list[int] | None:
    """Из UNO-измерения [page для каждой строки, ВКЛЮЧАЯ шапку] → размеры
    фрагментов (строки данных на каждой странице).

    pages[0] — страница шапки, pages[1:] — строки данных. Группируем
    данные по страницам: [9,9,9,10,10,11] → данные [9,9,10,10,11] →
    [2,2,1]. Это РЕАЛЬНОЕ распределение из движка LibreOffice — режем
    ровно там, где таблица переносится естественно. Возвращает None,
    если измерения нет или таблица на одной странице (резать нечего).
    """
    if not pages or len(pages) < 2:
        return None
    data_pages = [(p if p is not None and p >= 0 else None) for p in pages[1:]]
    groups = _page_groups(data_pages)  # None у строки → берём предыдущую
    if groups is None or len(groups) < 2:
        return None
    return groups


def split_long_tables_to_pages(doc, row_pages_by_table, table_captions) -> bool:
    """Режет нумерованные подписанные таблицы, переносящиеся через страницу.

    doc — загруженный python-docx Document.
    row_pages_by_table — {индекс таблицы в doc.tables: [номер страницы для
    каждой строки, ВКЛЮЧАЯ шапку]} из UNO-измерения (engine.uno_layout).
    Это ТОЧНАЯ пагинация от движка LibreOffice. Если для таблицы измерения
    нет (UNO недоступен/упал), таблицу не режем: продолжение без реальной
    страницы переноса создаёт ложные фрагменты.

    Возвращает True, если что-то менялось.
    """
    _remerge_continuations(doc)
    _normalize_continuation_labels(doc)

    try:
        captioned = _find_captioned_tables(doc, table_captions)
    except Exception:
        _logger.warning("table_continuation: сопоставление подписей не удалось",
                         exc_info=True)
        return False

    # Индексы таблиц в doc.tables снимаем ДО любой резки — после первой
    # резки число таблиц меняется, а UNO-измерение привязано к исходному
    # порядку.
    tbl_index: dict[int, int] = {}
    for idx, t in enumerate(doc.tables):
        tbl_index[id(t._tbl)] = idx

    _logger.info(
        "table_continuation: подписанных таблиц найдено: %d, UNO-измерений: %d",
        len(captioned), len(row_pages_by_table or {}),
    )
    changed = False
    for table, cap in captioned:
        try:
            if getattr(cap, "kind", None) != "table":
                continue
            number = str(getattr(cap, "new_number", "") or "").strip()
            if not number:
                _logger.info("table_continuation: пропуск — нет номера подписи")
                continue  # номер не выдумываем
            if _looks_decorative(table):
                _logger.info(
                    "table_continuation: таблица %s пропущена — декоративная",
                    number,
                )
                continue
            rows = list(table.rows)
            if len(rows) < 3:
                _logger.info(
                    "table_continuation: таблица %s пропущена — строк <3 (%d)",
                    number, len(rows),
                )
                continue
            if _has_straddling_vmerge(table):
                _logger.info(
                    "table_continuation: таблица %s пропущена — "
                    "вертикально объединённые ячейки (vMerge)",
                    number,
                )
                continue  # объединённые ячейки — не рискуем
            data_rows = rows[1:]
            ti = tbl_index.get(id(table._tbl))
            uno_pages = (row_pages_by_table or {}).get(ti) if ti is not None else None

            sizes: list[int] | None = _sizes_from_uno_pages(uno_pages)
            via = "UNO"
            if sizes is None:
                _logger.info(
                    "table_continuation: таблица %s пропущена — нет "
                    "валидного UNO-переноса: %s",
                    number,
                    uno_pages,
                )
                continue
            if not sizes or len(sizes) < 2:
                continue
            # Размеры должны суммироваться в число строк данных; хвост
            # _split_table добирает сам, но подстрахуемся от рассинхрона.
            if sum(sizes) != len(data_rows):
                _logger.info(
                    "table_continuation: таблица %s — sizes=%s sum!=%d, "
                    "хвост доберёт _split_table", number, sizes, len(data_rows),
                )
            _split_table(table, sizes, number)
            _logger.info(
                "table_continuation: таблица %s разрезана на %d фрагментов "
                "(%s): %s", number, len(sizes), via, sizes,
            )
            changed = True
        except Exception:
            _logger.warning(
                "table_continuation: пропуск таблицы из-за ошибки",
                exc_info=True,
            )
            continue
    if changed:
        apply_continuation_table_styles(doc, 0)
    return changed


# ── итеративная дорезка переполненных фрагментов ─────────────────────────────

_CAPTION_NUM_RE = re.compile(r"^\s*таблица\s+(\d+(?:\.\d+)?)\s*[–—-]", re.IGNORECASE)
_CONT_NUM_RE = re.compile(r"^\s*продолжение\s+таблицы\s*(\d+(?:\.\d+)?)", re.IGNORECASE)

# Если на первой странице таблицы влезает меньше этого числа строк данных,
# пробуем перенести таблицу на новую страницу (pageBreakBefore на абзац перед
# ней) и пересчитываем. Пустая страница с одной-единственной строкой
# выглядит намного хуже, чем небольшой отступ в конце предыдущей.
_MIN_FIRST_DATA_ROWS = 2


def _try_push_table_to_new_page(tbl_elem) -> bool:
    """Добавляет pageBreakBefore к абзацу перед таблицей, чтобы перенести её
    (вместе с заголовком/подписью) на новую страницу.

    Логика выбора целевого абзаца:
      * Находим ближайший непустой <w:p> перед таблицей (обычно это подпись
        «Таблица N»).
      * Если прямо перед ней стоит заголовок (outlineLvl в pPr), ставим
        pageBreakBefore на заголовок — он переезжает вместе с подписью и
        таблицей.
      * Если pageBreakBefore уже стоит — таблица уже начинается с новой
        страницы; дальнейший перенос бессмысленен.

    Возвращает True, если изменение сделано.
    """
    # Ближайший непустой <w:p> перед таблицей — ожидаем подпись «Таблица N».
    caption_el = None
    el = tbl_elem.getprevious()
    while el is not None:
        tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if tag == "p":
            if "".join(el.itertext()).strip():
                caption_el = el
                break
        elif tag == "tbl":
            return False  # другая таблица перед нашей — не трогаем
        el = el.getprevious()
    if caption_el is None:
        return False

    cap_pPr = caption_el.find(qn("w:pPr"))
    if cap_pPr is not None:
        pbb = cap_pPr.find(qn("w:pageBreakBefore"))
        if pbb is not None and pbb.get(qn("w:val"), "1") not in ("0", "false"):
            return False  # уже начинается с новой страницы

    # Если прямо перед подписью стоит заголовок — ставим разрыв на него,
    # чтобы заголовок не остался «сиротой» в конце предыдущей страницы.
    target = caption_el
    prev_el = caption_el.getprevious()
    if prev_el is not None:
        ptag = prev_el.tag.split("}")[-1] if "}" in prev_el.tag else prev_el.tag
        if ptag == "p" and "".join(prev_el.itertext()).strip():
            ppPr = prev_el.find(qn("w:pPr"))
            if ppPr is not None and ppPr.find(qn("w:outlineLvl")) is not None:
                target = prev_el  # заголовок переедет вместе с таблицей

    tgt_pPr = target.find(qn("w:pPr"))
    if tgt_pPr is None:
        tgt_pPr = OxmlElement("w:pPr")
        target.insert(0, tgt_pPr)
    existing = tgt_pPr.find(qn("w:pageBreakBefore"))
    if existing is not None:
        return False
    pbb_new = OxmlElement("w:pageBreakBefore")
    tgt_pPr.append(pbb_new)
    return True


def _table_number_before(tbl_elem) -> str | None:
    """Номер таблицы из ближайшего непустого <w:p> перед <w:tbl>.

    Перед первым фрагментом стоит подпись «Таблица N — …», перед каждым
    продолжением — наша надпись «Продолжение таблицы N». Из любой берём N.
    Если ближайший непустой абзац не про таблицу — None (номер не выдумываем).
    """
    el = tbl_elem.getprevious()
    while el is not None:
        tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if tag == "p":
            txt = "".join(el.itertext()).strip()
            if txt:
                m = _CONT_NUM_RE.match(txt) or _CAPTION_NUM_RE.match(txt)
                return m.group(1) if m else None
        elif tag == "tbl":
            return None  # перед нами другая таблица без подписи — стоп
        el = el.getprevious()
    return None


def _next_fragment_of_same_table(tbl_elem, number: str):
    """Следующий <w:tbl> — продолжение той же таблицы N, если оно идёт сразу
    за `tbl_elem` через надпись «Продолжение таблицы N». Иначе None.
    """
    el = tbl_elem.getnext()
    saw_label = False
    while el is not None:
        tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if tag == "p":
            txt = "".join(el.itertext()).strip()
            if txt:
                m = _CONT_NUM_RE.match(txt)
                if m and m.group(1) == number:
                    saw_label = True
                else:
                    return None  # чужой абзац между фрагментами
        elif tag == "tbl":
            return el if saw_label else None
        el = el.getnext()
    return None


def _move_overflow_to_next(tbl_elem, keep: int, next_tbl_elem) -> None:
    """Переносит строки данных сверх `keep` из `tbl_elem` в начало данных
    `next_tbl_elem` (сразу после его шапки), сохраняя порядок.

    Используется для балансировки: вместо того чтобы оставить одинокий
    «хвост» переполненного фрагмента на полупустой странице, прокатываем
    лишние строки в следующий фрагмент той же таблицы — там они лягут
    рядом со своими соседями.
    """
    trs = tbl_elem.findall(qn("w:tr"))
    data = trs[1:]  # без шапки
    overflow = data[keep:]
    if not overflow:
        return
    next_trs = next_tbl_elem.findall(qn("w:tr"))
    if not next_trs:
        return
    anchor = next_trs[0]  # шапка следующего фрагмента
    for tr in overflow:
        tbl_elem.remove(tr)
        anchor.addnext(tr)
        anchor = tr


def split_overflowing_fragments(doc, row_pages_by_table) -> bool:
    """Режет таблицы/фрагменты, переносящиеся через страницу — по ОДНОЙ
    (первой) точке переноса за вызов. Это основной механизм UNO-резки:
    вызывается итеративно, каждый раз по измерению ТЕКУЩЕГО документа.

    Почему итеративно по одной точке, а не сразу по всем группам: надпись
    «Продолжение таблицы N» + повтор шапки занимают место, которого не было
    при измерении голой таблицы. Если резать сразу по группам голой таблицы,
    фрагменты получаются неоптимальными (часть переполняется, часть пустует).
    Резка по одной точке реально измеренного документа кладёт в каждый
    фрагмент максимум строк, которые туда влезают.

    Номер для надписи берётся из подписи «Таблица N» (первый фрагмент) или
    из надписи «Продолжение таблицы N» (последующие). Декоративные таблицы и
    таблицы с вертикальным объединением ячеек пропускаются.

    row_pages_by_table — UNO-измерение ТЕКУЩЕГО документа {idx: [page_per_row]}.
    Возвращает True, если что-то разрезали (вызывающий перемерит и повторит).
    """
    _normalize_continuation_labels(doc)
    tables = list(doc.tables)
    changed = False
    for ti, table in enumerate(tables):
        try:
            pages = (row_pages_by_table or {}).get(ti)
            if not pages or len(pages) < 2:
                continue
            data_pages = [(p if p is not None and p >= 0 else None) for p in pages[1:]]
            groups = _page_groups(data_pages)
            if groups is None or len(groups) < 2:
                continue  # фрагмент помещается на одной странице — не трогаем
            if _looks_decorative(table) or _has_straddling_vmerge(table):
                continue  # титульные/бланковые и vMerge не режем
            number = _table_number_before(table._tbl)
            if not number:
                continue  # номер неизвестен — надпись делать не из чего
            data_rows = list(table.rows)[1:]
            first = groups[0]  # строк данных, влезающих на первую страницу
            rest = len(data_rows) - first
            if first < 1 or rest < 1:
                continue

            # Если на первой странице влезает слишком мало строк — переносим
            # таблицу (вместе с заголовком/подписью) на новую страницу.
            # На следующей итерации UNO перемерит: таблица стартует с чистого
            # листа, строк влезет больше или таблица уместится без разреза.
            #
            # Дополнительно: если вторая страница получает ВДВОЕ больше строк,
            # чем первая (groups[1] > first * 2), таблица явно началась в
            # середине страницы. Перенос исправит разрез [3,8] → [8,3].
            _very_uneven = len(groups) > 1 and groups[1] > first * 2
            if first < _MIN_FIRST_DATA_ROWS or _very_uneven:
                if _try_push_table_to_new_page(table._tbl):
                    _logger.info(
                        "table_continuation: таблица %s — %d/%d строк на "
                        "первой странице (очень неравномерный разрез), "
                        "перенос на новую страницу",
                        number, first, len(data_rows),
                    )
                    changed = True
                    continue  # пересчитаем на следующей итерации

            # Балансировка: если следующий фрагмент той же таблицы уже
            # существует — прокатываем лишние строки в его начало (а не
            # создаём из них одинокий хвост на полупустой странице). Иначе —
            # режем по точке переноса, создавая новый фрагмент с надписью.
            next_frag = _next_fragment_of_same_table(table._tbl, number)
            if next_frag is not None:
                _move_overflow_to_next(table._tbl, first, next_frag)
                _logger.info(
                    "table_continuation: таблица %s — %d строк прокатаны "
                    "в следующий фрагмент (оставлено %d)",
                    number, rest, first,
                )
            else:
                _split_table(table, [first, rest], number)
                _logger.info(
                    "table_continuation: таблица %s — разрез по точке "
                    "переноса [%d | %d]", number, first, rest,
                )
            changed = True
        except Exception:
            _logger.warning(
                "table_continuation: резка фрагмента не удалась",
                exc_info=True,
            )
            continue
    if changed:
        apply_continuation_table_styles(doc, 0)
    return changed
