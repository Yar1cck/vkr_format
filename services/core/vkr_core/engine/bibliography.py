"""Проверка списка использованных источников (ТЗ §7.6).

Проверки:
  B1. Каждой ссылке [N] в тексте должна соответствовать запись с этим же
      номером в разделе "Список использованных источников".
  B2. Записи в списке источников пронумерованы подряд: 1, 2, 3, ... .
  B3. У каждой записи должен быть хотя бы год и автор/название — короткие
      "обрубки" вызывают предупреждение.
"""

from __future__ import annotations

import re
from dataclasses import dataclass

from docx.document import Document as DocxDocument
from docx.oxml.ns import qn

from services.core.vkr_core.engine.citations import CitationUsage
from services.core.vkr_core.engine.detection import HeadingCandidate
from services.core.vkr_core.engine.stats import DocStats
from services.core.vkr_core.engine.violations import (
    SEVERITY_CRITICAL,
    SEVERITY_WARNING,
    PipelineViolation,
)
from services.core.vkr_core.models.enums import ViolationStatus

BIB_ENTRY_RE = re.compile(r"^(\d+)[.)]\s+(.*)$", re.DOTALL)
YEAR_RE = re.compile(r"\b(19|20)\d{2}\b")

# Признаки электронного ресурса или нормативного документа — для них
# отсутствие года не является нарушением (год может быть в дате обращения,
# а стандарт может не иметь традиционного года издания).
_URL_RE = re.compile(r"https?://|www\.", re.IGNORECASE)
_NORMATIVE_RE = re.compile(
    r"\bГОСТ\b|\bISO\b|\bRFC\b|\bСНиП\b|\bСП\s+\d|\bФЗ\b|\bПриказ\b|\bОСТ\b",
    re.IGNORECASE,
)

# Ключевые слова, по которым заголовок раздела распознаётся как список источников
# даже если точного совпадения с алиасами нет (опечатки, сокращения).
_BIB_KEYWORDS = frozenset({"источник", "литератур", "библиограф"})

# Признак строки-продолжения в group(2) BIB_ENTRY_RE: начинается с издательского
# тире/дефиса («– М.: Наука, 2020.»). Такой фрагмент — не самостоятельная запись.
_CONTINUATION_START_RE = re.compile(r"^[–—\-]")
_ENTRY_BODY_CONTINUATION_RE = re.compile(
    r"^(?:№|n\b|no\.?|с\.|стр\.|pp?\.|url\b|doi\b|\[электронный ресурс\])",
    re.IGNORECASE,
)

# Граница между двумя записями в одном абзаце: после знака препинания
# (точка, точка с запятой, закрывающая скобка), пробельные символы, затем
# новый маркер вида "N." или "N)". Lookbehind фиксированной длины — это
# требование re. Lookahead на маркер не съедает его, чтобы он остался у
# следующего куска.
#
# Используем \d{1,3} вместо \d+: номер записи — максимум трёхзначный (1–999),
# а год публикации всегда четырёхзначный (1900–2099). Это предотвращает ложный
# разрез внутри одной записи на ". YYYY." (типичный шаблон: «Журнал. 2024. №2.»).
_MULTI_ENTRY_SPLIT_RE = re.compile(r"(?<=[.;)\]])\s+(?=\d{1,3}[).]\s+)")

# Битый номер записи: диапазон вида «18-32», «18 – 32» или составной
# «18.5» вместо одиночного арабского числа (п.6.16 Порядка МИИГАиК —
# каждый источник нумеруется отдельным арабским числом). «18 Автор»
# (без точки) НЕ считаем битым — это валидно по п.6.16.
_MALFORMED_NUMBER_RE = re.compile(r"^\s*\d+\s*(?:[-–—]\s*\d+|[.,]\d)")


def _split_multi_entries(text: str) -> list[str]:
    """Делит текст абзаца на отдельные записи библиографии.

    Используется для случая, когда несколько записей слиплись в один абзац:
    студент вставил `1) Иванов... 2) Петров... 3) Сидоров...` без
    разделения на абзацы, либо с Shift+Enter (мягкий перенос — python-docx
    склеивает их через `\\n`).

    Возвращает список потенциальных записей. Дальше они должны быть
    проверены через `BIB_ENTRY_RE` — если хотя бы одна не прошла, разделение
    скорее всего ложное (например, наткнулись на «Том 1. ...» внутри
    одной записи), и абзац стоит обработать как одну запись.
    """
    parts: list[str] = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        for sub in _MULTI_ENTRY_SPLIT_RE.split(line):
            sub = sub.strip()
            if sub:
                parts.append(sub)
    return parts


def _looks_like_entry_body(text: str) -> bool:
    """Отличает начало записи от продолжения журнального описания.

    В списках источников часто встречается фрагмент «5. № 12-1 (81). С. ...»:
    цифра с точкой относится к году/переносу/склейке записи, а не к номеру
    источника. Без этой защиты такой хвост распознавался как новая запись №5
    и давал ложный разрыв нумерации.
    """
    return not _ENTRY_BODY_CONTINUATION_RE.match(text.strip())


@dataclass
class BibEntry:
    paragraph_index: int
    number: int
    text: str


def _norm(text: str) -> str:
    return " ".join(text.strip().lower().split())


def _is_references_heading(text: str, aliases: set[str]) -> bool:
    """Точное совпадение с алиасом, или fallback по ключевым словам (опечатки, сокращения)."""
    t = _norm(text)
    if t in aliases:
        return True
    # Нечёткий fallback: содержит хотя бы одно ключевое слово и достаточно короткий
    # (чтобы не ловить обычные абзацы с упоминанием «источников»).
    return len(t) < 80 and any(kw in t for kw in _BIB_KEYWORDS)


def _find_references_range(
    stats: DocStats,
    rules: dict,
    headings: list[HeadingCandidate],
) -> tuple[int, int] | None:
    aliases = {
        _norm(item) for item in rules.get("structural_elements", {}).get("references", [])
    }
    start_index: int | None = None
    end_index: int = len(stats.paragraphs)

    heading_indexes = [h.paragraph_index for h in headings]
    heading_indexes.sort()

    for i, ps in enumerate(stats.paragraphs):
        if not ps.is_empty and _is_references_heading(ps.stripped, aliases):
            start_index = i
            # Следующий заголовок после этого помечает конец раздела.
            for hi in heading_indexes:
                if hi > i:
                    end_index = hi
                    break
            break
    if start_index is None:
        return None
    return start_index + 1, end_index


def find_bibliography_start(
    stats: DocStats,
    rules: dict,
    headings: list[HeadingCandidate],
) -> int | None:
    """Возвращает индекс первого абзаца с содержимым списка источников или None."""
    bounds = _find_references_range(stats, rules, headings)
    return bounds[0] if bounds else None


def collect_bibliography(
    stats: DocStats,
    rules: dict,
    headings: list[HeadingCandidate],
) -> tuple[list[BibEntry], tuple[int, int] | None, list[tuple[int, str]]]:
    bounds = _find_references_range(stats, rules, headings)
    entries: list[BibEntry] = []
    malformed: list[tuple[int, str]] = []
    if bounds is None:
        return entries, None, malformed
    start, end = bounds
    for i in range(start, end):
        ps = stats.paragraphs[i]
        if ps.is_empty:
            continue

        # Сначала пытаемся разделить абзац на несколько записей. Это спасает
        # от ложного critical "источники не найдены", когда студент склеил
        # весь список в один абзац или использовал Shift+Enter.
        sub_entries = _split_multi_entries(ps.stripped)
        sub_matches = [BIB_ENTRY_RE.match(s) for s in sub_entries]
        if (
            len(sub_entries) > 1
            and all(m is not None for m in sub_matches)
            # Отклоняем ложный сплит «Т. 1. – М.:…»: у настоящей записи
            # group(2) начинается с автора/названия, а не с издательского тире.
            and not any(
                _CONTINUATION_START_RE.match(m.group(2).strip())
                or not _looks_like_entry_body(m.group(2))
                for m in sub_matches
            )
        ):
            for sub_match in sub_matches:
                try:
                    number = int(sub_match.group(1))
                except (ValueError, AttributeError):
                    continue
                entries.append(
                    BibEntry(
                        paragraph_index=i,
                        number=number,
                        text=sub_match.group(2).strip(),
                    )
                )
            continue

        # Fallback — одна запись на абзац (исходное поведение).
        match = BIB_ENTRY_RE.match(ps.stripped)
        if not match:
            # Автонумерация Word (w:numPr): видимое «1.» рисует поле
            # нумерации, в тексте абзаца цифры нет — буквальный BIB_ENTRY_RE
            # такую запись не ловит, и список «не распознавался». Внутри
            # раздела источников нумерованный абзац с текстом — это запись;
            # номер синтезируем по порядку (для чистого автосписка он точно
            # совпадает с тем, что рисует Word: 1, 2, 3, …). Так корректно
            # отрабатывают и B2 (сквозная нумерация), и B1 (сверка с [N]).
            if ps.numpr_num_id is not None and ps.stripped:
                entries.append(
                    BibEntry(
                        paragraph_index=i,
                        number=len(entries) + 1,
                        text=ps.stripped,
                    )
                )
                continue
            # Строка-продолжение длинной записи не помечается. Но если она
            # ПЫТАЕТСЯ быть записью с битым номером («18-32», «18.5») —
            # это нарушение п.6.16, иначе оно молча терялось.
            if _MALFORMED_NUMBER_RE.match(ps.stripped):
                malformed.append((i, ps.stripped))
            continue
        try:
            number = int(match.group(1))
        except ValueError:
            continue
        body = match.group(2).strip()
        if not _looks_like_entry_body(body):
            continue
        entries.append(BibEntry(paragraph_index=i, number=number, text=body))
    return entries, bounds, malformed


# Начало непронумерованной записи: первый значимый символ — заглавная буква
# (кириллица/латиница) или открывающая кавычка. Цифру В НАЧАЛЕ намеренно НЕ
# считаем: «18 Автор» (номер без точки, валиден по п.6.16) и «2024. …»
# (фрагмент с годом) не должны получать второй номер. Строчная буква, тире —
# это строка-продолжение длинной записи, её тоже не нумеруем.
_ENTRY_START_RE = re.compile(r"^[«\"“„(]?[A-ZА-ЯЁ]")

# Заголовок группы источников вида «I. Нормативные документы» или «II. Учебники»:
# латинские символы римской цифры, затем точка/скобка и пробел + непустой текст.
# Такие строки нельзя нумеровать как записи.
_ROMAN_GROUP_RE = re.compile(r"^[IVXLCDMivxlcdm]+[.)]\s+\S")


def _looks_like_entry_start(text: str) -> bool:
    return bool(_ENTRY_START_RE.match(text.strip()))


def _is_group_header(text: str) -> bool:
    return bool(_ROMAN_GROUP_RE.match(text.strip()))


def _prepend_number(paragraph, number: int) -> None:
    """Дописывает «N. » в начало абзаца отдельным run'ом."""
    run = paragraph.add_run(f"{number}. ")  # add_run кладёт в конец…
    r = run._element
    p = paragraph._p
    p.remove(r)  # …убираем и переставляем перед первым существующим run'ом.
    first_r = p.find(qn("w:r"))
    if first_r is not None:
        first_r.addprevious(r)
    else:
        p.append(r)


def apply_bibliography_numbering(
    doc: DocxDocument,
    stats: DocStats,
    rules: dict,
    headings: list[HeadingCandidate],
) -> list[int]:
    """Дописывает номера источникам без номера, продолжая существующую
    последовательность. Записи с уже проставленным номером (текстовым «N.» или
    Word-нумерацией w:numPr) не трогаются. Возвращает индексы изменённых абзацев.
    """
    bounds = _find_references_range(stats, rules, headings)
    if bounds is None:
        return []
    start, end = bounds
    paragraphs = doc.paragraphs
    fixed: list[int] = []
    counter = 0
    for i in range(start, min(end, len(paragraphs), len(stats.paragraphs))):
        ps = stats.paragraphs[i]
        if ps.is_empty:
            continue
        match = BIB_ENTRY_RE.match(ps.stripped)
        if match:
            try:
                counter = int(match.group(1))
            except ValueError:
                counter += 1
            continue
        if ps.numpr_num_id is not None:
            counter += 1
            continue
        if _looks_like_entry_start(ps.stripped) and not _is_group_header(ps.stripped):
            counter += 1
            _prepend_number(paragraphs[i], counter)
            fixed.append(i)
        # иначе — строка-продолжение, номер не дописываем.
    return fixed


def validate_bibliography(
    entries: list[BibEntry],
    citations: list[CitationUsage],
    bounds: tuple[int, int] | None,
    malformed: list[tuple[int, str]] | None = None,
) -> list[PipelineViolation]:
    violations: list[PipelineViolation] = []

    for para_idx, raw in malformed or []:
        violations.append(
            PipelineViolation(
                type="bibliography_entry_malformed",
                rule_reference="п. 6.16",
                description=(
                    "Запись списка источников имеет некорректный номер "
                    "(например, диапазон «18-32»). Каждый источник нумеруется "
                    "отдельным арабским числом (п.6.16 Порядка МИИГАиК)."
                ),
                status=ViolationStatus.manual_required,
                severity=SEVERITY_WARNING,
                paragraph_index=para_idx,
                original_text=raw,
            )
        )

    if bounds is None:
        violations.append(
            PipelineViolation(
                type="bibliography_missing",
                rule_reference="п. 6.16",
                description="Раздел «Список использованных источников» не найден.",
                status=ViolationStatus.manual_required,
                severity=SEVERITY_CRITICAL,
            )
        )
        return violations

    if not entries:
        violations.append(
            PipelineViolation(
                type="bibliography_entries_unrecognised",
                rule_reference="п. 6.16",
                description=(
                    "Записи источников не распознаны — возможно, нестандартный формат "
                    "оформления. Требуется ручная проверка."
                ),
                status=ViolationStatus.manual_required,
                severity=SEVERITY_WARNING,
            )
        )
        return violations

    # B2 — сквозная нумерация.
    expected = 1
    for entry in entries:
        if entry.number != expected:
            violations.append(
                PipelineViolation(
                    type="bibliography_numbering_gap",
                    rule_reference="п. 6.16",
                    description=(
                        f"Нарушена нумерация в списке источников: ожидался {expected}, "
                        f"найден {entry.number}."
                    ),
                    status=ViolationStatus.manual_required,
                    severity=SEVERITY_WARNING,
                    paragraph_index=entry.paragraph_index,
                    original_text=entry.text,
                )
            )
            expected = entry.number + 1
        else:
            expected += 1

    # B1 — у каждого номера-ссылки есть соответствующая запись.
    entry_numbers = {e.number for e in entries}
    cited_numbers: set[int] = set()
    for usage in citations:
        cited_numbers.update(usage.numbers)

    # Для каждого номера-ссылки — первый абзац и РЕАЛЬНЫЙ токен из текста
    # (например "[48-50]"), а не синтезированный "[50]": число могло прийти
    # из диапазона/группы, и литерального "[50]" в документе нет — тогда
    # подсветка/скролл не находили бы ничего.
    first_cite: dict[int, tuple[int, str]] = {}
    for usage in citations:
        for raw, nums in usage.groups:
            for n in nums:
                if n not in first_cite:
                    first_cite[n] = (usage.paragraph_index, raw)
        if not usage.groups:  # подстраховка для старых вызовов без groups
            for n in usage.numbers:
                first_cite.setdefault(n, (usage.paragraph_index, f"[{n}]"))

    for number in sorted(cited_numbers - entry_numbers):
        loc = first_cite.get(number)
        violations.append(
            PipelineViolation(
                type="citation_without_source",
                rule_reference="п. 6.16",
                description=(
                    f"Ссылка [{number}] встречается в тексте, но отсутствует в списке источников."
                ),
                status=ViolationStatus.manual_required,
                severity=SEVERITY_CRITICAL,
                paragraph_index=loc[0] if loc else None,
                original_text=loc[1] if loc else f"[{number}]",
            )
        )

    for entry_number in sorted(entry_numbers - cited_numbers):
        violations.append(
            PipelineViolation(
                type="source_not_cited",
                rule_reference="п. 6.16",
                description=(
                    f"Источник {entry_number} присутствует в списке, но не цитируется в тексте."
                ),
                status=ViolationStatus.manual_required,
                severity=SEVERITY_WARNING,
            )
        )

    # B3 — поверхностная проверка качества записей.
    # Электронные ресурсы (URL) и нормативные документы (ГОСТ, ISO, RFC…)
    # не требуют традиционного года издания — пропускаем их.
    for entry in entries:
        if _URL_RE.search(entry.text) or _NORMATIVE_RE.search(entry.text):
            continue
        if len(entry.text) < 20 or not YEAR_RE.search(entry.text):
            violations.append(
                PipelineViolation(
                    type="bibliography_entry_incomplete",
                    rule_reference="п. 6.16",
                    description=(
                        f"Запись {entry.number} в списке источников короче ожидаемой или не содержит года издания."
                    ),
                    status=ViolationStatus.manual_required,
                    severity=SEVERITY_WARNING,
                    paragraph_index=entry.paragraph_index,
                    original_text=entry.text,
                )
            )

    return violations
