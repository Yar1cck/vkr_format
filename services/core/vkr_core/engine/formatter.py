"""Форматирование страницы, шрифта и абзацев (ТЗ §6.1–6.3)."""

from __future__ import annotations

import logging
import re

from docx.document import Document as DocxDocument
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

from services.core.vkr_core.engine.detection import HeadingCandidate
from services.core.vkr_core.engine.heading_prefix import parse_heading_prefix
from services.core.vkr_core.engine.stats import DocStats
from services.core.vkr_core.engine.violations import SEVERITY_INFO, PipelineViolation
from services.core.vkr_core.models.enums import ViolationStatus

_logger = logging.getLogger(__name__)


def overwrite_paragraph_text(paragraph, new_text: str) -> None:
    """Перезаписывает текст абзаца, сохраняя форматирование первого run'а.

    paragraph.runs не видит <w:hyperlink>, но paragraph.text их склеивает —
    без удаления гиперссылок при перезаписи текст задваивался бы.
    """
    p = paragraph._p
    hyperlinks = p.findall(qn("w:hyperlink"))
    if hyperlinks:
        for h in hyperlinks:
            p.remove(h)
        _logger.warning(
            "Removed hyperlink(s) from paragraph during text rewrite; "
            "original anchor text is lost. Paragraph snippet: %r",
            (paragraph.text or "")[:80],
        )
    def _run_has_embedded_objects(run) -> bool:
        r = run._r
        return bool(
            r.find(qn("w:drawing")) is not None
            or r.find(qn("w:pict")) is not None
            or r.find(qn("w:object")) is not None
        )

    if paragraph.runs:
        target_run = None
        for run in paragraph.runs:
            if not _run_has_embedded_objects(run):
                target_run = run
                break
        if target_run is None:
            target_run = paragraph.add_run()
        target_run.text = new_text
        target_r = target_run._r
        for run in paragraph.runs:
            if run._r is target_r:
                continue
            if _run_has_embedded_objects(run):
                continue
            run.text = ""
    else:
        paragraph.add_run(new_text)

_NUM_PREFIX_RE = re.compile(r"^([\d.]+\s+)")
_CHAPTER_WORD_RE = re.compile(
    r"^(?:глава|раздел|часть|chapter|section|part)\s+(\d+(?:\.\d+)*)[.:\s]*",
    re.IGNORECASE,
)
# Ведущие пробелы: ASCII, таб, NBSP, Unicode — студенты часто вставляют вручную перед номером.
_LEADING_WS_RE = re.compile(r"^[\s  -   　]+")


def _lstrip_heading(text: str) -> str:
    return _LEADING_WS_RE.sub("", text)


_NUMBER_PREFIX_STRICT_RE = re.compile(r"^(\d+(?:\.\d+)*)[.)]?\s+")

_PAGE_SIZES_MM = {
    "A4": (210, 297),
    "A5": (148, 210),
    "Letter": (216, 279),
}

_ALIGNMENTS = {
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
}

_TABLE_CAPTION_EXEMPT_RE = re.compile(
    r"^(?:таблица|табл\.)\s+\d+(?:\.\d+)*(?:\s*[:–—-]\s*|\.\s+)\S",
    re.IGNORECASE,
)
_FIGURE_OR_LISTING_CAPTION_EXEMPT_RE = re.compile(
    r"^(?:рисунок|рис\.|листинг)\s+\d+(?:\.\d+)*(?:\s*[:–—-]\s*|\.\s+)\S",
    re.IGNORECASE,
)
_FORMULA_CAPTION_EXEMPT_RE = re.compile(
    r"^формула\s+\(?\d+(?:\.\d+)?\)?\b",
    re.IGNORECASE,
)
_TABLE_CONTINUATION_EXEMPT_RE = re.compile(
    r"^(?:продолжение|окончание)\s+таблицы\s+\d+(?:\.\d+)*\b",
    re.IGNORECASE,
)


def paragraph_alignment_from_rule(value: str | None) -> WD_ALIGN_PARAGRAPH:
    return _ALIGNMENTS.get(value or "justify", WD_ALIGN_PARAGRAPH.JUSTIFY)


def _next_visible_sibling_is_table(paragraph) -> bool:
    sibling = paragraph._element.getnext()
    while sibling is not None:
        if sibling.tag == qn("w:tbl"):
            return True
        if sibling.tag == qn("w:p"):
            text = "".join(t.text or "" for t in sibling.iter(qn("w:t"))).strip()
            if text:
                return False
        sibling = sibling.getnext()
    return False


def _is_body_style_exempt_caption(paragraph) -> bool:
    text = (paragraph.text or "").strip()
    if not text:
        return False
    if _TABLE_CONTINUATION_EXEMPT_RE.match(text):
        return True
    if _TABLE_CAPTION_EXEMPT_RE.match(text):
        return _next_visible_sibling_is_table(paragraph)
    return bool(
        _FIGURE_OR_LISTING_CAPTION_EXEMPT_RE.match(text)
        or _FORMULA_CAPTION_EXEMPT_RE.match(text)
    )


def _normalise_number_prefix(text: str) -> str:
    """Сворачивает "1. Title" / "1.1.2.  Title" / "1)\tTitle" → "1 Title" / "1.1.2 Title".

    Убирает финальную точку или скобку у ведущего номера и заменяет любые
    пробельные символы после него (табы, NBSP, серии пробелов) на один
    ASCII-пробел. Внутренние точки между уровнями иерархии ("1.1.2")
    сохраняются.
    """
    m = _NUMBER_PREFIX_STRICT_RE.match(text)
    if not m:
        return text
    number = m.group(1)
    rest = text[m.end():]
    return f"{number} {rest}" if rest else number


def _section_is_landscape(section) -> bool:
    """True, если секция уже альбомная — по атрибуту w:orient или по факту
    ширина > высоты (студенты иногда меняют размеры без атрибута orient)."""
    if section.orientation == WD_ORIENT.LANDSCAPE:
        return True
    return (
        section.page_width is not None
        and section.page_height is not None
        and section.page_width > section.page_height
    )


def apply_page_settings(doc: DocxDocument, rules: dict, skip_first_section: bool = False) -> None:
    page = rules["page_settings"]
    paper = page.get("paper", "A4")
    portrait_width_mm, portrait_height_mm = _PAGE_SIZES_MM.get(paper, _PAGE_SIZES_MM["A4"])
    default_landscape = page.get("orientation", "portrait") == "landscape"
    start_index = 1 if skip_first_section else 0
    for idx, section in enumerate(doc.sections):
        if idx < start_index:
            continue
        # Секция, сделанная альбомной вручную (широкие таблицы/приложения),
        # остаётся альбомной — не переворачиваем её обратно в книжную.
        landscape = default_landscape or _section_is_landscape(section)
        if landscape:
            width_mm, height_mm = portrait_height_mm, portrait_width_mm
            section.orientation = WD_ORIENT.LANDSCAPE
        else:
            width_mm, height_mm = portrait_width_mm, portrait_height_mm
            section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(width_mm)
        section.page_height = Mm(height_mm)
        section.left_margin = Mm(page["margin_left_mm"])
        section.right_margin = Mm(page["margin_right_mm"])
        section.top_margin = Mm(page["margin_top_mm"])
        section.bottom_margin = Mm(page["margin_bottom_mm"])


def disable_document_hyphenation(doc: DocxDocument) -> None:
    """Выставляет w:autoHyphenation val="false" в settings.xml.

    Явный false надёжнее удаления — перекрывает дефолт шаблона.
    consecutiveHyphenLimit убираем: без переносов он бессмысленен.
    """
    settings = doc.settings.element
    auto = settings.find(qn("w:autoHyphenation"))
    if auto is None:
        auto = OxmlElement("w:autoHyphenation")
        settings.append(auto)
    auto.set(qn("w:val"), "false")
    consec = settings.find(qn("w:consecutiveHyphenLimit"))
    if consec is not None:
        settings.remove(consec)


def _clear_footer_runs(footer) -> None:
    """Очищает колонтитул: удаляет таблицы и содержимое абзацев, оставляя w:pPr."""
    ftr = footer._element
    # Таблицы часто используют для размещения номеров страниц слева/справа.
    for tbl in list(ftr.findall(qn("w:tbl"))):
        ftr.remove(tbl)
    for p in list(ftr.iter(qn("w:p"))):
        for child in list(p):
            local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if local != "pPr":
                p.remove(child)


def _make_tnr_rPr() -> any:
    """Создаёт w:rPr с Times New Roman 14pt для использования в run'ах PAGE-поля."""
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), "Times New Roman")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "28")
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "28")
    rPr.append(szCs)
    return rPr


def _insert_page_field(paragraph) -> None:
    # Выравнивание через API и напрямую в XML — иначе наследование стиля перебивает.
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = paragraph._p.get_or_add_pPr()
    for _jc in list(pPr.findall(qn("w:jc"))):
        pPr.remove(_jc)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)

    # Стиль «Footer» может задавать шрифт, который LibreOffice подхватывает для текста поля.
    for _pStyle in list(pPr.findall(qn("w:pStyle"))):
        pPr.remove(_pStyle)

    # pPr/w:rPr — базовый шрифт абзаца; LibreOffice использует его для текста полей.
    for _prPr in list(pPr.findall(qn("w:rPr"))):
        pPr.remove(_prPr)
    pPr.append(_make_tnr_rPr())

    # Поле PAGE: пять run'ов по OOXML-стандарту.
    # LO берёт шрифт из result-placeholder (4-й run) — поэтому его w:rPr важен.
    p = paragraph._p

    def _make_field_run(fld_char_type: str | None = None, instr: str | None = None,
                        result_text: str | None = None) -> any:
        r = OxmlElement("w:r")
        r.append(_make_tnr_rPr())
        if fld_char_type is not None:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), fld_char_type)
            r.append(fc)
        if instr is not None:
            it = OxmlElement("w:instrText")
            it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            it.text = instr
            r.append(it)
        if result_text is not None:
            t = OxmlElement("w:t")
            t.text = result_text
            r.append(t)
        return r

    p.append(_make_field_run(fld_char_type="begin"))
    p.append(_make_field_run(instr=" PAGE "))
    p.append(_make_field_run(fld_char_type="separate"))
    p.append(_make_field_run(result_text="1"))   # placeholder — LO заменит при update_fields
    p.append(_make_field_run(fld_char_type="end"))


def _strip_sectpr_flags(section) -> None:
    """Удаляет w:titlePg и w:evenAndOddHeaders из sectPr.

    Нужно, чтобы все страницы (включая первую) использовали default-колонтитул.
    """
    sectPr = section._sectPr
    for tag in ("w:titlePg", "w:evenAndOddHeaders"):
        for el in list(sectPr.findall(qn(tag))):
            sectPr.remove(el)
    for tag in ("w:headerReference", "w:footerReference"):
        for ref in list(sectPr.findall(qn(tag))):
            if ref.get(qn("w:type")) in ("first", "even"):
                sectPr.remove(ref)


def _ensure_page_number_start(section, start: int = 1) -> None:
    """Устанавливает w:pgNumType start=<n>. Многие шаблоны несут start=0."""
    sectPr = section._sectPr
    for el in list(sectPr.findall(qn("w:pgNumType"))):
        sectPr.remove(el)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:start"), str(start))
    sectPr.append(pg)


def insert_pretoc_section_break(
    doc: DocxDocument,
    rules: dict,
    title_end_index: int = 0,
) -> bool:
    """Вставляет continuous sectPr перед СОДЕРЖАНИЕ, деля документ на две секции.

    Секция 0 (титул + задание) — пустой колонтитул. Секция 1+ (СОДЕРЖАНИЕ и далее) — PAGE-поле.
    Возвращает False, если СОДЕРЖАНИЕ не найдено или перед ним нет абзаца.

    title_end_index: индекс первого абзаца после конца титульного листа.
    Если между концом титула и СОДЕРЖАНИЕ нет содержательных абзацев (только
    пустые строки) — автоматически вставляется пустая страница, чтобы СОДЕРЖАНИЕ
    оказалось на стр. 3, а не стр. 2.
    """
    import copy as _copy

    from services.core.vkr_core.engine.toc import _find_contents_heading
    heading = _find_contents_heading(doc, rules)
    if heading is None:
        return False

    # Ищем ближайший <w:p> перед СОДЕРЖАНИЕ в body (пропускаем <w:tbl> и т.п.).
    elem = heading._p.getprevious()
    while elem is not None:
        _tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if _tag == "p":
            break
        elem = elem.getprevious()
    if elem is None:
        return False

    # continuous — разрыв без лишней страницы; СОДЕРЖАНИЕ само несёт pageBreakBefore.
    pPr = elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        elem.insert(0, pPr)

    # Не создаём дубликаты: если sectPr уже есть — пропускаем.
    if pPr.find(qn("w:sectPr")) is not None:
        return True

    # Проверяем: есть ли содержательные абзацы между концом титула и СОДЕРЖАНИЕ?
    # Если нет — вставляем пустую страницу (стр.2), чтобы TOC оказался на стр.3.
    paragraphs = doc.paragraphs
    toc_para_idx: int | None = None
    for _pi, _p in enumerate(paragraphs):
        if _p._p is heading._p:
            toc_para_idx = _pi
            break

    # Когда между концом титула и СОДЕРЖАНИЕ нет содержательных абзацев —
    # вставляем пустую страницу (стр.2) и используем nextPage-разрыв, чтобы
    # Section 1 (СОДЕРЖАНИЕ) гарантированно оказался на стр.3.
    # Если содержательный контент есть (задание и т.п.) — оставляем continuous:
    # СОДЕРЖАНИЕ само несёт pageBreakBefore из исходного документа.
    blank_inserted = False
    if toc_para_idx is not None:
        non_empty_between = sum(
            1 for _p in paragraphs[title_end_index:toc_para_idx]
            if (_p.text or "").strip()
        )
        if non_empty_between == 0:
            blank_p = OxmlElement("w:p")
            blank_r = OxmlElement("w:r")
            blank_br = OxmlElement("w:br")
            blank_br.set(qn("w:type"), "page")
            blank_r.append(blank_br)
            blank_p.append(blank_r)
            elem.addprevious(blank_p)
            blank_inserted = True

    sectPr = OxmlElement("w:sectPr")
    _sp = OxmlElement("w:type")
    # nextPage: разрыв секции сам переносит Section 1 на новую страницу.
    # continuous: Section 1 начинается на той же странице; СОДЕРЖАНИЕ полагается
    # на собственный pageBreakBefore из исходного документа.
    _sp.set(qn("w:val"), "nextPage" if blank_inserted else "continuous")
    sectPr.append(_sp)

    # Копируем pgSz/pgMar из главного sectPr, чтобы не сломать геометрию секции 0.
    try:
        main_sectPr = doc.element.body.find(qn("w:sectPr"))
        if main_sectPr is not None:
            for _tag in ("w:pgSz", "w:pgMar"):
                _el = main_sectPr.find(qn(_tag))
                if _el is not None:
                    sectPr.append(_copy.deepcopy(_el))
    except Exception:
        pass

    pPr.append(sectPr)
    return True


def apply_footer_page_number(doc: DocxDocument, with_blank_page_section: bool = False) -> None:
    """Вставляет центрированное PAGE-поле в нижние колонтитулы секций.

    with_blank_page_section=False (по умолчанию):
      Одна секция. Механизм "Different First Page": стр. 1 (титул) без номера,
      стр. 2+ показывают PAGE-поле.

    with_blank_page_section=True:
      Две секции (insert_pretoc_section_break уже создал разрыв между ними):
        Section 0 (стр. 1+, титул + задание) — нижний колонтитул пустой.
        Section 1+ (СОДЕРЖАНИЕ и далее) — PAGE-поле, без titlePg.
      Нумерация Section 1 не сбрасывается — счётчик продолжается с Section 0.
    """
    for i, section in enumerate(doc.sections):
        # Верхний колонтитул — всегда пустой.
        section.header.is_linked_to_previous = False
        _clear_footer_runs(section.header)

        if with_blank_page_section and i == 0:
            # Section 0: титул + пустая страница. Нижний колонтитул полностью пустой.
            section.footer.is_linked_to_previous = False
            _clear_footer_runs(section.footer)
            _strip_sectpr_flags(section)
            # Явно фиксируем start=1: некоторые шаблоны несут start=0 и сдвигают
            # нумерацию всего документа на единицу.
            _ensure_page_number_start(section, start=1)
            continue

        if i == 0 and not with_blank_page_section:
            _ensure_page_number_start(section, start=1)

        # Нижний колонтитул с PAGE-полем.
        footer = section.footer
        footer.is_linked_to_previous = False
        _clear_footer_runs(footer)
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        _insert_page_field(paragraph)

        if i == 0 and not with_blank_page_section:
            # Режим одной секции: Different First Page скрывает номер на титуле.
            sectPr = section._sectPr
            for tag in ("w:evenAndOddHeaders",):
                for el in list(sectPr.findall(qn(tag))):
                    sectPr.remove(el)
            for tag in ("w:headerReference", "w:footerReference"):
                for ref in list(sectPr.findall(qn(tag))):
                    if ref.get(qn("w:type")) == "even":
                        sectPr.remove(ref)
            for el in list(sectPr.findall(qn("w:titlePg"))):
                sectPr.remove(el)
            sectPr.append(OxmlElement("w:titlePg"))

            fph = section.first_page_header
            fph.is_linked_to_previous = False
            _clear_footer_runs(fph)
            fpf = section.first_page_footer
            fpf.is_linked_to_previous = False
            _clear_footer_runs(fpf)
        else:
            # Section 1+ в режиме blank_page_section, или любая не-первая секция:
            # убираем titlePg и сброс нумерации, чтобы счётчик продолжался с Section 0.
            _strip_sectpr_flags(section)
            if with_blank_page_section:
                sectPr = section._sectPr
                for pg_num in list(sectPr.findall(qn("w:pgNumType"))):
                    sectPr.remove(pg_num)


def _apply_heading_style(paragraph, level: int) -> bool:
    normalized_level = max(1, min(level, 3))
    candidates = (
        f"Heading {normalized_level}",
        f"Heading{normalized_level}",
        f"Заголовок {normalized_level}",
        f"Заголовок{normalized_level}",
    )
    for name in candidates:
        try:
            paragraph.style = name
            return True
        except KeyError:
            continue

    target_id = f"heading{normalized_level}"
    for style in paragraph.part.document.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        style_id = (style.style_id or "").replace(" ", "").lower()
        style_name = (style.name or "").replace(" ", "").lower()
        if style_id == target_id or style_name in {target_id, f"заголовок{normalized_level}"}:
            paragraph.style = style
            return True
    return False


def _clear_heading_indent_overrides(paragraph) -> None:
    """Очищает прямые XML-атрибуты left/hanging-отступа у заголовка."""
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return
    ind = pPr.find(qn("w:ind"))
    if ind is not None:
        for attr in ("w:left", "w:start", "w:hanging", "w:right", "w:end"):
            if ind.get(qn(attr)) is not None:
                del ind.attrib[qn(attr)]


def _remove_numpr(paragraph) -> None:
    """Удаляет w:numPr из pPr абзаца.

    Word рисует номер из numPr + TAB из numbering.xml — этим не управляешь из python-docx.
    Перед вызовом нужно перенести номер в текст через _ensure_number_in_text.
    """
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        pPr.remove(numPr)


_HAS_LEADING_NUMBER_RE = re.compile(r"^\s*\d+(?:\.\d+)*")


def _ensure_number_in_text(paragraph, derived_number: str | None) -> None:
    """Добавляет "{derived_number} " в начало текста, если числового префикса ещё нет.

    Нужно для заголовков с numPr: после _remove_numpr Word-номер пропадает,
    поэтому записываем его в первый run обычными символами.
    """
    if not derived_number:
        return
    if _HAS_LEADING_NUMBER_RE.match(paragraph.text or ""):
        return
    new_text = f"{derived_number} {(paragraph.text or '').lstrip()}"
    overwrite_paragraph_text(paragraph, new_text)


# Зазор до/после заголовка — реальный space before/after (а не имитация через
# line_spacing), одинаковый на всех уровнях — раздел, подраздел и пункт.
_HEADING_SPACING_PT = 6.0


def _force_heading_direct_format(paragraph, structural: bool, level: int = 1) -> None:
    """Применяет визуальное оформление заголовка напрямую, перекрывая стиль Word."""
    # Только сырые отступы — numPr оставляем, иначе авто-нумерация сломается.
    _clear_heading_indent_overrides(paragraph)

    pf = paragraph.paragraph_format
    # Межстрочный интервал заголовков — 1.5, как и во всём остальном тексте
    # (даже у раздела 1-го уровня — двойной интервал не используется).
    pf.line_spacing = 1.5
    _force_paragraph_spacing(
        paragraph, before_pt=_HEADING_SPACING_PT, after_pt=_HEADING_SPACING_PT
    )
    pf.widow_control = True
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)

    # Все заголовки строго по центру. Отступ первой строки не ставим — при центре выглядит криво.
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_alignment_xml(paragraph, "center")
    pf.first_line_indent = Cm(0)

    # Раздел (1) и подраздел (1.1) — полужирные; пункт (1.1.1) и глубже — обычным начертанием.
    bold = structural or level <= 2
    for run in paragraph.runs:
        run.font.bold = bold
        if not bold:
            _force_run_not_bold(run)
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"
        run.font.all_caps = False
        run.font.color.rgb = RGBColor(0, 0, 0)


def _apply_heading_fallback(paragraph, level: int, structural: bool) -> None:
    _force_heading_direct_format(paragraph, structural, level)


def _strip_paragraph_outline_level(paragraph) -> None:
    """Удаляет w:outlineLvl из pPr. Без этого body-абзацы с outlineLvl=0 из шаблона попадают в TOC."""
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return
    for olvl in list(pPr.findall(qn("w:outlineLvl"))):
        pPr.remove(olvl)


def _set_outline_level_xml(paragraph, level: int) -> None:
    """Записывает w:outlineLvl прямо в pPr.

    LO может сбросить стиль "Heading N" при пересохранении, но outlineLvl сохраняется —
    TOC-поле с флагом \\u подхватит заголовки по нему независимо от стиля.
    """
    normalized = max(1, min(level, 3))
    pPr = paragraph._p.get_or_add_pPr()
    for existing in list(pPr.findall(qn("w:outlineLvl"))):
        pPr.remove(existing)
    olvl = OxmlElement("w:outlineLvl")
    olvl.set(qn("w:val"), str(normalized - 1))
    pPr.append(olvl)


def _set_page_break_before(paragraph) -> None:
    """Устанавливает pageBreakBefore прямо в XML (надёжнее, чем API python-docx)."""
    pPr = paragraph._p.get_or_add_pPr()
    for existing in list(pPr.findall(qn("w:pageBreakBefore"))):
        pPr.remove(existing)
    pbr = OxmlElement("w:pageBreakBefore")
    pbr.set(qn("w:val"), "true")
    pPr.append(pbr)


def _strip_page_break_before(paragraph) -> None:
    """Удаляет w:pageBreakBefore из pPr."""
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return
    for pbr in list(pPr.findall(qn("w:pageBreakBefore"))):
        pPr.remove(pbr)


def _has_page_break_element(p_elem) -> bool:
    """True, если xml-элемент <w:p> уже вызывает разрыв страницы.

    Проверяет:
    - явный <w:br w:type='page'/> внутри параграфа,
    - <w:sectPr> внутри pPr параграфа (конец раздела = новая страница).
    """
    for br in p_elem.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
        return True
    return False


def _set_alignment_xml(paragraph, w_val: str) -> None:
    """Жёстко задаёт выравнивание абзаца через XML `w:jc`, обходя возможное
    наследование стиля, способное перекрыть alignment-сеттер python-docx.

    Значения w_val:
      "center"     — по центру (структурные заголовки)
      "both"       — по ширине (нумерованные заголовки, по предпочтению)
      "distribute" — распределение по символам (растягивает и одиночные строки)
    """
    pPr = paragraph._p.get_or_add_pPr()
    for existing in list(pPr.findall(qn("w:jc"))):
        pPr.remove(existing)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), w_val)
    pPr.append(jc)


_CAPS_HEADINGS = {
    "введение",
    "заключение",
    "содержание",
    "список использованных источников",
    "список литературы",
    "список использованной литературы",
    "список использованных источников и литературы",
    "список используемой литературы",
    "список используемых источников",
    "библиографический список",
    "реферат",
    "приложение",
    # Термины и определения / сокращения — те же структурные правила:
    # по центру, заглавными, страницу с новой страницы.
    "термины и определения",
    "перечень сокращений и обозначений",
    "список сокращений и условных обозначений",
    "перечень сокращений",
}


def _is_structural(text: str) -> bool:
    normalised = " ".join(text.strip().lower().split())
    return normalised in _CAPS_HEADINGS or any(normalised.startswith(k) for k in _CAPS_HEADINGS)


# Лейбл приложения («ПРИЛОЖЕНИЕ А», «Приложение Б», «ПРИЛОЖЕНИЕ А (обязательное)»).
_APPENDIX_LABEL_RE = re.compile(r"^\s*приложение\b", re.IGNORECASE)
# Строка-маркер обязательности под лейблом: «(обязательное)», «справочное» …
_APPENDIX_MARKER_RE = re.compile(
    r"^\s*\(?\s*(обязательн|справочн|рекомендуем|информацион)", re.IGNORECASE
)


def style_appendix_titles(doc: DocxDocument) -> None:
    """Строка после «ПРИЛОЖЕНИЕ X» — название приложения: по центру + полужирная.

    По требованию заказчика: НЕ важно, подпись ли это листинга/рисунка/таблицы
    или обычный текст — первый непустой абзац после лейбла приложения (и
    необязательная строка-маркер «(обязательное)») всегда оформляется как
    название приложения (TNR 14, по центру, полужирный, без outlineLvl —
    в Содержание не попадает).

    Поздний проход: вызывается ПОСЛЕ renumber_listings/figures/tables, чтобы
    перебить их левое выравнивание подписи, ставшей названием приложения.
    """
    paragraphs = doc.paragraphs
    n = len(paragraphs)

    def _next_nonempty(start: int) -> int:
        k = start
        while k < n and not paragraphs[k].text.strip():
            k += 1
        return k

    def _style(idx: int) -> None:
        para = paragraphs[idx]
        _strip_paragraph_outline_level(para)
        # Маркер/название сами могли попасть в apply_headings как отдельный
        # заголовок (например, если студент набрал их стилем "Заголовок 1") и
        # получить свой pageBreakBefore — это уводит их на страницу ПОСЛЕ
        # лейбла «ПРИЛОЖЕНИЕ X», хотя они должны идти сразу за ним.
        _strip_page_break_before(para)
        _force_heading_direct_format(para, True, 1)  # center + bold + TNR14

    def _style_title(idx: int) -> None:
        """Стилизует название (не маркер): дополнительно прописная первая буква."""
        para = paragraphs[idx]
        _strip_paragraph_outline_level(para)
        _strip_page_break_before(para)
        _force_heading_direct_format(para, True, 1)
        for run in para.runs:
            if run.text:
                if run.text[0].islower():
                    run.text = run.text[0].upper() + run.text[1:]
                break

    for i, p in enumerate(paragraphs):
        t = p.text.strip()
        if not t or len(t) >= 40 or not _APPENDIX_LABEL_RE.match(t):
            continue  # это не лейбл приложения (длинная фраза/проза)
        # Лейбл «ПРИЛОЖЕНИЕ А» — по правому краю (Порядок МИИГАиК п.5.8.1).
        _set_alignment_xml(p, "right")
        j = _next_nonempty(i + 1)
        if j >= n:
            continue
        # Необязательный маркер «(обязательное)/(справочное)» — тоже по центру.
        if _APPENDIX_MARKER_RE.match(paragraphs[j].text.strip()):
            _style(j)
            j = _next_nonempty(j + 1)
            if j >= n:
                continue
        nxt = paragraphs[j].text.strip()
        if not nxt or _is_structural(nxt):
            continue  # другой структурный заголовок — у приложения нет названия
        _style_title(j)


def _sentence_case_words(words: list[str]) -> list[str]:
    """Регистр предложения: первое слово с заглавной, остальные строчные.

    Одиночные заглавные буквы сохраняются ТОЛЬКО когда они идут после
    "приложение"/"часть"/"раздел" — это идентификаторы приложений ("А",
    "Б", "В"). Иначе одиночная "И", "А", "В" — это союз/предлог, его
    надо опустить в строчную (иначе "Теория И технологии" остаётся с
    заглавной "И" после нормализации заголовка из CAPS).
    """
    identifier_anchors = {"приложение", "часть", "раздел"}
    result: list[str] = []
    for i, word in enumerate(words):
        if i == 0:
            result.append(word[0].upper() + word[1:].lower() if word else word)
            continue
        prev = words[i - 1].lower() if i > 0 else ""
        is_identifier_after_anchor = (
            len(word) == 1
            and word.isalpha()
            and prev in identifier_anchors
        )
        if is_identifier_after_anchor:
            result.append(word.upper())  # "Приложение А" — А остаётся заглавной
        else:
            result.append(word.lower())
    return result


def _strip_chapter_word(text: str) -> str:
    """Преобразует 'Глава 1 Название' → '1 Название', 'Раздел 2.1' → '2.1' и т.п.

    Точка после номера не ставится — по ГОСТ между номером заголовка и
    текстом просто пробел.
    """
    stripped = text.strip()
    prefix = parse_heading_prefix(stripped)
    if prefix is None or prefix.kind != "chapter_word":
        return text
    number = prefix.number
    rest = prefix.title.strip()
    return f"{number} {rest}" if rest else number


def _normalise_heading_text(text: str) -> str:
    text = _lstrip_heading(text)
    text = _strip_chapter_word(text)
    text = _normalise_number_prefix(text)
    # Сворачиваем все внутренние пробельные последовательности (двойные
    # пробелы, забредшие табы или NBSP между словами) в один ASCII-пробел.
    text = re.sub(r"\s+", " ", text)
    stripped = text.rstrip(".")
    if _APPENDIX_LABEL_RE.match(stripped):
        # «Приложение А», не «ПРИЛОЖЕНИЕ А» — капс здесь не по норме.
        return " ".join(_sentence_case_words(stripped.split()))
    if _is_structural(stripped):
        return stripped.upper()

    # Отделяем ведущий числовой префикс ("1 ", "1.2 ") от слов.
    prefix_match = _NUM_PREFIX_RE.match(stripped)
    prefix = prefix_match.group(1) if prefix_match else ""
    rest = stripped[len(prefix):]

    if not rest:
        return stripped

    letters = [c for c in rest if c.isalpha()]
    if not letters:
        return stripped

    if all(c.isupper() for c in letters):
        # Слова целиком в верхнем регистре → регистр предложения.
        rest = " ".join(_sentence_case_words(rest.split()))
    elif rest[0].islower():
        # Начинается со строчной → делаем заглавной только первую букву.
        rest = rest[0].upper() + rest[1:]

    return prefix + rest


def apply_headings(
    doc: DocxDocument,
    headings: list[HeadingCandidate],
    stats: DocStats,
) -> set[int]:
    applied: set[int] = set()
    for candidate in headings:
        ps = stats.paragraphs[candidate.paragraph_index]
        if ps.index >= len(doc.paragraphs):
            continue
        paragraph = doc.paragraphs[candidate.paragraph_index]

        # numPr рисует номер + TAB, этим не управлять из python-docx — переводим в текст.
        if ps.numpr_ilvl is not None and ps.derived_number:
            _ensure_number_in_text(paragraph, ps.derived_number)
        _remove_numpr(paragraph)

        new_text = _normalise_heading_text(paragraph.text)
        if new_text != paragraph.text:
            overwrite_paragraph_text(paragraph, new_text)

        structural = _is_structural(new_text)
        if not _apply_heading_style(paragraph, candidate.level):
            _apply_heading_fallback(paragraph, candidate.level, structural)

        # Перекрываем встроенные стили Word (Calibri 16pt и т.д.) безусловно.
        _force_heading_direct_format(paragraph, structural, candidate.level)
        # outlineLvl прямо в pPr — страховка на случай если LO сбросит стиль Heading N.
        _set_outline_level_xml(paragraph, candidate.level)

        # Заголовки уровня 1 всегда начинаются с новой страницы.
        # pageBreakBefore сохраняем для совместимости с Word.
        if candidate.level == 1:
            _set_page_break_before(paragraph)

        applied.add(candidate.paragraph_index)

    return applied


def insert_l1_page_breaks(
    doc: DocxDocument,
    headings: list[HeadingCandidate],
    applied: set[int],
) -> None:
    """Вставляет явные разрывы страниц перед L1-заголовками.

    Вызывается ПОСЛЕ apply_body_style, чтобы вставка новых параграфов
    не сдвигала индексы, по которым apply_body_style обращается к stats.
    """
    l1_candidates = sorted(
        [c for c in headings if c.level == 1 and c.paragraph_index in applied],
        key=lambda c: c.paragraph_index,
        reverse=True,
    )
    for candidate in l1_candidates:
        idx = candidate.paragraph_index
        paragraph = doc.paragraphs[idx]
        prev_p = paragraph._p.getprevious()
        # pageBreakBefore уже есть (apply_headings), но LO иногда его теряет.
        # Явный <w:br type=page> в конец предыдущего абзаца — страховка без лишней пустой строки.
        if prev_p is None or prev_p.tag != qn("w:p"):
            continue
        if _has_page_break_element(prev_p):
            continue
        # Предыдущий абзац — лейбл «Приложение X» или маркер «(обязательное)»:
        # текущий кандидат на самом деле название приложения (свой pageBreakBefore
        # снимет style_appendix_titles позже). Разрыв здесь не нужен — иначе
        # название отрывается от лейбла на следующую страницу.
        prev_text = "".join(t.text or "" for t in prev_p.iter(qn("w:t"))).strip()
        if _APPENDIX_LABEL_RE.match(prev_text) or _APPENDIX_MARKER_RE.match(prev_text):
            continue
        _append_page_break_run(prev_p)


def _preserve_list_number_in_text(paragraph, derived_number: str | None) -> None:
    """Переносит авто-номер Word в начало текста как "N. ".

    _clear_paragraph_indent_overrides снимает numPr у всех body-абзацев,
    и Word-номер пропадает — переносим его заранее.
    Если есть гиперссылки — prefixим только первый w:r, не трогая их.
    """
    if not derived_number:
        return
    if _HAS_LEADING_NUMBER_RE.match(paragraph.text or ""):
        return
    p = paragraph._p
    hyperlinks = p.findall(qn("w:hyperlink"))
    if hyperlinks:
        # Находим первый w:r напрямую в w:p (не внутри гиперссылок).
        first_r = p.find(qn("w:r"))
        if first_r is not None:
            t = first_r.find(qn("w:t"))
            if t is not None:
                t.text = f"{derived_number}. {(t.text or '').lstrip()}"
            else:
                from lxml import etree
                t_el = etree.SubElement(first_r, qn("w:t"))
                t_el.text = f"{derived_number}. "
        else:
            # Нет ни одного run-а вне гиперссылок — добавляем новый run в начало.
            from lxml import etree
            new_r = etree.Element(qn("w:r"))
            new_t = etree.SubElement(new_r, qn("w:t"))
            new_t.text = f"{derived_number}. "
            p.insert(0, new_r)
        return
    new_text = f"{derived_number}. {(paragraph.text or '').lstrip()}"
    overwrite_paragraph_text(paragraph, new_text)


def _clear_paragraph_indent_overrides(paragraph) -> None:
    """Удаляет numPr и left/hanging-отступы из pPr абзаца.

    §6.2: у каждого body-абзаца отступ первой строки 1.25 см.
    Вызывать ПОСЛЕ _preserve_list_number_in_text — иначе Word-номер пропадёт.
    """
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return
    for numPr in list(pPr.findall(qn("w:numPr"))):
        pPr.remove(numPr)
    # Снимаем outlineLvl — иначе body-абзац с outlineLvl=0 из шаблона попадёт в TOC.
    for olvl in list(pPr.findall(qn("w:outlineLvl"))):
        pPr.remove(olvl)
    ind = pPr.find(qn("w:ind"))
    if ind is not None:
        for attr in ("w:left", "w:start", "w:hanging", "w:right", "w:end"):
            if ind.get(qn(attr)) is not None:
                del ind.attrib[qn(attr)]


def _force_paragraph_spacing(paragraph, before_pt: float = 0.0, after_pt: float = 0.0) -> None:
    """Патчит w:spacing напрямую и снимает autospacing-флаги.

    Python-docx сеттер не перебивает beforeAutospacing/afterAutospacing из стиля.
    """
    # 1 pt = 20 twips
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(int(round(before_pt * 20))))
    spacing.set(qn("w:after"), str(int(round(after_pt * 20))))
    spacing.attrib.pop(qn("w:beforeAutospacing"), None)
    spacing.attrib.pop(qn("w:afterAutospacing"), None)


def _effective_spacing_pt(paragraph, attr: str) -> float:
    """Возвращает space_before/space_after в pt с учётом цепочки стилей."""
    direct = getattr(paragraph.paragraph_format, attr)
    if direct is not None:
        return direct.pt
    style = getattr(paragraph, "style", None)
    while style is not None:
        val = getattr(style.paragraph_format, attr, None)
        if val is not None:
            return val.pt
        style = getattr(style, "base_style", None)
    return 0.0


def check_paragraph_spacing(
    doc: DocxDocument,
    stats: DocStats,
    skip_indexes: set[int],
    heading_indexes: set[int],
    rules: dict,
) -> list[PipelineViolation]:
    """Обнаруживает body-абзацы с неверным интервалом, отступом или alignment.

    По правилу §6.2 интервал между абзацами должен совпадать с межстрочным
    (добавочный space_before / space_after = 0). Проверяется эффективное
    значение — включая унаследованное из стиля абзаца (чаще всего именно
    так задаётся лишнее расстояние, например Word-дефолтные 8 pt after).
    Заголовки и настоящие подписи исключены.
    """
    style = rules.get("body_text_style", {})
    target_before = style.get("space_before_pt", 0)
    target_after = style.get("space_after_pt", 0)
    target_indent = float(style.get("first_line_indent_cm", 1.25))
    target_alignment = paragraph_alignment_from_rule(style.get("alignment"))

    offending: set[int] = set()
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx in skip_indexes or idx in heading_indexes:
            continue
        if idx >= len(stats.paragraphs):
            continue
        ps = stats.paragraphs[idx]
        if ps.in_table or not ps.stripped:
            continue
        if _is_body_style_exempt_caption(paragraph):
            continue

        before_pt = _effective_spacing_pt(paragraph, "space_before")
        after_pt = _effective_spacing_pt(paragraph, "space_after")
        if abs(before_pt - target_before) > 0.1 or abs(after_pt - target_after) > 0.1:
            offending.add(idx)
        if (
            ps.first_line_indent_cm is not None
            and abs(ps.first_line_indent_cm - target_indent) > 0.05
        ):
            offending.add(idx)
        if ps.alignment is not None and ps.alignment != target_alignment:
            offending.add(idx)

    if not offending:
        return []

    return [
        PipelineViolation(
            type="paragraph_spacing",
            rule_reference="Общие требования",
            description=(
                f"Форматирование абзацев основного текста отличается от правил "
                f"(затронуто абзацев: {len(offending)}). "
                "Выравнивание, абзацный отступ и интервалы приведены к настройкам."
            ),
            status=ViolationStatus.auto_fixed,
            severity=SEVERITY_INFO,
        )
    ]


def apply_body_style(
    doc: DocxDocument,
    rules: dict,
    stats: DocStats,
    skip_indexes: set[int],
    heading_indexes: set[int],
) -> None:
    style = rules["body_text_style"]
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx in skip_indexes or idx in heading_indexes:
            continue
        if idx >= len(stats.paragraphs):
            continue
        ps = stats.paragraphs[idx]
        if ps.in_table:
            continue

        # Сбрасываем заголовочный стиль и outlineLvl ДО ранних continue —
        # иначе шаблонный outlineLvl=0 затолкает body-абзацы в TOC.
        style_lower = (paragraph.style.name or "").strip().lower() if paragraph.style else ""
        if style_lower.startswith(("heading", "заголовок", "подзаголовок")):
            try:
                paragraph.style = "Normal"
            except KeyError:
                pass
        _strip_paragraph_outline_level(paragraph)

        # §6.2: нормализуем интервал через XML-патч — сеттер не перебивает autospacing из стиля.
        _force_paragraph_spacing(
            paragraph,
            before_pt=style["space_before_pt"],
            after_pt=style["space_after_pt"],
        )

        if _is_body_style_exempt_caption(paragraph):
            continue  # шрифт/отступы/выравнивание подписей — в captions.py

        # Абзацы с рисунками центрируем — ГОСТ требует центрирования иллюстраций.
        # Проверка ПЕРЕД is_empty-skip: абзац может содержать только <w:drawing>
        # без текста, тогда ps.stripped пустой и early-continue сработает раньше.
        # _set_alignment_xml дополнительно к сеттеру — защита от наследования
        # justify из стиля абзаца, способного перебить прямое форматирование.
        if _has_visual_content(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_alignment_xml(paragraph, "center")
            continue

        if not ps.stripped:
            continue

        for run in paragraph.runs:
            run.font.name = style["font_name"]
            run.font.size = Pt(style["font_size_pt"])
            run.font.color.rgb = RGBColor.from_string(style["font_color"])
            # Снимаем bold из шаблона (жирные пункты списков).
            # _force_run_not_bold добавляет w:bCs val=0 для надёжности.
            run.font.bold = False
            _force_run_not_bold(run)

        # §6.2: сначала переносим numPr-номер в текст, затем снимаем отступы из шаблона.
        if ps.numpr_ilvl is not None and ps.derived_number:
            _preserve_list_number_in_text(paragraph, ps.derived_number)
        _clear_paragraph_indent_overrides(paragraph)
        paragraph.paragraph_format.line_spacing = style["line_spacing"]
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.right_indent = Cm(0)
        paragraph.paragraph_format.first_line_indent = Cm(style["first_line_indent_cm"])
        paragraph.alignment = paragraph_alignment_from_rule(style.get("alignment"))


def _has_visual_content(para) -> bool:
    """Возвращает True, если абзац содержит изображение или встроенный объект."""
    p = para._element
    return bool(
        p.findall(".//" + qn("w:drawing"))
        or p.findall(".//" + qn("w:pict"))
        or p.findall(".//" + qn("w:object"))
    )


def _has_explicit_page_break(para) -> bool:
    """Возвращает True, если абзац содержит явный разрыв страницы (<w:br type=page>)."""
    for br in para._element.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def _has_section_break(p_elem) -> bool:
    """True, если в pPr абзаца есть <w:sectPr> (граница раздела). Такой
    абзац нельзя удалять — слияние разделов меняет поля/ориентацию."""
    pPr = p_elem.find(qn("w:pPr"))
    return pPr is not None and pPr.find(qn("w:sectPr")) is not None


def _append_page_break_run(p_elem) -> None:
    """Добавляет в конец абзаца run с <w:br w:type='page'/>. Содержимое
    абзаца не меняется — разрыв ставится ПОСЛЕ его текста, поэтому
    следующий абзац начинается с новой страницы без пустой строки."""
    new_rPr = OxmlElement("w:rPr")
    new_rFonts = OxmlElement("w:rFonts")
    new_rFonts.set(qn("w:ascii"), "Times New Roman")
    new_rFonts.set(qn("w:hAnsi"), "Times New Roman")
    new_rPr.append(new_rFonts)
    new_r = OxmlElement("w:r")
    new_r.append(new_rPr)
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    new_r.append(br)
    p_elem.append(new_r)


def _force_run_not_bold(run) -> None:
    """Явно снимает w:b и w:bCs val="0". bold=False иногда не трогает bCs."""
    rpr = run._r.get_or_add_rPr()
    for tag in ("w:b", "w:bCs"):
        el = rpr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rpr.append(el)
        el.set(qn("w:val"), "0")


# Подпись рисунка: "Рисунок N[.M] – <текст>". Используется для гарантии
# одной пустой строки после подписи (ТЗ §7.3, требование заказчика).
_FIGURE_CAPTION_RE = re.compile(
    r"^\s*рисунок\s+\d+(?:\.\d+)*\s*[–—\-]\s*\S", re.IGNORECASE
)


def _make_empty_paragraph():
    """Пустой <w:p> с Times New Roman — разделитель после таблиц/подписей."""
    new_p = OxmlElement("w:p")
    new_rPr = OxmlElement("w:rPr")
    new_rFonts = OxmlElement("w:rFonts")
    new_rFonts.set(qn("w:ascii"), "Times New Roman")
    new_rFonts.set(qn("w:hAnsi"), "Times New Roman")
    new_rPr.append(new_rFonts)
    pPr = OxmlElement("w:pPr")
    pPr.append(new_rPr)
    new_p.append(pPr)
    return new_p


def remove_empty_paragraphs(doc: DocxDocument, skip_indexes: set[int]) -> int:
    """Удаляет пустые абзацы за пределами skip_indexes. Абзацы только с
    изображением или явным разрывом страницы сохраняются.

    Особый случай: пустой абзац, идущий непосредственно за таблицей,
    сохраняется. Без этого разделителя соседние <w:tbl> склеиваются:
    Word рендерит их как одну таблицу с рваной шириной столбцов, а
    подпись следующей таблицы прибивается к предыдущей. Сохраняем
    только ПЕРВЫЙ пустой абзац после таблицы — повторные пустые
    дальше всё равно удаляем.
    """
    to_remove = []
    for idx, para in enumerate(doc.paragraphs):
        if idx in skip_indexes:
            continue
        if para.text.strip():
            continue
        if _has_visual_content(para):
            continue
        elem = para._element
        # Абзац с границей раздела не трогаем — иначе склеятся разделы.
        if _has_section_break(elem):
            continue
        prev = elem.getprevious()
        if prev is not None:
            prev_tag = prev.tag.split("}")[-1] if "}" in prev.tag else prev.tag
            if prev_tag == "tbl":
                # первый пустой абзац после таблицы — оставляем как разделитель
                continue
        if _has_explicit_page_break(para):
            # Пустой абзац-разрыв: сам по себе это лишняя пустая строка
            # (заголовки уже несут pageBreakBefore). Перенос страницы
            # сохраняем, перекладывая <w:br> в конец предыдущего абзаца;
            # если переложить некуда — оставляем абзац как есть.
            nxt = elem.getnext()
            nxt_breaks = (
                nxt is not None
                and nxt.tag == qn("w:p")
                and _has_page_break_element(nxt)
            )
            if not nxt_breaks:
                if prev is not None and prev.tag == qn("w:p"):
                    if not _has_page_break_element(prev):
                        _append_page_break_run(prev)
                else:
                    # переносить разрыв некуда — не удаляем абзац
                    continue
        to_remove.append(elem)
    for elem in to_remove:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
    return len(to_remove)


def ensure_separator_after_tables_and_figures(
    doc: DocxDocument, skip_indexes: set[int]
) -> int:
    """Гарантирует РОВНО одну пустую строку после каждой таблицы и после
    каждой подписи рисунка (ТЗ §7.3, требование заказчика).

    Запускается ПОСЛЕ remove_empty_paragraphs: лишние пустые уже убраны,
    поэтому здесь только дозаполняем недостающие разделители. После
    таблиц remove_empty_paragraphs уже сохраняет первый пустой абзац —
    добавляем разделитель лишь там, где его нет (две таблицы подряд,
    таблица перед заголовком и т.п.).
    """
    body = doc.element.body
    added = 0

    # 1. После таблиц.
    for tbl in list(body.findall(qn("w:tbl"))):
        nxt = tbl.getnext()
        if nxt is not None and nxt.tag == qn("w:p"):
            text = "".join(nxt.itertext()).strip()
            if not text and not _has_section_break(nxt):
                continue  # разделитель уже есть
        tbl.addnext(_make_empty_paragraph())
        added += 1

    # 2. После подписей рисунков.
    for idx, para in enumerate(doc.paragraphs):
        if idx in skip_indexes:
            continue
        if not _FIGURE_CAPTION_RE.match(para.text or ""):
            continue
        nxt = para._element.getnext()
        if nxt is not None and nxt.tag == qn("w:p"):
            text = "".join(nxt.itertext()).strip()
            if not text and not _has_section_break(nxt):
                continue  # разделитель уже есть
        para._element.addnext(_make_empty_paragraph())
        added += 1

    return added


def strip_chapter_word_everywhere(doc: DocxDocument, skip_indexes: set[int]) -> int:
    """Финальный защитный проход: переписывает любой короткий абзац,
    начинающийся с 'Глава N' / 'Раздел N' / 'Chapter N', в '<N> <остальное>'.

    Запускается после apply_headings — так если какой-то заголовок проскочил
    мимо детектора (или структура его run'ов помешала обычной перезаписи),
    префикс-слово всё равно срезается. Ограничено длиной абзаца, чтобы
    обычные предложения вида 'Глава 2 была посвящена…' не задеть.

    Попутно выставляет outlineLvl=0, чтобы _chapter_breaks_from_xml нашёл
    эти заголовки даже если apply_headings их пропустил.
    """
    changed = 0
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx in skip_indexes:
            continue
        text = paragraph.text
        stripped = text.strip()
        if not stripped or len(stripped) > 140:
            continue
        prefix = parse_heading_prefix(stripped)
        if prefix is None or prefix.kind != "chapter_word":
            continue
        new_text = _strip_chapter_word(stripped)
        if new_text != stripped:
            overwrite_paragraph_text(paragraph, new_text)
            changed += 1
        # Гарантируем outlineLvl=0 — иначе _chapter_breaks_from_xml пропустит
        # заголовки, которые apply_headings не обработал.
        pPr = paragraph._p.find(qn("w:pPr"))
        already_set = (
            pPr is not None
            and pPr.find(qn("w:outlineLvl")) is not None
            and pPr.find(qn("w:outlineLvl")).get(qn("w:val")) == "0"
        )
        if not already_set:
            _set_outline_level_xml(paragraph, 1)
    return changed


_FORMULA_NUM_TAIL_RE = re.compile(r"\s*\((\d+(?:\.\d+)?)\)\s*$")
_MM_TO_TWIPS = 1440 / 25.4


def fix_formula_number_alignment(
    doc: DocxDocument,
    formulas: list,
    rules: dict,
) -> int:
    """Выставляет правый таб-стоп и вставляет <w:tab/> перед (N) в формулах.

    Затрагивает только формулы с номером, у которых has_tab=False.
    Возвращает количество исправленных абзацев.
    """
    from copy import deepcopy

    page = rules.get("page_settings", {})
    paper = page.get("paper", "A4")
    width_mm, _ = _PAGE_SIZES_MM.get(paper, _PAGE_SIZES_MM["A4"])
    left_mm = page.get("margin_left_mm", 30)
    right_mm = page.get("margin_right_mm", 15)
    text_width_twips = int((width_mm - left_mm - right_mm) * _MM_TO_TWIPS)

    fixed = 0
    for formula in formulas:
        if formula.number is None or formula.has_tab:
            continue
        if formula.paragraph_index >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[formula.paragraph_index]

        # Добавляем правый таб-стоп, если его ещё нет.
        pPr = para._p.get_or_add_pPr()
        tabs_elem = pPr.find(qn("w:tabs"))
        if tabs_elem is None:
            tabs_elem = OxmlElement("w:tabs")
            pPr.append(tabs_elem)
        has_right = any(
            t.get(qn("w:val")) == "right" for t in tabs_elem.findall(qn("w:tab"))
        )
        if not has_right:
            tab_stop = OxmlElement("w:tab")
            tab_stop.set(qn("w:val"), "right")
            tab_stop.set(qn("w:pos"), str(text_width_twips))
            tabs_elem.append(tab_stop)

        # Перестраиваем текст: «текст формулы <w:tab/> (N)»
        full_text = para.text
        m = _FORMULA_NUM_TAIL_RE.search(full_text)
        if m is None:
            continue
        text_before = full_text[: m.start()].rstrip()
        num_str = m.group(1)

        # Сохраняем rPr первого run'а для переноса форматирования.
        runs = para.runs
        first_rpr = deepcopy(runs[0]._r.find(qn("w:rPr"))) if runs else None

        # Удаляем все run'ы абзаца.
        for r in para._p.findall(qn("w:r")):
            para._p.remove(r)

        def _make_run(
            rpr,
            text_value: str | None = None,
            is_tab: bool = False,
        ):
            r = OxmlElement("w:r")
            if rpr is not None:
                r.append(deepcopy(rpr))
            if is_tab:
                r.append(OxmlElement("w:tab"))
            else:
                t = OxmlElement("w:t")
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                t.text = text_value or ""
                r.append(t)
            return r

        if text_before:
            para._p.append(_make_run(first_rpr, text_before))
        para._p.append(_make_run(first_rpr, is_tab=True))
        para._p.append(_make_run(first_rpr, f"({num_str})"))

        fixed += 1
    return fixed


def normalize_spaces(
    doc: DocxDocument,
    skip_indexes: set[int],
    heading_indexes: set[int],
) -> None:
    """Сворачивает повторные пробелы и убирает ведущие/хвостовые пробелы в
    абзацах тела документа."""
    for idx, para in enumerate(doc.paragraphs):
        if idx in skip_indexes or idx in heading_indexes:
            continue
        runs = para.runs
        for i, run in enumerate(runs):
            if not run.text:
                continue
            text = re.sub(r" {2,}", " ", run.text)
            # Ведущий пробел убираем только у первого run'а.
            if i == 0:
                text = text.lstrip(" ")
            # Хвостовой пробел убираем только у последнего run'а.
            if i == len(runs) - 1:
                text = text.rstrip(" ")
            run.text = text
