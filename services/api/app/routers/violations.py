from __future__ import annotations

import asyncio
import copy
import logging
import re
import shutil
import tempfile
from pathlib import Path
from uuid import UUID

from docx import Document as WordDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, RGBColor
from docx.shared import Pt as DocxPt
from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, status
from pydantic import BaseModel
from sqlalchemy import delete as sa_delete
from sqlalchemy import select as sa_select
from sqlalchemy import update as sa_update
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from services.api.app.dependencies import get_current_user
from services.core.vkr_core.db.session import AsyncSessionLocal, get_db
from services.core.vkr_core.engine.formatter import (
    overwrite_paragraph_text,
    paragraph_alignment_from_rule,
)
from services.core.vkr_core.engine.toc import rebuild_toc_full
from services.core.vkr_core.engine.violation_labels import violation_type_label
from services.core.vkr_core.models import (
    Document,
    ProcessingReport,
    User,
    UserRole,
    ViolationRecord,
    ViolationStatus,
)
from services.core.vkr_core.schemas import (
    ViolationActionResponse,
    ViolationCommentRequest,
    ViolationCommentResponse,
    ViolationOut,
)
from services.core.vkr_core.services import (
    AccessDeniedError,
    NotFoundError,
    get_violation_or_404,
    list_violations,
)
from services.core.vkr_core.services.snapshot_service import create_snapshot
from services.core.vkr_core.services.storage_service import StorageService

_logger = logging.getLogger(__name__)


async def _regen_pdf_background(
    report_id: UUID,
    new_storage_path: str,
) -> None:
    # Поля оглавления уже пересчитаны (finalize_toc) до сохранения в storage,
    # поэтому здесь просто рендерим PDF из уже корректного .docx.
    storage = StorageService()
    local_path = storage.load_to_local_path(new_storage_path)
    try:
        from services.core.vkr_core.engine.preview import convert_docx_to_pdf
        new_pdf_local = await asyncio.to_thread(convert_docx_to_pdf, local_path)
        new_pdf_storage = storage.save_file(new_pdf_local, "previews", ".pdf")
        shutil.rmtree(new_pdf_local.parent, ignore_errors=True)
        async with AsyncSessionLocal() as session:
            await session.execute(
                sa_update(ProcessingReport)
                .where(ProcessingReport.id == report_id)
                .values(processed_pdf_storage_path=new_pdf_storage)
            )
            await session.commit()
    except Exception:
        _logger.warning("Фоновая регенерация PDF не удалась", exc_info=True)
    finally:
        if str(local_path) != new_storage_path:
            local_path.unlink(missing_ok=True)


async def _replace_caption_violations(
    session: AsyncSession,
    report_id: UUID,
    changes: list[tuple[int, str, str]],
) -> None:
    """Заменяет caption_renumbered-записи отчёта актуальными после пересчёта глав."""
    await session.execute(
        sa_delete(ViolationRecord).where(
            ViolationRecord.report_id == report_id,
            ViolationRecord.type == "caption_renumbered",
        )
    )
    for para_idx, orig_text, new_text in changes:
        session.add(ViolationRecord(
            report_id=report_id,
            type="caption_renumbered",
            rule_reference="п. 6.12",
            severity="info",
            description="Сквозная нумерация по порядку появления в тексте (см. Было/Стало).",
            status=ViolationStatus.manual_required,
            paragraph_index=para_idx,
            original_text=orig_text,
            fixed_text=new_text,
        ))


def _needs_full_table_split(
    pre_baked_nums: list[int | None] | None,
    post_baked_nums: list[int | None] | None,
) -> bool:
    """Нужен ли полный проход резки таблиц после быстрого finalize.

    Сигнал: после rebuild+fast-finalize изменились запечённые номера TOC.
    """
    if pre_baked_nums is None or post_baked_nums is None:
        return False
    return pre_baked_nums != post_baked_nums


async def _finalize_and_preview_bg(
    report_id: UUID,
    document_id: UUID,
    storage_path: str,
    rules_config: dict | None = None,
    renum_captions: bool = False,
) -> None:
    """Фоновая финализация TOC + превью (мгновенный отклик на правку).

    Хендлер сохраняет docx БЕЗ дорогого рендера и сразу отвечает; здесь
    (после ответа) запекаем номера страниц TOC, фиксируем TNR и кладём
    готовый PDF как превью. Document.processed_storage_path обновляется на
    финализированный файл, чтобы скачивание/превью отдавали корректную версию.

    rules_config — если передан, перестраивает строки оглавления (rebuild_toc_full)
    перед финализацией. Это позволяет вынести rebuild из HTTP-обработчика в фон
    при операциях, которые не меняют количество абзацев (fix_heading_number).

    renum_captions — если True, после перестройки TOC пересчитывает подписи рисунков,
    таблиц и листингов по актуальным chapter_breaks (нужно после исправления H1).
    """
    # Stale-check: если пользователь успел применить следующий фикс до того,
    # как мы начали работу, наш storage_path уже устарел. Записывать поверх
    # более свежего результата нельзя — тихо выходим.
    async with AsyncSessionLocal() as _stale_session:
        _cur_doc = await _stale_session.get(Document, document_id)
        if _cur_doc is None or _cur_doc.processed_storage_path != storage_path:
            _logger.warning(
                "_finalize_and_preview_bg: устарело — cur_path=%s != storage_path=%s",
                _cur_doc.processed_storage_path if _cur_doc else "None",
                storage_path,
            )
            return

    # Загружаем текущие caption_renumbered записи (paragraph_index → orig_text)
    # до изменения docx, чтобы сохранить связь с оригинальным текстом документа.
    _orig_map: dict[int, str] = {}
    if renum_captions:
        async with AsyncSessionLocal() as _cv_session:
            _cv_rows = await _cv_session.scalars(
                sa_select(ViolationRecord).where(
                    ViolationRecord.report_id == report_id,
                    ViolationRecord.type == "caption_renumbered",
                )
            )
            for _row in _cv_rows.all():
                if _row.paragraph_index is not None and _row.original_text:
                    _orig_map[_row.paragraph_index] = _row.original_text

    storage = StorageService()
    local_path = storage.load_to_local_path(storage_path)
    _caption_changes: list[tuple[int, str, str]] = []
    try:
        _logger.warning("_finalize_and_preview_bg: запускаем finalize local_path=%s", local_path)
        pre_nums: list[int | None] | None = None
        if rules_config is not None:
            def _rebuild_and_read_sync(path: Path, rules: dict) -> list[int | None]:
                from docx import Document as _WD

                from services.core.vkr_core.engine.toc import _read_baked_toc_nums
                _doc = _WD(str(path))
                pre_nums = _read_baked_toc_nums(_doc)
                rebuild_toc_full(_doc, rules)
                _doc.save(str(path))
                return pre_nums
            pre_nums = await asyncio.to_thread(_rebuild_and_read_sync, local_path, rules_config)

        if renum_captions and rules_config is not None:
            _caption_changes = await asyncio.to_thread(
                _rerenumber_captions_sync, local_path, rules_config, _orig_map
            )

        from services.core.vkr_core.engine.toc import (
            _read_baked_toc_nums,
            finalize_toc_with_preview,
        )
        _ok, pdf = await asyncio.to_thread(
            finalize_toc_with_preview, local_path,
            skip_table_split=True,
            known_page_nums=None,
        )
        if rules_config is not None and pdf is not None:
            post_nums = await asyncio.to_thread(
                lambda: _read_baked_toc_nums(WordDocument(str(local_path)))
            )
            if _needs_full_table_split(pre_nums, post_nums):
                try:
                    shutil.rmtree(pdf.parent, ignore_errors=True)
                except Exception:
                    pass

                def _rebuild_only_sync(path: Path, rules: dict) -> None:
                    from docx import Document as _WD
                    _doc = _WD(str(path))
                    rebuild_toc_full(_doc, rules)
                    _doc.save(str(path))

                await asyncio.to_thread(_rebuild_only_sync, local_path, rules_config)
                _ok, pdf = await asyncio.to_thread(
                    finalize_toc_with_preview,
                    local_path,
                    skip_table_split=False,
                    known_page_nums=None,
                )
        if pdf is not None:
            # finalize изменил docx (запекание номеров + TNR) — сохраняем его
            # и обновляем путь; PDF переиспользуем как превью (без 2-го рендера).
            new_docx = storage.save_file(local_path, "processed", ".docx")
            try:
                pdf_storage = storage.save_file(pdf, "previews", ".pdf")
            finally:
                shutil.rmtree(pdf.parent, ignore_errors=True)
            async with AsyncSessionLocal() as session:
                # Повторный stale-check перед записью: за время финализации
                # мог прийти ещё один фикс.
                result = await session.execute(
                    sa_update(Document)
                    .where(
                        Document.id == document_id,
                        Document.processed_storage_path == storage_path,
                    )
                    .values(processed_storage_path=new_docx)
                    .returning(Document.id)
                )
                if result.scalar() is None:
                    _logger.warning(
                        "_finalize_and_preview_bg: stale при записи docx "
                        "(другой фикс успел изменить processed_storage_path), pdf пропущен"
                    )
                    await session.rollback()
                    return
                await session.execute(
                    sa_update(ProcessingReport)
                    .where(ProcessingReport.id == report_id)
                    .values(processed_pdf_storage_path=pdf_storage)
                )
                if _caption_changes:
                    await _replace_caption_violations(session, report_id, _caption_changes)
                await session.commit()
                _logger.info("_finalize_and_preview_bg: PDF обновлён → %s", pdf_storage)
        else:
            # Нет TOC — docx не менялся; просто рендерим превью.
            _logger.info("_finalize_and_preview_bg: TOC нет, рендерим превью напрямую")
            from services.core.vkr_core.engine.preview import convert_docx_to_pdf
            p = await asyncio.to_thread(convert_docx_to_pdf, local_path)
            try:
                pdf_storage = storage.save_file(p, "previews", ".pdf")
            finally:
                shutil.rmtree(p.parent, ignore_errors=True)
            async with AsyncSessionLocal() as session:
                await session.execute(
                    sa_update(ProcessingReport)
                    .where(ProcessingReport.id == report_id)
                    .values(processed_pdf_storage_path=pdf_storage)
                )
                if _caption_changes:
                    await _replace_caption_violations(session, report_id, _caption_changes)
                await session.commit()
                _logger.info("_finalize_and_preview_bg: PDF (no-TOC) обновлён → %s", pdf_storage)
    except Exception:
        _logger.warning("Фоновая финализация TOC/превью не удалась", exc_info=True)
    finally:
        if str(local_path) != storage_path:
            local_path.unlink(missing_ok=True)


async def _lock_document_for_fix(db: AsyncSession, document_id: UUID) -> None:
    """Захватывает row-lock на documents.id до конца текущей транзакции.

    Защищает от race-condition при параллельных ручных правках одного
    документа: без lock'а двое пользователей могут одновременно загрузить
    одну и ту же версию processed.docx, оба сохранить свои новые версии
    под разными UUID, и второй UPDATE затрёт `processed_storage_path` —
    правка первого пользователя физически останется в storage, но БД о ней
    больше не знает. С lock'ом транзакции сериализуются: B ждёт коммита A,
    затем читает свежий path и применяет правку поверх.

    Сразу после lock'а вызывающий код должен `await db.refresh(document)`,
    чтобы получить актуальные данные (предыдущая транзакция могла обновить
    processed_storage_path после того, как мы загрузили document первый раз).
    """
    await db.execute(
        sa_select(Document.id)
        .where(Document.id == document_id)
        .with_for_update()
    )


async def _snapshot_before_fix(
    document,
    report,
    action_type: str,
    action_description: str,
    db: AsyncSession,
) -> None:
    """Создаёт снимок состояния документа перед мутирующей правкой.

    Вызывать ПОСЛЕ _lock_document_for_fix и db.refresh(document) —
    иначе processed_storage_path может быть устаревшим.
    """
    violations = list(await db.scalars(
        sa_select(ViolationRecord).where(ViolationRecord.report_id == report.id)
    ))
    await create_snapshot(document, report, violations, action_type, action_description, db)


async def _snapshot_violation_action(
    violation,
    action_type: str,
    action_description: str,
    db: AsyncSession,
) -> None:
    """История для немутирующих действий («Ознакомлен», «Отклонить»).

    Файл docx не меняется, поэтому снимок только фиксирует состояние замечаний
    (для undo) и само действие в ленте истории. Снимок лёгкий — без копии
    файла. Если у документа ещё нет processed_storage_path (снимок невозможен),
    тихо пропускаем: аудит-след не должен ронять простое действие.
    """
    report = await db.scalar(
        sa_select(ProcessingReport)
        .options(selectinload(ProcessingReport.document))
        .where(ProcessingReport.id == violation.report_id)
    )
    if report is None or report.document is None:
        return
    document = report.document
    await _lock_document_for_fix(db, document.id)
    await db.refresh(document)
    if not document.processed_storage_path:
        return
    await _snapshot_before_fix(document, report, action_type, action_description, db)


# Регулярки для ЗАМЕНЫ префикса целиком (через subn). Похожие имена в
# engine.heading_numbers — там они для ИЗВЛЕЧЕНИЯ номера (capture group =
# цифры). Здесь захватываем «префикс + хвост-пунктуация + пробел», чтобы
# subn полностью заменил его на new_prefix.
_REPLACEABLE_NUM_PREFIX_RE = re.compile(r"^(\d+(?:\.\d+)*[.)]?\s+)")
_REPLACEABLE_CHAPTER_WORD_PREFIX_RE = re.compile(
    r"^(?:глава|раздел|часть|chapter|section|part)\s+(?:\d+(?:\.\d+)*|[ivxlcdm]+)[.:\s]+",
    re.IGNORECASE,
)


def _replace_heading_prefix(para, new_prefix: str, original_number: str | None = None) -> bool:
    """Заменяет ведущий номер/префикс главы в абзаце-заголовке.

    Текст заголовка может быть одним run'ом или разбит на несколько; чтобы
    замена работала надёжно, склеиваем текст всех run'ов, делаем замену в
    склейке и записываем результат в run[0] (остальные очищаем — так
    сохраняется форматирование первого run'а).
    Возвращает True, если префикс был найден и заменён.
    """
    # paragraph.text включает текст из <w:hyperlink>. overwrite_paragraph_text
    # удалит hyperlink при записи, иначе текст склеился бы из нового и
    # старого hyperlink-text, дав дубликат.
    full = para.text or ""
    updated, count = _REPLACEABLE_NUM_PREFIX_RE.subn(new_prefix, full, count=1)
    if count == 0:
        updated, count = _REPLACEABLE_CHAPTER_WORD_PREFIX_RE.subn(new_prefix, full, count=1)
    if count == 0 and original_number:
        # Запасной вариант: если абзац начинается с записанного исходного
        # номера (возможно без точки/пробела) — делаем буквальную замену.
        stripped = full.lstrip()
        orig = original_number.strip()
        if stripped.startswith(orig):
            offset = len(full) - len(stripped)
            suffix = stripped[len(orig):].lstrip(". ") if stripped[len(orig):len(orig)+2].strip() in ("", ".", ")") else stripped[len(orig):]
            updated = full[:offset] + new_prefix + suffix
            count = 1
    if count == 0:
        return False
    overwrite_paragraph_text(para, updated)
    return True

router = APIRouter(prefix="/api/v1/violations", tags=["violations"])


class HeadingNumberFixRequest(BaseModel):
    chosen_number: str


class PromoteToHeadingRequest(BaseModel):
    chosen_number: str | None = None  # опционально; если задан — пишется префиксом
    level: int | None = None  # 1..3; если не задан — определяется по числу точек в номере


@router.get("/{document_id}", response_model=list[ViolationOut])
async def get_document_violations(
    document_id: UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> list[ViolationOut]:
    try:
        violations = await list_violations(document_id, current_user, db)
    except (NotFoundError, AccessDeniedError) as exc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=str(exc))
    return [ViolationOut.model_validate(v) for v in violations]


@router.patch("/{violation_id}/accept", response_model=ViolationActionResponse)
async def accept_violation(
    violation_id: UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> ViolationActionResponse:
    try:
        violation = await get_violation_or_404(violation_id, current_user, db)
    except (NotFoundError, AccessDeniedError) as exc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=str(exc))

    await _snapshot_violation_action(
        violation, "acknowledge",
        f"Ознакомлен: «{(violation.section_title or violation_type_label(violation.type)).strip()}»",
        db,
    )
    violation.status = ViolationStatus.accepted
    await db.commit()
    return ViolationActionResponse(violation_id=violation.id, status=violation.status)


@router.patch("/{violation_id}/reject", response_model=ViolationActionResponse)
async def reject_violation(
    violation_id: UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> ViolationActionResponse:
    try:
        violation = await get_violation_or_404(violation_id, current_user, db)
    except (NotFoundError, AccessDeniedError) as exc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=str(exc))

    await _snapshot_violation_action(
        violation, "reject",
        f"Замечание отклонено: «{(violation.section_title or violation.type).strip()}»",
        db,
    )
    violation.status = ViolationStatus.rejected
    await db.commit()
    return ViolationActionResponse(violation_id=violation.id, status=violation.status)


def _recheck_heading_numbers_sync(docx_local_path: Path, rules: dict) -> list:
    """Прогоняет обнаружение заголовков + проверку нумерации на docx-файле
    (синхронно, выполняется в отдельном потоке)."""
    from services.core.vkr_core.engine.detection import detect_headings
    from services.core.vkr_core.engine.heading_numbers import validate_heading_numbers
    from services.core.vkr_core.engine.stats import collect_stats
    from services.core.vkr_core.engine.title_page import detect_title_page_end
    from services.core.vkr_core.engine.toc import detect_toc_section_indexes

    doc = WordDocument(str(docx_local_path))
    stats = collect_stats(doc)
    title_end = detect_title_page_end(stats, rules)
    title_page_indexes = set(range(title_end))
    toc_section_indexes = detect_toc_section_indexes(doc, rules)
    skip_indexes = title_page_indexes | toc_section_indexes
    headings = detect_headings(stats, rules, skip_indexes)
    return validate_heading_numbers(headings, stats)


def _rerenumber_captions_sync(
    docx_path: Path,
    rules: dict,
    orig_map: dict[int, str],
) -> list[tuple[int, str, str]]:
    """Пересчитывает подписи рисунков/таблиц/листингов по chapter_breaks; сохраняет docx.

    orig_map: paragraph_index → исходный текст из оригинального документа (из БД).
    Возвращает [(paragraph_index, source_orig, new_text)] только для изменившихся.
    """
    from services.core.vkr_core.engine.captions import renumber_figures, renumber_tables
    from services.core.vkr_core.engine.formatter import strip_chapter_word_everywhere
    from services.core.vkr_core.engine.listings import renumber_listings
    from services.core.vkr_core.engine.pipeline import (
        _caption_chapter_breaks_for_mode,
        _chapter_breaks_from_xml,
    )
    from services.core.vkr_core.engine.stats import collect_stats
    from services.core.vkr_core.engine.title_page import detect_title_page_end
    from services.core.vkr_core.engine.toc import detect_toc_section_indexes

    doc = WordDocument(str(docx_path))
    stats = collect_stats(doc)
    title_end = detect_title_page_end(stats, rules)
    toc_idx = detect_toc_section_indexes(doc, rules)
    skip_idx = set(range(title_end)) | toc_idx

    strip_chapter_word_everywhere(doc, skip_idx)
    chapter_breaks = _chapter_breaks_from_xml(doc.paragraphs)
    figure_breaks = _caption_chapter_breaks_for_mode(rules, "figure", chapter_breaks)
    table_breaks = _caption_chapter_breaks_for_mode(rules, "table", chapter_breaks)
    listing_breaks = _caption_chapter_breaks_for_mode(rules, "listing", chapter_breaks)
    _, fig_changes = renumber_figures(doc.paragraphs, skip_idx, apply=True, chapter_breaks=figure_breaks)
    _, tbl_changes = renumber_tables(
        doc.paragraphs, skip_idx, apply=True,
        indent_cm=None,
        chapter_breaks=table_breaks,
    )
    _, listing_changes = renumber_listings(
        doc.paragraphs, skip_idx, apply=True, chapter_breaks=listing_breaks
    )
    doc.save(str(docx_path))

    result = []
    for idx, cur_text, new_text in fig_changes + tbl_changes + listing_changes:
        src_orig = orig_map.get(idx, cur_text)
        if src_orig != new_text:
            result.append((idx, src_orig, new_text))
    return result


@router.patch("/{violation_id}/fix-heading-number", response_model=ViolationActionResponse)
async def fix_heading_number(
    violation_id: UUID,
    body: HeadingNumberFixRequest,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> ViolationActionResponse:
    try:
        violation = await get_violation_or_404(violation_id, current_user, db)
    except (NotFoundError, AccessDeniedError) as exc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=str(exc))

    if violation.type != "heading_number_conflict":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Not a heading_number_conflict violation")

    chosen = (body.chosen_number or "").strip()
    allowed = violation.fix_options or []
    # Принимаем либо предложенные варианты (fix_options), либо любой
    # пользовательский номер вида "1", "1.2", "1.2.3" — он введён вручную
    # из inline-input'а на карточке конфликта.
    if chosen not in allowed and not re.fullmatch(r"\d+(?:\.\d+){0,2}", chosen):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Номер должен быть одним из предложенных вариантов или иметь вид 1, 1.2, 1.2.3",
        )
    body.chosen_number = chosen

    # "Оставить как есть" — пользователь выбрал исходный номер. Переписывать
    # ничего не нужно: помечаем конфликт как rejected (пользователь увидел
    # и согласился, подчёркивание убираем). Каскадную перепроверку
    # пропускаем — документ не изменился.
    if body.chosen_number == (violation.original_text or ""):
        violation.status = ViolationStatus.rejected
        violation.fixed_text = body.chosen_number
        await db.commit()
        return ViolationActionResponse(violation_id=violation.id, status=violation.status)

    if violation.paragraph_index is None:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Violation has no paragraph_index")

    # Загружаем отчёт + документ + нормативную версию.
    report_row = await db.scalar(
        sa_select(ProcessingReport)
        .options(
            selectinload(ProcessingReport.document).selectinload(Document.normative_version)
        )
        .where(ProcessingReport.id == violation.report_id)
    )
    if report_row is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")

    document = report_row.document
    if document is None or not document.processed_storage_path:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Processed document not found")

    # S-07: захватываем row-lock на документ до конца транзакции, чтобы
    # параллельная правка другого нарушения не затёрла processed_storage_path.
    await _lock_document_for_fix(db, document.id)
    await db.refresh(document)
    if not document.processed_storage_path:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Processed document not found")

    title = (violation.section_title or "").strip()
    desc = f"Номер заголовка «{title}» → {chosen}" if title else f"Номер заголовка → {chosen}"
    await _snapshot_before_fix(document, report_row, "fix_heading_number", desc, db)

    rules = document.normative_version.rules_config

    # Применяем правку текста в обработанном docx.
    storage = StorageService()
    local_path = storage.load_to_local_path(document.processed_storage_path)
    doc_obj = WordDocument(str(local_path))

    paragraphs = doc_obj.paragraphs
    chosen = body.chosen_number
    # По ГОСТ между номером заголовка и текстом — обычный пробел,
    # точка после номера не ставится.
    new_prefix = chosen.rstrip(".") + " "

    orig_number = violation.original_text

    def _try_replace_at(idx: int) -> bool:
        if idx < 0 or idx >= len(paragraphs):
            return False
        return _replace_heading_prefix(paragraphs[idx], new_prefix, original_number=orig_number)

    applied = False
    if violation.paragraph_index < len(paragraphs):
        applied = _try_replace_at(violation.paragraph_index)

    # Запасной вариант, когда paragraph_index устарел. Это происходит, если
    # фоновый таск финализации (_finalize_and_preview_bg) успел запустить
    # _remerge_continuations / split_overflowing_fragments и изменил количество
    # абзацев до того, как пользователь нажал «Исправить». Ищем заголовок по
    # тексту section_title по ВСЕМУ документу — сначала вблизи записанного
    # индекса (±40, быстро), затем по остатку.
    if not applied and violation.section_title:
        target = " ".join(violation.section_title.strip().lower().split())
        # Двухпроходный поиск: сначала окрестность ±40, затем весь документ.
        near_start = max(0, violation.paragraph_index - 40)
        near_end   = min(len(paragraphs), violation.paragraph_index + 41)
        search_order = (
            list(range(near_start, near_end))
            + [i for i in range(len(paragraphs)) if i < near_start or i >= near_end]
        )
        for idx in search_order:
            candidate = " ".join(paragraphs[idx].text.strip().lower().split())
            if candidate == target or (len(target) > 10 and target in candidate):
                if _try_replace_at(idx):
                    applied = True
                    break

    if not applied:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=(
                f"Не удалось применить правку: заголовок "
                f"«{(violation.section_title or violation.original_text or '')[:60]}» "
                "не найден в документе по индексу или тексту."
            ),
        )

    # Каскадный фикс: исправляем дочерние нарушения прямо сейчас, пока
    # paragraph_index актуальны (rebuild_toc_full ещё не сдвинул абзацы).
    # Алгоритм: корень «1» → «2» ⟹ ребёнок «1.1» → «2.1», «1.2» → «2.2».
    _cascade_fixed_indices: list[int] = []
    _root_orig_parts = (violation.original_text or "").split(".")
    _new_root_parts = chosen.split(".")
    if _root_orig_parts != [""] and _new_root_parts != _root_orig_parts:
        _cascade_rows = list((await db.scalars(
            sa_select(ViolationRecord).where(
                ViolationRecord.report_id == violation.report_id,
                ViolationRecord.type == "heading_number_conflict",
                ViolationRecord.caused_by_index == violation.paragraph_index,
            )
        )).all())
        for _child in _cascade_rows:
            if not _child.original_text or _child.paragraph_index is None:
                continue
            _child_parts = _child.original_text.split(".")
            _new_child_num = ".".join(_new_root_parts + _child_parts[len(_root_orig_parts):]).rstrip(".")
            if 0 <= _child.paragraph_index < len(paragraphs):
                if _replace_heading_prefix(
                    paragraphs[_child.paragraph_index],
                    _new_child_num + " ",
                    original_number=_child.original_text,
                ):
                    _cascade_fixed_indices.append(_child.paragraph_index)

    # rebuild_toc_full перенесён в _finalize_and_preview_bg (передаём rules_config):
    # fix_heading_number не меняет количество абзацев — только текст заголовка,
    # поэтому _recheck_heading_numbers_sync ниже даёт корректные paragraph_index
    # даже без предварительного перестроения строк TOC.

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        updated_path = Path(tmp.name)
    doc_obj.save(str(updated_path))
    # Финализацию TOC (рендер ~6 с) выносим в фон — мгновенный отклик.
    new_storage_path = storage.save_file(updated_path, "processed", ".docx")
    document.processed_storage_path = new_storage_path

    updated_path.unlink(missing_ok=True)

    # Перезапускаем проверку нумерации заголовков на обновлённом документе (каскад).
    # В S3-режиме load_to_local_path создаёт tmp-файл — удаляем после use.
    new_docx_local = storage.load_to_local_path(new_storage_path)
    try:
        new_violations = await asyncio.to_thread(
            _recheck_heading_numbers_sync, new_docx_local, rules
        )
    finally:
        if str(new_docx_local) != new_storage_path:
            new_docx_local.unlink(missing_ok=True)

    # Сохраняем решения "оставить как есть", принятые ранее: собираем
    # section_title нарушений, которые пользователь явно отклонил, и до того,
    # как удалить старые строки, помечаем совпадающие новые строки как rejected.
    prior_rejected_rows = await db.scalars(
        sa_select(ViolationRecord).where(
            ViolationRecord.report_id == violation.report_id,
            ViolationRecord.type == "heading_number_conflict",
            ViolationRecord.status == ViolationStatus.rejected,
        )
    )
    rejected_titles = {
        (row.section_title or "").strip().lower()
        for row in prior_rejected_rows.all()
        if row.section_title
    }

    # Заменяем ВСЕ нарушения heading_number_conflict в этом отчёте, включая
    # только что исправленное — оно заменяется результатами перевалидации.
    # Аудит-след зафиксирован снапшотом (_snapshot_before_fix), накапливать
    # auto_fixed-записи не нужно.
    violation_id_saved = violation.id
    await db.execute(
        sa_delete(ViolationRecord).where(
            ViolationRecord.report_id == violation.report_id,
            ViolationRecord.type == "heading_number_conflict",
        )
    )

    # Пользователь исправил номер — подтвердил что это заголовок.
    # Снимаем heading_confirm для этого абзаца + каскадных детей (если были
    # исправлены автоматически), чтобы убрать подчёркивание в PDF.
    _confirm_indices = (
        ([violation.paragraph_index] if violation.paragraph_index is not None else [])
        + _cascade_fixed_indices
    )
    if _confirm_indices:
        await db.execute(
            sa_update(ViolationRecord)
            .where(
                ViolationRecord.report_id == violation.report_id,
                ViolationRecord.type == "heading_confirm",
                ViolationRecord.paragraph_index.in_(_confirm_indices),
                ViolationRecord.status == ViolationStatus.manual_required,
            )
            .values(status=ViolationStatus.accepted)
        )

    await db.flush()

    for v in new_violations:
        entry_status = v.status
        title_norm = (v.section_title or "").strip().lower()
        if title_norm and title_norm in rejected_titles:
            entry_status = ViolationStatus.rejected
        db.add(
            ViolationRecord(
                report_id=violation.report_id,
                type=v.type,
                rule_reference=v.rule_reference,
                severity=v.severity,
                description=v.description,
                paragraph_index=v.paragraph_index,
                section_title=v.section_title,
                original_text=v.original_text,
                fix_options=v.fix_options,
                caused_by_index=v.caused_by_index,
                detector_signals=v.detector_signals,
                status=entry_status,
            )
        )

    await db.commit()
    # Исправление H1-заголовка меняет chapter_breaks → нужно пересчитать подписи.
    _is_h1_fix = "." not in chosen
    background_tasks.add_task(
        _finalize_and_preview_bg,
        report_row.id,
        document.id,
        new_storage_path,
        rules_config=rules,
        renum_captions=_is_h1_fix,
    )
    return ViolationActionResponse(violation_id=violation_id_saved, status=ViolationStatus.auto_fixed)


@router.patch("/{violation_id}/revert-caption", response_model=ViolationActionResponse)
async def revert_caption(
    violation_id: UUID,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> ViolationActionResponse:
    """Откатывает перенумерацию подписи — восстанавливает оригинальный текст."""
    try:
        violation = await get_violation_or_404(violation_id, current_user, db)
    except (NotFoundError, AccessDeniedError) as exc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=str(exc))

    if violation.type != "caption_renumbered":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Not a caption_renumbered violation")

    if not violation.original_text or not violation.fixed_text:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Violation missing original_text or fixed_text")

    report_row = await db.scalar(
        sa_select(ProcessingReport)
        .options(selectinload(ProcessingReport.document))
        .where(ProcessingReport.id == violation.report_id)
    )
    if report_row is None or report_row.document is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")
    document = report_row.document
    if not document.processed_storage_path:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Processed document not found")

    # S-07: row-lock на документ для сериализации параллельных правок.
    await _lock_document_for_fix(db, document.id)
    await db.refresh(document)
    if not document.processed_storage_path:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Processed document not found")

    fixed = (violation.fixed_text or "").strip()
    orig  = (violation.original_text or "").strip()
    await _snapshot_before_fix(
        document, report_row, "revert_caption",
        f"Откат подписи: «{fixed}» → «{orig}»",
        db,
    )

    storage = StorageService()
    local_path = storage.load_to_local_path(document.processed_storage_path)
    doc_obj = WordDocument(str(local_path))

    target_text = violation.fixed_text.strip()
    original_text = violation.original_text.strip()

    reverted = False
    for para in doc_obj.paragraphs:
        if para.text.strip() == target_text:
            para.text = original_text
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = DocxPt(14)
            reverted = True
            break
    if not reverted:
        # Подпись могла измениться между сохранением нарушения и нажатием
        # «откатить» (например, пользователь вручную правил текст в Word).
        # Не падаем — нарушение всё равно пометим как rejected, но в логах
        # должен остаться след, чтобы было видно «откат не применился».
        _logger.warning(
            "revert_caption: подпись «%s» не найдена в документе",
            target_text,
        )

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        updated_path = Path(tmp.name)
    doc_obj.save(str(updated_path))

    new_storage_path = storage.save_file(updated_path, "processed", ".docx")
    document.processed_storage_path = new_storage_path
    updated_path.unlink(missing_ok=True)

    violation.status = ViolationStatus.rejected
    await db.commit()
    background_tasks.add_task(_regen_pdf_background, report_row.id, new_storage_path)

    return ViolationActionResponse(violation_id=violation.id, status=violation.status)


class DemoteHeadingRequest(BaseModel):
    # "text" — пользователь сказал, что это обычный абзац.
    # "list" — пункт нумерованного списка. По варианту (ii) обе ветки
    # делают одинаковую OXML-мутацию (снимаем стиль заголовка, ставим
    # body-стиль, текст и префикс "1.2" сохраняются как есть).
    # Различие фиксируется только в violation.fixed_text для аудита.
    as_kind: str = "text"


_DEMOTE_KINDS = {"text", "list"}


def _strip_page_break_before(paragraph) -> None:
    """Удаляет w:pageBreakBefore из pPr — заголовки L1 при apply_headings
    получают этот атрибут; если мы понижаем заголовок до тела, его нужно
    убрать, иначе в исправленном файле абзац будет всегда начинаться с
    новой страницы."""
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return
    for existing in list(pPr.findall(qn("w:pageBreakBefore"))):
        pPr.remove(existing)


_LIST_NUM_PREFIX_RE = re.compile(r"^(\d+(?:\.\d+){0,2})[.)]?\s*")


def _ensure_list_dot_after_number(paragraph) -> bool:
    """Гарантирует, что числовой префикс пункта списка завершается точкой:
    "1.2 Foo"  → "1.2. Foo"
    "1) Foo"   → "1. Foo"
    "1.2. Foo" — без изменений.

    Возвращает True, если текст был изменён.
    """
    full = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
    m = _LIST_NUM_PREFIX_RE.match(full)
    if not m:
        return False
    digits = m.group(1)
    rest = full[m.end():]
    desired = f"{digits}. {rest.lstrip()}" if rest else f"{digits}."
    if full == desired:
        return False
    overwrite_paragraph_text(paragraph, desired)
    return True


def _apply_body_paragraph_style(paragraph, rules: dict) -> None:
    """Применяет body-стилизацию к одному абзацу — те же параметры, что
    apply_body_style, но точечно (без полной перепрогонки документа).

    Сбрасывает стиль на Normal, убирает outline_level и pageBreakBefore,
    выставляет шрифт TNR 14pt, межстрочный 1.5, выравнивание justify,
    first_line_indent 1.25 см. numPr (если был) не трогаем — формально
    варианты "обычный текст" и "пункт списка" по варианту (ii) идентичны
    на уровне OXML.
    """
    style = rules.get("body_text_style", {})
    font_name = style.get("font_name", "Times New Roman")
    font_size = float(style.get("font_size_pt", 14))
    line_spacing = float(style.get("line_spacing", 1.5))
    first_line_indent = float(style.get("first_line_indent_cm", 1.25))
    font_color = style.get("font_color", "000000")

    try:
        paragraph.style = "Normal"
    except KeyError:
        pass

    # Снимаем атрибуты pPr, унаследованные от стиля заголовка.
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is not None:
        for olvl in list(pPr.findall(qn("w:outlineLvl"))):
            pPr.remove(olvl)
    _strip_page_break_before(paragraph)

    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = DocxPt(float(style.get("space_before_pt", 0)))
    pf.space_after = DocxPt(float(style.get("space_after_pt", 0)))
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.first_line_indent = Cm(first_line_indent)
    paragraph.alignment = paragraph_alignment_from_rule(style.get("alignment"))

    rgb = RGBColor.from_string(font_color)
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = DocxPt(font_size)
        run.font.bold = False
        run.font.italic = False
        run.font.underline = False
        run.font.all_caps = False
        run.font.color.rgb = rgb


def _p_has_outline_level(paragraph) -> bool:
    pPr = paragraph._p.find(qn("w:pPr"))
    return pPr is not None and pPr.find(qn("w:outlineLvl")) is not None


def _p_is_toc_entry(paragraph) -> bool:
    """True, если абзац — строка оглавления (наша PAGEREF-запись или Word-TOC).

    Нужно, чтобы fallback-поиск по section_title не «целился» в саму запись
    оглавления (её текст совпадает с заголовком), а находил реальный
    заголовок в теле документа.
    """
    from services.core.vkr_core.engine.toc import _is_toc_entry
    try:
        return _is_toc_entry(paragraph)
    except Exception:
        return "_VkrToc_" in paragraph._p.xml or "PAGEREF" in paragraph._p.xml


def _resolve_heading_target(paragraphs, violation) -> int | None:
    """Находит индекс абзаца-заголовка для ручной правки.

    Сначала доверяем violation.paragraph_index; если текст там не совпадает с
    section_title — ищем в окне ±20. В обоих случаях ПРОПУСКАЕМ абзацы-записи
    оглавления (их текст дублирует заголовок) и ПРЕДПОЧИТАЕМ абзац, у которого
    сейчас есть w:outlineLvl (т.е. это действительно заголовок, а не уже
    понижённый/случайно совпавший абзац тела).
    """
    pidx = violation.paragraph_index
    in_range = pidx is not None and 0 <= pidx < len(paragraphs)
    target_norm = (
        " ".join(violation.section_title.strip().lower().split())
        if violation.section_title
        else None
    )

    def _match(text: str) -> bool:
        c = " ".join(text.strip().lower().split())
        return c == target_norm or (
            len(target_norm) > 10 and target_norm in c
        )

    # 1. Прямое попадание по paragraph_index (не запись TOC; текст совпадает
    #    либо section_title не задан).
    if in_range and not _p_is_toc_entry(paragraphs[pidx]):
        if target_norm is None or _match(paragraphs[pidx].text):
            return pidx

    if target_norm is None or pidx is None:
        # Без section_title больше ничего проверить нельзя — доверяем индексу.
        return pidx if in_range else None

    # 2. Поиск по тексту в окне ±20: пропускаем записи TOC, приоритет —
    #    абзацу с outlineLvl (это и есть настоящий заголовок).
    start = max(0, pidx - 20)
    end = min(len(paragraphs), pidx + 21)
    window_fallback: int | None = None
    for idx in range(start, end):
        p = paragraphs[idx]
        if _p_is_toc_entry(p) or not _match(p.text):
            continue
        if _p_has_outline_level(p):
            return idx
        if window_fallback is None:
            window_fallback = idx
    if window_fallback is not None:
        return window_fallback

    # 3. Индексы могли сильно сдвинуться (после rebuild_toc_full) — ищем по
    #    всему документу заголовок (с outlineLvl), совпадающий по тексту.
    for idx, p in enumerate(paragraphs):
        if _p_is_toc_entry(p):
            continue
        if _p_has_outline_level(p) and _match(p.text):
            return idx

    # 4. Последний шанс — доверяем сырому индексу (как в прежней логике),
    #    лишь бы это не была сама строка оглавления.
    if in_range and not _p_is_toc_entry(paragraphs[pidx]):
        return pidx
    return None


# ── Откат auto_fixed нарушений ────────────────────────────────────────────────

_REVERTABLE_AUTO_FIXED_TYPES = {
    "heading_recovered",
    "forbidden_heading",
    "long_table_split_warning",
    "table_mass_bold_removed",
    "citation_separator_missing",
    "paragraph_spacing",
    "font_size_inconsistency",
}

# Типы, для которых нужен оригинальный docx чтобы знать исходное состояние.
_REVERT_NEEDS_ORIGINAL = {
    "citation_separator_missing",
    "paragraph_spacing",
    "font_size_inconsistency",
}


def _find_table_after_para_idx(doc, para_idx: int):
    """Первая таблица, идущая в body сразу после абзаца с индексом para_idx."""
    body = doc.element.body
    p_count = -1
    hit = False
    for elem in body.iterchildren():
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag == "p":
            p_count += 1
            if p_count == para_idx:
                hit = True
        elif tag == "tbl" and hit:
            for tbl in doc.tables:
                if tbl._tbl is elem:
                    return tbl
            break
    return None


def _revert_auto_fix_sync(
    local_path: Path,
    original_local: Path | None,
    violation_type: str,
    paragraph_index: int | None,
    original_text: str | None,
    rules: dict,
) -> Path:
    """Откатывает одно auto_fixed нарушение; возвращает путь к новому docx.

    Финализация TOC (для heading_recovered) выполняется вызывающим в фоне.
    """
    doc = WordDocument(str(local_path))
    paras = doc.paragraphs

    if violation_type == "heading_recovered":
        if paragraph_index is not None and paragraph_index < len(paras):
            _apply_body_paragraph_style(paras[paragraph_index], rules)

    elif violation_type == "forbidden_heading":
        orig = (original_text or "").strip()
        if paragraph_index is not None and orig and paragraph_index < len(paras):
            overwrite_paragraph_text(paras[paragraph_index], orig)

    elif violation_type == "long_table_split_warning":
        if paragraph_index is not None:
            tbl = _find_table_after_para_idx(doc, paragraph_index)
            if tbl and tbl.rows:
                trPr = tbl.rows[0]._tr.find(qn("w:trPr"))
                if trPr is not None:
                    for el in list(trPr.findall(qn("w:tblHeader"))):
                        trPr.remove(el)

    elif violation_type == "table_mass_bold_removed":
        if paragraph_index is not None:
            tbl = _find_table_after_para_idx(doc, paragraph_index)
            if tbl and len(tbl.rows) > 1:
                for row in tbl.rows[1:]:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.bold = True

    elif violation_type == "citation_separator_missing" and original_local:
        orig_doc = WordDocument(str(original_local))
        orig_paras = orig_doc.paragraphs
        if (
            paragraph_index is not None
            and paragraph_index < len(paras)
            and paragraph_index < len(orig_paras)
        ):
            dst_p = paras[paragraph_index]._p
            for r in list(dst_p.findall(qn("w:r"))):
                dst_p.remove(r)
            for r in orig_paras[paragraph_index]._p.findall(qn("w:r")):
                dst_p.append(copy.deepcopy(r))

    elif violation_type == "paragraph_spacing" and original_local:
        orig_doc = WordDocument(str(original_local))
        # strict=False — после remove_empty_paragraphs длины могут различаться;
        # сравнение идёт по короткому, остаток молча игнорируется.
        for para, orig_para in zip(paras, orig_doc.paragraphs, strict=False):
            orig_pPr = orig_para._p.find(qn("w:pPr"))
            pPr = para._p.find(qn("w:pPr"))
            orig_spacing = orig_pPr.find(qn("w:spacing")) if orig_pPr is not None else None
            if orig_spacing is None:
                if pPr is not None:
                    for el in list(pPr.findall(qn("w:spacing"))):
                        pPr.remove(el)
            else:
                if pPr is None:
                    pPr = OxmlElement("w:pPr")
                    para._p.insert(0, pPr)
                for el in list(pPr.findall(qn("w:spacing"))):
                    pPr.remove(el)
                pPr.append(copy.deepcopy(orig_spacing))

    elif violation_type == "font_size_inconsistency" and original_local:
        orig_doc = WordDocument(str(original_local))
        for para, orig_para in zip(paras, orig_doc.paragraphs, strict=False):
            for run, orig_run in zip(para.runs, orig_para.runs, strict=False):
                if orig_run.font.size is not None:
                    run.font.size = orig_run.font.size

    if violation_type == "heading_recovered":
        rebuild_toc_full(doc, rules)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        updated_path = Path(tmp.name)
    doc.save(str(updated_path))
    # Финализация TOC (для heading_recovered) — в фоне у вызывающего.
    return updated_path


@router.patch("/{violation_id}/revert", response_model=ViolationActionResponse)
async def revert_auto_fix(
    violation_id: UUID,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> ViolationActionResponse:
    """Откатывает авто-исправление к исходному состоянию документа.

    Поддерживаемые типы: heading_recovered, forbidden_heading,
    long_table_split_warning, table_mass_bold_removed,
    citation_separator_missing, paragraph_spacing, font_size_inconsistency.
    После отката violation.status = rejected.
    """
    try:
        violation = await get_violation_or_404(violation_id, current_user, db)
    except (NotFoundError, AccessDeniedError) as exc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=str(exc))

    if violation.status != ViolationStatus.auto_fixed:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Откат доступен только для автоматически исправленных нарушений",
        )
    if violation.type not in _REVERTABLE_AUTO_FIXED_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Откат нарушения типа «{violation.type}» не поддерживается",
        )

    report_row = await db.scalar(
        sa_select(ProcessingReport)
        .options(
            selectinload(ProcessingReport.document).selectinload(Document.normative_version)
        )
        .where(ProcessingReport.id == violation.report_id)
    )
    if report_row is None or report_row.document is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Отчёт не найден")
    document = report_row.document
    if not document.processed_storage_path:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Обработанный документ не найден")

    await _lock_document_for_fix(db, document.id)
    await db.refresh(document)
    if not document.processed_storage_path:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Обработанный документ не найден")

    await _snapshot_before_fix(
        document, report_row, "revert_auto_fix",
        f"Откат «{violation_type_label(violation.type)}»",
        db,
    )

    rules = document.normative_version.rules_config
    storage = StorageService()
    local_path = storage.load_to_local_path(document.processed_storage_path)

    original_local: Path | None = None
    if violation.type in _REVERT_NEEDS_ORIGINAL:
        original_local = storage.load_to_local_path(document.original_storage_path)

    try:
        updated_path = await asyncio.to_thread(
            _revert_auto_fix_sync,
            local_path,
            original_local,
            violation.type,
            violation.paragraph_index,
            violation.original_text,
            rules,
        )
    finally:
        if original_local is not None and str(original_local) != document.original_storage_path:
            original_local.unlink(missing_ok=True)

    new_storage_path = storage.save_file(updated_path, "processed", ".docx")
    document.processed_storage_path = new_storage_path

    updated_path.unlink(missing_ok=True)

    # heading_recovered: убран один заголовок — перевалидируем соседей.
    if violation.type == "heading_recovered":
        new_docx_local = storage.load_to_local_path(new_storage_path)
        try:
            new_violations = await asyncio.to_thread(
                _recheck_heading_numbers_sync, new_docx_local, rules
            )
        finally:
            if str(new_docx_local) != new_storage_path:
                new_docx_local.unlink(missing_ok=True)

        await db.execute(
            sa_delete(ViolationRecord).where(
                ViolationRecord.report_id == violation.report_id,
                ViolationRecord.type == "heading_number_conflict",
            )
        )
        await db.flush()
        for v in new_violations:
            db.add(ViolationRecord(
                report_id=violation.report_id,
                type=v.type,
                rule_reference=v.rule_reference,
                severity=v.severity,
                description=v.description,
                paragraph_index=v.paragraph_index,
                section_title=v.section_title,
                original_text=v.original_text,
                fix_options=v.fix_options,
                caused_by_index=v.caused_by_index,
                detector_signals=v.detector_signals,
                status=v.status,
            ))

    violation.status = ViolationStatus.rejected
    await db.commit()
    if violation.type == "heading_recovered":
        # TOC изменился — финализация (рендер) + превью в фоне.
        background_tasks.add_task(
            _finalize_and_preview_bg, report_row.id, document.id, new_storage_path
        )
    else:
        background_tasks.add_task(
            _regen_pdf_background, report_row.id, new_storage_path,
        )

    return ViolationActionResponse(violation_id=violation.id, status=violation.status)


@router.patch("/{violation_id}/demote-heading-to-text", response_model=ViolationActionResponse)
async def demote_heading_to_text(
    violation_id: UUID,
    body: DemoteHeadingRequest,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> ViolationActionResponse:
    """Понижает абзац-«заголовок-с-конфликтом-нумерации» в обычный абзац.

    Применимо только к heading_number_conflict: пользователь увидел
    предполагаемый заголовок и решил, что это не заголовок, а обычная
    строка тела (как абзац или пункт списка). Текст и префикс ("1.2 ...")
    сохраняем; снимаем заголовочный стиль, outline_level и pageBreakBefore;
    ставим Normal + body-стиль (TNR 14pt, justify, first_line_indent 1.25см).
    После этого каскадно перевалидируем нумерацию заголовков — соседние
    конфликты могут изменить ожидаемый номер, потому что один заголовок
    из последовательности исчез.
    """
    try:
        violation = await get_violation_or_404(violation_id, current_user, db)
    except (NotFoundError, AccessDeniedError) as exc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=str(exc))

    # heading_number_conflict — «это заголовок, но номер спорный»;
    # heading_confirm / possible_missed_heading — «возможно, это заголовок».
    # Во всех трёх случаях пользователь решает «это НЕ заголовок» и хочет
    # вернуть абзац к обычному телу. Для confirm-типов каскад по нумерации
    # заголовков не нужен — там просто гасим само нарушение.
    _DEMOTABLE_TYPES = {
        "heading_number_conflict", "heading_confirm", "possible_missed_heading",
    }
    if violation.type not in _DEMOTABLE_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Понижение в обычный текст применимо только к {sorted(_DEMOTABLE_TYPES)}",
        )

    kind = (body.as_kind or "text").strip().lower()
    if kind not in _DEMOTE_KINDS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"as_kind должен быть одним из {sorted(_DEMOTE_KINDS)}",
        )

    if violation.paragraph_index is None:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="У нарушения нет paragraph_index")

    # Быстрый путь. heading_confirm / possible_missed_heading в новой модели —
    # это абзац, который УЖЕ обычный текст тела (детектор сознательно его не
    # принял как заголовок). «Это не заголовок» здесь — лишь подтверждение
    # дефолта: содержимое docx менять не нужно, поэтому НЕ открываем docx и
    # НЕ запускаем фоновую 2-проходную перерисовку превью (цена правки
    # платится только при явном promote-to-heading). Но действие пользователя
    # всё равно фиксируем в истории снапшотом — он лёгкий (только ссылки на
    # пути + JSON статусов, без копии файла) и даёт undo.
    # (Для kind="list" нужна реальная OXML-мутация — точка после номера, —
    #  поэтому туда не сворачиваем; heading_number_conflict — настоящий
    #  заголовок, тоже идёт обычным путём с каскадом.)
    if kind == "text" and violation.type in {
        "heading_confirm", "possible_missed_heading",
    }:
        report_row = await db.scalar(
            sa_select(ProcessingReport)
            .options(selectinload(ProcessingReport.document))
            .where(ProcessingReport.id == violation.report_id)
        )
        if report_row is None or report_row.document is None:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Отчёт не найден")
        document = report_row.document
        await _lock_document_for_fix(db, document.id)
        await db.refresh(document)
        title = (violation.section_title or "").strip()
        await _snapshot_before_fix(
            document, report_row, "acknowledge",
            f"Это не заголовок (оставлено текстом): «{title}»",
            db,
        )
        await db.execute(
            sa_update(ViolationRecord)
            .where(
                ViolationRecord.report_id == violation.report_id,
                ViolationRecord.type.in_(
                    ["heading_confirm", "possible_missed_heading"]
                ),
                ViolationRecord.paragraph_index == violation.paragraph_index,
                ViolationRecord.status == ViolationStatus.manual_required,
            )
            .values(status=ViolationStatus.accepted)
        )
        await db.commit()
        return ViolationActionResponse(
            violation_id=violation.id, status=ViolationStatus.auto_fixed
        )

    report_row = await db.scalar(
        sa_select(ProcessingReport)
        .options(
            selectinload(ProcessingReport.document).selectinload(Document.normative_version)
        )
        .where(ProcessingReport.id == violation.report_id)
    )
    if report_row is None or report_row.document is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Отчёт не найден")
    document = report_row.document
    if not document.processed_storage_path:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Обработанный документ не найден")

    # S-07: row-lock на документ для сериализации параллельных правок.
    await _lock_document_for_fix(db, document.id)
    await db.refresh(document)
    if not document.processed_storage_path:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Обработанный документ не найден")

    title = (violation.section_title or "").strip()
    kind_label = "обычный текст" if kind == "text" else "пункт списка"
    await _snapshot_before_fix(
        document, report_row, "demote_heading",
        f"Это не заголовок: «{title}» → {kind_label}",
        db,
    )

    rules = document.normative_version.rules_config

    storage = StorageService()
    local_path = storage.load_to_local_path(document.processed_storage_path)
    doc_obj = WordDocument(str(local_path))
    paragraphs = doc_obj.paragraphs

    # Поиск целевого абзаца: пропускаем записи оглавления (их текст дублирует
    # заголовок) и предпочитаем абзац с outlineLvl — иначе при сдвиге индексов
    # после rebuild_toc_full повторное «это не заголовок» мутировало строку
    # TOC, а реальный заголовок оставался и снова попадал в оглавление.
    target = _resolve_heading_target(paragraphs, violation)

    if target is None:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Не удалось найти абзац: возможно, документ уже изменён",
        )

    paragraph = paragraphs[target]
    _apply_body_paragraph_style(paragraph, rules)
    # Для пункта списка по ГОСТ — после номера обязательно точка ("1.2.").
    # Для обычного абзаца текста префикс не трогаем (вдруг это часть фразы
    # типа "1.2 раза больше").
    if kind == "list":
        _ensure_list_dot_after_number(paragraph)

    rebuild_toc_full(doc_obj, rules)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        updated_path = Path(tmp.name)
    doc_obj.save(str(updated_path))
    # Финализацию TOC (рендер ~6 с) выносим в фон — мгновенный отклик.
    new_storage_path = storage.save_file(updated_path, "processed", ".docx")
    document.processed_storage_path = new_storage_path
    updated_path.unlink(missing_ok=True)

    violation_id_saved = violation.id

    if violation.type == "heading_number_conflict":
        # Каскад: убран один заголовок → у соседних может смениться ожидаемый
        # номер. Перепроверяем; повторно используется уже существующий хелпер.
        new_docx_local = storage.load_to_local_path(new_storage_path)
        try:
            new_violations = await asyncio.to_thread(
                _recheck_heading_numbers_sync, new_docx_local, rules
            )
        finally:
            if str(new_docx_local) != new_storage_path:
                new_docx_local.unlink(missing_ok=True)

        prior_rejected_rows = await db.scalars(
            sa_select(ViolationRecord).where(
                ViolationRecord.report_id == violation.report_id,
                ViolationRecord.type == "heading_number_conflict",
                ViolationRecord.status == ViolationStatus.rejected,
            )
        )
        rejected_titles = {
            (row.section_title or "").strip().lower()
            for row in prior_rejected_rows.all()
            if row.section_title
        }
        # Пользователь явно пометил абзац как НЕ заголовок — даже если
        # каскадная перевалидация снова его подхватит (текст начинается с
        # числа), нарушение должно оставаться скрытым.
        if violation.section_title:
            rejected_titles.add(violation.section_title.strip().lower())

        # Удаляем все heading_number_conflict включая текущее — аналогично
        # fix-heading-number, накопление auto_fixed-записей не нужно.
        await db.execute(
            sa_delete(ViolationRecord).where(
                ViolationRecord.report_id == violation.report_id,
                ViolationRecord.type == "heading_number_conflict",
            )
        )

        # Абзац понижен до тела — heading_confirm для него не актуален.
        if violation.paragraph_index is not None:
            await db.execute(
                sa_update(ViolationRecord)
                .where(
                    ViolationRecord.report_id == violation.report_id,
                    ViolationRecord.type == "heading_confirm",
                    ViolationRecord.paragraph_index == violation.paragraph_index,
                    ViolationRecord.status == ViolationStatus.manual_required,
                )
                .values(status=ViolationStatus.accepted)
            )

        await db.flush()

        for v in new_violations:
            entry_status = v.status
            title_norm = (v.section_title or "").strip().lower()
            if title_norm and title_norm in rejected_titles:
                entry_status = ViolationStatus.rejected
            db.add(
                ViolationRecord(
                    report_id=violation.report_id,
                    type=v.type,
                    rule_reference=v.rule_reference,
                    severity=v.severity,
                    description=v.description,
                    paragraph_index=v.paragraph_index,
                    section_title=v.section_title,
                    original_text=v.original_text,
                    fix_options=v.fix_options,
                    caused_by_index=v.caused_by_index,
                    detector_signals=v.detector_signals,
                    status=entry_status,
                )
            )
    else:
        # heading_confirm / possible_missed_heading: «это не заголовок».
        # Каскад нумерации не нужен — гасим само нарушение и все
        # confirm-карточки на том же абзаце.
        if violation.paragraph_index is not None:
            await db.execute(
                sa_update(ViolationRecord)
                .where(
                    ViolationRecord.report_id == violation.report_id,
                    ViolationRecord.type.in_(
                        ["heading_confirm", "possible_missed_heading"]
                    ),
                    ViolationRecord.paragraph_index == violation.paragraph_index,
                    ViolationRecord.status == ViolationStatus.manual_required,
                )
                .values(status=ViolationStatus.accepted)
            )
        else:
            await db.execute(
                sa_update(ViolationRecord)
                .where(ViolationRecord.id == violation.id)
                .values(status=ViolationStatus.accepted)
            )

    await db.commit()
    background_tasks.add_task(
        _finalize_and_preview_bg, report_row.id, document.id, new_storage_path
    )

    return ViolationActionResponse(violation_id=violation_id_saved, status=ViolationStatus.auto_fixed)


_PROMOTABLE_TYPES = {"possible_missed_heading", "heading_confirm"}


@router.patch("/{violation_id}/promote-to-heading", response_model=ViolationActionResponse)
async def promote_to_heading(
    violation_id: UUID,
    body: PromoteToHeadingRequest,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> ViolationActionResponse:
    """Оформляет указанный абзац как полноценный заголовок.

    Используется для нарушений типа "possible_missed_heading" (LOW-кандидат,
    которого детектор не принял автоматически) и "heading_confirm" (MEDIUM-
    кандидат, ждавший подтверждения). Пользователь может задать номер
    заголовка вручную (например, "2.1.3") — он будет дописан в начало
    текста, а уровень определится по количеству точек, если параметр
    `level` не задан явно.
    """
    try:
        violation = await get_violation_or_404(violation_id, current_user, db)
    except (NotFoundError, AccessDeniedError) as exc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=str(exc))

    if violation.type not in _PROMOTABLE_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Этот тип нарушения нельзя оформить как заголовок: {violation.type}",
        )

    if violation.paragraph_index is None:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="У нарушения нет paragraph_index — невозможно найти абзац",
        )

    chosen_number = (body.chosen_number or "").strip()
    if chosen_number and not re.fullmatch(r"\d+(?:\.\d+){0,2}", chosen_number):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Номер должен иметь вид 1, 1.2 или 1.2.3",
        )

    if body.level is not None:
        if not 1 <= body.level <= 3:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="level должен быть в диапазоне 1..3",
            )
        level = body.level
    elif chosen_number:
        level = chosen_number.count(".") + 1
    else:
        level = 1

    # Загружаем документ, на котором будем применять оформление.
    report_row = await db.scalar(
        sa_select(ProcessingReport)
        .options(
            selectinload(ProcessingReport.document).selectinload(Document.normative_version)
        )
        .where(ProcessingReport.id == violation.report_id)
    )
    if report_row is None or report_row.document is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Отчёт не найден")
    document = report_row.document
    if not document.processed_storage_path:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Обработанный документ не доступен",
        )

    # S-07: row-lock на документ для сериализации параллельных правок.
    await _lock_document_for_fix(db, document.id)
    await db.refresh(document)
    if not document.processed_storage_path:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Обработанный документ не доступен",
        )

    title = (violation.section_title or "").strip()
    num_part = f" {chosen_number}" if chosen_number else ""
    await _snapshot_before_fix(
        document, report_row, "promote_heading",
        f"Оформить «{title}»{num_part} как заголовок",
        db,
    )

    rules = document.normative_version.rules_config

    storage = StorageService()
    local_path = storage.load_to_local_path(document.processed_storage_path)
    doc_obj = WordDocument(str(local_path))
    paragraphs = doc_obj.paragraphs
    if violation.paragraph_index >= len(paragraphs):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="paragraph_index указывает за пределы документа",
        )
    # Поиск целевого абзаца с пропуском записей оглавления (их текст дублирует
    # заголовок) — иначе при сдвиге индексов promote мог оформить как заголовок
    # саму строку TOC вместо абзаца в теле документа.
    _ptarget = _resolve_heading_target(paragraphs, violation)
    if _ptarget is None:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Не удалось найти абзац: возможно, документ уже изменён",
        )
    paragraph = paragraphs[_ptarget]

    # Применяем номер: если задан — пишем префиксом, иначе оставляем текст
    # как есть (заголовок без номера, типа "Введение"/"Заключение").
    from services.core.vkr_core.engine.formatter import _normalise_heading_text

    original_text = paragraph.text or ""
    normalized_heading_text = _normalise_heading_text(original_text)
    if chosen_number:
        prefix_re = re.compile(r"^\s*\d+(?:\.\d+){0,2}[.)]?\s+")
        stripped_text = prefix_re.sub("", normalized_heading_text)
        new_text = f"{chosen_number} {stripped_text.lstrip()}".rstrip()
        overwrite_paragraph_text(paragraph, new_text)
    elif normalized_heading_text != original_text:
        overwrite_paragraph_text(paragraph, normalized_heading_text)

    # Применяем визуальное оформление и outline_level. Используем те же
    # хелперы, что и pipeline → стиль гарантированно совпадает.
    from services.core.vkr_core.engine.formatter import (
        _apply_heading_fallback,
        _apply_heading_style,
        _force_heading_direct_format,
        _is_structural,
        _set_outline_level_xml,
        _set_page_break_before,
    )

    text_for_check = paragraph.text or ""
    structural = _is_structural(text_for_check)
    if not _apply_heading_style(paragraph, level):
        _apply_heading_fallback(paragraph, level, structural)
    _force_heading_direct_format(paragraph, structural, level)
    _set_outline_level_xml(paragraph, level)
    if level == 1:
        _set_page_break_before(paragraph)

    rebuild_toc_full(doc_obj, rules)

    # Сохраняем обновлённый docx; финализацию TOC (рендер ~6 с) — в фон.
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        updated_path = Path(tmp.name)
    doc_obj.save(str(updated_path))
    new_storage_path = storage.save_file(updated_path, "processed", ".docx")
    document.processed_storage_path = new_storage_path
    updated_path.unlink(missing_ok=True)

    # Помечаем нарушение как исправленное.
    if chosen_number:
        violation.fixed_text = f"{chosen_number} {paragraph.text.replace(chosen_number, '', 1).lstrip()}".strip()
    else:
        violation.fixed_text = paragraph.text.strip()
    violation.status = ViolationStatus.auto_fixed

    # Добавлен новый заголовок → ожидаемые номера соседей могли сдвинуться.
    # Перепроверяем нумерацию и заменяем все heading_number_conflict-записи.
    recheck_local = storage.load_to_local_path(new_storage_path)
    try:
        new_hn_violations = await asyncio.to_thread(
            _recheck_heading_numbers_sync, recheck_local, rules
        )
    finally:
        if str(recheck_local) != new_storage_path:
            recheck_local.unlink(missing_ok=True)

    prior_rejected = await db.scalars(
        sa_select(ViolationRecord).where(
            ViolationRecord.report_id == violation.report_id,
            ViolationRecord.type == "heading_number_conflict",
            ViolationRecord.status == ViolationStatus.rejected,
        )
    )
    rejected_titles_promote = {
        (row.section_title or "").strip().lower()
        for row in prior_rejected.all()
        if row.section_title
    }

    await db.execute(
        sa_delete(ViolationRecord).where(
            ViolationRecord.report_id == violation.report_id,
            ViolationRecord.type == "heading_number_conflict",
        )
    )
    await db.flush()

    for v in new_hn_violations:
        entry_status = v.status
        title_norm = (v.section_title or "").strip().lower()
        if title_norm and title_norm in rejected_titles_promote:
            entry_status = ViolationStatus.rejected
        db.add(ViolationRecord(
            report_id=violation.report_id,
            type=v.type,
            rule_reference=v.rule_reference,
            severity=v.severity,
            description=v.description,
            paragraph_index=v.paragraph_index,
            section_title=v.section_title,
            original_text=v.original_text,
            fix_options=v.fix_options,
            caused_by_index=v.caused_by_index,
            detector_signals=v.detector_signals,
            status=entry_status,
        ))

    await db.commit()
    background_tasks.add_task(
        _finalize_and_preview_bg, report_row.id, document.id, new_storage_path
    )

    return ViolationActionResponse(violation_id=violation.id, status=violation.status)


@router.patch("/{violation_id}/comment", response_model=ViolationCommentResponse)
async def set_violation_comment(
    violation_id: UUID,
    body: ViolationCommentRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> ViolationCommentResponse:
    """Руководитель/декан/админ добавляет или очищает комментарий к нарушению."""
    if current_user.role not in {UserRole.supervisor, UserRole.dean, UserRole.admin}:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Комментировать нарушения может только руководитель, декан или админ",
        )
    try:
        violation = await get_violation_or_404(
            violation_id,
            current_user,
            db,
            allow_supervisor=True,
        )
    except (NotFoundError, AccessDeniedError) as exc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=str(exc))

    violation.supervisor_comment = (body.comment or "").strip() or None
    await db.commit()
    return ViolationCommentResponse(violation_id=violation.id, supervisor_comment=violation.supervisor_comment)
